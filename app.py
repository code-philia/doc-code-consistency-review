import os
import time
from flask import Flask, json, render_template, request, jsonify, send_file
import sqlite3
import json as pyjson
import socket
from utils import get_all_files_with_relative_paths, parse_markdown, split_code, count_lines_of_code, convert_doc_to_markdown, get_filename_without_extension,\
    replace_text_in_docx, generate_issue_content, include_related_blocks
from agent import query_generated_requirement, query_related_code, query_review_result, query_flow_chart, query_related_requirement
from rag_chroma import rag_engine
from doc_block import chunk_markdown
import random
import string
from datetime import datetime, timedelta
from werkzeug.utils import secure_filename
import uuid
from docx import Document
import io
import shutil
import re
import zipfile

from code_block import get_all_code_blocks

import logging
import sys
from utils import parse_programming_rules, parse_issue_reports, format_rules_for_rag, format_issues_for_rag, read_docx_text
from agent import smart_parse_doc
# 配置日志
logging.basicConfig(
    level = logging.INFO,
    format='%(message)s'
    )

logger = logging.getLogger(__name__)

# 定义全局历史文件路径
HISTORY_FILE = 'history.json'
MAX_HISTORY_ITEMS = 15 # 最多记录15条历史

# 定义testdata目录路径
TESTDATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'testdata')

app = Flask(__name__)

# templates
@app.route('/')
def index():
    """Render the welcome page"""
    return render_template('welcome.html')

@app.route('/welcome')
def welcome():
    """Render the welcome page"""
    return render_template('welcome.html')

@app.route('/semi-automatic')
def semi_automatic():
    """Render the semi-automatic mode page"""
    return render_template('semi-automatic.html')

@app.route('/project')
def project():
    """Render the project page"""
    return render_template('project.html')

@app.route('/annotation')
def annotation():
    """Render the annotation page"""
    return render_template('annotation.html')

@app.route('/templates/flowchart-viewer.html')
def flowchart_viewer_template():
    """Serve the flowchart viewer template"""
    return send_file('templates/flowchart-viewer.html', mimetype='text/html')

# project
@app.route('/project/create', methods=['POST'])
def create_project():
    data = request.json
    creation_type = data.get('creationType', 'blank')
    project_name = data.get('projectName')
    project_location = data.get('projectLocation')

    if not project_name or not project_location:
        return jsonify({"status": "error", "message": "项目名称和路径不能为空。"}), 400

    if creation_type == 'blank':
        return create_blank_project(project_name, project_location)
    elif creation_type == 'folder':
        return create_project_from_folder(project_name, project_location)
    else:
        return jsonify({"status": "error", "message": "无效的创建类型。"}), 400


def create_blank_project(project_name, project_location):
    """处理创建空白项目的逻辑"""
    project_path = os.path.join(project_location, project_name)
    if os.path.exists(project_path):
        return jsonify({"status": "error", "message": f"项目文件夹 '{project_name}' 已存在于目标位置。"}), 400

    try:
        code_repo_path = os.path.join(project_path, 'code_repo')
        doc_repo_path = os.path.join(project_path, 'doc_repo')
        os.makedirs(code_repo_path, exist_ok=True)
        os.makedirs(doc_repo_path, exist_ok=True)

        metadata = {
            "project_name": project_name,
            "project_location": project_location,
            "create_time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()),
            "code_repo": code_repo_path,
            "doc_repo": doc_repo_path,
            "code_files": [],
            "doc_files": [],
            "code_scale": 0,
            "code_file_lines": {}
        }
        
        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)
        
        update_history(project_name, project_path)
        init_project_db(project_path)
        auto_load_rag_db(project_path)
        
        return jsonify({"status": "success", "project_path": project_path}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"创建目录或文件时出错: {e}"}), 500


def create_project_from_folder(project_name, folder_path):
    """处理从现有文件夹创建项目的逻辑"""
    project_path = folder_path # 项目路径就是用户选择的文件夹
    if not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "提供的路径不是一个有效的文件夹。"}), 400

    code_repo_path = os.path.join(project_path, 'code_repo')
    doc_repo_path = os.path.join(project_path, 'doc_repo')

    if not os.path.isdir(code_repo_path) or not os.path.isdir(doc_repo_path):
        return jsonify({"status": "error", "message": "文件夹结构不符合要求，必须包含 'code_repo' 和 'doc_repo' 子目录。"}), 400

    if os.path.exists(os.path.join(project_path, 'metadata.json')):
        return jsonify({"status": "error", "message": "该文件夹已包含 'metadata.json'，似乎已是一个项目。"}), 400

    try:
        code_files = get_all_files_with_relative_paths(code_repo_path, type ='code')
        doc_files = get_all_files_with_relative_paths(doc_repo_path, type ='doc')
        
        # 检查 doc_repo 目录下是否有 docx 文件，如果有则进行格式转换
        has_docx = False
        for root, dirs, files in os.walk(doc_repo_path):
            for file in files:
                if file.endswith('.docx'):
                    has_docx = True
                    break
            if has_docx:
                break
        
        if has_docx:
            convert_doc_to_markdown(doc_repo_path)
            # 转换后重新获取文档文件列表
            doc_files = get_all_files_with_relative_paths(doc_repo_path, type ='doc')
        
        code_file_lines = {}
        total_loc = 0
        for file in code_files:
            loc = count_lines_of_code(os.path.join(code_repo_path, file))
            code_file_lines[file] = loc
            total_loc += loc

        metadata = {
            "project_name": project_name,
            "project_location": os.path.dirname(project_path), # 存储其父目录
            "create_time": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()),
            "code_repo": code_repo_path,
            "doc_repo": doc_repo_path,
            "code_files": code_files,
            "doc_files": doc_files,
            "code_scale": total_loc,
            "code_file_lines": code_file_lines
        }
        
        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        update_history(project_name, project_path)
        init_project_db(project_path)
        auto_load_rag_db(project_path)
        
        return jsonify({"status": "success", "project_path": project_path}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"扫描文件夹或生成元数据时出错: {e}"}), 500
    

def update_history(project_name, project_path):
    """读取、更新并写回项目历史记录"""
    history = []
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            try:
                history = json.load(f)
            except json.JSONDecodeError:
                history = [] # 如果文件内容损坏，则重置

    # 检查项目是否已在历史中，如果在则移除旧条目
    history = [item for item in history if item.get('path') != project_path]

    # 添加新条目到列表顶部
    new_entry = {
        "name": project_name,
        "path": project_path,
        "last_opened": datetime.now().isoformat() # 使用ISO 8601格式的时间戳
    }
    history.insert(0, new_entry)

    # 限制历史记录的长度
    history = history[:MAX_HISTORY_ITEMS]

    # 写回文件
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, indent=4, ensure_ascii=False)


def get_db_path(project_path):
    return os.path.join(project_path, 'project.db')


def get_db_conn(project_path):
    db_path = get_db_path(project_path)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute('PRAGMA foreign_keys=ON')
        conn.execute('PRAGMA journal_mode=WAL')
    except Exception:
        pass
    return conn


def init_project_db(project_path):
    db_path = get_db_path(project_path)
    os.makedirs(project_path, exist_ok=True)
    conn = sqlite3.connect(db_path)
    try:
        conn.execute('PRAGMA foreign_keys=ON')
        conn.execute('PRAGMA journal_mode=WAL')
        conn.execute(
            'CREATE TABLE IF NOT EXISTS alignments ('
            'id TEXT PRIMARY KEY,'
            'name TEXT,'
            'isReviewed INTEGER DEFAULT 0,'
            'reviewThoughts TEXT,'
            'docRanges TEXT NOT NULL DEFAULT "[]",'
            'codeRanges TEXT NOT NULL DEFAULT "[]",'
            'createdAt TEXT DEFAULT CURRENT_TIMESTAMP,'
            'updatedAt TEXT)'
        )
        conn.execute(
            'CREATE TABLE IF NOT EXISTS issues ('
            'id INTEGER PRIMARY KEY AUTOINCREMENT,'
            'displayId TEXT UNIQUE,'
            'alignmentId TEXT NOT NULL,'
            'severity TEXT,'
            'title TEXT,'
            'content TEXT,'
            'status TEXT,'
            'relatedDocFile TEXT,'
            'relatedRequirementId TEXT,'
            'briefRequirement TEXT,'
            'briefCode TEXT,'
            'createdAt TEXT DEFAULT CURRENT_TIMESTAMP,'
            'updatedAt TEXT,'
            'FOREIGN KEY(alignmentId) REFERENCES alignments(id) ON DELETE CASCADE)'
        )
        conn.execute('CREATE INDEX IF NOT EXISTS idx_issues_alignmentId ON issues(alignmentId)')
    finally:
        conn.close()


def auto_load_rag_db(project_path):
    """自动加载项目下的第一个RAG知识库"""
    rag_db_root = os.path.join(project_path, "rag_database")
    if os.path.exists(rag_db_root) and os.path.isdir(rag_db_root):
        # 获取所有子目录
        subdirs = [d for d in os.listdir(rag_db_root) if os.path.isdir(os.path.join(rag_db_root, d))]
        if subdirs:
            # 默认加载第一个（按字母序）
            first_db = sorted(subdirs)[0]
            try:
                print(f"[AutoLoad] Found RAG DBs: {subdirs}, loading: {first_db}")
                rag_engine.initialize(project_path, db_name=first_db)
            except Exception as e:
                print(f"[AutoLoad] Failed to load RAG DB '{first_db}': {e}")


def import_json_to_db(project_path):
    conn = get_db_conn(project_path)
    try:
        cur = conn.cursor()
        cur.execute('SELECT COUNT(1) AS c FROM alignments')
        row = cur.fetchone()
        need_import_alignments = (row['c'] == 0)
        results_dir = os.path.join(project_path, 'results')
        if need_import_alignments and os.path.isdir(results_dir):
            for filename in os.listdir(results_dir):
                if filename.endswith('.json'):
                    fp = os.path.join(results_dir, filename)
                    try:
                        with open(fp, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                        if isinstance(data, dict):
                            for aid, alignment in data.items():
                                try:
                                    cur.execute(
                                        'INSERT INTO alignments(id,name,isReviewed,reviewThoughts,docRanges,codeRanges,createdAt,updatedAt)'
                                        ' VALUES (?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)'
                                        ' ON CONFLICT(id) DO UPDATE SET name=excluded.name,isReviewed=excluded.isReviewed,reviewThoughts=excluded.reviewThoughts,docRanges=excluded.docRanges,codeRanges=excluded.codeRanges,updatedAt=CURRENT_TIMESTAMP',
                                        (
                                            alignment.get('id') or aid,
                                            alignment.get('name'),
                                            1 if alignment.get('isReviewed') else 0,
                                            alignment.get('reviewThoughts') or '',
                                            pyjson.dumps(alignment.get('docRanges') or []),
                                            pyjson.dumps(alignment.get('codeRanges') or [])
                                        )
                                    )
                                except Exception:
                                    continue
                        conn.commit()
                    except Exception:
                        continue
        cur.execute('SELECT COUNT(1) AS c FROM issues')
        total_issues = cur.fetchone()['c']
        need_import_issues = (total_issues == 0)
        try:
            cur.execute("SELECT COUNT(1) AS c FROM issues WHERE IFNULL(relatedDocFile,'')='' OR IFNULL(briefRequirement,'')='' OR IFNULL(briefCode,'')='' ")
            missing_fields = cur.fetchone()['c']
            if missing_fields and missing_fields > 0:
                need_import_issues = True
        except Exception:
            pass
        issues_file = os.path.join(project_path, 'issues.json')
        if need_import_issues and os.path.exists(issues_file):
            try:
                with open(issues_file, 'r', encoding='utf-8') as f:
                    issues = json.load(f)
                for issue in issues or []:
                    try:
                        cur.execute(
                            'INSERT INTO issues(displayId,alignmentId,severity,title,content,status,relatedDocFile,relatedRequirementId,briefRequirement,briefCode,createdAt,updatedAt) '
                            'VALUES (?,?,?,?,?,?,?,?,?,?,?,?) '
                            'ON CONFLICT(displayId) DO UPDATE SET '
                            'alignmentId=excluded.alignmentId, '
                            'severity=excluded.severity, '
                            'title=excluded.title, '
                            'content=excluded.content, '
                            'status=excluded.status, '
                            'relatedDocFile=excluded.relatedDocFile, '
                            'relatedRequirementId=excluded.relatedRequirementId, '
                            'briefRequirement=excluded.briefRequirement, '
                            'briefCode=excluded.briefCode, '
                            'createdAt=excluded.createdAt, '
                            'updatedAt=excluded.updatedAt',
                            (
                                issue['id'],
                                issue['alignmentId'],
                                issue['level'],
                                issue['summary'],
                                issue['description'],
                                issue['status'] or 'unconfirmed',
                                issue['relatedDocFile'],
                                issue['relatedRequirementId'],
                                issue['briefRequirement'],
                                issue['briefCode'],
                                issue['createdDate'] or datetime.now().isoformat(),
                                issue['updatedDate'] or datetime.now().isoformat()
                            )
                        )
                    except Exception:
                        continue
                conn.commit()
            except Exception:
                pass
    finally:
        conn.close()


@app.route('/project/history', methods=['GET'])
def get_project_history():
    """获取最近打开的项目列表"""
    if not os.path.exists(HISTORY_FILE):
        return jsonify([])
    
    with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
        try:
            history = json.load(f)
            return jsonify(history)
        except json.JSONDecodeError:
            return jsonify([])

@app.route('/project/history', methods=['DELETE'])
def delete_project_history():
    """删除指定的历史记录"""
    data = request.json
    project_path = data.get('path')
    
    if not project_path:
        return jsonify({"status": "error", "message": "项目路径不能为空"}), 400
    
    if not os.path.exists(HISTORY_FILE):
        return jsonify({"status": "error", "message": "历史记录文件不存在"}), 404
    
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            history = json.load(f)
        
        # 查找并删除指定的历史记录
        original_length = len(history)
        history = [item for item in history if item.get('path') != project_path]
        
        if len(history) == original_length:
            return jsonify({"status": "error", "message": "未找到指定的历史记录"}), 404
        
        # 写回文件
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, indent=4, ensure_ascii=False)
        
        return jsonify({"status": "success", "message": "历史记录删除成功"})
    
    except json.JSONDecodeError:
        return jsonify({"status": "error", "message": "历史记录文件格式错误"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": f"删除历史记录时出错: {e}"}), 500

@app.route('/project/delete', methods=['POST'])
def delete_project():
    """删除项目目录和历史记录"""
    data = request.json
    project_path = data.get('path')
    
    if not project_path:
        return jsonify({"status": "error", "message": "项目路径不能为空"}), 400
    
    if not os.path.exists(project_path):
        return jsonify({"status": "error", "message": "项目路径不存在"}), 404
    
    try:
        # 删除项目目录
        shutil.rmtree(project_path)
        
        # 从历史记录中删除项目条目
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                history = json.load(f)
            
            # 过滤掉要删除的项目
            history = [item for item in history if item.get('path') != project_path]
            
            # 写回历史记录文件
            with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
                json.dump(history, f, indent=4, ensure_ascii=False)
        
        return jsonify({"status": "success", "message": "项目删除成功"})
        
    except PermissionError:
        return jsonify({"status": "error", "message": "没有权限删除项目文件"}), 403
    except Exception as e:
        return jsonify({"status": "error", "message": f"删除项目时出错: {str(e)}"}), 500

@app.route('/project/open', methods=['POST'])
def open_project():
    """当用户打开一个项目时，更新其历史记录"""
    data = request.json
    project_name = data.get('name') or data.get('projectName')
    project_path = data.get('path') or data.get('projectPath')
    if not project_name or not project_path:
        return jsonify({"status": "error", "message": "项目信息不完整"}), 400
    
    # 可以在此添加校验，确保项目路径真实存在
    if not os.path.exists(project_path):
         return jsonify({"status": "error", "message": "项目路径不存在，可能已被移动或删除"}), 404

    update_history(project_name, project_path)
    try:
        init_project_db(project_path)
        import_json_to_db(project_path)
        auto_load_rag_db(project_path)
    except Exception:
        pass
    return jsonify({"status": "success"})


@app.route('/project/import', methods=['POST'])
def import_project():
    """验证一个现有项目文件夹的结构是否有效"""
    data = request.json
    project_path = data.get('path')

    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": f"提供的路径不是一个有效的文件夹。"}), 400
    
    # 检查必需的文件和文件夹
    code_repo_path = os.path.join(project_path, 'code_repo')
    doc_repo_path = os.path.join(project_path, 'doc_repo')
    metadata_file = os.path.join(project_path, 'metadata.json')

    if not os.path.isdir(code_repo_path):
        return jsonify({"status": "error", "message": "文件夹内缺少 'code_repo' 子目录。"}), 400
    if not os.path.isdir(doc_repo_path):
        return jsonify({"status": "error", "message": "文件夹内缺少 'doc_repo' 子目录。"}), 400
    if not os.path.isfile(metadata_file):
        return jsonify({"status": "error", "message": "文件夹内缺少 'metadata.json' 文件。"}), 400
    
    # 读取 metadata.json 以获取项目名称
    try:
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        project_name = metadata.get('project_name')
        if not project_name:
            return jsonify({"status": "error", "message": "'metadata.json' 文件中缺少 'project_name' 字段。"}), 400
        
        update_history(project_name, project_path)

        # 如为原始版本（不存在 project.db），初始化数据库并迁移旧数据
        try:
            db_path = get_db_path(project_path)
            if not os.path.exists(db_path):
                init_project_db(project_path)
                import_json_to_db(project_path)
        except Exception:
            pass

        # 验证成功，返回项目信息
        project_data = {
            "name": project_name,
            "path": project_path
        }
        return jsonify({"status": "success", "project": project_data})

    except (json.JSONDecodeError, Exception) as e:
        return jsonify({"status": "error", "message": f"读取 'metadata.json' 文件失败: {e}"}), 500

@app.route('/project/metadata', methods=['GET'])
def get_project_metadata():
    """根据项目路径获取元数据"""
    project_path = request.args.get('path')
    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "无效的项目路径。"}), 400

    metadata_file = os.path.join(project_path, 'metadata.json')
    if not os.path.isfile(metadata_file):
        return jsonify({"status": "error", "message": "项目元数据文件不存在。"}), 404

    try:
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        return jsonify({"status": "success", "metadata": metadata}), 200

    except (json.JSONDecodeError, Exception) as e:
        return jsonify({"status": "error", "message": f"读取元数据文件失败: {e}"}), 500

@app.route('/project/upload-files', methods=['POST'])
def upload_files():
    try:
        project_path = request.form.get('path')
        file_type = request.form.get('fileType')  # 'doc' or 'code'
        files = request.files.getlist('files')

        if not all([project_path, file_type, files]):
            return jsonify({"status": "error", "message": "请求参数不完整。"}), 400

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        if file_type == 'code':
            code_repo_path = metadata.get('code_repo')
            for file in files:
                # 保留包含中文的原始相对路径
                relative_path = file.filename.replace('\\', '/')

                # 安全检查：防止目录遍历攻击 (e.g., ../../secret.txt)
                # 1. 路径拼接后进行规范化
                dest_path = os.path.abspath(os.path.join(code_repo_path, relative_path))
                # 2. 确保目标路径仍然在 code_repo 目录内
                if not dest_path.startswith(os.path.abspath(code_repo_path)):
                    return jsonify({"status": "error", "message": f"检测到不安全的路径: {relative_path}"}), 400

                # 创建目标目录并保存文件
                dest_dir = os.path.dirname(dest_path)
                os.makedirs(dest_dir, exist_ok=True)
                file.save(dest_path)

            # 更新元数据
            metadata['code_files'] = get_all_files_with_relative_paths(code_repo_path, type='code')
            code_file_lines = {}
            total_loc = 0
            for f in metadata['code_files']:
                loc = count_lines_of_code(os.path.join(code_repo_path, f))
                code_file_lines[f] = loc
                total_loc += loc
            metadata['code_scale'] = total_loc
            metadata['code_file_lines'] = code_file_lines

        elif file_type == 'annotation':
            # 1. 既然 project_path 是根目录，我们在这里手动拼接 annotations
            target_dir = os.path.join(project_path, 'annotations')
            
            # 2. 如果文件夹不存在，自动创建
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)

            # 3. 保存文件
            for file in files:
                # 只取文件名，防止路径包含多余信息
                filename = os.path.basename(file.filename)
                save_dest = os.path.join(target_dir, filename)
                file.save(save_dest)
                print(f"标注文件已保存至: {save_dest}")

        elif file_type == 'doc':
            doc_repo_path = metadata.get('doc_repo')
            has_docx = False
            for file in files:
                # 直接使用原始文件名，仅取最后的文件名部分，天然防止了目录遍历
                filename = os.path.basename(file.filename)
                if filename.endswith(('.md', '.docx')):
                    file.save(os.path.join(doc_repo_path, filename))
                    if filename.endswith('.docx'):
                        has_docx = True
            
            if has_docx:
                convert_doc_to_markdown(doc_repo_path)
            
            metadata['doc_files'] = get_all_files_with_relative_paths(doc_repo_path, type='doc')

        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        return jsonify({"status": "success"}), 200

    except Exception as e:
        print(f"Error during file upload: {e}")
        return jsonify({"status": "error", "message": f"服务器处理文件上传时出错: {e}"}), 500


@app.route('/project/upload-folder', methods=['POST'])
def upload_folder():
    """处理文件夹上传功能"""
    try:
        # 获取上传的文件和文件夹名称
        files = request.files.getlist('files')
        paths = request.form.getlist('paths')
        folder_name = request.form.get('folderName')
        
        if not files or not folder_name:
            return jsonify({"status": "error", "message": "没有接收到文件或文件夹名称"}), 400
        
        # 确保testdata目录存在
        os.makedirs(TESTDATA_DIR, exist_ok=True)
        
        # 创建目标文件夹路径
        target_folder_path = os.path.join(TESTDATA_DIR, folder_name)
        
        # 如果目标文件夹已存在，添加时间戳后缀
        if os.path.exists(target_folder_path):
            timestamp = int(time.time())
            target_folder_path = os.path.join(TESTDATA_DIR, f"{folder_name}_{timestamp}")
        
        # 创建目标文件夹
        os.makedirs(target_folder_path, exist_ok=True)
        
        # 保存所有文件，保持目录结构
        for file, relative_path in zip(files, paths):
            if not file.filename:
                continue
                
            # 移除文件夹名称前缀，获取相对路径
            if relative_path.startswith(folder_name + '/'):
                file_relative_path = relative_path[len(folder_name) + 1:]
            else:
                file_relative_path = relative_path
            
            # 构建完整的目标文件路径
            target_file_path = os.path.join(target_folder_path, file_relative_path)
            
            # 安全检查：确保目标路径在目标文件夹内
            target_file_path = os.path.abspath(target_file_path)
            if not target_file_path.startswith(os.path.abspath(target_folder_path)):
                return jsonify({"status": "error", "message": f"检测到不安全的路径: {relative_path}"}), 400
            
            # 创建目标目录
            target_dir = os.path.dirname(target_file_path)
            os.makedirs(target_dir, exist_ok=True)
            
            # 保存文件
            file.save(target_file_path)
        
        return jsonify({
            "status": "success", 
            "message": f"文件夹 '{folder_name}' 上传成功",
            "serverPath": target_folder_path,
            "folderName": os.path.basename(target_folder_path)
        }), 200
        
    except Exception as e:
        print(f"Error during folder upload: {e}")
        return jsonify({"status": "error", "message": f"文件夹上传失败: {str(e)}"}), 500


@app.route('/project/file-content', methods=['GET'])
def get_file_content():
    """根据项目路径、文件名和文件类型获取文件内容"""
    project_path = request.args.get('path')
    filename = request.args.get('filename')
    file_type = request.args.get('type') # 'doc' or 'code'

    if not project_path or not filename or not file_type:
        return jsonify({"status": "error", "message": "缺少必要的参数"}), 400

    repo_map = {
        'doc': 'doc_repo',
        'code': 'code_repo'
    }

    if file_type not in repo_map:
        return jsonify({"status": "error", "message": "无效的文件类型"}), 400

    try:
        # 获取项目元数据以确定文件仓库路径
        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        if file_type == 'code':
            repo_path = metadata.get(repo_map[file_type])
            file_path = os.path.join(repo_path, filename)
        else: # 'doc'
            if filename.endswith('.md'):
                repo_path = metadata.get(repo_map[file_type])
                file_path = os.path.join(repo_path, filename)
            else: # docx类型，读取转换后的md文件
                file_name_prefix = filename.split('.')[0]
                file_path = os.path.join(project_path, 'doc_repo_converted', file_name_prefix, file_name_prefix + '.md')
            
        if not os.path.exists(file_path):
            return jsonify({"status": "error", "message": "文件未找到"}), 404
        
        # 读取文件内容
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return jsonify({"status": "success", "content": content}), 200
    
    except Exception as e:
        return jsonify({"status": "error", "message": f"读取文件内容时出错: {e}"}), 500



@app.route('/api/requirement-decomposition', methods=['POST'])
def requirement_decomposition():
    """处理需求分解请求 - 保存到JSONL文件"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        
        if not project_path:
            return jsonify({'status':'error', 'message': '缺少项目路径'})

        annotations_dir = os.path.join(project_path, 'annotations')
        if not os.path.exists(annotations_dir):
            return jsonify({'status':'error', 'message': '标注目录不存在'})

        json_files = [f for f in os.listdir(annotations_dir) if f.endswith('.json')]
        if not json_files:
            return jsonify({'status':'error', 'message': '标注结果文件不存在'})

        annotation_file = os.path.join(annotations_dir, json_files[0])
        with open(annotation_file, 'r', encoding='utf-8') as f:
            annotation_data = json.load(f)
        
        # 解析docFiles建立映射
        doc_files = annotation_data.get('docFiles', [])
        doc_id_to_name = {doc['id']: doc['name'] for doc in doc_files}
        
        # 解析annotations构建需求点
        annotations = annotation_data.get('annotations', [])
        processed_count = 0
        
        req_blocks = []
        
        for annotation in annotations:
            doc_ranges = annotation.get('docRanges', [])
            
            if not doc_ranges:
                doc_ranges = annotation.get('documentRanges', [])
                if not doc_ranges:
                    continue
                
            # 获取文档ID
            document_id = doc_ranges[0].get('documentId')
            doc_name = doc_id_to_name.get(document_id)
            
            if not doc_name:
                continue
                
            # 构建需求块对象 (扁平化，每个docRange作为一个独立的块)
            for doc_range in doc_ranges:
                req_block = {
                    'filename': doc_name,
                    'documentId': doc_name,
                    'content': doc_range.get('content'),
                    'start': doc_range.get('start'),
                    'end': doc_range.get('end')
                }
                req_blocks.append(req_block)
                processed_count += 1
        
        # 保存到 doc_block_repo/doc_blocks.jsonl
        doc_block_repo = os.path.join(project_path, 'doc_block_repo')
        os.makedirs(doc_block_repo, exist_ok=True)
        doc_blocks_file = os.path.join(doc_block_repo, 'doc_blocks.jsonl')
        
        # 如果文件存在，先删除
        if os.path.exists(doc_blocks_file):
            os.remove(doc_blocks_file)
        
        with open(doc_blocks_file, 'w', encoding='utf-8') as f:
            for block in req_blocks:
                f.write(json.dumps(block, ensure_ascii=False) + '\n')
        
        return jsonify({
            'status': 'success',
            'message': '需求分解完成，结果已保存',
            'processedCount': processed_count
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/auto-markdown-split', methods=['POST'])
def auto_markdown_split():
    """处理自动Markdown分解请求 - 保存到JSONL文件"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        
        if not project_path:
            return jsonify({'status':'error', 'message': '缺少项目路径'})

        # 获取项目中的文档文件
        doc_repo_path = os.path.join(project_path, 'doc_repo')
        if not os.path.exists(doc_repo_path):
            return jsonify({'status':'error', 'message': '文档目录不存在'})

        # 查找所有markdown文件
        md_files = []
        for root, dirs, files in os.walk(doc_repo_path):
            for file in files:
                if file.lower().endswith('.md'):
                    md_files.append(os.path.join(root, file))
        
        if not md_files:
            return jsonify({'status':'error', 'message': '未找到Markdown文档'})

        processed_count = 0
        req_blocks = []
        
        # 处理每个markdown文件
        for md_file in md_files:
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 分解markdown内容
            doc_name = os.path.basename(md_file)
            blocks = chunk_markdown(doc_name, md_content)
            
            if not blocks:
                continue
            
            # 构建需求块
            for block_info in blocks:
                req_block = {
                    'filename': doc_name,
                    'documentId': doc_name,
                    'content': block_info['content'],
                    'start': block_info['start'],
                    'end': block_info['end']
                }
                req_blocks.append(req_block)
                processed_count += 1
        
        # 保存到 doc_block_repo/doc_blocks.jsonl
        doc_block_repo = os.path.join(project_path, 'doc_block_repo')
        os.makedirs(doc_block_repo, exist_ok=True)
        doc_blocks_file = os.path.join(doc_block_repo, 'doc_blocks.jsonl')
        
        # 如果文件存在，先删除
        if os.path.exists(doc_blocks_file):
            os.remove(doc_blocks_file)
        
        with open(doc_blocks_file, 'w', encoding='utf-8') as f:
            for block in req_blocks:
                f.write(json.dumps(block, ensure_ascii=False) + '\n')
        
        return jsonify({
            'status': 'success',
            'message': '自动分解完成，结果已保存',
            'processedCount': processed_count
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/project/export-issues-download', methods=['POST'])
def export_issues_download():
    """导出所有问题单到一个docx文件"""
    try:
        data = request.json
        issues = data.get('issues', [])
        form_data = data.get('formData', {})
        project_path = data.get('projectPath', '')
        
        template_path = os.path.join(os.path.dirname(__file__), 'templates', '问题单模板.docx')
        
        if not issues:
            return jsonify({'status': 'error', 'message': '没有问题单可导出'})
        
        # 创建临时目录存储文件
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir, exist_ok=True)
        
        # 生成docx文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"问题单导出_{form_data.get('issueId', 'BBB')}_{timestamp}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)
        
        # 检查是否提供了DOCX模板路径
        if template_path and os.path.exists(template_path):
            current_date = datetime.now().strftime("%Y%m%d")
            issue_categories = form_data.get('issueCategories', [])
            
            # 将英文级别转换为中文的映射
            level_mapping = {
                'high': '重大',
                'medium': '严重', 
                'low': '一般'
            }
            
            # 处理第一个问题单作为基础文档
            first_issue = issues[0]
            merged_doc = Document(template_path)
            
            replacements = {}
            # 替换页码信息
            replacements["CURRENT"] = "1"
            replacements["TOTAL"] = str(len(issues))
            
            replacements["AAAAA软件"] = form_data.get('productName', '')
            replacements["BBBBB"] = f"{form_data.get('issueId', '')}_1"
            replacements["CCCCC"] = form_data.get('productId', '')
            replacements["DDDDD"] = form_data.get('discoveryMethod', '')
            replacements["EEEEE"] = form_data.get('issueTracking', '')
            replacements["GGGGG"] = current_date
            
            # 处理问题类别
            for category in ['设计', '编码', '测试', '文档', '数据', '其他']:
                if category in issue_categories:
                    replacements[f"□{category}"] = f"■{category}"
            
            # 处理问题级别
            issue_level = first_issue.get('level', '')
            chinese_level = level_mapping.get(issue_level.lower(), issue_level)
            
            for level in ['重大', '严重', '一般']:
                if level == chinese_level:
                    replacements[f"□{level}"] = f"■{level}"
            
            replacements["CONTENTCONTENT"] = first_issue.get('description', '')
            
            # 替换第一个文档的占位符
            replace_text_in_docx(merged_doc, replacements)
            
            # 处理剩余的问题单
            for i, issue in enumerate(issues[1:], 2):
                # 添加分页符
                merged_doc.add_page_break()
                
                # 为每个问题单加载新的模板并填充
                temp_doc = Document(template_path)
                
                replacements = {}
                # 替换页码信息
                replacements["CURRENT"] = str(i)
                replacements["TOTAL"] = str(len(issues))
                
                replacements["AAAAA软件"] = form_data.get('productName', '')
                replacements["BBBBB"] = f"{form_data.get('issueId', '')}_{i}"
                replacements["CCCCC"] = form_data.get('productId', '')
                replacements["DDDDD"] = form_data.get('discoveryMethod', '')
                replacements["EEEEE"] = form_data.get('issueTracking', '')
                replacements["GGGGG"] = current_date
                
                # 处理问题类别
                for category in ['设计', '编码', '测试', '文档', '数据', '其他']:
                    if category in issue_categories:
                        replacements[f"□{category}"] = f"■{category}"
                
                # 处理问题级别
                issue_level = issue.get('level', '')
                chinese_level = level_mapping.get(issue_level.lower(), issue_level)
                
                for level in ['重大', '严重', '一般']:
                    if level == chinese_level:
                        replacements[f"□{level}"] = f"■{level}"
                
                replacements["CONTENTCONTENT"] = issue.get('description', '')
                
                # 替换模板中的占位符
                replace_text_in_docx(temp_doc, replacements)
                
                # 直接拼接填充好的页面内容到合并文档
                for element in temp_doc.element.body:
                    merged_doc.element.body.append(element)
            
            # 保存合并后的文档
            merged_doc.save(docx_path)
        else:
            # 使用文本格式导出（备用方案）
            content = ""
            for i, issue in enumerate(issues, 1):
                content += f"问题单 {i}/{len(issues)}\n"
                content += generate_issue_content(issue, form_data)
                content += "\n" + "="*50 + "\n\n"
            
            # 创建一个简单的docx文档
            doc = Document()
            doc.add_paragraph(content)
            doc.save(docx_path)
        
        return jsonify({
            'status': 'success', 
            'message': f'成功生成包含 {len(issues)} 个问题单的docx文件',
            'docxFile': docx_filename
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/project/download-file/<filename>', methods=['GET'])
def download_file(filename):
    """下载临时文件并在下载后删除，支持zip文件"""
    try:
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
        file_path = os.path.join(temp_dir, filename)
        
        if not os.path.exists(file_path):
            return jsonify({'status': 'error', 'message': '文件不存在'}), 404
        
        # 读取文件内容到内存
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        # 立即删除临时文件
        try:
            os.remove(file_path)
            print(f"已删除临时文件: {file_path}")
        except Exception as e:
            print(f"删除临时文件失败: {e}")
        
        # 根据文件类型设置MIME类型
        if filename.endswith('.zip'):
            mimetype = 'application/zip'
        elif filename.endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'text/plain'
        
        # 从内存返回文件
        return send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

#####################################
# 核心流程函数
#####################################
@app.route('/api/align-requirement-to-project', methods=['POST'])
def align_requirement_to_project():
    """
    为单个需求点在项目中查找相关代码并返回codeRanges格式的结果
    """
    data = request.json
    doc_ranges = data.get('docRanges', [])
    project_path = data.get('projectPath', '')
    
    doc_name = doc_ranges[0]['filename']
    random_flag = False

    if '协议' in doc_name:
        random_flag = True

    
    # 拼接所有docRanges的content作为requirement_text
    requirement_text = '\n\n'.join([doc_range.get('content', '') for doc_range in doc_ranges if doc_range.get('content')])
    if not requirement_text or not project_path:
        return jsonify({"status": "error", "message": "缺少需求内容或项目路径参数"}), 400
    
    try:
        # 获取项目中所有代码文件
        code_repo_path = os.path.join(project_path, 'code_repo')
        code_block_base_path = os.path.join(project_path, 'code_block_repo')
        all_files = get_all_files_with_relative_paths(code_repo_path, 'code')

        # 为代码进行分块或读取分块结果
        all_code_blocks = get_all_code_blocks(code_repo_path, all_files, code_block_base_path)
        # 调用对齐函数获取相关代码
        related_code = query_related_code(requirement_text, all_code_blocks, random_flag, block_limit=50)

        # 检查并添加 related_id 对应的代码块
        related_code = include_related_blocks(related_code, all_code_blocks)
        
        # 转换为codeRanges格式
        code_ranges = []
        for code_block in related_code:
            # 获取原始代码内容（不带行号）
            file_path = os.path.join(code_repo_path, code_block['file'])
            if os.path.exists(file_path):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    original_content = f.read()
                    lines = original_content.splitlines(keepends=True)  # 保留换行符
                    
                    # 提取指定行范围的内容
                    start_line = max(1, code_block['range'][0])
                    end_line = min(len(lines), code_block['range'][1])
                    
                    if start_line <= end_line:
                        # 计算字符偏移量
                        char_start = sum(len(line) for line in lines[:start_line-1])
                        char_end = sum(len(line) for line in lines[:end_line])
                        
                        # 提取内容（不保留换行符用于显示）
                        range_content = '\n'.join([line.rstrip('\n\r') for line in lines[start_line-1:end_line]])
                        
                        code_ranges.append({
                            'filename': code_block['file'],
                            'start': char_start,  # 字符偏移量
                            'end': char_end,      # 字符偏移量
                            'content': range_content,
                            'documentId': code_block['file'],
                            'startLine': start_line,
                            'endLine': end_line
                        })
        
        return jsonify({
            "status": "success",
            "codeRanges": code_ranges
        })
        
    except Exception as e:
        return jsonify({
            "status": "error", 
            "message": f"对齐过程中出错: {str(e)}"
        }), 500


@app.route('/api/generate-flowchart', methods=['POST'])
def generate_flowchart():
    try:
        data = request.get_json()
        code_content = data.get('codeContent')
        
        if not code_content:
            return jsonify({"status": "error", "message": "Missing code content"}), 400
        
        mermaid_code = query_flow_chart(code_content)

        return jsonify({
            "status": "success",
            "mermaidCode": mermaid_code
        })

    except Exception as e:
        print(f"Error generating flowchart: {str(e)}")
        return jsonify({"status": "error", "message": f"Failed to generate flowchart: {str(e)}"}), 500

@app.route('/api/generate-reverse-requirement', methods=['POST'])
def generate_reverse_requirement():
    try:
        data = request.get_json()
        requirement_content = data.get('requirementContent')
        code_content = data.get('codeContent')
        
        if not code_content:
            return jsonify({"status": "error", "message": "Missing code content"}), 400
        
        # 构建代码块列表，格式与现有函数兼容
        code_blocks = []
        if isinstance(code_content, list):
            for code_block in code_content:
                code_blocks.append({
                    'filename': code_block.get('filename', 'unknown'),
                    'content': code_block.get('content', '')
                })
        else:
            # 如果是字符串，创建单个代码块
            code_blocks.append({
                'filename': 'code',
                'content': code_content
            })
        
        # 调用LLM生成需求，传入参考需求内容
        generated_requirement = query_generated_requirement(code_blocks, requirement_content or "")
        
        # 调用LLM生成流程图
        mermaid_code = query_flow_chart(code_content if isinstance(code_content, str) else 
                                       '\n\n'.join([block.get('content', '') for block in code_content]))

        return jsonify({
            "status": "success",
            "generatedRequirement": generated_requirement,
            "mermaidCode": mermaid_code
        })

    except Exception as e:
        print(f"Error generating reverse requirement: {str(e)}")
        return jsonify({"status": "error", "message": f"Failed to generate reverse requirement: {str(e)}"}), 500

@app.route('/api/review-alignment', methods=['POST'])
def review_alignment():
    data = request.json
    project_path = data.get('projectPath')
    doc_file = data.get('docFile')
    alignment = data.get('alignment')

    if not all([project_path, doc_file, alignment]):
        return jsonify({"status": "error", "message": "Missing required parameters"}), 400

    init_project_db(project_path)

    # 1. 调用 agent 获取审查结果
    review_process, issue = query_review_result(
        alignment.get('docRanges', []),
        alignment.get('codeRanges', [])
    )

    # 2. 更新对齐关系
    alignment['isReviewed'] = True
    alignment['reviewThoughts'] = review_process

    if isinstance(issue, list):
        issues_list = [x for x in issue if isinstance(x, dict)]
    elif isinstance(issue, dict):
        issues_list = [issue]
    else:
        issues_list = []

    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()

        cur.execute(
            'UPDATE alignments SET isReviewed=1, reviewThoughts=?, updatedAt=CURRENT_TIMESTAMP WHERE id=?',
            (alignment.get('reviewThoughts') or '', alignment.get('id'))
        )

        if issues_list:
            cur.execute("SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%'")
            used = set()
            for r in cur.fetchall():
                disp = r['displayId']
                if disp and disp.startswith('ISSUE-'):
                    try:
                        used.add(int(disp.split('-')[1]))
                    except Exception:
                        pass
            next_number = (max(used) + 1) if used else 1

            brief_req = alignment.get('docRanges', [{}])[0].get('content', '') or ''
            brief_code = alignment.get('codeRanges', [{}])[0].get('content', '') or ''

            for one in issues_list:
                display_id = f"ISSUE-{next_number:03d}"
                next_number += 1

                severity = one.get('level') or one.get('severity')
                title = one.get('summary') or one.get('title') or ''
                content = one.get('description') or one.get('content') or ''
                status = one.get('status') or 'unconfirmed'

                cur.execute(
                    'INSERT INTO issues(displayId,alignmentId,severity,title,content,status,relatedDocFile,relatedRequirementId,briefRequirement,briefCode,createdAt,updatedAt) '
                    'VALUES (?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
                    (
                        display_id,
                        alignment.get('id'),
                        severity,
                        title,
                        content,
                        status,
                        doc_file,
                        alignment.get('id'),
                        brief_req,
                        brief_code
                    )
                )

        conn.commit()
        conn.close()
    except Exception as e:
        try:
            conn.rollback()
            conn.close()
        except Exception:
            pass
        return jsonify({"status": "error", "message": f"Failed to save review result: {str(e)}"}), 500

    return jsonify({"status": "success", "createdIssues": len(issues_list)})


@app.route('/api/clear-alignment-review', methods=['POST'])
def clear_alignment_review():
    """Clear review state for a single alignment and delete its related issues (DB)."""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        alignment_id = data.get('alignmentId')
        if not all([project_path, alignment_id]):
            return jsonify({"status": "error", "message": "Missing required parameters"}), 400
        init_project_db(project_path)
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('UPDATE alignments SET isReviewed=0, reviewThoughts="" WHERE id=?', (alignment_id,))
        cur.execute('DELETE FROM issues WHERE alignmentId=?', (alignment_id,))
        removed = cur.rowcount
        conn.commit()
        conn.close()
        return jsonify({"status": "success", "removedIssues": removed})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})




def get_filename_without_extension(filename):
    """去掉文件名的扩展名"""
    return os.path.splitext(filename)[0]

@app.route('/project/alignments', methods=['GET'])
def get_alignments():
    """按文件筛选获取对齐关系，支持 doc 或 code 文件"""
    project_path = request.args.get('path')
    file_path = request.args.get('file')
    kind = request.args.get('kind', 'doc')
    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径参数。"}), 400

    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        
        if file_path:
            target = file_path
            col = 'docRanges' if kind == 'doc' else 'codeRanges'
            query = f"SELECT id,name,isReviewed,reviewThoughts,docRanges,codeRanges FROM alignments WHERE EXISTS (SELECT 1 FROM json_each({col}) jr WHERE json_extract(jr.value,'$.documentId') = ?)"
            cur.execute(query, (target,))
        else:
            query = "SELECT id,name,isReviewed,reviewThoughts,docRanges,codeRanges FROM alignments"
            cur.execute(query)
            
        rows = cur.fetchall()
        result = {}
        for r in rows:
            alignment = {
                'id': r['id'],
                'name': r['name'],
                'isReviewed': bool(r['isReviewed']),
                'reviewThoughts': r['reviewThoughts'] or '',
                'docRanges': pyjson.loads(r['docRanges'] or '[]'),
                'codeRanges': pyjson.loads(r['codeRanges'] or '[]')
            }
            result[alignment['id']] = alignment
        conn.close()
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"读取对齐数据失败: {e}"}), 500

@app.route('/project/alignments', methods=['POST'])
def add_alignment():
    """添加或更新对齐关系到数据库"""
    project_path = request.args.get('path')
    new_alignment = request.json
    if not project_path or not new_alignment or 'id' not in new_alignment:
        return jsonify({"status": "error", "message": "缺少项目路径或无效的对齐数据。"}), 400

    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute(
            'INSERT INTO alignments(id,name,isReviewed,reviewThoughts,docRanges,codeRanges,createdAt,updatedAt) '
            'VALUES (?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP) '
            'ON CONFLICT(id) DO UPDATE SET name=excluded.name,isReviewed=excluded.isReviewed,reviewThoughts=excluded.reviewThoughts,docRanges=excluded.docRanges,codeRanges=excluded.codeRanges,updatedAt=CURRENT_TIMESTAMP',
            (
                new_alignment.get('id'),
                new_alignment.get('name'),
                1 if new_alignment.get('isReviewed') else 0,
                new_alignment.get('reviewThoughts') or '',
                pyjson.dumps(new_alignment.get('docRanges') or []),
                pyjson.dumps(new_alignment.get('codeRanges') or [])
            )
        )
        conn.commit()
        conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"写入对齐数据失败: {e}"}), 500


@app.route('/project/alignment', methods=['DELETE'])
def delete_alignment():
    """删除一个对齐关系（数据库）"""
    project_path = request.args.get('path')
    alignment_id = request.args.get('id')
    if not all([project_path, alignment_id]):
        return jsonify({"status": "error", "message": "缺少项目路径或对齐ID参数。"}), 400
    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('DELETE FROM alignments WHERE id=?', (alignment_id,))
        conn.commit()
        conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"删除对齐项时出错: {e}"}), 500


@app.route('/project/issues', methods=['GET'])
def get_issues():
    try:
        project_path = request.args.get('path')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('SELECT id,displayId,alignmentId,severity,title,content,status,relatedDocFile,relatedRequirementId,briefRequirement,briefCode,createdAt,updatedAt FROM issues ORDER BY id ASC')
        rows = cur.fetchall()
        issues = []
        for r in rows:
            issues.append({
                'id': str(r['id']),
                'displayId': r['displayId'],
                'alignmentId': r['alignmentId'],
                'level': r['severity'],
                'summary': r['title'] or '',
                'description': r['content'] or '',
                'status': r['status'] or 'unconfirmed',
                'relatedDocFile': r['relatedDocFile'] or '',
                'relatedRequirementId': r['relatedRequirementId'] or '',
                'briefRequirement': r['briefRequirement'] or '',
                'briefCode': r['briefCode'] or '',
                'createdDate': r['createdAt'],
                'updatedDate': r['updatedAt']
            })
        conn.close()
        return jsonify({'status': 'success', 'data': issues})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/project/issues', methods=['POST'])
def add_issue():
    try:
        project_path = request.args.get('path')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        issue_data = request.json or {}
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        # 分配展示ID
        if not issue_data.get('displayId'):
            cur.execute("SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%'")
            used = set()
            for r in cur.fetchall():
                disp = r['displayId']
                if disp and disp.startswith('ISSUE-'):
                    try:
                        used.add(int(disp.split('-')[1]))
                    except Exception:
                        pass
            next_number = (max(used) + 1) if used else 1
            issue_data['displayId'] = f"ISSUE-{next_number:03d}"
        cur.execute(
            'INSERT INTO issues(displayId,alignmentId,severity,title,content,status,createdAt,updatedAt) '
            'VALUES (?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
            (
                issue_data.get('displayId'),
                issue_data.get('alignmentId'),
                issue_data.get('level') or issue_data.get('severity'),
                issue_data.get('title'),
                issue_data.get('description') or issue_data.get('content'),
                issue_data.get('status')
            )
        )
        conn.commit()
        conn.close()
        return jsonify({'status': 'success', 'message': '问题单添加成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/project/issues/<issue_id>', methods=['PUT'])
def update_issue(issue_id):
    try:
        project_path = request.args.get('path')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        issue_data = request.json or {}
        init_project_db(project_path)
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('UPDATE issues SET severity=?, title=?, content=?, status=?, updatedAt=CURRENT_TIMESTAMP WHERE id=?', (
            issue_data.get('level') or issue_data.get('severity'),
            issue_data.get('title'),
            issue_data.get('description') or issue_data.get('content'),
            issue_data.get('status'),
            issue_id
        ))
        conn.commit()
        conn.close()
        return jsonify({'status': 'success', 'message': '问题单更新成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/project/issues/<issue_id>', methods=['DELETE'])
def delete_issue(issue_id):
    try:
        project_path = request.args.get('path')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        init_project_db(project_path)
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('DELETE FROM issues WHERE id=?', (issue_id,))
        conn.commit()
        conn.close()
        return jsonify({'status': 'success', 'message': '问题单删除成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/clear-project-results', methods=['POST'])
def clear_project_results():
    """清空项目的所有结果：需求片段、对齐结果、审查结果、问题单结果"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})
        
        if not os.path.exists(project_path):
            return jsonify({'status': 'error', 'message': '项目路径不存在'})
        
        init_project_db(project_path)
        conn = get_db_conn(project_path)
        conn.execute('DELETE FROM issues')
        conn.execute('DELETE FROM alignments')
        conn.commit()
        conn.close()
        
        return jsonify({
            'status': 'success',
            'message': '所有结果已清空'
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/clear-code-ranges', methods=['POST'])
def clear_code_ranges():
    """清空项目中的所有对齐关系"""
    data = request.json
    project_path = data.get('projectPath')
    
    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径参数"}), 400
    
    try:
        init_project_db(project_path)
        conn = get_db_conn(project_path)
        try:
            cur = conn.cursor()
            cur.execute('DELETE FROM alignments')
            conn.commit()
        finally:
            conn.close()
        
        return jsonify({"status": "success", "message": "已清空项目的对齐关系"})
        
    except Exception as e:
        return jsonify({"status": "error", "message": f"清空对齐关系失败: {str(e)}"}), 500


@app.route('/api/clear-review-results', methods=['POST'])
def clear_review_results():
    """清空项目的审查结果：问题单结果，并重置对齐关系的审查状态"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})
        
        if not os.path.exists(project_path):
            return jsonify({'status': 'error', 'message': '项目路径不存在'})
        
        conn = get_db_conn(project_path)
        try:
            conn.execute('DELETE FROM issues')
            conn.execute('UPDATE alignments SET isReviewed=0, reviewThoughts=""')
            conn.commit()
        finally:
            conn.close()
        
        return jsonify({
            'status': 'success',
            'message': '审查结果已清空'
        })
        
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/project/issue/update', methods=['POST'])
def update_issue_content():
    data = request.json
    project_path = data.get('path')
    issue_id = data.get('issueId')
    new_description = data.get('description')
    new_status = data.get('status')
    new_level = data.get('level')  # 添加问题级别参数

    if not all([project_path, issue_id]):
        return jsonify({"status": "error", "message": "缺少项目路径或问题单ID"}), 400

    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('UPDATE issues SET content=?, status=?, severity=?, updatedAt=CURRENT_TIMESTAMP WHERE id=?', (
            new_description,
            new_status,
            new_level,
            issue_id
        ))
        if cur.rowcount == 0:
            conn.close()
            return jsonify({"status": "error", "message": "未找到指定ID的问题单"}), 404
        conn.commit()
        conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"更新问题单失败: {str(e)}"}), 500
     
@app.route('/api/code-decomposition', methods=['POST'])
def code_decomposition():
    """将代码分块为仅包含代码范围的对齐关系并保存到文件"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        code_repo_path = os.path.join(project_path, 'code_repo')
        if not os.path.exists(code_repo_path):
            return jsonify({'status': 'error', 'message': '代码目录不存在'})

        all_files = get_all_files_with_relative_paths(code_repo_path, 'code')
        code_block_base_path = os.path.join(project_path, 'code_block_repo')
        os.makedirs(code_block_base_path, exist_ok=True)
        
        # 删除现有的 code_blocks.jsonl
        code_blocks_file = os.path.join(code_block_base_path, 'code_blocks.jsonl')
        if os.path.exists(code_blocks_file):
            os.remove(code_blocks_file)
            
        all_code_blocks = get_all_code_blocks(code_repo_path, all_files, code_block_base_path)

        # 仅确保文件生成，不写入数据库
        processed_count = len(all_code_blocks)

        return jsonify({'status': 'success', 'message': '代码分解完成，结果已保存', 'processedCount': processed_count})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/project/alignment-by-id', methods=['GET'])
def get_alignment_by_id():
    project_path = request.args.get('path')
    alignment_id = request.args.get('id')
    if not project_path or not alignment_id:
        return jsonify({'status': 'error', 'message': '缺少项目路径或对齐ID'}), 400
    try:
        conn = get_db_conn(project_path)
        cur = conn.cursor()
        cur.execute('SELECT id,name,isReviewed,reviewThoughts,docRanges,codeRanges FROM alignments WHERE id=?', (alignment_id,))
        r = cur.fetchone()
        conn.close()
        if not r:
            return jsonify({'status': 'error', 'message': '未找到对齐关系'}), 404
        alignment = {
            'id': r['id'],
            'name': r['name'],
            'isReviewed': bool(r['isReviewed']),
            'reviewThoughts': r['reviewThoughts'] or '',
            'docRanges': pyjson.loads(r['docRanges'] or '[]'),
            'codeRanges': pyjson.loads(r['codeRanges'] or '[]')
        }
        return jsonify({'status': 'success', 'data': alignment})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})



@app.route('/api/get-doc-blocks', methods=['GET'])
def get_doc_blocks():
    """获取需求分块列表"""
    try:
        project_path = request.args.get('projectPath')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})
            
        doc_block_base_path = os.path.join(project_path, 'doc_block_repo')
        doc_block_file_path = os.path.join(doc_block_base_path, 'doc_blocks.jsonl')
        
        if not os.path.exists(doc_block_file_path):
             return jsonify({'status': 'success', 'data': []})
             
        doc_blocks = []
        with open(doc_block_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    doc_blocks.append(json.loads(line.strip()))
        
        return jsonify({'status': 'success', 'data': doc_blocks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@app.route('/api/get-code-blocks', methods=['GET'])
def get_code_blocks():
    """获取代码分块列表 (原始格式)"""
    try:
        project_path = request.args.get('projectPath')
        filename = request.args.get('filename') # 可选：按文件名筛选
        
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})
            
        code_block_base_path = os.path.join(project_path, 'code_block_repo')
        code_block_file_path = os.path.join(code_block_base_path, 'code_blocks.jsonl')
        
        if not os.path.exists(code_block_file_path):
             return jsonify({'status': 'success', 'data': []})
             
        code_blocks = []
        with open(code_block_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    try:
                        block = json.loads(line.strip())
                        # 如果提供了文件名，则进行筛选
                        if filename:
                            if block.get('file') == filename:
                                code_blocks.append(block)
                        else:
                            code_blocks.append(block)
                    except json.JSONDecodeError:
                        continue
        
        return jsonify({'status': 'success', 'data': code_blocks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


def _find_first_file_by_basename(base_dir: str, basename: str):
    if not base_dir or not basename or not os.path.isdir(base_dir):
        return None
    direct_path = os.path.join(base_dir, basename)
    if os.path.exists(direct_path):
        return direct_path
    for root, _, files in os.walk(base_dir):
        if basename in files:
            return os.path.join(root, basename)
    return None


def _read_text_file(file_path: str):
    if not file_path or not os.path.exists(file_path):
        return None
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _read_doc_raw_content(project_path: str, filename: str):
    doc_repo_path = os.path.join(project_path, 'doc_repo')
    metadata_file = os.path.join(project_path, 'metadata.json')
    if os.path.exists(metadata_file):
        try:
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = pyjson.load(f)
            if metadata.get('doc_repo'):
                doc_repo_path = metadata.get('doc_repo')
        except Exception:
            pass

    if filename and filename.lower().endswith('.md'):
        file_path = _find_first_file_by_basename(doc_repo_path, os.path.basename(filename))
        return _read_text_file(file_path)

    if filename:
        prefix = filename.split('.')[0]
        converted_md = os.path.join(project_path, 'doc_repo_converted', prefix, prefix + '.md')
        converted_content = _read_text_file(converted_md)
        if converted_content is not None:
            return converted_content

        file_path = _find_first_file_by_basename(doc_repo_path, os.path.basename(filename))
        return _read_text_file(file_path)

    return None


def _offset_to_line_numbers(text: str, start: int, end: int):
    if text is None:
        return 1, 1
    n = len(text)
    try:
        s = max(0, min(int(start), n))
    except Exception:
        s = 0
    try:
        e = max(0, min(int(end), n))
    except Exception:
        e = s
    start_line = text.count('\n', 0, s) + 1
    end_line = text.count('\n', 0, e) + 1
    return start_line, end_line


def _line_range_to_char_offsets(text: str, start_line: int, end_line: int):
    if text is None:
        return 0, 0
    lines = text.splitlines(keepends=True)
    if not lines:
        return 0, 0
    s = max(1, int(start_line or 1))
    e = max(s, int(end_line or s))
    s = min(s, len(lines))
    e = min(e, len(lines))
    char_start = sum(len(line) for line in lines[:s - 1])
    char_end = sum(len(line) for line in lines[:e])
    return char_start, char_end


def _compact_title_from_text(text: str, max_len: int = 24):
    if not text:
        return '未命名'
    t = re.sub(r'\s+', ' ', text).strip()
    t = re.sub(r'[`*_>#-]+', '', t).strip()
    if len(t) <= max_len:
        return t or '未命名'
    return t[:max_len].rstrip() + '…'


@app.route('/api/get-requirement-chunks', methods=['GET'])
def get_requirement_chunks():
    try:
        project_path = request.args.get('projectPath')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400

        doc_block_file_path = os.path.join(project_path, 'doc_block_repo', 'doc_blocks.jsonl')
        if not os.path.exists(doc_block_file_path):
            return jsonify({'status': 'success', 'data': []})

        chunks = []
        with open(doc_block_file_path, 'r', encoding='utf-8') as f:
            for i, line in enumerate(f):
                if not line.strip():
                    continue
                try:
                    block = pyjson.loads(line.strip())
                except Exception:
                    continue

                filename = block.get('filename') or block.get('documentId') or ''
                content = block.get('content') or ''
                start = block.get('start') if block.get('start') is not None else 0
                end = block.get('end') if block.get('end') is not None else 0

                raw_doc = _read_doc_raw_content(project_path, filename)
                start_line, end_line = _offset_to_line_numbers(raw_doc, start, end)

                doc_range = {
                    'documentId': filename,
                    'filename': filename,
                    'start': start,
                    'end': end,
                    'content': content,
                    'startLine': start_line,
                    'endLine': end_line
                }

                chunk_id = f"auto_req_{uuid.uuid4().hex}"
                chunk_name = _compact_title_from_text(content, 24)
                chunks.append({
                    'id': chunk_id,
                    'name': chunk_name,
                    'isReviewed': False,
                    'reviewThoughts': '',
                    'docRanges': [doc_range],
                    'codeRanges': []
                })

        chunks.sort(key=lambda x: ((x.get('docRanges') or [{}])[0].get('filename') or '', (x.get('docRanges') or [{}])[0].get('start') or 0))
        return jsonify({'status': 'success', 'data': chunks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/get-code-chunks', methods=['GET'])
def get_code_chunks():
    try:
        project_path = request.args.get('projectPath')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400

        code_block_file_path = os.path.join(project_path, 'code_block_repo', 'code_blocks.jsonl')
        if not os.path.exists(code_block_file_path):
            return jsonify({'status': 'success', 'data': []})

        code_repo_path = os.path.join(project_path, 'code_repo')
        chunks = []
        with open(code_block_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if not line.strip():
                    continue
                try:
                    block = pyjson.loads(line.strip())
                except Exception:
                    continue

                file_rel = block.get('file') or ''
                rng = block.get('range') or []
                if not (isinstance(rng, list) and len(rng) == 2):
                    continue
                start_line = int(rng[0])
                end_line = int(rng[1])
                content = block.get('code') or block.get('content') or ''

                abs_code_path = os.path.join(code_repo_path, file_rel) if file_rel else None
                raw_code = _read_text_file(abs_code_path)
                char_start, char_end = _line_range_to_char_offsets(raw_code, start_line, end_line)

                code_range = {
                    'documentId': file_rel,
                    'filename': file_rel,
                    'start': char_start,
                    'end': char_end,
                    'content': content,
                    'startLine': start_line,
                    'endLine': end_line
                }

                chunk_id = f"auto_code_{uuid.uuid4().hex}"
                chunk_name = f"{file_rel}:{start_line}-{end_line}" if file_rel else f"代码块:{start_line}-{end_line}"
                chunks.append({
                    'id': chunk_id,
                    'name': chunk_name,
                    'isReviewed': False,
                    'reviewThoughts': '',
                    'docRanges': [],
                    'codeRanges': [code_range]
                })

        chunks.sort(key=lambda x: ((x.get('codeRanges') or [{}])[0].get('filename') or '', (x.get('codeRanges') or [{}])[0].get('startLine') or 0))
        return jsonify({'status': 'success', 'data': chunks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@app.route('/api/align-code-to-requirement', methods=['POST'])
def align_code_to_requirement():
    """为单个代码块在项目中查找相关需求"""
    try:
        data = request.json
        code_ranges = data.get('codeRanges', [])
        project_path = data.get('projectPath', '')
        
        if not code_ranges or not project_path:
             return jsonify({"status": "error", "message": "缺少代码内容或项目路径参数"}), 400

        # 获取需求块
        doc_block_base_path = os.path.join(project_path, 'doc_block_repo')
        doc_block_file_path = os.path.join(doc_block_base_path, 'doc_blocks.jsonl')
        
        if not os.path.exists(doc_block_file_path):
             return jsonify({"status": "success", "docRanges": []}) # 没有需求文件，无法对齐
             
        requirements = []
        with open(doc_block_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    # 提取简化的需求信息用于匹配
                    block = json.loads(line.strip())
                    requirements.append({
                        "filename": block.get('filename'),
                        "content": block.get('content')
                    })
        
        code_content = code_ranges[0].get('content', '')
        
        # 调用LLM
        related_reqs = query_related_requirement(code_content, requirements, random_flag=False, block_limit=50)
        
        # 转换结果为docRanges
        doc_ranges = []
        
        # 加载完整需求信息以获取位置信息
        full_requirements_map = {} # (filename, content_hash) -> full_req
        with open(doc_block_file_path, 'r', encoding='utf-8') as f:
             for line in f:
                if line.strip():
                    block = json.loads(line.strip())
                    # 简单使用 content 前50个字符作为key的一部分，实际应更严谨
                    key = f"{block.get('filename')}_{block.get('content')[:50]}" 
                    full_requirements_map[key] = block

        for item in related_reqs:
            # item: {'filename': ..., 'content': ..., 'similarity': ...}
            # 尝试找回原始位置信息
            key = f"{item.get('filename')}_{item.get('content')[:50]}"
            original_doc_range = full_requirements_map.get(key)
            
            if original_doc_range:
                doc_ranges.append(original_doc_range)
            else:
                # 如果找不到原始信息（不太可能，除非LLM修改了内容），则构造一个基本的
                doc_ranges.append({
                    'filename': item.get('filename'),
                    'documentId': item.get('filename'),
                    'content': item.get('content'),
                    'start': 0, # 未知
                    'end': 0    # 未知
                })

        return jsonify({
            "status": "success",
            "docRanges": doc_ranges
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": f"对齐过程中出错: {str(e)}"}), 500

# 1. 获取标注文件列表
@app.route('/api/files/list-annotations', methods=['GET'])
def list_annotation_files():
    project_path = request.args.get('projectPath')
    
    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径"})
    
    # 拼接 annotations 目录路径
    ann_dir = os.path.join(project_path, 'annotations')
    
    files = []
    # 检查目录是否存在
    if os.path.exists(ann_dir) and os.path.isdir(ann_dir):
        # 遍历目录，只获取 .json 文件
        files = [f for f in os.listdir(ann_dir) if f.endswith('.json')]
        # 可选：按修改时间排序，让最新的排在前面
        # files.sort(key=lambda x: os.path.getmtime(os.path.join(ann_dir, x)), reverse=True)
    
    return jsonify({"status": "success", "files": files})

# 2. 构建知识库 (调用 rag_chroma)
@app.route('/api/rag/build', methods=['POST'])
def build_rag_db():
    data = request.json
    project_path = data.get('projectPath')
    annotation_file = data.get('annotationFile') 
    db_name = data.get('dbName', 'default_rag')
    kb_type = data.get('kbType', 'other') # [新增] 获取知识库类型

    if not project_path or not annotation_file:
        return jsonify({"status": "error", "message": "参数缺失"})

    try:
        rag_engine.initialize(project_path, db_name)
    except Exception as e:
        return jsonify({"status": "error", "message": f"初始化失败: {str(e)}"})

    full_path = os.path.join(project_path, 'annotations', annotation_file)
    if not os.path.exists(full_path):
        return jsonify({"status": "error", "message": f"找不到文件: {annotation_file}"})

    try:
        json_data = None
        
        # === Word 文档处理逻辑 (kbType 分流) ===
        if annotation_file.lower().endswith('.docx'):
            print(f"[RAG] 解析文档: {annotation_file}, 类型: {kb_type}")
            
            # A. 编程规则
            if kb_type == 'rule':
                raw_rules = parse_programming_rules(full_path) # 优先用 utils 解析
                if not raw_rules:
                    print("[RAG] 正则解析为空，尝试 LLM 兜底...")
                    doc_text = read_docx_text(full_path)
                    raw_rules = smart_parse_doc(doc_text, type='rule') # 兜底
                if raw_rules:
                    json_data = format_rules_for_rag(raw_rules)
            
            # B. 问题单
            elif kb_type == 'issue':
                raw_issues = parse_issue_reports(full_path) # 优先用 utils 解析
                if not raw_issues:
                    print("[RAG] 表格解析为空，尝试 LLM 兜底...")
                    doc_text = read_docx_text(full_path)
                    raw_issues = smart_parse_doc(doc_text, type='issue') # 兜底
                if raw_issues:
                    json_data = format_issues_for_rag(raw_issues)
            
            # C. 其他/典型案例 (直接 LLM)
            elif kb_type in ['case', 'other']:
                doc_text = read_docx_text(full_path)
                # 简单复用 rule 提取逻辑作为通用提取
                raw_data = smart_parse_doc(doc_text, type='rule') 
                if raw_data:
                    json_data = format_rules_for_rag(raw_data)

            # 保存解析结果并构建
            if json_data:
                temp_json = full_path + ".parsed.json"
                with open(temp_json, 'w', encoding='utf-8') as f:
                    json.dump(json_data, f, ensure_ascii=False)
                
                result = rag_engine.build_from_json(temp_json)
                try: os.remove(temp_json) 
                except: pass
                return jsonify(result)
            else:
                return jsonify({"status": "error", "message": "文档解析失败，未能提取有效数据"})

        # === 原有 JSON 逻辑 ===
        elif annotation_file.lower().endswith('.json'):
            result = rag_engine.build_from_json(full_path)
            return jsonify(result)
            
        else:
            return jsonify({"status": "error", "message": "不支持的文件格式"})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": str(e)}), 500

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
TEST_DATA_DIR = os.path.join(PROJECT_ROOT, 'testdata')

if not os.path.exists(TEST_DATA_DIR):
    os.makedirs(TEST_DATA_DIR)

# === 路由定义 ===
@app.route('/review')  # 建议路径也改一下，避免冲突
def kb_review():       # <--- 重点：把这里改成 kb_review 或其他名字
    # 获取现有知识库列表
    # 注意：这里需要引入 os 和 PROJECT_ROOT
    kb_root = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rag_database")
    existing_kbs = []
    if os.path.exists(kb_root):
        existing_kbs = [d for d in os.listdir(kb_root) if os.path.isdir(os.path.join(kb_root, d))]
    
    return render_template('index.html', kbs=existing_kbs)

# 在 app.py 中添加/修改以下代码

@app.route('/api/list-testdata', methods=['GET'])
def list_testdata():
    """列出 testdata 目录下的所有文件"""
    if not os.path.exists(TEST_DATA_DIR):
        return jsonify({"status": "success", "files": []})
    
    files = [f for f in os.listdir(TEST_DATA_DIR) if os.path.isfile(os.path.join(TEST_DATA_DIR, f)) and not f.startswith('.')]
    return jsonify({"status": "success", "files": files})

@app.route('/api/list-kbs', methods=['GET'])
def list_kbs():
    """列出已有的知识库"""
    kb_root = os.path.join(PROJECT_ROOT, "rag_database")
    kbs = []
    if os.path.exists(kb_root):
        kbs = [d for d in os.listdir(kb_root) if os.path.isdir(os.path.join(kb_root, d))]
    return jsonify({"status": "success", "kbs": kbs})

@app.route('/preview', methods=['POST'])
def preview_file():
    doc_type = request.form.get('doc_type')
    use_server_file = request.form.get('use_server_file') == 'true'
    
    target_path = ""

    try:
        if use_server_file:
            # 模式 A: 使用服务器上的 testdata 文件
            filename = request.form.get('filename')
            target_path = os.path.join(TEST_DATA_DIR, filename)
            if not os.path.exists(target_path):
                return jsonify({"status": "error", "message": "文件不存在"}), 404
        else:
            # 模式 B: 上传新文件
            if 'file' not in request.files:
                return jsonify({"error": "没有上传文件"}), 400
            file = request.files['file']
            target_path = os.path.join(TEST_DATA_DIR, file.filename)
            file.save(target_path)
        
        # --- 开始解析 ---
        parsed_data = []
        
        if doc_type == 'rule':
            raw_rules = parse_programming_rules(target_path, debug=False)
            for r in raw_rules:
                # 1. 组合用于展示和检索的完整文本
                # 把“描述”、“违背示例”、“遵循示例”拼在一起
                combined_content = f"【规则描述】\n{r.get('description', '')}\n"
                
                if r.get('violation_code', '').strip():
                    combined_content += f"\n【❌ 违背示例】\n{r['violation_code']}"
                
                if r.get('compliance_code', '').strip():
                    combined_content += f"\n【✅ 遵循示例】\n{r['compliance_code']}"

                parsed_data.append({
                    "id": r.get('id', '未知ID'),
                    "summary": r.get('description', '')[:50] + "...",
                    "content": combined_content,  # <--- 现在这里包含了完整代码！
                    "full_data": r,
                    "type": "编程规则"
                })
                
        elif doc_type == 'issue':
            raw_issues = parse_issue_reports(target_path, debug=False)
            for i in raw_issues:
                # 增强：确保字段存在，防止前端显示空白
                desc = i.get('desc', '无描述')
                opinion = i.get('opinion', '无处理意见')
                
                parsed_data.append({
                    "id": i.get('id', '未知ID'),
                    "summary": desc[:50] + "...",
                    # 组合详细内容用于展示
                    "content": f"【问题描述】\n{desc}\n\n【处理意见】\n{opinion}\n\n【追踪ID】\n{i.get('trace_id','')}",
                    "full_data": i,
                    "type": "问题单"
                })
        
        return jsonify({"status": "success", "data": parsed_data})
        
    except Exception as e:
        print(f"Preview Error: {e}") # 打印报错以便调试
        return jsonify({"status": "error", "message": str(e)})

@app.route('/commit', methods=['POST'])
def commit_knowledge():
    """接收前端确认的数据 -> 写入 Chroma"""
    data = request.json
    kb_name = data.get('kb_name')     # 用户填写的库名，例如 "v2.0_issue_db"
    selected_items = data.get('items') # 用户勾选的数据
    # 获取前端传来的项目路径，如果没有则用默认的
    project_path = data.get('projectPath', PROJECT_ROOT)
    
    if not kb_name or not selected_items:
        return jsonify({"status": "error", "message": "参数不完整"}), 400
        
    try:
        # 使用传入的 project_path 初始化
        rag_engine.initialize(project_path=project_path, db_name=kb_name)
        
        # 2. 转换数据格式适配 add_manual_data
        # 前端发来的 items 里的 full_data 对应后端的 meta
        items_to_add = []
        for item in selected_items:
            items_to_add.append({
                "id": item['id'],
                "content": item['content'], # 检索文本
                "meta": item['full_data']   # 完整元数据
            })
            
        # 3. 调用我们在 rag_chroma.py 里新加的方法
        result = rag_engine.add_manual_data(items_to_add)
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Commit Error: {e}")
        return jsonify({"status": "error", "message": str(e)})

# 3. 获取列表 (改为 POST 或者带参数的 GET)
# 为了方便传路径，建议用 POST，或者 GET ?projectPath=...
@app.route('/api/rag/list-dbs', methods=['POST']) 
def list_rag_dbs():
    data = request.json
    project_path = data.get('projectPath')
    
    if not project_path:
        return jsonify({"status": "error", "message": "未提供项目路径"})

    # 这里的路径和 rag_chroma.py 里对应
    kb_root = os.path.join(project_path, "rag_database")
    
    dbs = []
    if os.path.exists(kb_root):
        dbs = [d for d in os.listdir(kb_root) if os.path.isdir(os.path.join(kb_root, d))]
    
    # 如果没有建立过，返回空或者 default
    if not dbs: 
        dbs = [] 

    return jsonify({"status": "success", "dbs": dbs})


# 4. 切换数据库
@app.route('/api/rag/switch', methods=['POST'])
def switch_rag_db():
    data = request.json
    project_path = data.get('projectPath') # 必传
    db_name = data.get('dbName')

    if not project_path or not db_name:
         return jsonify({"status": "error", "message": "参数缺失"})

    try:
        # 传入项目路径进行切换
        rag_engine.initialize(project_path, db_name)
        return jsonify({"status": "success", "message": f"已切换至 {db_name}"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

def find_available_port(start_port):
    port = start_port
    while True:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            if s.connect_ex(('0.0.0.0', port)) != 0:
                return port
            port += 1

if __name__ == '__main__':
    start_port = 5055
    available_port = find_available_port(start_port)
    app.run(host='0.0.0.0', port=available_port, debug=True)
