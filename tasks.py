import json
import os
import time
import shutil
import redis
from collections import defaultdict
from datetime import datetime

import pandas as pd
from celery import Celery
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from werkzeug.datastructures import FileStorage

from app import create_app
from app.agent import query_review_result, query_review_result_by_feedback, query_codefile_from_abstract, \
    query_related_code, query_related_code_graph_rerank, query_generated_requirement, query_flow_chart, query_related_requirement
from app.db import (
    get_db_celery,
    append_missing_doc_blocks,
    append_missing_code_blocks,
)
from app.rag_chroma import rag_engine
from app.utils import get_all_files_with_relative_paths, include_related_blocks, replace_text_in_docx, \
    generate_issue_content
from app.call_graph import ensure_project_call_graph, query_function_graph, resolve_code_block_to_function
from callgraph.text_encoding import read_source_file
from app.alignment_config import ALIGN_FILE_ABSTRACT_BATCH_LIMIT, CALL_GRAPH_ALIGN_DEPTH
from app.views import (
    logger,
    get_abstracts_from_sqlite,
    generate_abstract,
    save_abstract_to_db,
    filter_non_abstract_files,
    _get_doc_blocks_for_matching,
    _match_doc_ranges_from_related_reqs,
    _get_or_build_code_blocks_for_file, SECRET_LEVEL_MAP, do_upload_files_logic,
)


celery = Celery(
    'app',
    broker='redis://localhost:6379/0',
    backend='redis://localhost:6379/0'
)

celery.conf.update(
    send_events=False,
    worker_send_task_events=False,
    task_serializer='json',
    accept_content=['json'],
    result_serializer='json',
    timezone='Asia/Shanghai',
    enable_utc=True
)

app = create_app()


class ContextTask(celery.Task):
    def __call__(self, *args, **kwargs):
        with app.app_context():
            return self.run(*args, **kwargs)


celery.Task = ContextTask

def _normalize_kb_type_for_use(raw_type):
    kb_type = (raw_type or "other").strip()
    if kb_type in ("rule", "coding_rule", "checklist"):
        return "rule"
    if kb_type in ("issue", "history_issue"):
        return "issue"
    if kb_type in ("align", "history_align"):
        return "align"
    return "other"

def _safe_first_range_field(ranges, field, default=''):
    if isinstance(ranges, list) and ranges:
        first = ranges[0] or {}
        if isinstance(first, dict):
            return first.get(field, default) or default
    return default	

def _safe_int(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _code_block_file(block):
    return block.get('file') or block.get('filename') or block.get('documentId') or ''


def _code_block_range(block):
    line_range = block.get('range')
    if isinstance(line_range, (list, tuple)) and len(line_range) == 2:
        return [_safe_int(line_range[0]), _safe_int(line_range[1])]
    start_line = _safe_int(block.get('startLine') or block.get('start_line'))
    end_line = _safe_int(block.get('endLine') or block.get('end_line') or start_line)
    return [start_line, end_line]


def _code_block_key(block):
    line_range = _code_block_range(block)
    return (_code_block_file(block), line_range[0], line_range[1])


def _ranges_intersect(start_a, end_a, start_b, end_b):
    return max(_safe_int(start_a), _safe_int(start_b)) <= min(_safe_int(end_a), _safe_int(end_b))


def _dedupe_code_blocks(blocks):
    result = []
    seen = set()
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        key = _code_block_key(block)
        if not key[0] or key[1] <= 0 or key[2] <= 0 or key in seen:
            continue
        seen.add(key)
        result.append(block)
    return result


def _match_related_items_to_blocks(related_items, all_code_blocks):
    matched = []
    for item in related_items or []:
        if not isinstance(item, dict):
            continue
        file_name = _code_block_file(item)
        start_line, end_line = _code_block_range(item)
        if not file_name or start_line <= 0 or end_line <= 0:
            continue

        exact = None
        containing = None
        intersecting = None
        for block in all_code_blocks or []:
            if _code_block_file(block) != file_name:
                continue
            block_start, block_end = _code_block_range(block)
            if block_start == start_line and block_end == end_line:
                exact = block
                break
            if containing is None and block_start <= start_line and block_end >= end_line:
                containing = block
            if intersecting is None and _ranges_intersect(block_start, block_end, start_line, end_line):
                intersecting = block

        selected = exact or containing or intersecting
        if selected:
            enriched = dict(selected)
            if item.get('similarity') is not None:
                enriched['similarity'] = item.get('similarity')
            if item.get('role'):
                enriched['role'] = item.get('role')
            if item.get('reason'):
                enriched['reason'] = item.get('reason')
            matched.append(enriched)
        else:
            matched.append(dict(item))
    return _dedupe_code_blocks(matched)


def _expand_code_blocks_with_call_graph(project_path, project_id, seed_blocks):
    if not project_path or not seed_blocks:
        return []

    try:
        graph_result = ensure_project_call_graph(project_path)
        if not graph_result.get('payload'):
            logger.info(f"调用图不可用，跳过对齐扩展: {graph_result.get('message')}")
            return []
    except Exception as exc:
        logger.info(f"调用图检查/构建失败，跳过对齐扩展: {exc}")
        return []

    expanded = []
    seen_functions = set()
    for block in seed_blocks:
        file_name = _code_block_file(block)
        start_line, end_line = _code_block_range(block)
        if not file_name or start_line <= 0 or end_line <= 0:
            continue
        try:
            function_id = resolve_code_block_to_function(project_path, project_id, file_name, start_line, end_line)
            if not function_id or function_id in seen_functions:
                continue
            seen_functions.add(function_id)
            graph = query_function_graph(
                project_path,
                function_id,
                max_depth=CALL_GRAPH_ALIGN_DEPTH,
                direction="both",
            )
            for code_range in graph.get('code_ranges') or []:
                candidate = {
                    'file': code_range.get('filename') or code_range.get('file') or code_range.get('documentId') or file_name,
                    'filename': code_range.get('filename') or code_range.get('file') or code_range.get('documentId') or file_name,
                    'documentId': code_range.get('documentId') or code_range.get('filename') or file_name,
                    'range': [code_range.get('startLine'), code_range.get('endLine')],
                    'startLine': code_range.get('startLine'),
                    'endLine': code_range.get('endLine'),
                    'start': code_range.get('start'),
                    'end': code_range.get('end'),
                    'content': code_range.get('content') or '',
                    'code': code_range.get('content') or '',
                    'type': code_range.get('type') or 'function',
                    'name': code_range.get('name') or '',
                    'source': 'call_graph',
                }
                expanded.append(candidate)
        except Exception as exc:
            logger.info(f"调用图扩展单个代码块失败 {file_name}:{start_line}-{end_line}: {exc}")
            continue
    return _dedupe_code_blocks(expanded)


def _read_code_range_from_file(project_path, block):
    code_repo_path = os.path.join(project_path, 'code_repo')
    file_name = _code_block_file(block)
    start_line, end_line = _code_block_range(block)
    if not file_name or start_line <= 0 or end_line <= 0:
        return None

    file_path = os.path.join(code_repo_path, file_name)
    if not os.path.exists(file_path):
        return None

    original_content = read_source_file(file_path)
    lines = original_content.splitlines(keepends=True)
    if not lines:
        return None

    safe_start = max(1, start_line)
    safe_end = min(len(lines), max(safe_start, end_line))
    char_start = sum(len(line) for line in lines[:safe_start - 1])
    char_end = sum(len(line) for line in lines[:safe_end])
    range_content = '\n'.join(line.rstrip('\n\r') for line in lines[safe_start - 1:safe_end])

    return {
        'filename': file_name,
        'start': char_start,
        'end': char_end,
        'content': range_content,
        'documentId': file_name,
        'startLine': safe_start,
        'endLine': safe_end,
        'name': block.get('name') or '',
        'type': block.get('type') or '',
    }


def _code_blocks_to_code_ranges(project_path, blocks):
    code_ranges = []
    seen = set()
    for block in blocks or []:
        code_range = _read_code_range_from_file(project_path, block)
        if not code_range:
            continue
        key = (code_range['filename'], code_range['startLine'], code_range['endLine'])
        if key in seen:
            continue
        seen.add(key)
        code_ranges.append(code_range)
    return code_ranges
	
@celery.task(name="user_bp.task.add")
def add(x, y):
    time.sleep(10)
    return x + y


@celery.task(bind=True)
def abstract_code_from_project_task(self, params, code_file_path, user_id):
    #print('*********1111111*******')
    try:
        project_id = params.get('project_id')
        project_path = params.get('projectPath', '')
        # 从SQLite读取数据
        # print('从SQLite读取数据!!!!!!!!!!!!!!!')
        df = get_abstracts_from_sqlite(project_id)
        # print(df.head())
        # 排除无关文件夹/目录
        exclude_folders = ['.git', '.idea']
        # 基于文件名后缀，指定文件类型
        include_files = ['.py', '.c', '.cpp', '.h', '.hpp', '.java', '.html', '.vhd', '.v', '.sv']
        # 遍历文件夹
        file_abstract = {}
        for root, dirs, files in os.walk(code_file_path):
            dirs[:] = [d for d in dirs if d not in exclude_folders]
            total = len(files)
            for i, file in enumerate(files, 1):
                self.update_state(
                    state="PROGRESS",
                    meta={
                        'current': i,
                        'total': total,
                        'name': f'正在摘要:{file}',
                        'status': f'任务进行中{i}/{total}...'
                    }
                )
                if os.path.splitext(file)[1] in include_files:
                    file_path = os.path.join(root, file)
                    # 文件不能是0KB的空文件
                    if os.path.getsize(file_path) == 0:
                        continue
                    # 构建相对路径    
                    rel_path = os.path.relpath(file_path, code_file_path)

                    if not df.empty:
                        # 先看数据库里有没有已经生成好的代码摘要
                        row_data = df[df['filename'] == rel_path]

                        # 数据库有该代码文件的摘要
                        if not row_data.empty:
                            logger.info('数据库有该代码文件的摘要')
                            abstract_data = row_data['abstract'].values[0]
                            file_abstract[rel_path] = abstract_data
                        # 数据库没有该代码文件的摘要
                        else:
                            logger.info('数据库没有该代码文件的摘要')
                            file_path = os.path.join(root, file)
                            codefile_abstract = generate_abstract(file_path)
                            file_abstract[rel_path] = codefile_abstract

                            # save_abstract_to_db(project_path, file, codefile_abstract, project_id)
                            save_abstract_to_db(project_path, rel_path, codefile_abstract, project_id, user_id)

                    # 数据库里代码摘要这张表是空的，需要新生成
                    else:
                        logger.info('数据库里代码摘要这张表是空的')
                        file_path = os.path.join(root, file)
                        codefile_abstract = generate_abstract(file_path)
                        file_abstract[rel_path] = codefile_abstract
                        # save_abstract_to_db(project_path, file, codefile_abstract, project_id)
                        save_abstract_to_db(project_path, rel_path, codefile_abstract, project_id, user_id)

        self.update_state(
            state="SUCCESS",
            meta={
                'status': f'摘要生成完毕'
            }
        )
        logger.info('代码文件的摘要已生成')
        return {'status': True, 'message': '任务执行成功', 'data': file_abstract}

    except Exception as e:
        self.update_state(
            state="FAILURE",
            message=f"摘要过程中出错{e}"
        )
        logger.error(f"生成代码摘要出错: {str(e)}", exc_info=True)
        return {'status': False, 'message': str(e)}


@celery.task(bind=True)
def review_alignment_task(self, project_path, project_id, user_id, files, prompt_type=None, reviewed_count=None):
    try:
        result = []
        if isinstance(files, dict):
            result.extend(
                {"doc_file": doc_file, "alignment": item}
                for doc_file, alignments in files.items()
                for item in (alignments or [])
            )

        if not result:
            raise ValueError("缺少可审查的数据：既没有需求-代码对齐，也没有代码块列表")

        # print('result===============', result)
        total = len(result)
        # print('reviewed_count===============',reviewed_count)
        if not reviewed_count or not isinstance(reviewed_count, int):
            reviewed_count = 0
        for i, item in enumerate(result, reviewed_count):
            alignment = item['alignment']
            self.update_state(
                state="PROGRESS",
                meta={
                    'current': i,
                    'total': total,
                    'name': alignment['name'],
                    'status': f'任务进行中{i}/{total}...'
                }
            )
            # 获取选定的 knowledge base
            selected_rule_kbs = []
            selected_issue_kbs = []

            try:
                metadata_file = os.path.join(project_path, 'metadata.json')
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                        selected_kbs = metadata.get('selected_kbs', [])
                        selected_rule_kbs = [kb['name'] for kb in selected_kbs if _normalize_kb_type_for_use(kb.get('type')) == 'rule']
                        selected_issue_kbs = [kb['name'] for kb in selected_kbs if _normalize_kb_type_for_use(kb.get('type')) == 'issue']
            except Exception as e:
                logger.error(str(e), exc_info=True)
            # 检索上下文
            retrieved_rules = []
            retrieved_issues = []

            doc_ranges = alignment.get('docRanges', [])
            code_ranges = alignment.get('codeRanges', [])

            # 构造查询文本
            query_text = ""
            if doc_ranges:
                query_text += doc_ranges[0].get('content', '') + "\n"
            if code_ranges:
                query_text += code_ranges[0].get('content', '')

            # 检索规则
            for kb_name in selected_rule_kbs:
                collection = rag_engine.get_collection('rule', kb_name)
                if collection:
                    results = collection.query(query_texts=[query_text], n_results=3)
                    if results and results['documents']:
                        for doc in results['documents'][0]:
                            retrieved_rules.append(doc)

            # 检索问题单
            for kb_name in selected_issue_kbs:
                collection = rag_engine.get_collection('issue', kb_name)
                if collection:
                    results = collection.query(query_texts=[query_text], n_results=3)
                    if results and results['documents']:
                        for doc in results['documents'][0]:
                            retrieved_issues.append(doc)

            # 1. 调用 agent 获取审查结果
            review_process, issue = query_review_result(
                doc_ranges,
                code_ranges,
                rules=retrieved_rules,
                issues=retrieved_issues,
                user_id=user_id,
                project_path=project_path,
                prompt_type=prompt_type
            )

            # 2. 更新对齐关系
            alignment['isReviewed'] = True
            alignment['reviewThoughts'] = review_process

            if isinstance(issue, list):
                issues_list = [x for x in issue if isinstance(x, dict)]
            elif isinstance(issue, dict):
                issues_list = [issue]
            else:
                issues_list = []
            
            # 需求反生成
            # generated_requirement, mermaid_code = gen_requirement(doc_ranges, code_ranges)
            generated_requirement, mermaid_code = '', ''

            conn = get_db_celery()
            cur = conn.cursor()
            try:

                cur.execute(
                    'INSERT INTO alignments(id,user_id,project_id,name,isReviewed,reviewThoughts,docRanges,codeRanges,'
                    'createdAt,updatedAt,GenReq,GenMermaid,is_code_review) '
                    'VALUES(%s,%s,%s,%s,1,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP,%s,%s,%s) '
                    'ON DUPLICATE KEY UPDATE '
                    'isReviewed=1,'
                    'reviewThoughts=VALUES(reviewThoughts),'
                    'GenReq=VALUES(GenReq),'
                    'GenMermaid=VALUES(GenMermaid),'
                    'updatedAt=CURRENT_TIMESTAMP',
                    (
                        alignment.get('id'),
                        user_id,
                        project_id,
                        alignment.get('name') or '',
                        alignment.get('reviewThoughts') or '',
                        json.dumps(doc_ranges or []),
                        json.dumps(code_ranges or []),
                        generated_requirement or '',
                        mermaid_code or '',
                        0 if not prompt_type else 1
                    )
                )

                if issues_list:
                    cur.execute(
                        f"SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%' and project_id={project_id}")
                    used = set()
                    for r in cur.fetchall():
                        disp = r['displayId']
                        if disp and disp.startswith('ISSUE-'):
                            try:
                                used.add(int(disp.split('-')[1]))
                            except Exception as e:
                                logger.error(str(e), exc_info=True)

                    next_number = (max(used) + 1) if used else 1

                    doc_ranges = alignment.get('docRanges', [])
                    if doc_ranges:
                        content = doc_ranges[0].get('content', '')
                    else:
                        content = ''
                    brief_req = content or ''
                    brief_code = _safe_first_range_field(alignment.get('codeRanges', []), 'content')
                    related_doc_file = (
                        item.get('doc_file')
                        or _safe_first_range_field(alignment.get('docRanges', []), 'filename')
                        or _safe_first_range_field(alignment.get('codeRanges', []), 'filename')
                    )
                    for one in issues_list:
                        display_id = f"ISSUE-{next_number:03d}"
                        next_number += 1

                        severity = one.get('level') or one.get('severity')
                        title = one.get('summary') or one.get('title') or ''
                        content = one.get('description') or one.get('content') or ''
                        status = one.get('status') or 'unconfirmed'

                        cur.execute(
                            'INSERT INTO issues(user_id,project_id,displayId,alignmentId,severity,title,content,status,'
                            'relatedDocFile,relatedRequirementId,briefRequirement,briefCode,category,createdAt,updatedAt) '
                            'VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
                            (
                                user_id,
                                project_id,
                                display_id,
                                alignment.get('id'),
                                severity,
                                title,
                                content,
                                status,
                                item['doc_file'],
                                alignment.get('id'),
                                brief_req,
                                brief_code,
                                alignment.get('align_type') if alignment.get('align_type') else 'codeReview'
                            )
                        )

                conn.commit()
            except Exception as e:
                conn.rollback()
                logger.error(f"写入数据库失败: {str(e)}", exc_info=True)
                return {"status": "error", "message": f"写入数据库失败: {str(e)}"}
            finally:
                conn.close()
    except Exception as e:
        self.update_state(
            state="FAILURE",
            message=f"审查出错{e}"
        )
        logger.error(f"审查失败2: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to save review result: {str(e)}") from e

    self.update_state(
        state="SUCCESS",
        meta={}
    )


# @celery.task
# def review_alignment_addprompt_task(project_path, alignment, project_id, user_id, doc_file=None, user_prompt=None):
#     # 获取选定的 knowledge base
#     selected_rule_kbs = []
#     selected_issue_kbs = []
#     try:
#         metadata_file = os.path.join(project_path, 'metadata.json')
#         if os.path.exists(metadata_file):
#             with open(metadata_file, 'r', encoding='utf-8') as f:
#                 metadata = json.load(f)
#                 selected_kbs = metadata.get('selected_kbs', [])
#                 selected_rule_kbs = [kb['name'] for kb in selected_kbs if _normalize_kb_type_for_use(kb.get('type')) == 'rule']
#                 selected_issue_kbs = [kb['name'] for kb in selected_kbs if _normalize_kb_type_for_use(kb.get('type')) == 'issue']
#     except Exception as e:
#         logger.error(str(e), exc_info=True)
#
#     # 检索上下文
#     retrieved_rules = []
#     retrieved_issues = []
#
#     doc_ranges = alignment.get('docRanges', [])
#     code_ranges = alignment.get('codeRanges', [])
#     reviewThoughts = alignment.get('reviewThoughts', [])
#
#     # 构造查询文本
#     query_text = ""
#     if doc_ranges:
#         query_text += doc_ranges[0].get('content', '') + "\n"
#     if code_ranges:
#         query_text += code_ranges[0].get('content', '')
#
#     # 检索规则
#     for kb_name in selected_rule_kbs:
#         collection = rag_engine.get_collection('rule', kb_name)
#         if collection:
#             results = collection.query(query_texts=[query_text], n_results=3)
#             if results and results['documents']:
#                 for doc in results['documents'][0]:
#                     retrieved_rules.append(doc)
#
#     # 检索问题单
#     for kb_name in selected_issue_kbs:
#         collection = rag_engine.get_collection('issue', kb_name)
#         if collection:
#             results = collection.query(query_texts=[query_text], n_results=3)
#             if results and results['documents']:
#                 for doc in results['documents'][0]:
#                     retrieved_issues.append(doc)
#
#     # 1. 调用 agent 获取审查结果
#     review_process, issue = query_review_result_by_feedback(
#         doc_ranges,
#         code_ranges,
#         reviewThoughts,
#         user_prompt,
#         rules=retrieved_rules,
#         issues=retrieved_issues,
#         user_id=user_id,
#         project_path=project_path
#     )
#
#     # 2. 更新对齐关系
#     alignment['isReviewed'] = True
#     alignment['reviewThoughts'] = review_process
#
#     if isinstance(issue, list):
#         issues_list = [x for x in issue if isinstance(x, dict)]
#     elif isinstance(issue, dict):
#         issues_list = [issue]
#     else:
#         issues_list = []
#
#     # 需求反生成
#     generated_requirement, mermaid_code = gen_requirement(doc_ranges, code_ranges)
#
#     conn = get_db_celery()
#     cur = conn.cursor()
#
#     try:
#         cur.execute(
#             'UPDATE alignments SET isReviewed=1, reviewThoughts=%s, GenReq=%s, GenMermaid=%s, updatedAt=CURRENT_TIMESTAMP '
#             'WHERE id=%s and project_id=%s',
#             (alignment.get('reviewThoughts') or '', generated_requirement or '', mermaid_code or '', alignment.get('id'), project_id)
#         )
#
#         if issues_list:
#             cur.execute(f"SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%' and project_id={project_id}")
#             used = set()
#             for r in cur.fetchall():
#                 disp = r['displayId']
#                 if disp and disp.startswith('ISSUE-'):
#                     try:
#                         used.add(int(disp.split('-')[1]))
#                     except Exception as e:
#                         logger.error(str(e), exc_info=True)
#
#             next_number = (max(used) + 1) if used else 1
#
#             brief_req = _safe_first_range_field(alignment.get('docRanges', []), 'content')
#             brief_code = _safe_first_range_field(alignment.get('codeRanges', []), 'content')
#             related_doc_file = (
#                 doc_file
#                 or _safe_first_range_field(alignment.get('docRanges', []), 'filename')
#                 or _safe_first_range_field(alignment.get('codeRanges', []), 'filename')
#             )
#
#             for one in issues_list:
#                 display_id = f"ISSUE-{next_number:03d}"
#                 next_number += 1
#
#                 severity = one.get('level') or one.get('severity')
#                 title = one.get('summary') or one.get('title') or ''
#                 content = one.get('description') or one.get('content') or ''
#                 status = one.get('status') or 'unconfirmed'
#
#                 cur.execute(
#                     'INSERT INTO issues(user_id,project_id,displayId,alignmentId,severity,title,content,status,'
#                     'relatedDocFile,relatedRequirementId,briefRequirement,briefCode,createdAt,updatedAt) '
#                     'VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
#                     (
#                         user_id,
#                         project_id,
#                         display_id,
#                         alignment.get('id'),
#                         severity,
#                         title,
#                         content,
#                         status,
#                         related_doc_file,
#                         alignment.get('id'),
#                         brief_req,
#                         brief_code
#                     )
#                 )
#
#         conn.commit()
#     except Exception as e:
#         conn.rollback()
#         logger.error(f"审查失败 user_prompt2: {str(e)}", exc_info=True)
#         return {"status": "error", "message": f"Failed to save review result: {str(e)}"}
#     finally:
#         conn.close()
#
#     return {"status": "success", "createdIssues": len(issues_list)}


@celery.task(bind=True)
def align_requirement_to_project_task(self, abstract, params, user_id):

    project_path = params.get('projectPath', '')
    project_id = params.get('project_id')
    chunks = params.get('requirements')
    y_align = params.get('y_align')
    total = len(chunks)
    try:

        for i, chunk in enumerate(chunks, y_align):

            self.update_state(
                state="PROGRESS",
                meta={
                    'current': i,
                    'total': total,
                    'name': chunk['name'],
                    'status': f'任务进行中{i}/{total}...'
                }
            )
            print(f"正在执行==文件名{chunk['name']} {i}/{len(chunks)}")
            doc_ranges = chunk['docRanges']
            # 获取项目中所有代码文件
            code_repo_path = os.path.join(project_path, 'code_repo')
            all_files = get_all_files_with_relative_paths(code_repo_path, 'code')

            # 拼接所有docRanges的content作为requirement_text
            requirement_text = '\n\n'.join(
                [doc_range.get('content', '') for doc_range in doc_ranges if doc_range.get('content')])
            if not requirement_text or not project_path:
                return {"status": "error", "message": "缺少需求内容或项目路径参数"}

            # 如果有多个代码文件，执行检索代码摘要
            file_abstract = abstract.get('data', abstract) if isinstance(abstract, dict) else abstract
            if len(all_files) > 1:
                # 基于需求，利用大模型检索代码摘要，先定位代码文件
                # 调用llm
                file_name_list = query_codefile_from_abstract(requirement_text, file_abstract)
                
                # 解析异常，返回空列表时
                # 过滤掉含有乱码的代码摘要（作为被定位的代码文件防止遗漏），重新调用大模型定位代码文件
                FILE_MAX_LIMIT = ALIGN_FILE_ABSTRACT_BATCH_LIMIT
                if not file_name_list:
                    filter_non_file, filter_file_abstract, file_cnt = filter_non_abstract_files(file_abstract)
                    
                    # 代码摘要数量小于阈值时，可以直接调用
                    if file_cnt <=FILE_MAX_LIMIT:
                        # 调用llm
                        file_name_list = query_codefile_from_abstract(requirement_text, filter_file_abstract)
                        #print(file_name_list)
                    
                    # 可能由于代码摘要过多，影响大模型分析理解而报错
                    else:
                        # 对代码摘要分批次处理
                        file_name_list = []
                        file_cnt = 0
                        batch_file_abstract = {}
                        for key, value in filter_file_abstract.items():
                            file_cnt += 1
                            batch_file_abstract[key] = value
                            if file_cnt >= FILE_MAX_LIMIT:
                                file_cnt = 0
                                batch_file_name_list = query_codefile_from_abstract(requirement_text, batch_file_abstract)
                                file_name_list += batch_file_name_list
                    
                    # 谨防遗漏，将有摘要是乱码的代码文件全部放入候选区
                    if not file_name_list:                    
                        file_name_list += filter_non_file
                    #print(file_name_list)
                    
            # 如果只有一个代码文件，就不检索代码摘要
            else:
                file_name_list = all_files

            # 尝试初始化RAG引擎，以便在agent中使用
            # try:
            # rag_engine.initialize(project_path)
            # except Exception as e:
            # print(f"[Align] RAG initialize failed: {e}")

            code_ranges = []

            # 遍历经过代码摘要筛选的代码文件
            for file_name in file_name_list:
                # print('22222222222222222222222222')
                # 假设 file_name 是一个字典，例如：{"filename": "main.py"}
                if isinstance(file_name, dict):
                    file_name = file_name.get("filename", file_name.get("file", ""))

                # 确保 file_name 是字符串
                # print('file_name, (str, bytes, os.PathLike):', file_name, type(file_name))
                if not isinstance(file_name, (str, bytes, os.PathLike)):
                    continue
            
            
                all_code_blocks = _get_or_build_code_blocks_for_file(project_path, file_name, project_id)
                # print('all_code_blocks:', all_code_blocks)
                if not all_code_blocks:
                    continue
                
                # 先用现有语义/LLM链路找种子块，再用调用图扩展候选，最后二次保守筛选。
                seed_related_code = query_related_code(
                    requirement_text,
                    all_code_blocks,
                    block_limit=50,
                    user_id=user_id,
                    project_path=project_path
                )
                
                try:
                    try:
                        seed_blocks = include_related_blocks(seed_related_code, all_code_blocks)
                    except Exception as exc:
                        logger.info(f"related_id 扩展失败，使用基础匹配结果: {exc}")
                        seed_blocks = _match_related_items_to_blocks(seed_related_code, all_code_blocks)

                    seed_blocks = _dedupe_code_blocks(seed_blocks)
                    graph_blocks = _expand_code_blocks_with_call_graph(project_path, project_id, seed_blocks)
                    candidate_blocks = _dedupe_code_blocks(seed_blocks + graph_blocks)

                    final_blocks = []
                    if graph_blocks:
                        reranked_code = query_related_code_graph_rerank(
                            requirement_text,
                            seed_blocks,
                            candidate_blocks
                        )
                        final_blocks = _match_related_items_to_blocks(reranked_code, candidate_blocks)

                    if not final_blocks:
                        final_blocks = seed_blocks

                    code_ranges.extend(_code_blocks_to_code_ranges(project_path, final_blocks))
                                    
                except Exception as e:
                    print(f"add related_code failed: {e}") 
                    
            if code_ranges:
                chunk['codeRanges'] = code_ranges
                add_alignment_data(project_path, chunk, project_id, user_id)

        self.update_state(
            state="SUCCESS",
            meta={}
        )

    except Exception as e:
        self.update_state(
            state="FAILURE",
            message=f"对齐过程中出错{e}"
        )
        logger.error(f'对齐过程中出错:{str(e)}', exc_info=True)
        raise RuntimeError(f'对齐过程中出错:{str(e)}') from e


@celery.task(bind=True)
def align_code_to_requirements_task(self, project_path, code_blocks, project_id, user_id, y_align):
    total = len(code_blocks)
    try:
        all_doc_blocks, _, blocks_by_file = _get_doc_blocks_for_matching(project_path, project_id)
        for i, code_block in enumerate(code_blocks, y_align):
            self.update_state(
                state="PROGRESS",
                meta={
                    'current': i,
                    'total': total,
                    'name': code_block['name'],
                    'status': f'任务进行中{i}/{total}...'
                }
            )
            print(f"正在执行==文件名{code_block['name']} {i}/{total}")
            code_ranges = code_block['codeRanges']
            # 获取选定的 align 类型知识库
            selected_align_kbs = []
            try:
                metadata_file = os.path.join(project_path, 'metadata.json')
                if os.path.exists(metadata_file):
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                        selected_kbs = metadata.get('selected_kbs', [])
                        selected_align_kbs = [kb['name'] for kb in selected_kbs if _normalize_kb_type_for_use(kb.get('type')) == 'align']
            except Exception:
                pass

            if not all_doc_blocks:
                code_block['docRanges'] = []
                add_alignment_data(project_path, code_block, project_id, user_id)
                continue

            # 如果没有选择任何 align 知识库，使用原来的 LLM 逻辑
            if not selected_align_kbs:

                code_content = '\n\n'.join(
                    [code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

                # 调用LLM
                related_reqs = query_related_requirement(
                    code_content,
                    all_doc_blocks,
                    block_limit=50,
                    user_id=user_id,
                    project_path=project_path
                )

                doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

                code_block['docRanges'] = doc_ranges
                add_alignment_data(project_path, code_block, project_id, user_id)
                continue

            # 如果选择了 align 知识库，使用 RAG 进行检索
            try:
                rag_engine.initialize()  # 确保初始化
            except Exception as e:
                print(f"[Align] RAG initialize failed: {e}")

            code_content = '\n\n'.join(
                [code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

            all_retrieved_items = []
            for kb_name in selected_align_kbs:
                # 检索 'align' 类型的知识库
                # 注意：align 知识库里存储的是历史对齐数据
                # 这里的检索策略可以是：用代码去搜相关的历史对齐，然后把历史对齐中的文档部分作为推荐

                # 暂时假设 rag_chroma 提供了 query 接口，如果没有需要添加
                # 这里先模拟直接调用 collection.query
                collection = rag_engine.get_collection('align', kb_name)
                if collection:
                    results = collection.query(
                        query_texts=[code_content],
                        n_results=5  # Top 5 per KB
                    )

                    if results and results['documents']:
                        for i, doc in enumerate(results['documents'][0]):
                            meta = results['metadatas'][0][i]
                            # 历史对齐数据通常包含 code_text 和 doc_text (query_text)
                            # 我们需要提取其中的 doc 部分
                            # 在 build_from_json 中，document 是 doc_text, meta 中有 code_text
                            # 但我们现在是用 code 去搜 doc，所以 doc 正好是 document
                            all_retrieved_items.append({
                                'content': doc,
                                'score': results['distances'][0][i] if results['distances'] else 0,
                                'meta': meta
                            })

            # 对所有结果排序
            all_retrieved_items.sort(key=lambda x: x['score'])  # distance 越小越好
            top_items = all_retrieved_items[:5]

            # 构造返回结果
            # 注意：这里返回的是参考的历史对齐文档内容，而不是当前项目中的具体需求块
            # 前端可能需要展示这些参考内容供用户选择，或者作为提示
            # 但目前的接口契约是返回 docRanges (当前项目的需求块)
            # 这是一个逻辑断层：历史对齐是"参考"，而不是"直接结果"
            # 如果要用历史对齐来辅助定位当前项目的需求，需要两步：
            # 1. 检索历史对齐 -> 得到相关的历史需求描述
            # 2. 用历史需求描述去匹配当前项目的需求块 (类似 query_related_requirement)

            history_doc_contents = [item['content'] for item in top_items]
            combined_history_content = "\n".join(history_doc_contents)

            # 使用历史需求内容 + 代码内容 共同作为 Query 去查询当前项目需求
            enhanced_query = f"Code:\n{code_content}\n\nRelated History Requirements:\n{combined_history_content}"

            # 调用LLM (使用增强后的 Query)
            related_reqs = query_related_requirement(
                enhanced_query,
                all_doc_blocks,
                block_limit=50,
                project_path=project_path
            )

            doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

            code_block['docRanges'] = doc_ranges
            add_alignment_data(project_path, code_block, project_id, user_id)

        self.update_state(
            state="SUCCESS",
            meta={}
        )

    except Exception as e:
        self.update_state(
            state="FAILURE",
            message=f"对齐过程中出错{e}"
        )
        logger.error(f'对齐 代码=>需求 过程中出错:{str(e)}', exc_info=True)
        raise RuntimeError(f'对齐 代码=>需求 过程中出错:{str(e)}') from e


def add_alignment_data(project_path, new_alignment, project_id, user_id):
    if not project_path or not new_alignment or 'id' not in new_alignment:
        return {"status": "error", "message": "缺少项目路径或无效的对齐数据。"}

    code_ranges = new_alignment['codeRanges']
    # --- 自动创建块的逻辑 ---
    try:
        doc_ranges = new_alignment.get('docRanges', [])
        if doc_ranges:
            append_missing_doc_blocks(project_path, doc_ranges)

        if code_ranges:
            append_missing_code_blocks(project_path, code_ranges)

    except Exception as e:
        # print(f"Error auto-creating blocks: {e}")
        # 即使块创建失败，也不应该阻止对齐关系的保存，但最好记录日志
        logger.error(f"Error auto-creating blocks: {str(e)}", exc_info=True)


    conn = get_db_celery()
    cur = conn.cursor()
    try:
        # conn = get_db_conn(project_path)
        cur.execute(
            '''
            INSERT INTO alignments(id, user_id, project_id, name, isReviewed, reviewThoughts, docRanges, codeRanges, createdAt, updatedAt, is_code_review, is_alignment) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, %s, 1) 
            ON DUPLICATE KEY UPDATE 
                name = VALUES(name),
                isReviewed = VALUES(isReviewed),
                reviewThoughts = VALUES(reviewThoughts),
                docRanges = VALUES(docRanges),
                codeRanges = VALUES(codeRanges),
                updatedAt = CURRENT_TIMESTAMP,
                is_alignment = 1
            ''',
            (
                new_alignment.get('id'),
                user_id,
                project_id,
                new_alignment.get('name'),
                1 if new_alignment.get('isReviewed') else 0,
                new_alignment.get('reviewThoughts') or '',
                json.dumps(new_alignment.get('docRanges') or []),
                json.dumps(code_ranges or []),
                0
                # generated_requirement or ''
                # mermaid_code or ''
            )
        )
        conn.commit()
        # conn.close()
        return {"status": "success"}
    except Exception as e:
        conn.rollback()
        logger.error(f"写入对齐数据失败:{str(e)}", exc_info=True)
        return {"status": "error", "message": f"写入对齐数据失败: {e}"}
    finally:
        conn.close()


def gen_requirement(doc_ranges, code_ranges):
    generated_requirement = ''
    flowchart_code = ''

    # 反生成需求+流程图
    try:
        requirement_content = doc_ranges
        code_content = code_ranges

        # if not code_content:
        #     return jsonify({"status": "error", "message": "Missing code content"}), 400

        # 构建代码块列表，格式与现有函数兼容
        code_blocks = []
        if isinstance(code_content, list):
            for code_block in code_content:
                code_blocks.append({
                    'filename': code_block.get('filename', 'unknown'),
                    'content': code_block.get('content', '')
                })
        else:
            # 如果是字符串，创建单个代码块
            code_blocks.append({
                'filename': 'code',
                'content': code_content
            })

        # 调用LLM生成需求，传入参考需求内容
        generated_requirement = query_generated_requirement(code_blocks, requirement_content or "")

        # 调用LLM生成流程图
        try:
            flowchart_code = query_flow_chart(code_content if isinstance(code_content, str) else
                                              '\n\n'.join([block.get('content', '') for block in code_content]))
        except Exception as flowchart_error:
            logger.error(f"Error generating flowchart: {str(flowchart_error)}", exc_info=True)
            flowchart_code = ''

    except Exception as e:
        # print(f"Error generating reverse requirement: {str(e)}")
        logger.error(f"Error generating reverse requirement: {str(e)}", exc_info=True)
        # return jsonify({"status": "error", "message": f"Failed to generate reverse requirement: {str(e)}"}), 500
    
    return generated_requirement, flowchart_code


@celery.task(bind=True)
def gen_requirement_task(self, alignments, total, generated):
    db = get_db_celery()
    cursor = db.cursor()
    try:
        for i, alignment in enumerate(alignments, generated):
            # if alignment['GenReq']:
            #     continue
            self.update_state(
                state="PROGRESS",
                meta={
                    'current': i,
                    'total': total,
                    'name': alignment['name'],
                    'status': f'任务进行中{i}/{total}...'
                }
            )
            code_ranges = json.loads(alignment['codeRanges'])
            doc_ranges = json.loads(alignment['docRanges'])
            # 需求反生成
            generated_requirement, mermaid_code = gen_requirement(doc_ranges, code_ranges)

            cursor.execute("""
                UPDATE alignments
                SET GenReq = %s,
                    GenMermaid = %s
                WHERE id = %s
            """, (generated_requirement, mermaid_code, alignment['id']))
            db.commit()
        self.update_state(
            state="SUCCESS",
            meta={}
        )
    except Exception as e:
        self.update_state(
            state="FAILURE",
            message=f"需求反生成失败:{e}"
        )
        db.rollback()
        logger.error(f"需求反生成失败:{str(e)}", exc_info=True)
    finally:
        db.close()


@celery.task(bind=True)
def export_issues_task(self, data, docx_path):
    db = get_db_celery()
    cursor = db.cursor()
    try:

        issues = data.get('issues', [])
        form_data = data.get('formData', {})
        secret_level = data.get('secret_level')
        template_path = os.path.join(os.path.dirname(__file__), 'templates/', '问题单模板.docx')
        # 创建临时目录存储文件
        temp_dir = os.path.join(os.path.dirname(__file__), 'app/', 'temp_exports')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir, exist_ok=True)

        # 检查是否提供了DOCX模板路径
        if template_path and os.path.exists(template_path):
            current_date = datetime.now().strftime("%Y%m%d")
            issue_categories = form_data.get('issueCategories', [])

            # 将英文级别转换为中文的映射
            level_mapping = {
                'high': '重大',
                'medium': '严重',
                'low': '一般'
            }

            # 处理第一个问题单作为基础文档
            first_issue = issues[0]
            merged_doc = Document(template_path)

            export_prompt = SECRET_LEVEL_MAP[secret_level]
            p = merged_doc.paragraphs[0].insert_paragraph_before(export_prompt) if merged_doc.paragraphs \
                else merged_doc.add_paragraph(export_prompt)
            p.alignment = 0
            run = p.runs[0]
            set_run_font(run)

            replacements = {}
            # 替换页码信息
            replacements["CURRENT"] = "1"
            replacements["TOTAL"] = str(len(issues))

            replacements["AAAAA软件"] = form_data.get('productName', '')
            replacements["BBBBB"] = f"{form_data.get('issueId', '')}_1"
            replacements["CCCCC"] = form_data.get('productId', '')
            replacements["DDDDD"] = form_data.get('discoveryMethod', '')
            #replacements["EEEEE"] = form_data.get('issueTracking', '')
            replacements["EEEEE"] = f"{form_data.get('issueTracking', '')}_1"
            
            replacements["GGGGG"] = current_date

            # 处理问题类别
            for category in ['需求问题', '设计问题', '编码问题', '测试问题', '文档问题', '数据问题', '其他问题']:
                if category in issue_categories:
                    replacements[f"□{category}"] = f"■{category}"

            # 处理问题级别
            issue_level = first_issue.get('level', '')
            chinese_level = level_mapping.get(issue_level.lower(), issue_level)

            for level in ['重大', '严重', '一般']:
                if level == chinese_level:
                    replacements[f"□{level}"] = f"■{level}"

            replacements["CONTENTCONTENT"] = first_issue.get('description', '')

            # 替换第一个文档的占位符
            replace_text_in_docx(merged_doc, replacements, 'issue')
            
            # 【新增代码开始】
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            # 创建分页符元素 <w:br w:type="page"/>
            page_break = OxmlElement('w:br')
            page_break.set(qn('w:type'), 'page')
            merged_doc.element.body.append(page_break)

            cursor.execute("""
                        UPDATE export_tasks SET status = 'processing'
                        WHERE task_id = %s
                    """, (self.request.id,))
            db.commit()
            
            
            # 处理剩余的问题单
            for i, issue in enumerate(issues[1:], 2):
                # 添加分页符
                #merged_doc.add_page_break()

                # 为每个问题单加载新的模板并填充
                temp_doc = Document(template_path)

                replacements = {}
                # 替换页码信息
                replacements["CURRENT"] = str(i)
                replacements["TOTAL"] = str(len(issues))

                replacements["AAAAA软件"] = form_data.get('productName', '')
                replacements["BBBBB"] = f"{form_data.get('issueId', '')}_{i}"
                replacements["CCCCC"] = form_data.get('productId', '')
                replacements["DDDDD"] = form_data.get('discoveryMethod', '')
                replacements["EEEEE"] = f"{form_data.get('issueTracking', '')}_{i}"#form_data.get('issueTracking', '')
                
                replacements["GGGGG"] = current_date

                # 处理问题类别
                for category in ['需求问题', '设计问题', '编码问题', '测试问题', '文档问题', '数据问题', '其他问题']:
                    if category in issue_categories:
                        replacements[f"□{category}"] = f"■{category}"

                # 处理问题级别
                issue_level = issue.get('level', '')
                chinese_level = level_mapping.get(issue_level.lower(), issue_level)

                for level in ['重大', '严重', '一般']:
                    if level == chinese_level:
                        replacements[f"□{level}"] = f"■{level}"

                replacements["CONTENTCONTENT"] = issue.get('description', '')

                # 替换模板中的占位符
                replace_text_in_docx(temp_doc, replacements, 'issue')

                # 直接拼接填充好的页面内容到合并文档
                for element in temp_doc.element.body:
                    merged_doc.element.body.append(element)
                    
                # 添加分页符
                page_break = OxmlElement('w:br')
                page_break.set(qn('w:type'), 'page')
                merged_doc.element.body.append(page_break)
                
            # 保存合并后的文档
            merged_doc.save(docx_path)
        else:
            # 使用文本格式导出（备用方案）
            content = ""
            for i, issue in enumerate(issues, 1):
                content += f"问题单 {i}/{len(issues)}\n"
                content += generate_issue_content(issue, form_data)
                content += "\n" + "=" * 50 + "\n\n"

            # 创建一个简单的docx文档
            doc = Document()
            doc.add_paragraph(content)
            doc.save(docx_path)

        cursor.execute("""
            UPDATE export_tasks SET status = 'success', completed_at = NOW()
            WHERE task_id = %s
        """, (self.request.id, ))
        db.commit()
    except Exception as e:
        cursor.execute("""
            UPDATE export_tasks SET status = 'failure', completed_at = NOW(), error_msg = %s
            WHERE task_id = %s
        """, (str(e), self.request.id))
        db.commit()
        logger.error(f"生成问题单文件失败:{str(e)}", exc_info=True)
    finally:
        db.close()


def set_run_font(run, font_name='SimSun', font_size=Pt(10), color=None):
    """
    设置 run 字体，解决跨平台乱码
    """

    east_asia_map = {
        'Microsoft YaHei': '微软雅黑',
        'SimSun': '宋体(中文正体)',
        'SimHei': '黑体',
        'KaiTi': '楷体',
        'FangSong': '仿宋'
    }

    run.font.name = font_name
    run.font.size = font_size
    if color:
        run.font.color.rgb = color

    east_asia = east_asia_map.get(font_name, font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)

def _update_snapshot(task_id, **fields):
    """更新 user_task_snapshot，fields 支持 state / title / current_progress / is_running"""
    if not fields:
        return
    sets = ', '.join(f'{k}=%s' for k in fields)
    db = get_db_celery()
    cursor = db.cursor()
    cursor.execute(
        f"UPDATE user_task_snapshot SET {sets} WHERE task_id=%s",
        (*fields.values(), task_id)
    )
    db.commit()    


# db2 专门用来放锁，和你的 broker(db0)、结果(db1) 分开
redis_client = redis.Redis(host='127.0.0.1', port=6379, db=2, decode_responses=True)


@celery.task(bind=True)
def upload_files_task(self, project_path, file_type, parseDocMethod, temp_files):
    task_id = self.request.id
    handles = []
    # 锁的 key 按项目区分：同项目的任务互相排队，不同项目互不干扰
    lock_key = f'upload_lock:{project_path}'

    try:
        # ===== 包装 FileStorage（不用锁，各任务各转各的） =====
        files = []
        for item in temp_files:
            fh = open(item['temp_path'], 'rb')
            handles.append(fh)
            files.append(FileStorage(stream=fh, filename=item['origin_name']))

        # 排队时给用户个提示，不然卡片一直显示"排队等待中"也不知道在等啥
        _update_snapshot(task_id, state='PROCESSING', title='同项目有任务在处理，等待中...')

        # ===== 加锁：同一项目同一时刻只有一个任务在处理 =====
        # timeout：锁最长持有时间，防任务崩溃锁永远不释放（按你最坏的处理时长估，宁大勿小）
        # blocking_timeout：排队最长等待时间，超时就抛异常放弃
        with redis_client.lock(lock_key, timeout=7200, blocking_timeout=7200):
            # ===== 真正动共享资源（metadata.json、项目目录）的部分才放锁里 =====
            do_upload_files_logic(project_path, file_type, files, parseDocMethod, task=self)
        # with 结束自动释放锁

        _update_snapshot(task_id, state='SUCCESS', title='处理完成', is_running=0)
        return {'message': '处理完成'}

    except redis.exceptions.LockError:
        # 排队超时（前面任务跑了 2 小时还没完，基本就是出问题了）
        _update_snapshot(task_id, state='FAILURE', title='排队超时，请重试', is_running=0)
        raise

    except Exception as e:
        _update_snapshot(task_id, state='FAILURE', title=str(e), is_running=0)
        raise

    finally:
        for fh in handles:
            fh.close()
        if temp_files:
            shutil.rmtree(os.path.dirname(temp_files[0]['temp_path']), ignore_errors=True)

