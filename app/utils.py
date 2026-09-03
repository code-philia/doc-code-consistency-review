import os
import markdown
from bs4 import BeautifulSoup
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import io
import traceback
import uuid
import olefile
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.document import Document as _Document

from app.docx_utils import freeze_numbering
from callgraph.text_encoding import read_source_file
from doc2md import docToMd
import random
import sys
import zipfile
import xml.etree.ElementTree as ET

from typing import List, Dict, Optional, Tuple

_ILLEGAL_XML_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff\ufdd0-\ufddf\ufffe\uffff]')

def sanitize_xml(text):
    """过滤 XML 非法字符"""
    if not text:
        return text
    return _ILLEGAL_XML_CHARS.sub('', str(text))

def count_lines_of_code(filepath):
    """一个简单的代码行数统计函数，忽略空行"""
    try:
        return len([line for line in read_source_file(filepath).splitlines() if line.strip()])
    except OSError:
        # 如果文件无法读取或解码，则计为0
        return 0

def parse_markdown(md_text):
    """
    解析Markdown文本，提取需求点、表格和公式。
    """
    
    # 转换Markdown为HTML
    html = markdown.markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')
    
    requirements = []
    
    current_context = []
    grouped_content = ""

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'table']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if grouped_content.strip():
                requirements.append({
                    "type": "描述文本",
                    "id": f"text_{len(requirements)}",
                    "content": grouped_content.strip(),
                    "context": " > ".join(current_context)
                })
                grouped_content = ""
            # 更新标题上下文
            current_context = current_context[:int(element.name[1]) - 1] + [element.get_text()]
        elif element.name in ['p', 'li']:
            # 将段落和列表项内容累积到当前上下文
            grouped_content += element.get_text() + "\n"
        elif element.name == 'table':
            # 如果有未处理的内容，添加到需求点
            if grouped_content.strip():
                requirements.append({
                    "type": "描述文本",
                    "id": f"text_{len(requirements)}",
                    "content": grouped_content.strip(),
                    "context": " > ".join(current_context)
                })
                grouped_content = ""
            # 添加整个表格需求
            table_id = f"table_{len(requirements)}"
            requirements.append({
                "type": "表格",
                "id": table_id,
                "content": str(element),
                "context": " > ".join(current_context)
            })
            headers = [th.get_text() for th in element.find_all('th')]
            for i, row in enumerate(element.find_all('tr')[1:]):  # 跳过表头行
                cells = [td.get_text() for td in row.find_all('td')]
                requirements.append({
                    "type": "表格行",
                    "id": f"{table_id}_row_{i}",
                    "content": dict(zip(headers, cells)),
                    "context": " > ".join(current_context + [table_id])
                })

    # 如果有未处理的内容，添加到需求点
    if grouped_content.strip():
        requirements.append({
            "type": "描述文本",
            "id": f"text_{len(requirements)}",
            "content": grouped_content.strip(),
            "context": " > ".join(current_context)
        })

    # 解析公式
    formula_pattern = r'\$(.*?)\$|\$\$(.*?)\$\$'
    formulas = re.findall(formula_pattern, md_text, re.DOTALL)
    for k, formula_pair in enumerate(formulas):
        formula = formula_pair[0] if formula_pair[0] else formula_pair[1]
        requirements.append({
            "type": "公式",
            "id": f"formula_{k}",
            "content": f"${formula.strip()}$" if formula_pair[0] else f"$$ {formula.strip()} $$",
            "context": " > ".join(current_context)
        })
    
    return requirements


def split_code(filename, content, max_length=10000):
    """
    优化后的代码分块函数：
    1. 在分块前为每行代码添加行号
    2. 尽可能填充每个块直到接近最大长度
    3. 不拆分完整代码结构（函数/类等）
    4. 保持行完整性
    
    参数:
        filename: 文件名
        content: 代码内容
        max_length: 最大token长度限制
        
    返回:
        分块列表，每个元素包含:
        - filename: 文件名
        - start_line: 起始行号
        - end_line: 结束行号
        - content: 块内容
    """
    # 添加行号到每行代码
    lines = content.splitlines(keepends=True)
    numbered_lines = [f"{i + 1}: {line}" for i, line in enumerate(lines)]
    
    # encoder = tiktoken.get_encoding("cl100k_base")
    encoder = None # Mock encoder
    line_token_counts = [estimate_tokens(encoder, line) for line in numbered_lines]
    
    # 识别完整代码结构
    protected_blocks = identify_protected_blocks(content)
    
    chunks = []
    current_chunk = []
    current_token_count = 0
    current_start = 0  # 当前块起始行索引
    
    i = 0
    while i < len(numbered_lines):
        line = numbered_lines[i]
        token_count = line_token_counts[i]
        line_num = i + 1
        
        # 检查当前行是否属于某个受保护块
        block = find_enclosing_block(line_num, protected_blocks)
        
        # 情况1：遇到受保护块
        if block:
            block_start, block_end = block
            block_lines = numbered_lines[block_start-1:block_end]
            block_token_count = sum(line_token_counts[block_start-1:block_end])
            
            # 情况1a：当前块为空，直接添加整个受保护块
            if not current_chunk:
                chunks.append(create_chunk(filename, block_start, block_end, block_lines))
                i = block_end
                continue
            
            # 情况1b：添加受保护块会超出限制，先提交当前块
            elif current_token_count + block_token_count > max_length:
                chunks.append(create_chunk(filename, current_start+1, i, current_chunk))
                current_chunk = block_lines
                current_token_count = block_token_count
                current_start = block_start - 1
                i = block_end
            
            # 情况1c：可以添加到当前块
            else:
                current_chunk.extend(block_lines)
                current_token_count += block_token_count
                i = block_end
        
        # 情况2：普通行，添加后会超出限制
        elif current_token_count + token_count > max_length and current_chunk:
            chunks.append(create_chunk(filename, current_start+1, i, current_chunk))
            current_chunk = [line]
            current_token_count = token_count
            current_start = i
            i += 1
        
        # 情况3：可以添加到当前块
        else:
            current_chunk.append(line)
            current_token_count += token_count
            i += 1
    
    # 添加最后一个块
    if current_chunk:
        chunks.append(create_chunk(filename, current_start+1, len(numbered_lines), current_chunk))
    
    return chunks

def identify_protected_blocks(content):
    """识别需要保护的代码块范围（起始行，结束行）"""
    blocks = []
    
    # 函数定义
    for match in re.finditer(r'\b[\w:<>]+\s+\w+\s*\([^)]*\)\s*\{', content):
        start_line = content[:match.start()].count('\n') + 1
        end_line = find_matching_brace(content, match.end()-1)
        if end_line > 0:
            blocks.append((start_line, end_line))
    
    # 类/结构体定义
    for match in re.finditer(r'\b(class|struct)\s+\w+\s*\{', content):
        start_line = content[:match.start()].count('\n') + 1
        end_line = find_matching_brace(content, match.end()-1)
        if end_line > 0:
            blocks.append((start_line, end_line))
    
    # 命名空间
    for match in re.finditer(r'\bnamespace\s+\w+\s*\{', content):
        start_line = content[:match.start()].count('\n') + 1
        end_line = find_matching_brace(content, match.end()-1)
        if end_line > 0:
            blocks.append((start_line, end_line))
    
    return blocks

def find_enclosing_block(line_num, blocks):
    """检查行是否属于某个受保护块"""
    for start, end in blocks:
        if start <= line_num <= end:
            return (start, end)
    return None

def estimate_tokens(encoder, line):
    # Fallback when tiktoken is not available: approx 1 token = 4 chars
    return len(line) // 4 + 1

def find_matching_brace(content, open_pos):
    """找到匹配的闭括号行号"""
    stack = 1
    pos = open_pos + 1
    while pos < len(content) and stack > 0:
        if content[pos] == '{':
            stack += 1
        elif content[pos] == '}':
            stack -= 1
        pos += 1
    return content[:pos].count('\n') + 1 if stack == 0 else -1

def create_chunk(filename, start, end, lines):
    """创建分块字典"""
    return {
        "filename": filename,
        "start_line": start,
        "end_line": end,
        "content": "".join(lines)
    }
    

def get_all_files_with_relative_paths(base_path, type = 'code'):
    """递归遍历目录，获取所有文件的相对路径"""
    all_files = []
    adb_files = []
    # 从FPGA代码中去除'.adb', '.ads'
    flag_FPGA = False
    include_files_FPGA = ['.vhd', '.v', '.sv']
    include_files_Ada = ['.adb', '.ads']
    # 基于文件名后缀，指定文件类型
    for root, _, files in os.walk(base_path):
        for file in files:
            if type == 'code' and not (file.endswith('.py') or file.endswith('.java') or file.endswith('.cpp') or file.endswith('.hpp') or file.endswith('.js') or file.endswith('.c') or file.endswith('.h') or file.endswith('.vhd') or file.endswith('.v') or file.endswith('.sv') or file.endswith('.adb') or file.endswith('.ads')):
                continue
            if type == 'doc' and not (file.endswith('.docx') or file.endswith('.md')):
                continue
            if os.path.splitext(file)[1] in include_files_FPGA:
                flag_FPGA = True
            # 从FPGA代码中去除'.adb', '.ads'
            if flag_FPGA and os.path.splitext(file)[1] in include_files_Ada:
                continue
            relative_path = os.path.relpath(os.path.join(root, file), base_path)
            
            if os.path.splitext(file)[1] in include_files_Ada:
                adb_files.append(relative_path)
            else:
                all_files.append(relative_path)
    if not flag_FPGA:
        all_files.extend(adb_files)
    return all_files

def convert_doc_to_markdown(doc_repo_path, parseDocMethod):
    converted_repo_path = os.path.join(os.path.dirname(doc_repo_path), "doc_repo_converted")
    os.makedirs(converted_repo_path, exist_ok=True)
    
    for root, _, files in os.walk(doc_repo_path):
        for file in files:
            if not file.endswith('.docx'):
                continue
            
            file_name_prefix = os.path.splitext(file)[0]
            converted_md_path = os.path.join(converted_repo_path, file_name_prefix, file_name_prefix + '.md')
            
            if os.path.exists(converted_md_path):
                continue
            #freeze_numbering(os.path.join(root, file))
            docToMd.convertDocToMarkdown(os.path.join(root, file), converted_repo_path, parseDocMethod)
            
def convert_docfile_to_markdown(doc_file_path, doc_repo_path, parseDocMethod):
    converted_repo_path = os.path.join(os.path.dirname(doc_repo_path), "doc_repo_converted")
    os.makedirs(converted_repo_path, exist_ok=True)
    docToMd.convertDocToMarkdown(doc_file_path, converted_repo_path, parseDocMethod)


    
def get_filename_without_extension(filepath):
    """获取不带扩展名的文件名"""
    return os.path.splitext(os.path.basename(filepath))[0]

# def chunk_list(lst, limit):
#     result = []
#     for i in range(0, len(lst), limit):
#         result.append(lst[i:i+limit])
#     return result

def chunk_list(code_blocks, max_chunk_size):
    all_chunks = []
    current_chunk = []
    max_chunk_size = 30000

    # 尝试随机打乱list中元素顺序
    # if random_flag:
        # random.shuffle(code_blocks)
    
    for block in code_blocks:
        block_str_len = len(str(block))
        
        # 某个块本身就大于max_chunk_size
        if block_str_len >= max_chunk_size:
            # 当前有未保存的chunk，先保存
            if len(current_chunk) > 0:
                all_chunks.append(current_chunk)
                current_chunk = []
            
            # 将这个大块独立为一个chunk，最后到模型输入处截断
            all_chunks.append([block])

        # 正常块，先尝试放入
        potential_chunk = current_chunk + [block]
        potential_chunk_len = len(str(potential_chunk))

        if potential_chunk_len < max_chunk_size:
            current_chunk.append(block)
        else: #放入后超了max_chunk_size，放入新的chunk
            if len(current_chunk) > 0:
                all_chunks.append(current_chunk)
            current_chunk = [block]

    if len(current_chunk) > 0:
        all_chunks.append(current_chunk)

    print(f"共有{len(code_blocks)}个代码块，分{len(all_chunks)}次询问模型")
    return all_chunks



def replace_text_in_docx(doc, replacements, export_type):
    """在DOCX文档中替换文本，保持原有格式"""
    
    def replace_text_in_runs(paragraph, old_text, new_text):
        """在段落的runs中替换文本，保持格式"""
        modified_runs = []
        full_text = paragraph.text
        if old_text not in full_text:
            return False
        
        # 找到替换位置
        start_pos = full_text.find(old_text)
        end_pos = start_pos + len(old_text)
        
        # 遍历runs，找到包含目标文本的runs
        current_pos = 0
        runs_to_modify = []
        
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            # 检查这个run是否与目标文本有重叠
            if run_start < end_pos and run_end > start_pos:
                runs_to_modify.append({
                    'index': i,
                    'run': run,
                    'start': run_start,
                    'end': run_end,
                    'text': run.text
                })
            
            current_pos = run_end
        
        if not runs_to_modify:
            return False
        
        # 执行替换
        # 如果替换文本完全在一个run内
        if len(runs_to_modify) == 1:
            run_info = runs_to_modify[0]
            run = run_info['run']
            relative_start = start_pos - run_info['start']
            relative_end = end_pos - run_info['start']
            
            new_run_text = (run.text[:relative_start] + 
                           new_text + 
                           run.text[relative_end:])
            run.text = sanitize_xml(new_run_text)
            modified_runs.append(run)
        else:
            # 替换文本跨越多个runs
            for i, run_info in enumerate(runs_to_modify):
                run = run_info['run']
                
                if i == 0:  # 第一个run
                    relative_start = start_pos - run_info['start']
                    run.text = run.text[:relative_start] + new_text
                    modified_runs.append(run)
                elif i == len(runs_to_modify) - 1:  # 最后一个run
                    relative_end = end_pos - run_info['start']
                    run.text = run.text[relative_end:]
                else:  # 中间的runs
                    run.text = ""
        
        return modified_runs

    all_modified_runs = []
    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                modified = replace_text_in_runs(paragraph, old_text, new_text)
                all_modified_runs.extend(modified)
    
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            modified = replace_text_in_runs(paragraph, old_text, new_text)
                            all_modified_runs.extend(modified)

    # 统一给被修改的run设置字体
    seen = set()
    for run in all_modified_runs:
        if id(run) not in seen:
            seen.add(id(run))
            # 设置中文字体
            run.font.name = 'Times New Roman'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            if export_type == 'issue':
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), '宋体(中文正文)')
            else:
                rFonts.set(qn('w:ascii'), '等线(西文正文)')
                rFonts.set(qn('w:hAnsi'), '等线(西文正文)')
                rFonts.set(qn('w:eastAsia'), '等线(中文正文)')


def generate_issue_content(issue, form_data):
    """生成问题单文件内容"""
    content = []
    
    # 表单信息
    content.append("=" * 50)
    content.append("问题单信息")
    content.append("=" * 50)
    content.append(f"被测产品名称: {form_data.get('productName', '')}")
    content.append(f"问题单标识: {form_data.get('issueId', '')}")
    content.append(f"被测产品标识: {form_data.get('productId', '')}")
    content.append(f"发现手段: {form_data.get('discoveryMethod', '')}")
    content.append(f"问题追踪: {form_data.get('issueTracking', '')}")
    content.append(f"问题类别: {', '.join(form_data.get('issueCategories', []))}")
    content.append(f"问题级别: {issue.get('level', '')}")
    content.append("")
    
    # 问题单详细信息
    content.append("=" * 50)
    content.append("问题详情")
    content.append("=" * 50)
    content.append(f"问题ID: {issue.get('id', '')}")
    content.append(f"问题摘要: {issue.get('summary', '')}")
    content.append(f"问题级别: {issue.get('level', '')}")
    content.append(f"相关文档: {issue.get('relatedDocFile', '')}")
    content.append(f"相关代码: {issue.get('relatedCodeFile', '')}")
    content.append(f"创建时间: {issue.get('createdDate', '')}")
    content.append(f"状态: {issue.get('status', '')}")
    content.append("")
    
    # 问题描述
    content.append("=" * 50)
    content.append("问题描述")
    content.append("=" * 50)
    description = issue.get('description', '暂无描述')
    content.append(description)
    content.append("")
    
    return "\n".join(content)

def include_related_blocks(related_code, all_code_blocks):
    """
    检查查询结果中每个 code_block 的 related_id，将缺失的相关 block 加入结果
    
    参数:
        related_code: 查询返回的相关代码块列表，格式为 [{"file": "xxx", "range": [start, end]}, ...]
        all_code_blocks: 所有代码块的列表，包含完整的代码块信息
        
    返回:
        包含相关块的完整代码块列表
    """
    def normalize_related_line_range(value):
        if isinstance(value, dict):
            start = value.get('startLine') or value.get('start_line') or value.get('start')
            end = value.get('endLine') or value.get('end_line') or value.get('end') or start
            return int(start), int(end)
        if isinstance(value, (list, tuple)):
            if len(value) >= 2:
                return int(value[0]), int(value[1])
            if len(value) == 1:
                return int(value[0]), int(value[0])
        return int(value), int(value)

    # 创建一个字典，以便快速查找代码块
    block_dict = {}
    for block in all_code_blocks:
        # 使用文件名和行号范围作为唯一标识
        key = (block['file'], tuple(block['range']))
        block_dict[key] = block
    
    # 首先根据 file 和 range 匹配到完整的代码块
    matched_blocks = []
    for related_item in related_code:
        file_name = related_item['file']
        range_info = related_item['range']
        key = (file_name, tuple(range_info))
        
        if key in block_dict:
            matched_blocks.append(block_dict[key])
        else:
            # 如果精确匹配失败，尝试找到包含该范围的代码块
            for block in all_code_blocks:
                if (block['file'] == file_name and 
                    block['range'][0] <= range_info[0] and 
                    block['range'][1] >= range_info[1]):
                    matched_blocks.append(block)
                    break
                
                # 如果精确匹配失败，尝试找到小于该范围的代码块
                elif (block['file'] == file_name and 
                    block['range'][0] >= range_info[0] and 
                    block['range'][1] <= range_info[1]):
                    matched_blocks.append(block)

    
    # 创建结果集合，避免重复
    result_blocks = []
    added_keys = set()
    
    # 添加匹配到的代码块
    for block in matched_blocks:
        key = (block['file'], tuple(block['range']))
        if key not in added_keys:
            result_blocks.append(block)
            added_keys.add(key)
    
    # 仅检查匹配到的代码块的一层 related_id，不递归
    for matched_block in matched_blocks:
        # 检查当前块的 related_id
        if 'related_id' in matched_block and matched_block['related_id']:
            for related_id in matched_block['related_id']:
                # 在所有代码块中查找对应的 related_id
                for block in all_code_blocks:
                    if str(block.get('id')) == str(related_id):
                        # key = (block['file'], tuple(block['range']))
                        # if key not in added_keys:
                            
                        # 按照related_range添加具体的代码片段
                        related_range_map = matched_block.get('related_range') or {}
                        related_value = related_range_map.get(str(block['id']))
                        if related_value is None:
                            related_value = block.get('range')
                        try:
                            related_start, related_end = normalize_related_line_range(related_value)
                        except Exception:
                            related_start, related_end = block['range'][0], block['range'][1]

                        related_start = max(int(block['range'][0]), int(related_start))
                        related_end = min(int(block['range'][1]), int(related_end))
                        if related_start > related_end:
                            related_start, related_end = int(block['range'][0]), int(block['range'][1])

                        temp_block = {}
                        temp_block['file'] = block['file']
                        temp_block['range'] = [related_start, related_end]

                        key = (temp_block['file'], tuple(temp_block['range']))
                        if key in added_keys:
                            break

                        temp_block['type'] = block['type']
                        # # 按照\n切分
                        temp_list = block['code'].split('\n')
                                      
                        pos_start = temp_block['range'][0] - block['range'][0]
                        pos_end = temp_block['range'][1] - block['range'][0]
                        
                        temp_block['code'] = '\n'.join(temp_list[pos_start:pos_end+1])
                        
                        
                        temp_block['id'] = block['id']
                        temp_block['related_id'] = []
                        temp_block['related_id'].append(matched_block['id'])
                        temp_block['related_range'] = {str(matched_block['id']): matched_block.get('range')}
                        
                        
                        # print(key)
                        # print(temp_block)
                        result_blocks.append(temp_block)
                        
                        # print(block)
                        # sys.exit()
                        # result_blocks.append(block)
                        added_keys.add(key)
                        break
    
    return result_blocks

def split_markdown_to_blocks(md_content):
    """
    将Markdown内容按标题分解为块，返回包含内容和字符偏移量的块信息
    参考split_markdown.py的实现逻辑
    """
    original_content = md_content
    target_flag = '# <div class="page"></div>'
    content_start_offset = 0
    
    # 如果存在目标标志，从标志后开始处理
    if target_flag in md_content:
        parts = md_content.split(target_flag, 1)
        if len(parts) > 1:
            content_start_offset = len(parts[0]) + len(target_flag)
            md_content = parts[1]
    
    # 匹配所有标题
    heading_pattern = re.compile(r'^(#{1,}) (.*)$', re.M)
    headings = [(match.start(), len(match.group(1)), match.group(0))
                for match in heading_pattern.finditer(md_content)]
    
    blocks = []
    n = len(headings)
    
    # 处理每个标题间的内容，筛选叶子标题
    for i in range(n):
        cur_pos, cur_level, cur_heading = headings[i]
        # 确定当前标题的结束位置
        next_pos = headings[i+1][0] if (i+1 < n) else len(md_content)
        # 提取当前标题到下一个标题之间的内容
        content = md_content[cur_pos:next_pos].strip()
        
        if '\n' not in content:  # 只有标题没有内容，不保存
            continue
            
        # 计算在原始文档中的字符偏移量
        actual_start = content_start_offset + cur_pos
        actual_end = content_start_offset + cur_pos + len(content)
        
        blocks.append({
            'content': content,
            'start': actual_start,
            'end': actual_end
        })
    
    return blocks

class DocxNativeParser:
    """
    【通用兼容版解析器】
    保留这个强大的解析器，因为它能把 VML/文本框 里的代码“挖”出来变成普通的 para 段落。
    这样后续的 parse_programming_rules 才能像处理普通文字一样处理代码。
    """
    def __init__(self, filepath, debug=False):
        self.filepath = filepath
        self.debug = debug
        self.elements = [] 
        if self.debug:
            print(f"\n[Parser] 🚀 开始解析文件: {filepath}")
        self._parse()

    def _get_tag_name(self, element):
        return element.tag.split('}')[-1]

    def _extract_text_from_node_recursive(self, node):
        text_parts = []
        tag = self._get_tag_name(node)

        # 核心：遇到文本框 (txbxContent) -> 强制提取为独立代码块
        if tag == 'txbxContent':
            code_lines = []
            for child in node:
                if self._get_tag_name(child) == 'p':
                    line_text = self._get_pure_text(child)
                    if line_text.strip():
                        code_lines.append(line_text)
            if code_lines:
                # 前后加换行，确保独立
                return ["\n" + "\n".join(code_lines) + "\n"]
            return []

        if tag == 't' and node.text:
            text_parts.append(node.text)
        
        for child in node:
            text_parts.extend(self._extract_text_from_node_recursive(child))
        return text_parts

    def _get_pure_text(self, node):
        text = ""
        tag = self._get_tag_name(node)
        if tag == 't' and node.text:
            text += node.text
        for child in node:
            text += self._get_pure_text(child)
        return text

    def _process_paragraph_node(self, p_node):
        parts = self._extract_text_from_node_recursive(p_node)
        return "".join(parts)

    def _parse(self):
        try:
            with zipfile.ZipFile(self.filepath, 'r') as z:
                xml_content = z.read('word/document.xml')
            
            tree = ET.fromstring(xml_content)
            body = None
            for elem in tree.iter():
                if self._get_tag_name(elem) == 'body':
                    body = elem
                    break
            
            if body is None:
                return

            para_count = 0
            table_count = 0

            for child in body:
                tag = self._get_tag_name(child)
                
                # 处理段落
                if tag == 'p': 
                    text = self._process_paragraph_node(child)
                    if text.strip(): 
                        self.elements.append({'type': 'para', 'text': text})
                        para_count += 1
                
                # 处理表格 (问题单需要这个)
                elif tag == 'tbl':
                    table_rows = []
                    for row in child:
                        if self._get_tag_name(row) != 'tr': continue
                        row_cells = []
                        for cell in row:
                            if self._get_tag_name(cell) != 'tc': continue
                            cell_text_parts = []
                            for p in cell:
                                if self._get_tag_name(p) == 'p':
                                    p_text = self._process_paragraph_node(p)
                                    if p_text: cell_text_parts.append(p_text)
                            full_cell_text = "\n".join(cell_text_parts)
                            row_cells.append(full_cell_text.strip())
                        table_rows.append(row_cells)
                    
                    if table_rows:
                        self.elements.append({'type': 'table', 'rows': table_rows})
                        table_count += 1
            
            if self.debug:
                print(f"[Parser] ✅ 解析完成. 段落: {para_count}, 表格: {table_count}")

        except Exception as e:
            print(f"[Parser] ❌ 解析 XML 出错: {e}")

    def get_full_text(self):
        lines = []
        for el in self.elements:
            if el['type'] == 'para':
                lines.append(el['text'])
            elif el['type'] == 'table':
                for row in el['rows']:
                    lines.append(" | ".join(row))
        return "\n".join(lines)


def read_docx_text(path):
    parser = DocxNativeParser(path)
    return parser.get_full_text()


# ==========================================
# 1. 编程规则解析 (GJB8114逻辑)
# ==========================================
def parse_programming_rules(docx_path, debug=True):
    print("="*60)
    print(f"🛠️  正在解析编程规则: {docx_path}")
    
    # 使用增强版 Parser，这样文本框里的代码会被提取为 text
    parser = DocxNativeParser(docx_path, debug=debug)
    rules = []
    
    current_rule = {}
    capturing_code = None 
    
    # 使用你提供的正则：匹配 R-1-1-1
    rule_id_pattern = re.compile(r'(R-\d+-\d+-\d+([A-Z]?))') 
    
    # 关键词
    violation_keys = ["违背示例", "错误示例", "违背", "不符合", "违背 1", "违背 2", "提示 1", "提示 2"]
    compliance_keys = ["遵循示例", "正确示例", "遵循", "符合"]

    for i, el in enumerate(parser.elements):
        # 按照你的要求，规则解析只看段落
        if el['type'] != 'para': continue
            
        text = el['text'].strip()
        if not text: continue
        
        # 1. 尝试匹配规则 ID
        match = rule_id_pattern.search(text)
        
        if match:
            if current_rule:
                rules.append(current_rule)
                if debug: print(f"   💾 [保存上一条] ID: {current_rule['id']}")
            
            rule_id = match.group(1)
            #description = text
            next_el = (parser.elements)[i+1]
            description = next_el['text'].strip()
            # 取最后一行，防止出现多余的代码
            description = description.splitlines()
            description = description[-1]
            # 过滤掉异常字符
            description = description.replace("违背示例：","") 
            #print(description)
            
            #if debug: print(f"🟢 [发现新规则] ID: {rule_id} | 描述: {description[:20]}...")
            if debug: print(f"🟢 [发现新规则] ID: {rule_id} | 描述: {description}.")

            current_rule = {
                "id": rule_id,
                "description": description, 
                "violation_code": "",
                "compliance_code": ""
            }
            capturing_code = None
            continue
            
        if not current_rule: continue

        # 2. 识别示例标记
        # 注意：Parser 可能会把代码提取在“违背示例”同一行的后面
        # 所以我们得检查这一行除去关键词后，是否还有残留内容（代码）
        
        is_mode_switch = False
        
        # 检查违背
        for k in violation_keys:
            if k in text:
                capturing_code = 'violation'
                if debug: print(f"   ⚠️  [进入模式] 违背示例")
                # 【重要修改】检查同行的残留代码
                code_part = text.replace(k, "", 1).replace(":", "").replace("：", "").strip()
                if len(code_part) > 1: # 如果剩下的内容够长，说明代码跟在后面了
                    current_rule['violation_code'] += code_part + "\n"
                is_mode_switch = True
                break
        
        if is_mode_switch: continue

        # 检查遵循
        for k in compliance_keys:
            if k in text:
                capturing_code = 'compliance'
                if debug: print(f"   ✅ [进入模式] 遵循示例")
                # 【重要修改】检查同行的残留代码
                code_part = text.replace(k, "", 1).replace(":", "").replace("：", "").strip()
                if len(code_part) > 1:
                    current_rule['compliance_code'] += code_part + "\n"
                is_mode_switch = True
                break
        
        if is_mode_switch: continue
            
        # 3. 提取代码 (只有处于模式中才提取)
        if capturing_code == 'violation':
            current_rule['violation_code'] += text + "\n"
        elif capturing_code == 'compliance':
            current_rule['compliance_code'] += text + "\n"
            
    if current_rule:
        rules.append(current_rule)
        
    print(f"🏁 规则解析结束. 共提取到 {len(rules)} 条规则.")
    return rules


# ==========================================
# 2. 问题单解析 (保持当前版本，依赖表格)
# ==========================================
def parse_issue_reports(docx_path, debug=True):
    if debug:
        print("="*60)
        print(f"🛠️  正在解析问题单: {docx_path}")

    # 问题单依赖表格结构
    parser = DocxNativeParser(docx_path, debug=debug)
    
    all_issues = []     
    current_issue = {}  
    
    key_map = {
        "问题描述": "desc", "问题内容": "desc",
        "处理意见": "opinion", "改正措施": "opinion", "修改建议": "opinion",
        "问题追踪": "trace_id", "用例标识": "trace_id",
        "问题单标识": "id", "问题标识": "id", "问题单号": "id"
    }

    for el in parser.elements:
        # 只看表格
        if el['type'] != 'table': continue
            
        table_rows = el['rows']
        
        for row in table_rows:
            for i in range(len(row)):
                cell_text = row[i].replace(" ", "").replace(":", "").replace("：", "").strip()
                
                # 检查 Key
                found_key = None
                for k_word, k_field in key_map.items():
                    if k_word in cell_text and len(cell_text) < 20: 
                        found_key = k_field
                        break
                
                if found_key:
                    if found_key == "id":
                        if current_issue and ("id" in current_issue or "desc" in current_issue):
                            all_issues.append(current_issue)
                            current_issue = {} 

                    # 提取 Value (优先取右侧单元格)
                    value = ""
                    if i + 1 < len(row):
                        value = row[i+1].strip()
                    
                    # 简单的防错位检查
                    is_next_cell_key = False
                    clean_val = value.replace(" ", "").replace(":", "")
                    for k in key_map.keys():
                        if k in clean_val and len(clean_val) < 20:
                            is_next_cell_key = True
                            break
                    
                    if value and not is_next_cell_key:
                        # 清理 Value 中的 Key 前缀
                        for k_word in key_map.keys():
                            if value.startswith(k_word):
                                value = value.replace(k_word, "", 1).lstrip("：: ")
                        
                        if found_key in current_issue:
                            current_issue[found_key] += "\n" + value
                        else:
                            current_issue[found_key] = value

    if current_issue.get("desc") or current_issue.get("opinion") or current_issue.get("id"):
        if "id" not in current_issue:
            current_issue["id"] = f"issue_autogen_{len(all_issues)+1}"
        all_issues.append(current_issue)

    return all_issues

# ==========================================
# 3. 典型案例解析
# ==========================================    

# ============ 可配置项 ============
CASE_HEADING_LEVEL = 3          # 案例标题级别（导航里的第3级）
SECTION_HEADING_LEVEL = 4       # 小节标题级别（1 概述 / 2 机理分析 ...）
DISCARD_SECTION_KEYWORD = '启示'  # 小节标题包含该词 → 整节丢弃
NS_OFFICE = 'urn:schemas-microsoft-com:office:office'

SECTION_KEYWORDS = ['概述', '机理分析', '纠正措施', '启示和建议']

# ============ 基础工具 ============
def iter_block_items(parent):
    """按文档实际顺序遍历段落和表格"""
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def get_heading_level(para, doc=None):
    """
    返回标题级别数字，非标题返回 None。
    逻辑：
    1. 优先匹配样式名称 (Heading 3, 标题 3 等)。
    2. 如果样式名不匹配，检查段落的大纲级别 (Outline Level)。
       - 大纲级别 0 = 正文 (通常不视为标题，除非用户特意设为标题)
       - 大纲级别 1-9 = 对应 Level 1-9
       - 大纲级别 10+ = 正文 (None)
    """
    # 1. 优先检查样式名称
    name = (para.style.name or '').strip()
    m = re.match(r'^(?:heading|标题)\s*(\d+)$', name, re.IGNORECASE)
    if m and m is not None:     
        return int(m.group(1))
    
    # 2. 如果样式名不匹配，检查大纲级别
    if doc is None:
        return None

    # 获取样式对象
    style = para.style
    # 检查样式是否有定义 (有些内置样式可能没有独立的 style element)
    if style is None or style._element is None:
        return None

    # 解析样式定义的 XML
    style = para.style
    if not style or not style._element:
        return None

    style_elm = style._element
    style_name = (para.style.name or '').strip()
    # 打印调试信息：查看样式定义的 XML 结构
    #from lxml import etree
    #xml_str = etree.tostring(style_elm, encoding='utf-8').decode('utf-8')
    #print(f"样式 '{style_name}' 的 XML:\n{xml_str[:500]}...\n")

    # 在样式定义的 XML 中搜索 outlineLevel
    # 遍历所有子节点
    for elem in style_elm.iter():
        if elem.tag == qn('w:outlineLevel'):
            val_str = elem.get(qn('w:val'))
            if val_str:
                try:
                    level = int(val_str)
                    # 验证级别范围
                    if 1 <= level <= 9:
                        return level
                except ValueError:
                    pass

    # --- 策略 4: 兜底逻辑 ---
    if style_name == "RC标题":
        # 假设 RC标题 是一级标题
        return 1
    
    return None

def extract_textbox_text(para):
    """提取段落中嵌入的文本框文字（python-docx 默认读不到文本框）"""
    lines = []
    for txbx in para._p.findall('.//' + qn('w:txbxContent')):
        for p in txbx.findall(qn('w:p')):
            line = ''.join(t.text or '' for t in p.findall('.//' + qn('w:t')))
            lines.append(line)
    return '\n'.join(lines).strip()


def table_to_text(table):
    """表格转纯文本（正文里的表格才会用到）"""
    return '\n'.join(
        ' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells)
        for row in table.rows
    )


def extract_from_ole10native(raw):
    """从 Ole10Native 流里抠原始文件： vsdx 是 zip(PK 头), 老 vsd 是 OLE 复合文件 (DOCF 头)"""
    idx = raw.find(b'PK\x03\x04')
    if idx != -1:
        return raw[idx:], '.vsdx'
    idx = raw.find(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')
    if idx != -1:
        return raw[idx:], '.vsd'
    return raw, '.bin'


def save_ole_blob(blob, save_dir, name_prefix):
    """保存 OLE 对象，Visio 对象尽量还原成 .vsd/.vsdx，否则存 .bin"""
    os.makedirs(save_dir, exist_ok=True)
    data, ext = blob, '.bin'
    try:

        bio = io.BytesIO(blob)
        if olefile.isOleFile(bio):
            ole = olefile.OleFileIO(bio)
            streams = ['/'.join(s) for s in ole.listdir()]
            # print('OLE streams:', streams)

            # 嵌入的 Visio 数据，就是完整 vsd
            if any('VisioDocument' in s or 'Visio Document' in s for s in streams):
                # 嵌入的 Visio 对象: 整个OLE 容器本身就是 vsd 格式, 直接存
                data, ext = blob, '.vsd'
            elif 'Package' in streams:           # 少数情况内嵌完整包
                pkg = ole.openstream('Package').read()
                data, ext = pkg, ('.vsdx' if pkg[:2] == b'PK' else '.bin')
            else:
                # Ole10Native 包装: 抠出里面的原始文件
                native = next((s for s in streams if s.endswith('Ole10Native')), None)
                if native:
                    data, ext = extract_from_ole10native(ole.openstream(native).read())
            ole.close()
    except ImportError:
        traceback.print_exc()
        pass  # 没装 olefile，按 .bin 原样保存
    except Exception as e:
        print(f'OLE 对象解析失败，按 .bin 保存: {e}')

    fname = f'{name_prefix}{ext}'
    with open(os.path.join(save_dir, fname), 'wb') as f:
        f.write(data)
    return fname


def extract_ole_objects(para, doc_part, save_dir, doc_tag, counter):
    """提取段落中的 OLE 嵌入对象（Visio），返回 (文件名列表, 新counter)"""
    saved = []
    for ole in para._p.findall('.//{%s}OLEObject' % NS_OFFICE):
        rid = ole.get(qn('r:id'))
        if not rid:
            continue
        try:
            blob = doc_part.rels[rid].target_part.blob
        except KeyError:
            continue
        progid = ole.get('ProgID', '')
        kind = 'visio' if 'visio' in progid.lower() else 'ole'
        fname = save_ole_blob(blob, save_dir, f'{doc_tag}_{kind}_{counter}')
        saved.append(fname)
        counter += 1
    return saved, counter


# ============ 主解析函数 ============
def parse_typical_cases(source, visio_dir='visio_files'):
    """
    解析案例集 docx
    source: 文件路径 或 文件流（Flask 上传直接传 file.stream）
    visio_dir: Visio 对象保存目录
    返回: [{'title': 案例标题, 'content': 内容, 'visio_files': [...]}, ...]
    """
    doc = Document(source)
    doc_tag = uuid.uuid4().hex[:6]   # 防止批量解析时 Visio 文件重名覆盖
    cases = []
    current = None              # 当前案例
    started = False             # 是否已进入第一个案例（之前的内容全部跳过：目录/封面/篇章标题）
    before_keywords = False     # 是否在 案例标题 ~ 关键词 之间
    skip_section = False        # 是否在"启示"小节里
    current_section = ''        # 当前小节标题文本
    ole_counter = 1

    for block in iter_block_items(doc):

        if isinstance(block, Paragraph):
            level = get_heading_level(block)

            # --- 案例标题：开启新案例 ---
            if level == CASE_HEADING_LEVEL:
                if current:
                    cases.append(current)
                current = {'title': block.text.strip(), 'parts': [], 'visio_files': []}
                started, before_keywords = True, True
                skip_section, current_section = False, ''
                continue

            if not started:
                continue  # 目录、封面、篇章大标题全部跳过

            # --- 小节标题 ---
            if level == SECTION_HEADING_LEVEL:
                sec = block.text.strip()
                before_keywords = False
                current_section = sec
                if DISCARD_SECTION_KEYWORD in sec:
                    skip_section = True        # 启示小节：整节丢弃
                    continue
                skip_section = False
                current['parts'].append(sec)   # 其他小节标题保留进 content
                continue

            # --- 其他级别标题（上篇/第一部分等）忽略 ---
            if level is not None:
                continue

            if skip_section:
                continue

            # --- 普通段落 ---
            text = block.text.strip()
            if text:
                if '【关键词】' in text:
                    before_keywords = False
                current['parts'].append(text)   # 简介、关键词、正文、图注都走这里

            # 文本框（代码段）：按所在小节命名为 违背示例 / 遵循示例
            tb_text = extract_textbox_text(block)
            if tb_text:
                if '机理分析' in current_section:
                    current['parts'].append('违背示例')
                elif '纠正措施' in current_section:
                    current['parts'].append('遵循示例')
                # 其他小节的文本框：不加标题，直接放内容
                current['parts'].append(tb_text)

            # Visio OLE 对象：保存文件 + content 留占位
            saved, ole_counter = extract_ole_objects(block, doc.part, visio_dir, doc_tag, ole_counter)
            for fname in saved:
                current['parts'].append(f'[Visio对象: {fname}]')
                current['visio_files'].append(fname)

        elif isinstance(block, Table):
            if not started or skip_section:
                continue
            if before_keywords:
                continue                        # 标题和关键词之间的表格不要
            current['parts'].append(table_to_text(block))

    if current:
        cases.append(current)

    for c in cases:
        c['content'] = '\n'.join(c.pop('parts'))
    return cases


def probe(file_path):
    doc = Document(file_path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            print(f'{p.style.name:<15} | {t[:40]}')


# ============ 新解析函数：parse_typical_cases_new ============
def parse_typical_cases_new(source, visio_dir='visio_files_new'):
    """
    专门处理备用逻辑文档的解析函数
    逻辑：
    1. 先扫描文档，确认是否存在标准 Level 3 标题。
       - 如果存在：直接返回空列表（说明这是标准文档，应由 parse_typical_cases 处理）。
       - 如果不存在：继续判断是 B1 还是 B2 模式。
    2. 场景 B1 (全一级): Level 1 标题中，靠关键词区分案例和小节。
    3. 场景 B2 (一级案例+二级小节): Level 1 是案例，Level 2 是小节。
    """
    doc = Document(source)
    doc_tag = uuid.uuid4().hex[:6]
    
    # --- 第一步：探测文档结构 ---
    has_l3 = False
    has_l1 = False
    has_l2 = False

    # 快速扫描，只检查标题级别
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block, doc)
            if level == 3: has_l3 = True
            elif level == 1: has_l1 = True
            elif level == 2: has_l2 = True
    
    # 确定解析模式
    mode = None
    case_level = None
    section_level = None
    use_keywords = False

    if has_l1 and not has_l2:
        # 场景 B1: 只有 Level 1
        mode = 'B1_AllLevel1'
        case_level = 1
        section_level = 1
        use_keywords = True
    elif has_l1 and has_l2:
        # 场景 B2: Level 1 和 Level 2 共存
        mode = 'B2_L1Case_L2Section'
        case_level = 1
        section_level = 2
        use_keywords = False
    else:
        print(f"[{os.path.basename(source)}] 无法识别备用逻辑结构 (L1:{has_l1}, L2:{has_l2})")
        return []

    print(f"[{os.path.basename(source)}] 启用备用逻辑解析 -> 解析模式: {mode}")

    # --- 第二步：正式解析 ---
    cases = []
    current = None
    started = False
    before_keywords = False
    skip_section = False
    current_section = ''
    ole_counter = 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block, doc)
            text = block.text.strip()
            #print(text,level)
            # --- 标题处理逻辑 ---
            if level is not None:
                is_case = False
                is_section = False
                
                if use_keywords and level == case_level:
                    # 场景 B1: 所有标题都是 L1，靠关键词区分
                    text_type = classify_text_type(text)
                    if text_type == 'case':
                        is_case = True
                    elif text_type == 'section':
                        is_section = True
                    
                else:
                    # 场景 B2: 直接按层级判断
                    if level == case_level:
                        is_case = True
                    elif level == section_level:
                        is_section = True

                # 处理案例标题
                if is_case:

                    if current:
                        cases.append(current)
                    current = {'title': text, 'parts': [], 'visio_files': []}
                    started, before_keywords = True, True
                    skip_section, current_section = False, ''
                    continue
                
                # 处理小节标题
                if is_section:
                    if not started:
                        continue # 案例还没开始，忽略小节
                    before_keywords = False
                    current_section = text
                    if DISCARD_SECTION_KEYWORD in text:
                        skip_section = True
                        continue
                    skip_section = False
                    current['parts'].append(text)
                    continue
            
            # --- 非标题段落处理 ---
            if not started:
                continue
            
            if skip_section:
                continue

            # 普通段落
            if text:
                if '【关键词】' in text:
                    before_keywords = False
                if current:
                    current['parts'].append(text)

            # 文本框
            tb_text = extract_textbox_text(block)
            if tb_text and current:
                if '机理分析' in current_section:
                    current['parts'].append('违背示例')
                elif '纠正措施' in current_section:
                    current['parts'].append('遵循示例')
                current['parts'].append(tb_text)

            # Visio OLE
            saved, ole_counter = extract_ole_objects(block, doc.part, visio_dir, doc_tag, ole_counter)
            for fname in saved:
                if current:
                    current['parts'].append(f'[Visio对象: {fname}]')
                    current['visio_files'].append(fname)

        elif isinstance(block, Table):
            if not started or skip_section:
                continue
            if before_keywords:
                continue
            if current:
                current['parts'].append(table_to_text(block))

    if current:
        cases.append(current)

    for c in cases:
        c['content'] = '\n'.join(c.pop('parts'))
    
    return cases
            
 
# ============ 结构探测函数 ============
def detect_document_structure(doc):
    """
    扫描文档，判断标题层级结构，返回解析策略配置
    策略定义：
    1. 'standard': 存在 Level 3 (案例) 和 Level 4 (小节)
    2. 'level1_all': 最高级是 Level 1，且没有 Level 2/3。靠关键词区分案例和小节。
    3. 'level1_l2': 存在 Level 1 (案例) 和 Level 2 (小节)。
    """
    levels_found = set()
    has_l1 = False
    has_l2 = False
    has_l3 = False
    
    # 第一次扫描：统计存在的标准标题级别
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block)
            if level:
                levels_found.add(level)
                if level == 1: has_l1 = True
                if level == 2: has_l2 = True
                if level == 3: has_l3 = True
    
    # 策略判断
    if has_l3:
        # 标准模式：Level 3 = 案例，Level 4 = 小节
        return 'standard'
    
    if has_l1 and not has_l2 and not has_l3:
        # 场景 B1：所有标题都是 Level 1
        # 需要靠关键词区分：包含“案例/问题”的是案例，包含“概述/机理”的是小节
        return 'level1_all'
    
    if has_l1 and has_l2 and not has_l3:
        # 场景 B2：Level 1 = 案例，Level 2 = 小节
        return 'level1_l2'
    
    # 兜底：如果既没有 L3，也没有 L1，或者结构混乱
    return 'unknown'

# ============ 辅助函数：判断文本类型 (用于场景 B1) ============
def classify_text_type(text):
    """
    根据关键词判断文本是案例标题还是小节标题
    """
    text_lower = text.lower()
    
    # 优先判断小节 
    for kw in SECTION_KEYWORDS:
        if kw in text:
            return 'section'

    # 默认视为案例标题
    return 'case'


# ==========================================
# 4. 必查清单解析
# ==========================================    
def parse_check_lists(docx_path: str) -> List[Dict]:
    """
    通用代码审查清单解析器
    自动识别文档中的表格结构和ID前缀，无需硬编码关键词。
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        return [{"error": f"无法读取文件: {str(e)}"}]

    results = []
    
     # 1. 动态定位目标表格 (终极版：定位最后一行表头)
    checklist_data = []
    target_table = None
    id_prefix = ""
    header_row_idx = -1

    # 遍历所有表格
    for table in doc.tables:
        if not table.rows:
            continue
        
        # 扫描整个表格，寻找包含“标识”和“审查项”的行
        # 寻找的是“最后一行”包含关键词的行，以防上面有干扰行
        last_valid_header_row = -1
        temp_clean_headers = []
        temp_idx_id = -1
        temp_idx_item = -1
        temp_idx_type = -1

        for r_idx in range(len(table.rows)):
            row = table.rows[r_idx]
            # 提取该行所有单元格文本
            raw_cells = [cell.text.strip() for cell in row.cells]
            
            # 清洗：去除空格，转为小写
            clean_row = [c.replace(' ', '').lower() for c in raw_cells]
            
            # 检查是否包含关键列 (兼容“测试项”)
            has_id = "标识" in clean_row
            has_item = "审查项" in clean_row or "测试项" in clean_row
            
            if has_id and has_item:
                # 记录这一行作为候选表头
                last_valid_header_row = r_idx
                temp_clean_headers = clean_row
                try:
                    temp_idx_id = clean_row.index("标识")
                    if "审查项" in clean_row:
                        temp_idx_item = clean_row.index("审查项")
                    else:
                        temp_idx_item = clean_row.index("测试项")
                    temp_idx_type = clean_row.index("类型") if "类型" in clean_row else -1
                except ValueError:
                    continue

        # 如果找到了表头行
        if last_valid_header_row != -1:
            header_row_idx = last_valid_header_row
            clean_headers = temp_clean_headers
            idx_id = temp_idx_id
            idx_item = temp_idx_item
            idx_type = temp_idx_type
            target_table = table
            
            # 提取 ID 前缀 (从表头行的下一行开始找，跳过可能的空行)
            # 我们需要找到第一个非空的数据行
            data_start_found = False
            for r_idx in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[r_idx]
                cells = [cell.text.strip() for cell in row.cells]
                # 检查是否有有效数据 (ID列非空)
                if len(cells) > idx_id and cells[idx_id]:
                    sample_id = cells[idx_id]
                    match = re.match(r'^([A-Za-z]+\d*)', sample_id)
                    if match:
                        id_prefix = match.group(1)
                    data_start_found = True
                    break
            
            if not id_prefix:
                # 如果第一行数据也没提取到，尝试在后续所有数据行中找
                for r_idx in range(header_row_idx + 1, len(table.rows)):
                    row = table.rows[r_idx]
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) > idx_id and cells[idx_id]:
                        sample_id = cells[idx_id]
                        match = re.match(r'^([A-Za-z]+\d*)', sample_id)
                        if match:
                            id_prefix = match.group(1)
                            break

            # 提取表格数据 (从表头行的下一行开始)
            for r_idx in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[r_idx]
                cells = [cell.text.strip() for cell in row.cells]
                
                # 跳过空行
                if not any(cells):
                    continue
                
                # 确保行长度足够
                if len(cells) > max(idx_id, idx_item):
                    id_val = cells[idx_id]
                    item_val = cells[idx_item]
                    
                    # 过滤空行和备注行
                    if id_val and not id_val.startswith("注："):
                        type_val = ""
                        if idx_type != -1 and len(cells) > idx_type:
                            type_val = cells[idx_type]
                        
                        checklist_data.append({
                            "category": type_val, # 类型
                            "id": id_val, # 标识
                            "content": item_val, # 审查项
                            "desc": "" # 具体内容
                        })
            
            # 如果在这个表格里找到了数据，就停止遍历其他表格
            if checklist_data:
                break 

    if not target_table:
        print(f"文件 {file_path.name} 中未找到包含“标识”和“审查项”的表格。")
        return None
    
    if not id_prefix:
        print(f"文件 {file_path.name} 中无法识别出 ID 前缀。")

    # 2. 动态定位“规则简要说明”章节
    section_start_index = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "规则简要说明" in text or "简要说明" in text:
            section_start_index = i
            break
            
    if section_start_index == -1:
        return [{"error": f"未在文档中找到“规则简要说明”章节，请检查文档结构"}]

    # 3. 动态解析说明内容 (严格边界模式 - 修复版)
    # 【步骤 1：准备 ID 查找表】
    target_ids = {}
    for item in checklist_data:
        raw_id = item["id"]
        norm_id = raw_id.replace(' ', '').upper()
        target_ids[norm_id] = raw_id

    if not target_ids:
        return results

    # 【步骤 2：定义查找函数】
    def find_id_in_text(text, id_mapping):
        clean_text = text.replace(' ', '').upper()
        for norm_id in id_mapping:
            if norm_id in clean_text:
                return id_mapping[norm_id], norm_id
        return None, None

    id_to_content = {}
    current_id = None
    current_content_lines = []
    
    # 遍历段落
    for i in range(section_start_index + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            continue
            
        # 查找是否包含任何目标 ID
        matched_original_id, matched_norm_id = find_id_in_text(text, target_ids)
        
        if matched_original_id:
            # 只要发现新 ID，无论它是什么（父级、子级、新系列）
            # 1. 立即保存上一个 ID 的内容（如果存在）
            if current_id and current_content_lines:
                id_to_content[current_id] = "\n".join(current_content_lines)
                # 清空列表，准备新内容
                current_content_lines = []
            
            # 2. 更新当前 ID
            current_id = matched_original_id
            
            # 3. 提取当前行中 ID 之后的内容作为新 ID 的起始内容
            # 如果 ID 后面没有内容，current_content_lines 保持为空
            try:
                escaped_id = re.escape(matched_original_id)
                match_obj = re.search(escaped_id, text, re.IGNORECASE)
                if match_obj:
                    content_after = text[match_obj.end():].strip()
                    if content_after:
                        current_content_lines.append(content_after)
                    # 如果 content_after 为空，说明这一行只是标题，不添加任何内容到列表中
                else:
                    pass
            except Exception:
                pass

            continue
        
        # 如果当前正在收集内容 (且当前段落不包含任何 ID)
        if current_id:
            # 过滤干扰行
            if "问题单" in text or "XXX检查" in text or "注：" in text:
                continue
            
            # 将当前段落作为内容追加
            current_content_lines.append(text)

    # 保存最后一个 ID 的内容
    if current_id and current_content_lines:
        id_to_content[current_id] = "\n".join(current_content_lines)

    # 4. 关联数据
    for item in checklist_data:
        item_id = item["id"]
        normalized_table_id = item_id.replace(' ', '').upper()
        
        content = id_to_content.get(item_id, "")
        
        # 兜底逻辑：如果表格里有子项 (FXKZCC1_1)，但说明里只写了父项 (FXKZCC1)
        if not content:
            # 提取基础 ID (去掉 _数字)
            base = item_id
            if '_' in item_id and item_id.split('_')[-1].isdigit():
                base = '_'.join(item_id.split('_')[:-1])
            
            if base in id_to_content:
                content = id_to_content[base]
        
        if '【' in content:
            content = '\n'.join(content.splitlines()[1:])
        item["desc"] = content
        results.append(item)

    return results
    
# 格式化函数 (保持不变)
def format_rules_for_rag(rules):
    formatted = []
    for idx, r in enumerate(rules):
        combined_code = ""
        if r['violation_code'].strip():
            combined_code += f"// [违背示例]\n{r['violation_code']}\n"
        if r['compliance_code'].strip():
            combined_code += f"// [遵循示例]\n{r['compliance_code']}"
        if not combined_code.strip():
            combined_code = "// 文档中未提取到明确示例代码"

        formatted.append({
            "id": r.get('id', f"rule_{idx}"),
            "category": "编程规则",
            "docRanges": [{"content": r['description'], "documentId": "rules_doc"}],
            "codeRanges": [{"content": combined_code, "documentId": "rules_code"}]
        })
    return {"annotations": formatted}

def format_issues_for_rag(issues):
    formatted = []
    for idx, i in enumerate(issues):
        search_text = f"【问题描述】\n{i.get('desc', '')}\n\n【处理意见】\n{i.get('opinion', '')}"
        func_name = i.get('trace_id', 'Unknown_Trace_ID')
        code_text = f"// 相关追踪标识/用例ID: {func_name}\n// (此处暂存标识，后续可扩展为提取具体函数代码)"
        
        formatted.append({
            "id": i.get('id', f"issue_{idx}"),
            "category": "问题单",
            "docRanges": [{"content": search_text, "documentId": "issue_doc"}],
            "codeRanges": [{"content": code_text, "documentId": func_name}]
        })
    return {"annotations": formatted}
    
def format_cases_for_rag(cases):
    formatted = []
    for idx, i in enumerate(cases):
        search_text = f"【问题描述】\n{i.get('content', '')}\n【简要说明】\n{i.get('desc', '')}"
        formatted.append({
            "id": i.get('id'),
            "category": "必查清单",
            "docRanges": [{"content": search_text, "documentId": "case_doc"}],
            "codeRanges": [{"content": "", "documentId": ""}]
        })
    return {"annotations": formatted}
