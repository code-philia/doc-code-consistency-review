import os
import threading

from celery import chain
from docx.shared import Pt, RGBColor
from flask_login import login_manager, login_required, current_user

from utils.kb_permission import kb_perm_required
from utils.word_parser import preprocess_files
from .db import (
    DB_CONFIG,
    get_db,
    get_db_celery,
    get_project_id_by_path,
    resolve_project_id,
    get_doc_blocks_by_project,
    get_paginated_doc_blocks_by_project,
    get_code_blocks_by_project,
    get_paginated_code_blocks_by_project,
    save_code_blocks_for_file,
    replace_doc_blocks,
    replace_code_blocks,
    append_missing_doc_blocks,
    append_missing_code_blocks,
    _default_doc_block_name,
    _default_code_block_name,
    _get_next_block_id
)

os.environ["OPENBLAS_NUM_THREADS"] = "128"
os.environ["OMP_NUM_THREADS"]="128"

import time
import traceback
from .project import project_access, get_project_id_by_name
from flask import Flask, json, render_template, request, jsonify, send_file, Blueprint, current_app
import sqlite3
import json as pyjson
import socket
from .alignment_config import (
    ALIGN_FILE_ABSTRACT_BATCH_LIMIT,
    CALL_GRAPH_ALIGN_DEPTH,
    CALL_GRAPH_MIN_SEED_SIMILARITY,
    CALL_GRAPH_SEED_LIMIT,
)
from .utils import get_all_files_with_relative_paths, parse_markdown, split_code, count_lines_of_code, convert_doc_to_markdown, get_filename_without_extension,\
    replace_text_in_docx, generate_issue_content, include_related_blocks
from callgraph.text_encoding import read_source_file
from .agent import query_generated_requirement, query_related_code, query_related_code_graph_rerank, query_review_result, query_flow_chart, query_related_requirement, query_code_abstract, query_codefile_abstract, query_codefile_from_abstract
from .agent import query_related_code_by_feedback, query_review_result_by_feedback, query_related_requirement_by_feedback
from .rag_chroma import rag_engine
from .doc_block import chunk_markdown
import random
import string
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from werkzeug.utils import secure_filename
import uuid
from docx import Document
import io
import shutil
import re
import zipfile

from .code_block import get_codefile_blocks
from .call_graph import (
    build_selection_code_range,
    build_project_call_graph,
    code_block_to_code_range,
    ensure_project_call_graph,
    find_first_intersecting_code_block,
    get_call_graph_metadata,
    mark_call_graph_building,
    query_multiple_function_graphs,
    query_function_graph,
    resolve_called_functions_in_line_range,
    resolve_code_block_to_function,
)

import logging
import chromadb
import sys
from io import BytesIO
import pandas as pd
from collections import defaultdict
from sqlalchemy import create_engine
import pymysql

from .utils import parse_programming_rules, parse_issue_reports, format_rules_for_rag, format_issues_for_rag, read_docx_text
from .utils import convert_docfile_to_markdown
from .agent import smart_parse_doc

# 配置日志
handler = logging.StreamHandler(sys.stdout)
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)

logger = logging.getLogger(__name__)
logger.addHandler(handler)
logger.setLevel(logging.INFO)

ALIGNMENT_SIDEBAR_PAGE_SIZE = 100
BLOCK_SIDEBAR_PAGE_SIZE = 100

SECRET_LEVEL_MAP = {
    'public': '公开',
    'internal': '内部',
    'secret': '秘密★10年',
    'confidential': '机密★20年'
}

SECRET_LEVEL_EXTENSION_MAP = {
    'public': '公开',
    'internal': '内部',
    'secret': '秘密',
    'confidential': '机密'
}


def project_now_str():
    return datetime.now(ZoneInfo("Asia/Shanghai")).strftime('%Y-%m-%d %H:%M:%S')


def _resolve_doc_block_name(block_data):
    return (
        (block_data.get('name') or '').strip()
        or _default_doc_block_name(block_data.get('content') or '', 24)
    )


def _resolve_code_block_name(block_data):
    return (
        (block_data.get('name') or '').strip()
        or _default_code_block_name(block_data.get('code') or block_data.get('content') or '', 60)
    )

def _call_graph_status_message(metadata):
    if not isinstance(metadata, dict):
        return '调用图不可用'
    status = metadata.get('status') or 'unavailable'
    error_message = metadata.get('error_message') or ''
    mapping = {
        'ready': '调用图已就绪',
        'building': '调用图正在后台重建',
        'failed': '调用图构建失败',
        'stale': '调用图需要重建',
        'unavailable': '当前项目不可用调用图',
    }
    base = mapping.get(status, '调用图状态未知')
    return f'{base}: {error_message}' if error_message else base
 
 
def _start_call_graph_rebuild(project_path):
    mark_call_graph_building(project_path)
 
    def _worker():
        try:
            build_project_call_graph(project_path)
        except Exception:
            traceback.print_exc()
 
    threading.Thread(target=_worker, name='call-graph-rebuild', daemon=True).start()
    
    
def _safe_int(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _build_doc_block_lookup(doc_blocks):
    llm_blocks = []
    blocks_by_file = defaultdict(list)
    for block in doc_blocks:
        filename = block.get('filename') or block.get('documentId') or ''
        start = _safe_int(block.get('start'))
        end = _safe_int(block.get('end'))
        content = block.get('content') or ''
        normalized = {
            'id': _safe_int(block.get('id')),
            'name': block.get('name') or block.get('type') or _compact_title_from_text(content, 24),
            'type': block.get('type') or '',
            'filename': filename,
            'documentId': filename,
            'content': content,
            'start': start,
            'end': end
        }
        llm_blocks.append({
            'file': filename,
            'range': [start, end],
            'content': content
        })
        blocks_by_file[filename].append(normalized)
    return llm_blocks, blocks_by_file


def _get_doc_blocks_for_matching(project_path, project_id=None):
    doc_blocks = get_doc_blocks_by_project(project_path, project_id)
    llm_blocks, blocks_by_file = _build_doc_block_lookup(doc_blocks)
    return llm_blocks, doc_blocks, blocks_by_file


def _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file):
    doc_ranges = []
    seen = set()
    for req in related_reqs or []:
        req_range = req.get('range') or [0, 0]
        req_start = _safe_int(req_range[0] if len(req_range) >= 1 else 0)
        req_end = _safe_int(req_range[1] if len(req_range) >= 2 else req_start)
        target_file = req.get('file', 'default')
        candidates = blocks_by_file.get(target_file, [])
        for block in candidates:
            block_start = _safe_int(block.get('start'))
            block_end = _safe_int(block.get('end'))
            if block_start <= req_start and block_end >= req_end:
                key = (block.get('filename') or '', block_start, block_end)
                if key not in seen:
                    doc_ranges.append(block)
                    seen.add(key)
    return doc_ranges


def _code_block_file(block):
    return block.get('file') or block.get('filename') or block.get('documentId') or ''


def _code_block_range(block):
    line_range = block.get('range')
    if isinstance(line_range, (list, tuple)) and len(line_range) == 2:
        return [_safe_int(line_range[0]), _safe_int(line_range[1])]
    start_line = _safe_int(block.get('startLine') or block.get('start_line'))
    end_line = _safe_int(block.get('endLine') or block.get('end_line') or start_line)
    return [start_line, end_line]


def _code_block_key(block):
    start_line, end_line = _code_block_range(block)
    return (_code_block_file(block), start_line, end_line)


def _code_ranges_intersect(start_a, end_a, start_b, end_b):
    return max(_safe_int(start_a), _safe_int(start_b)) <= min(_safe_int(end_a), _safe_int(end_b))


def _safe_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _dedupe_code_blocks(blocks):
    result = []
    seen = set()
    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        key = _code_block_key(block)
        if not key[0] or key[1] <= 0 or key[2] <= 0 or key in seen:
            continue
        seen.add(key)
        result.append(block)
    return result


def _trim_seed_related_code_items(items, *, max_items=None, min_similarity=None):
    max_items = CALL_GRAPH_SEED_LIMIT if max_items is None else max(1, int(max_items))
    min_similarity = CALL_GRAPH_MIN_SEED_SIMILARITY if min_similarity is None else float(min_similarity)
    scored = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        similarity = _safe_float(item.get('similarity'), 0.0)
        if similarity < min_similarity:
            continue
        scored.append((similarity, item))
    scored.sort(key=lambda pair: pair[0], reverse=True)
    return [item for _, item in scored[:max_items]]


def _match_related_items_to_code_blocks(related_items, all_code_blocks):
    matched = []
    for item in related_items or []:
        if not isinstance(item, dict):
            continue
        file_name = _code_block_file(item)
        start_line, end_line = _code_block_range(item)
        if not file_name or start_line <= 0 or end_line <= 0:
            continue

        exact = None
        containing = None
        intersecting = None
        for block in all_code_blocks or []:
            if _code_block_file(block) != file_name:
                continue
            block_start, block_end = _code_block_range(block)
            if block_start == start_line and block_end == end_line:
                exact = block
                break
            if containing is None and block_start <= start_line and block_end >= end_line:
                containing = block
            if intersecting is None and _code_ranges_intersect(block_start, block_end, start_line, end_line):
                intersecting = block

        selected = exact or containing or intersecting
        if selected:
            enriched = dict(selected)
            for field in ('similarity', 'role', 'reason'):
                if item.get(field) is not None:
                    enriched[field] = item.get(field)
            matched.append(enriched)
        else:
            matched.append(dict(item))
    return _dedupe_code_blocks(matched)


def _expand_code_blocks_with_call_graph_for_alignment(project_path, project_id, seed_blocks):
    if not project_path or not seed_blocks:
        return []

    try:
        graph_result = ensure_project_call_graph(project_path)
        if not graph_result.get('payload'):
            logger.info(f"调用图不可用，跳过重新对齐扩展: {graph_result.get('message')}")
            return []
    except Exception as exc:
        logger.info(f"调用图检查/构建失败，跳过重新对齐扩展: {exc}")
        return []

    expanded = []
    seen_functions = set()
    for block in seed_blocks:
        file_name = _code_block_file(block)
        start_line, end_line = _code_block_range(block)
        if not file_name or start_line <= 0 or end_line <= 0:
            continue
        try:
            function_id = resolve_code_block_to_function(project_path, project_id, file_name, start_line, end_line)
            if not function_id or function_id in seen_functions:
                continue
            seen_functions.add(function_id)
            graph = query_function_graph(
                project_path,
                function_id,
                max_depth=CALL_GRAPH_ALIGN_DEPTH,
                direction="both",
            )
            for code_range in graph.get('code_ranges') or []:
                code_file = code_range.get('filename') or code_range.get('file') or code_range.get('documentId') or file_name
                expanded.append({
                    'file': code_file,
                    'filename': code_file,
                    'documentId': code_file,
                    'range': [code_range.get('startLine'), code_range.get('endLine')],
                    'startLine': code_range.get('startLine'),
                    'endLine': code_range.get('endLine'),
                    'start': code_range.get('start'),
                    'end': code_range.get('end'),
                    'content': code_range.get('content') or '',
                    'code': code_range.get('content') or '',
                    'type': code_range.get('type') or 'function',
                    'name': code_range.get('name') or '',
                    'source': 'call_graph',
                })
        except Exception as exc:
            logger.info(f"调用图扩展单个代码块失败 {file_name}:{start_line}-{end_line}: {exc}")
            continue
    return _dedupe_code_blocks(expanded)


def _code_blocks_to_code_ranges(project_path, blocks):
    code_repo_path = os.path.join(project_path, 'code_repo')
    code_ranges = []
    seen = set()
    for block in blocks or []:
        file_name = _code_block_file(block)
        start_line, end_line = _code_block_range(block)
        if not file_name or start_line <= 0 or end_line <= 0:
            continue
        file_path = os.path.join(code_repo_path, file_name)
        if not os.path.exists(file_path):
            continue
        original_content = read_source_file(file_path)
        lines = original_content.splitlines(keepends=True)
        if not lines:
            continue
        safe_start = max(1, start_line)
        safe_end = min(len(lines), max(safe_start, end_line))
        char_start = sum(len(line) for line in lines[:safe_start - 1])
        char_end = sum(len(line) for line in lines[:safe_end])
        key = (file_name, safe_start, safe_end)
        if key in seen:
            continue
        seen.add(key)
        code_ranges.append({
            'filename': file_name,
            'start': char_start,
            'end': char_end,
            'content': '\n'.join(line.rstrip('\n\r') for line in lines[safe_start - 1:safe_end]),
            'documentId': file_name,
            'startLine': safe_start,
            'endLine': safe_end,
            'name': block.get('name') or '',
            'type': block.get('type') or '',
        })
    return code_ranges


def _get_or_build_code_blocks_for_file(project_path, file_name, project_id=None):
    existing_blocks = get_code_blocks_by_project(project_path, project_id, filename=file_name)
    if existing_blocks:
        return existing_blocks

    code_repo_path = os.path.join(project_path, 'code_repo')
    file_path = os.path.join(code_repo_path, file_name)
    if not os.path.exists(file_path):
        return []

    computed_blocks = get_codefile_blocks(code_repo_path, file_name)
    if not computed_blocks:
        return []
    return save_code_blocks_for_file(project_path, file_name, computed_blocks, project_id)

# 定义全局历史文件路径
HISTORY_FILE = 'app/history.json'
MAX_HISTORY_ITEMS = 15 # 最多记录15条历史

# 定义testdata目录路径
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
TESTDATA_DIR = os.path.join(PROJECT_ROOT, r'../testdata')
# NEW_PROJECT_ROOT = os.path.dirname(PROJECT_ROOT)
# NEW_TESTDATA_DIR = os.path.join(NEW_PROJECT_ROOT, '/testdata')
# print('NEW_TESTDATA_DIR',NEW_TESTDATA_DIR)
# print('NEW_PROJECT_ROOT',NEW_PROJECT_ROOT)

# app = Flask(__name__)
bp = Blueprint('main', __name__)


# templates
@bp.route('/')
@login_required
def index():
    """Render the welcome page"""
    #return render_template('login.html') # 切换成需要登录的页面
    return render_template('welcome.html')


@bp.route('/welcome')
@login_required
def welcome():
    """Render the welcome page"""
    return render_template('welcome.html')

@bp.route('/semi-automatic')
def semi_automatic():
    """Render the semi-automatic mode page"""
    return render_template('semi-automatic.html')

@bp.route('/project')
def project():
    """Render the project page"""
    return render_template('project.html')

@bp.route('/annotation')
def annotation():
    """Render the annotation page"""
    return render_template('annotation.html')

@bp.route('/templates/flowchart-viewer.html')
def flowchart_viewer_template():
    """Serve the flowchart viewer template"""
    return send_file('../templates/flowchart-viewer.html', mimetype='text/html')

# login
@bp.route('/get-ip', methods=['GET'])
def get_ip():
    ip = request.headers.get('X-Real-IP', request.remote_addr)
    return jsonify({"ip": ip})

# @bp.route('/login/ip', methods=['POST'])
# def validate_ip():
#     data = request.get_json()
#     ip = data.get('ip')
#
#     if not ip:
#         return jsonify({"success": False, "message": "IP 无效"})
#
#     # 数据库查询
#     #user = next((u for u in users if u["ip"] == ip), None)
#     user = read_ip_mapping_from_db(ip)
#
#     if user:
#         return jsonify({"success": True, "message": "IP 登录成功"})
#     else:
#         return jsonify({"success": False, "message": "IP 未授权，请联系管理员"})

def read_ip_mapping_from_db(IP):
    # 检查数据库文件是否存在
    user = ''
    DB_PATH = './user_data/user_info.db'
    if not os.path.exists(DB_PATH):
        print(f"数据库文件 {DB_PATH} 不存在。")
        return user

    # 连接数据库
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row  # 允许通过列名访问数据
    cursor = conn.cursor()

    try:
        # 查询 ip_mapping 表
        cursor.execute("SELECT username, ip FROM ip_mapping")
        rows = cursor.fetchall()

        if not rows:
            print("数据库中没有找到任何 IP 映射记录。")
            return user

        #print(f"共找到 {len(rows)} 条 IP 映射记录：")
        #print("-" * 50)
        for row in rows:
            #username = row['username']
            if row['ip'] == IP:
                user = row['username']
                conn.close()
                return user

    except Exception as e:
        print(f"读取数据库时出错: {e}")
    finally:
        conn.close()
        return user

def read_user_from_db(username, password):
    # 检查数据库文件是否存在
    user = ''
    DB_PATH = './user_data/user_info.db'
    if not os.path.exists(DB_PATH):
        print(f"数据库文件 {DB_PATH} 不存在。")
        return user

    # 连接数据库
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row  # 允许通过列名访问数据
    cursor = conn.cursor()

    try:
        # 查询 ip_mapping 表
        cursor.execute("SELECT username, password FROM user_pass")
        rows = cursor.fetchall()

        if not rows:
            print("数据库中没有找到任何用户记录。")
            return user

        #print(f"共找到 {len(rows)} 条用户记录：")
        #print("-" * 50)
        for row in rows:
            if row['username'] == username and row['password'] == password:
                user = username
                conn.close()
                return user

    except Exception as e:
        print(f"读取数据库时出错: {e}")
    finally:
        conn.close()
        return user



# project
@bp.route('/project/create', methods=['POST'])
def create_project():
    data = request.json
    creation_type = data.get('creationType', 'blank')
    project_name = data.get('projectName')
    project_location = data.get('projectLocation')
    project_id = data.get('project_id')
    parseDocMethod = data.get('parseDocMethod')
    project_secret_level = data.get('projectSecretLevel')
    #print(f'parseDocMethod=======:{parseDocMethod}')
    
    code = project_location.split('_')[-1]
    #project_name += code
    # print(f'project_name=======:{project_name}')

    if not project_name or not project_location:
        return jsonify({"status": "error", "message": "项目名称和路径不能为空。"}), 400

    if creation_type == 'blank':
        return create_blank_project(project_name, project_location, parseDocMethod,project_secret_level)
    elif creation_type == 'folder':
        return create_project_from_folder(project_name, project_location, parseDocMethod, project_secret_level,project_id=project_id)
    else:
        return jsonify({"status": "error", "message": "无效的创建类型。"}), 400


def create_blank_project(project_name, project_location, parseDocMethod, project_secret_level):
    """处理创建空白项目的逻辑"""
    project_path = os.path.join(project_location, project_name)
    if os.path.exists(project_path):
        return jsonify({"status": "error", "message": f"项目文件夹 '{project_name}' 已存在于目标位置。"}), 400

    try:
        now_str = project_now_str()
        code_repo_path = os.path.join(project_path, 'code_repo')
        doc_repo_path = os.path.join(project_path, 'doc_repo')
        os.makedirs(code_repo_path, exist_ok=True)
        os.makedirs(doc_repo_path, exist_ok=True)

        metadata = {
            "project_name": project_name,
            "project_location": project_location,
            "create_time": now_str,
            "code_repo": code_repo_path,
            "doc_repo": doc_repo_path,
            "code_files": [],
            "doc_files": [],
            "code_scale": 0,
            "code_file_lines": {}
        }

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        # update_history(project_name, project_path)
        # init_project_db(project_path)
        auto_load_rag_db(project_path)
        sql = f"""
            insert into project(user_id,last_opened,name,path,create_time,update_time, project_secret_level) 
            values({current_user.user_id}, "{now_str}", "{project_name}", 
            "{project_path}", "{now_str}", "{now_str}", "{project_secret_level}");
            """
        # print('sql:', sql)

        db = get_db()
        c = db.cursor()
        c.execute(sql)
        new_id = c.lastrowid

        return jsonify({"status": "success", "project_path": project_path, 'new_id': new_id}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"创建目录或文件时出错: {e}"}), 500


def create_project_from_folder(project_name, folder_path, parseDocMethod, project_secret_level, project_id=None):
    """处理从现有文件夹创建项目的逻辑"""
    project_path = folder_path # 项目路径就是用户选择的文件夹
    if not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "提供的路径不是一个有效的文件夹。"}), 400

    code_repo_path = os.path.join(project_path, 'code_repo')
    doc_repo_path = os.path.join(project_path, 'doc_repo')

    if not os.path.isdir(code_repo_path) or not os.path.isdir(doc_repo_path):
        return jsonify({"status": "error", "message": "文件夹结构不符合要求，必须包含 'code_repo' 和 'doc_repo' 子目录。"}), 400

    if os.path.exists(os.path.join(project_path, 'metadata.json')):
        return jsonify({"status": "error", "message": "该文件夹已包含 'metadata.json'，似乎已是一个项目。"}), 400

    try:
        now_str = project_now_str()
        code_files = get_all_files_with_relative_paths(code_repo_path, type ='code')
        doc_files = get_all_files_with_relative_paths(doc_repo_path, type ='doc')

        # 检查 doc_repo 目录下是否有 docx 文件，如果有则进行格式转换
        has_docx = False
        for root, dirs, files in os.walk(doc_repo_path):
            for file in files:
                if file.endswith('.docx'):
                    has_docx = True
                    break
            if has_docx:
                break

        if has_docx:
            convert_doc_to_markdown(doc_repo_path, parseDocMethod)
            # 转换后重新获取文档文件列表
            doc_files = get_all_files_with_relative_paths(doc_repo_path, type ='doc')

        code_file_lines = {}
        total_loc = 0
        for file in code_files:
            loc = count_lines_of_code(os.path.join(code_repo_path, file))
            code_file_lines[file] = loc
            total_loc += loc

        metadata = {
            "project_name": project_name,
            "project_location": os.path.dirname(project_path), # 存储其父目录
            "create_time": now_str,
            "code_repo": code_repo_path,
            "doc_repo": doc_repo_path,
            "code_files": code_files,
            "doc_files": doc_files,
            "code_scale": total_loc,
            "code_file_lines": code_file_lines
        }

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        # update_history(project_name, project_path)
        # init_project_db(project_path)
        auto_load_rag_db(project_path)

        db = get_db()
        c = db.cursor()
        if project_id:
            sql = """
                update project
                set name=%s, path=%s, update_time=%s
                where project_id=%s and user_id=%s
            """
            params = (project_name, project_path, now_str, project_id, current_user.user_id)
            c.execute(sql, params)
            new_id = project_id
        else:
            sql = """
                insert into project(user_id,last_opened,name,path,create_time,update_time,project_secret_level)
                values(%s, %s, %s, %s, %s, %s, %s)
            """
            params = (current_user.user_id, now_str, project_name, project_path, now_str, now_str, project_secret_level)
            c.execute(sql, params)
            new_id = c.lastrowid

        return jsonify({"status": "success", "project_path": project_path, "project_name": project_name, "new_id": new_id}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"扫描文件夹或生成元数据时出错: {e}"}), 500


# def update_history(project_name, project_path):
#     """读取、更新并写回项目历史记录"""
#     history = []
#     if os.path.exists(HISTORY_FILE):
#         with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
#             try:
#                 history = json.load(f)
#             except json.JSONDecodeError:
#                 history = [] # 如果文件内容损坏，则重置
#
#     # 检查项目是否已在历史中，如果在则移除旧条目
#     history = [item for item in history if item.get('path') != project_path]
#
#     # 添加新条目到列表顶部
#     new_entry = {
#         "name": project_name,
#         "path": project_path,
#         "last_opened": datetime.now().isoformat() # 使用ISO 8601格式的时间戳
#     }
#     history.insert(0, new_entry)
#
#     # 限制历史记录的长度
#     history = history[:MAX_HISTORY_ITEMS]
#
#     # 写回文件
#     with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
#         json.dump(history, f, indent=4, ensure_ascii=False)


def get_db_path(project_path):
    return os.path.join(project_path, 'project.db')


def get_db_conn(project_path):
    db_path = get_db_path(project_path)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    try:
        conn.execute('PRAGMA foreign_keys=ON')
        conn.execute('PRAGMA journal_mode=WAL')
    except Exception:
        pass
    return conn


# def init_project_db(project_path):
#     db_path = get_db_path(project_path)
#     os.makedirs(project_path, exist_ok=True)
#     conn = sqlite3.connect(db_path)
#     conn.row_factory = sqlite3.Row  # Ensure we can access columns by name
#     try:
#         conn.execute('PRAGMA foreign_keys=ON')
#         conn.execute('PRAGMA journal_mode=WAL')
#
#         # 代码摘要
#         conn.execute(
#             'CREATE TABLE IF NOT EXISTS abstracts ('
#             'filename TEXT PRIMARY KEY,'
#             'abstract TEXT NOT NULL DEFAULT "[]",'
#             'createdAt TEXT DEFAULT CURRENT_TIMESTAMP,'
#             'updatedAt TEXT)'
#         )
#
#         conn.execute(
#             'CREATE TABLE IF NOT EXISTS alignments ('
#             'id TEXT PRIMARY KEY,'
#             'name TEXT,'
#             'isReviewed INTEGER DEFAULT 0,'
#             'reviewThoughts TEXT,'
#             'docRanges TEXT NOT NULL DEFAULT "[]",'
#             'codeRanges TEXT NOT NULL DEFAULT "[]",'
#             'GenReq TEXT,'
#             'GenMermaid TEXT,'
#             'createdAt TEXT DEFAULT CURRENT_TIMESTAMP,'
#             'updatedAt TEXT)'
#         )
#
#         # --- Migration: Check and add missing columns ---
#         cur = conn.cursor()
#         cur.execute("PRAGMA table_info(alignments)")
#         existing_columns = {row['name'] for row in cur.fetchall()}
#
#         if 'GenReq' not in existing_columns:
#             print(f"[DB Migration] Adding GenReq column to {db_path}")
#             conn.execute("ALTER TABLE alignments ADD COLUMN GenReq TEXT")
#
#         if 'GenMermaid' not in existing_columns:
#             print(f"[DB Migration] Adding GenMermaid column to {db_path}")
#             conn.execute("ALTER TABLE alignments ADD COLUMN GenMermaid TEXT")
#
#         conn.execute(
#             'CREATE TABLE IF NOT EXISTS issues ('
#             'id INTEGER PRIMARY KEY AUTOINCREMENT,'
#             'displayId TEXT UNIQUE,'
#             'alignmentId TEXT NOT NULL,'
#             'severity TEXT,'
#             'title TEXT,'
#             'content TEXT,'
#             'status TEXT,'
#             'relatedDocFile TEXT,'
#             'relatedRequirementId TEXT,'
#             'briefRequirement TEXT,'
#             'briefCode TEXT,'
#             'createdAt TEXT DEFAULT CURRENT_TIMESTAMP,'
#             'updatedAt TEXT,'
#             'FOREIGN KEY(alignmentId) REFERENCES alignments(id) ON DELETE CASCADE)'
#         )
#         conn.execute('CREATE INDEX IF NOT EXISTS idx_issues_alignmentId ON issues(alignmentId)')
#     finally:
#         conn.close()


def auto_load_rag_db(project_path):
    """自动加载项目下的RAG知识库"""
    try:
        # Initialize standard KBs (align_rules, user_manuals, code)
        rag_engine.initialize(project_path)
    except Exception as e:
        print(f"[AutoLoad] Failed to load RAG DBs: {e}")


# def import_json_to_db(project_path):
#     conn = get_db_conn(project_path)
#     try:
#         cur = conn.cursor()
#         cur.execute('SELECT COUNT(1) AS c FROM alignments')
#         row = cur.fetchone()
#         need_import_alignments = (row['c'] == 0)
#         results_dir = os.path.join(project_path, 'results')
#         if need_import_alignments and os.path.isdir(results_dir):
#             for filename in os.listdir(results_dir):
#                 if filename.endswith('.json'):
#                     fp = os.path.join(results_dir, filename)
#                     try:
#                         with open(fp, 'r', encoding='utf-8') as f:
#                             data = json.load(f)
#                         if isinstance(data, dict):
#                             for aid, alignment in data.items():
#                                 try:
#                                     cur.execute(
#                                         'INSERT INTO alignments(id,name,isReviewed,reviewThoughts,docRanges,codeRanges,createdAt,updatedAt)'
#                                         ' VALUES (%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)'
#                                         ' ON CONFLICT(id) DO UPDATE SET name=excluded.name,isReviewed=excluded.isReviewed,reviewThoughts=excluded.reviewThoughts,docRanges=excluded.docRanges,codeRanges=excluded.codeRanges,updatedAt=CURRENT_TIMESTAMP',
#                                         (
#                                             alignment.get('id') or aid,
#                                             alignment.get('name'),
#                                             1 if alignment.get('isReviewed') else 0,
#                                             alignment.get('reviewThoughts') or '',
#                                             pyjson.dumps(alignment.get('docRanges') or []),
#                                             pyjson.dumps(alignment.get('codeRanges') or [])
#                                         )
#                                     )
#                                 except Exception:
#                                     continue
#                         conn.commit()
#                     except Exception:
#                         continue
#         cur.execute('SELECT COUNT(1) AS c FROM issues')
#         total_issues = cur.fetchone()['c']
#         need_import_issues = (total_issues == 0)
#         try:
#             cur.execute("SELECT COUNT(1) AS c FROM issues WHERE IFNULL(relatedDocFile,'')='' OR IFNULL(briefRequirement,'')='' OR IFNULL(briefCode,'')='' ")
#             missing_fields = cur.fetchone()['c']
#             if missing_fields and missing_fields > 0:
#                 need_import_issues = True
#         except Exception:
#             pass
#         issues_file = os.path.join(project_path, 'issues.json')
#         if need_import_issues and os.path.exists(issues_file):
#             try:
#                 with open(issues_file, 'r', encoding='utf-8') as f:
#                     issues = json.load(f)
#                 for issue in issues or []:
#                     try:
#                         cur.execute(
#                             'INSERT INTO issues(displayId,alignmentId,severity,title,content,status,relatedDocFile,relatedRequirementId,briefRequirement,briefCode,createdAt,updatedAt) '
#                             'VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) '
#                             'ON CONFLICT(displayId) DO UPDATE SET '
#                             'alignmentId=excluded.alignmentId, '
#                             'severity=excluded.severity, '
#                             'title=excluded.title, '
#                             'content=excluded.content, '
#                             'status=excluded.status, '
#                             'relatedDocFile=excluded.relatedDocFile, '
#                             'relatedRequirementId=excluded.relatedRequirementId, '
#                             'briefRequirement=excluded.briefRequirement, '
#                             'briefCode=excluded.briefCode, '
#                             'createdAt=excluded.createdAt, '
#                             'updatedAt=excluded.updatedAt',
#                             (
#                                 issue['id'],
#                                 issue['alignmentId'],
#                                 issue['level'],
#                                 issue['summary'],
#                                 issue['description'],
#                                 issue['status'] or 'unconfirmed',
#                                 issue['relatedDocFile'],
#                                 issue['relatedRequirementId'],
#                                 issue['briefRequirement'],
#                                 issue['briefCode'],
#                                 issue['createdDate'] or datetime.now().isoformat(),
#                                 issue['updatedDate'] or datetime.now().isoformat()
#                             )
#                         )
#                     except Exception:
#                         continue
#                 conn.commit()
#             except Exception:
#                 pass
#     finally:
#         conn.close()


# @bp.route('/project/history', methods=['GET'])
# def get_project_history():
#     """获取最近打开的项目列表"""
#     if not os.path.exists(HISTORY_FILE):
#         return jsonify([])
#
#     with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
#         try:
#             history = json.load(f)
#             return jsonify(history)
#         except json.JSONDecodeError:
#             return jsonify([])

# @bp.route('/project/history', methods=['DELETE'])
# def delete_project_history():
#     """删除指定的历史记录"""
#     data = request.json
#     project_path = data.get('path')
#
#     if not project_path:
#         return jsonify({"status": "error", "message": "项目路径不能为空"}), 400
#
#     if not os.path.exists(HISTORY_FILE):
#         return jsonify({"status": "error", "message": "历史记录文件不存在"}), 404
#
#     try:
#         with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
#             history = json.load(f)
#
#         # 查找并删除指定的历史记录
#         original_length = len(history)
#         history = [item for item in history if item.get('path') != project_path]
#
#         if len(history) == original_length:
#             return jsonify({"status": "error", "message": "未找到指定的历史记录"}), 404
#
#         # 写回文件
#         with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
#             json.dump(history, f, indent=4, ensure_ascii=False)
#
#         return jsonify({"status": "success", "message": "历史记录删除成功"})
#
#     except json.JSONDecodeError:
#         return jsonify({"status": "error", "message": "历史记录文件格式错误"}), 500
#     except Exception as e:
#         return jsonify({"status": "error", "message": f"删除历史记录时出错: {e}"}), 500
#
# @bp.route('/project/delete', methods=['POST'])
# def delete_project():
#     """删除项目目录和历史记录"""
#     data = request.json
#     project_path = data.get('path')
#
#     if not project_path:
#         return jsonify({"status": "error", "message": "项目路径不能为空"}), 400
#
#     if not os.path.exists(project_path):
#         return jsonify({"status": "error", "message": "项目路径不存在"}), 404
#
#     try:
#         # 删除项目目录
#         shutil.rmtree(project_path)
#
#         # 从历史记录中删除项目条目
#         if os.path.exists(HISTORY_FILE):
#             with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
#                 history = json.load(f)
#
#             # 过滤掉要删除的项目
#             history = [item for item in history if item.get('path') != project_path]
#
#             # 写回历史记录文件
#             with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
#                 json.dump(history, f, indent=4, ensure_ascii=False)
#
#         return jsonify({"status": "success", "message": "项目删除成功"})
#
#     except PermissionError:
#         return jsonify({"status": "error", "message": "没有权限删除项目文件"}), 403
#     except Exception as e:
#         return jsonify({"status": "error", "message": f"删除项目时出错: {str(e)}"}), 500

@bp.route('/project/open', methods=['POST'])
def open_project():
    """当用户打开一个项目时，更新其历史记录"""
    data = request.json
    project_name = data.get('name') or data.get('projectName')
    project_path = data.get('path') or data.get('projectPath')
    project_id = data.get('project_id')
    if not project_id:
        project_id = get_project_id_by_name(project_name)
    print('project/open project_id:', project_id)

    if not project_name or not project_path:
        return jsonify({"status": "error", "message": "项目信息不完整"}), 400

    # Handle relative paths
    # PROJECT_ROOT is essentially the directory of views.py
    PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))

    if not os.path.isabs(project_path):
        # 1. Try resolving relative to PROJECT_ROOT
        abs_path = os.path.join(PROJECT_ROOT, project_path)
        if os.path.exists(abs_path):
            project_path = abs_path
        else:
            # 2. Try resolving relative to uploads (if applicable)
            abs_path_uploads = os.path.join(PROJECT_ROOT, 'uploads', project_path)
            if os.path.exists(abs_path_uploads):
                project_path = abs_path_uploads

    # 可以在此添加校验，确保项目路径真实存在
    if not os.path.exists(project_path):
         return jsonify({"status": "error", "message": f"项目路径不存在: {project_path}"}), 404

    project_access(project_id)
    try:
        # init_project_db(project_path)
        # import_json_to_db(project_path)
        auto_load_rag_db(project_path)
    except Exception:
        pass
    return jsonify({"status": "success", "path": project_path, 'project_id': project_id})


@bp.route('/project/import', methods=['POST'])
def import_project():
    """验证一个现有项目文件夹的结构是否有效"""
    data = request.json
    project_path = data.get('path')
    project_id = data.get('project_id')

    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": f"提供的路径不是一个有效的文件夹。"}), 400

    # 检查必需的文件和文件夹
    code_repo_path = os.path.join(project_path, 'code_repo')
    doc_repo_path = os.path.join(project_path, 'doc_repo')
    metadata_file = os.path.join(project_path, 'metadata.json')

    if not os.path.isdir(code_repo_path):
        return jsonify({"status": "error", "message": "文件夹内缺少 'code_repo' 子目录。"}), 400
    if not os.path.isdir(doc_repo_path):
        return jsonify({"status": "error", "message": "文件夹内缺少 'doc_repo' 子目录。"}), 400
    if not os.path.isfile(metadata_file):
        return jsonify({"status": "error", "message": "文件夹内缺少 'metadata.json' 文件。"}), 400

    # 读取 metadata.json 以获取项目名称
    try:
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        project_name = metadata.get('project_name')
        if not project_name:
            return jsonify({"status": "error", "message": "'metadata.json' 文件中缺少 'project_name' 字段。"}), 400

        # update_history(project_name, project_path)

        # 如为原始版本（不存在 project.db），初始化数据库并迁移旧数据
        try:
            db_path = get_db_path(project_path)
            # 总是调用 init_project_db 以确保表结构最新（自动迁移）
            # init_project_db(project_path)
            # import_json_to_db 内部会检查表是否为空，所以安全调用
            # import_json_to_db(project_path)
        except Exception:
            pass

        # 验证成功，返回项目信息
        project_data = {
            "name": project_name,
            "path": project_path
        }
        return jsonify({"status": "success", "project": project_data})

    except (json.JSONDecodeError, Exception) as e:
        return jsonify({"status": "error", "message": f"读取 'metadata.json' 文件失败: {e}"}), 500

# @bp.route('/project/recent-projects', methods=['GET'])
# def get_recent_projects():
#     """获取最近打开的项目列表"""
#     history = []
#     if os.path.exists(HISTORY_FILE):
#         with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
#             try:
#                 history = json.load(f)
#             except json.JSONDecodeError:
#                 history = []
#     return jsonify({"status": "success", "recentProjects": history})

@bp.route('/project/metadata', methods=['GET'])
def get_project_metadata():
    """根据项目路径获取元数据"""
    project_path = request.args.get('path')
    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "无效的项目路径。"}), 400

    metadata_file = os.path.join(project_path, 'metadata.json')
    if not os.path.isfile(metadata_file):
        return jsonify({"status": "error", "message": "项目元数据文件不存在。"}), 404

    try:
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        def normalize_kb_type_for_storage(raw_type):
            """统一 selected_kbs 中的类型到系统内部类型。"""
            kb_type = (raw_type or "other").strip()
            if kb_type in ["rule", "coding_rule", "checklist"]:
                return "rule"
            if kb_type in ["issue", "history_issue"]:
                return "issue"
            if kb_type in ["align", "history_align"]:
                return "align"
            return "other"

        def kb_exists(kb_root, kb_name, kb_type):
            # 兼容当前展平结构：rag_database/<kb_name>
            flat_path = os.path.join(kb_root, kb_name)
            if os.path.exists(flat_path):
                return True
            # 兼容历史分类型目录结构
            legacy_folder_map = {
                'rule': 'rule_knowledge_base',
                'issue': 'issue_knowledge_base',
                'align': 'align_knowledge_base',
                'other': 'other_knowledge_base'
            }
            legacy_folder = legacy_folder_map.get(kb_type, 'other_knowledge_base')
            legacy_path = os.path.join(kb_root, legacy_folder, kb_name)
            return os.path.exists(legacy_path)

        # 校验 selected_kbs 是否有效
        if 'selected_kbs' in metadata and metadata['selected_kbs']:
            kb_root = os.path.join(PROJECT_ROOT, "../rag_database")
            valid_kbs = []

            has_changes = False
            for kb in metadata['selected_kbs']:
                # kb should be {name, type}
                kb_name = kb.get('name')
                kb_type = kb.get('type')

                if not kb_name or not kb_type:
                    has_changes = True
                    continue

                normalized_type = normalize_kb_type_for_storage(kb_type)
                normalized_kb = {'name': kb_name, 'type': normalized_type}
                if normalized_type != kb_type:
                    has_changes = True

                if kb_exists(kb_root, kb_name, normalized_type):
                    valid_kbs.append(normalized_kb)
                else:
                    has_changes = True

            if has_changes:
                metadata['selected_kbs'] = valid_kbs
                # Write back changes
                with open(metadata_file, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=4, ensure_ascii=False)
                    
        metadata['call_graph'] = get_call_graph_metadata(project_path)
        
        return jsonify({"status": "success", "metadata": metadata}), 200

    except (json.JSONDecodeError, Exception) as e:
        return jsonify({"status": "error", "message": f"读取元数据文件失败: {e}"}), 500

@bp.route('/project/save-kbs', methods=['POST'])
def save_project_kbs():
    """保存项目选用的知识库"""
    data = request.json
    project_path = data.get('projectPath')
    selected_kbs = data.get('selectedKbs') # List of {name, type}

    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "无效的项目路径"}), 400

    metadata_file = os.path.join(project_path, 'metadata.json')
    if not os.path.isfile(metadata_file):
        return jsonify({"status": "error", "message": "metadata.json 不存在"}), 404

    try:
        def normalize_kb_type_for_storage(raw_type):
            kb_type = (raw_type or "other").strip()
            if kb_type in ["rule", "coding_rule", "checklist"]:
                return "rule"
            if kb_type in ["issue", "history_issue"]:
                return "issue"
            if kb_type in ["align", "history_align"]:
                return "align"
            return "other"

        normalized_selected_kbs = []
        seen = set()
        for kb in (selected_kbs or []):
            kb_name = (kb or {}).get('name')
            kb_type = normalize_kb_type_for_storage((kb or {}).get('type'))
            if not kb_name:
                continue
            key = (kb_name, kb_type)
            if key in seen:
                continue
            seen.add(key)
            normalized_selected_kbs.append({'name': kb_name, 'type': kb_type})

        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        metadata['selected_kbs'] = normalized_selected_kbs

        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        return jsonify({"status": "success", "message": "知识库配置已保存"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@bp.route('/project/upload-files-async', methods=['POST'])
@login_required
def upload_files_async():
    project_path = request.form.get('path')
    project_id = request.form.get('project_id')
    file_type = request.form.get('fileType')
    files = request.files.getlist('files')
    parseDocMethod = request.form.get('parseDocMethod')

    if not all([project_path, project_id, file_type, files]):
        return jsonify({"status": "error", "message": "请求参数不完整。"}), 400

    try:
        # ===== 中转：文件先落临时目录，任务里再转回 FileStorage =====
        temp_dir = os.path.join('./', uuid.uuid4().hex)
        os.makedirs(temp_dir, exist_ok=True)

        temp_files = []
        for f in files:
            # 文件名用纯 uuid（f.filename 可能是 webkitRelativePath 含 '/'，直接拼会出错）
            # 原始相对路径存在 origin_name 里，任务里还原
            save_path = os.path.join(temp_dir, uuid.uuid4().hex)
            f.save(save_path)
            temp_files.append({'temp_path': save_path, 'origin_name': f.filename})

        # ===== 提交 Celery 任务 =====
        from tasks import upload_files_task
        task = upload_files_task.delay(project_path, file_type, parseDocMethod, temp_files)

        # ===== 插状态表 =====
        title = temp_files[0]['origin_name'] if len(temp_files) == 1 else f'{len(temp_files)} 个文件'
        db = get_db()
        cursor = db.cursor()
        cursor.execute("""
            INSERT INTO user_task_snapshot
                (user_id, project_id, task_id, task_type, task_category, title, state, is_running)
            VALUES (%s, %s, %s, 'upload', 'upload', %s, 'PENDING', 1)
        """, (current_user.user_id, project_id, task.id, title))
        db.commit()

        return jsonify({"status": "success", "task_id": task.id}), 200

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


def do_upload_files_logic(project_path, file_type, files, parseDocMethod, task=None):
    files = preprocess_files(files)

    try:

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        if file_type == 'code':
            code_repo_path = metadata.get('code_repo')
            for i, file in enumerate(files):
                if task:
                    # print(i, '======================')

                    from tasks import _update_snapshot
                    _update_snapshot(task.request.id, state='PROCESSING',
                                     title=f'正在处理 {file.filename} ({i}/{len(files)})',
                                     current_progress=int(i / len(files) * 100))
                    # time.sleep(10)
                # 保留包含中文的原始相对路径
                relative_path = file.filename.replace('\\', '/')

                # 安全检查：防止目录遍历攻击 (e.g., ../../secret.txt)
                # 1. 路径拼接后进行规范化
                dest_path = os.path.abspath(os.path.join(code_repo_path, relative_path))
                # 2. 确保目标路径仍然在 code_repo 目录内
                if not dest_path.startswith(os.path.abspath(code_repo_path)):
                    return jsonify({"status": "error", "message": f"检测到不安全的路径: {relative_path}"}), 400

                # 创建目标目录并保存文件
                dest_dir = os.path.dirname(dest_path)
                os.makedirs(dest_dir, exist_ok=True)
                file.save(dest_path)

            # 更新元数据
            metadata['code_files'] = get_all_files_with_relative_paths(code_repo_path, type='code')
            code_file_lines = {}
            total_loc = 0
            for f in metadata['code_files']:
                loc = count_lines_of_code(os.path.join(code_repo_path, f))
                code_file_lines[f] = loc
                total_loc += loc
            metadata['code_scale'] = total_loc
            metadata['code_file_lines'] = code_file_lines

        elif file_type == 'annotation':
            # 1. 既然 project_path 是根目录，我们在这里手动拼接 annotations
            target_dir = os.path.join(project_path, 'annotations')

            # 2. 如果文件夹不存在，自动创建
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)

            # 3. 保存文件
            for file in files:
                # 只取文件名，防止路径包含多余信息
                filename = os.path.basename(file.filename)
                save_dest = os.path.join(target_dir, filename)
                file.save(save_dest)
                print(f"标注文件已保存至: {save_dest}")

        elif file_type == 'doc':
            doc_repo_path = metadata.get('doc_repo')
            for i, file in enumerate(files):

                filename = os.path.basename(file.filename)
                if task:
                    from tasks import _update_snapshot
                    _update_snapshot(task.request.id, state='PROCESSING',
                                     title=f'正在处理 {filename} ({i}/{len(files)})',
                                     current_progress=int(i / len(files) * 100))
                # time.sleep(10)
                if filename.endswith(('.md', '.docx')):
                    doc_file_path = os.path.join(doc_repo_path, filename)
                    # print(doc_file_path)
                    # 去除文件名中的空格，重新保存
                    new_filename = process_word_file(doc_file_path)
                    if new_filename:
                        filename = new_filename
                    # print("新保存的文件名！！！！！！！！！！！")
                    # print(doc_file_path)
                    doc_file_path = os.path.join(doc_repo_path, filename)
                    file.save(doc_file_path)
                    if filename.endswith('.docx'):
                        convert_docfile_to_markdown(doc_file_path, doc_repo_path, parseDocMethod)

            metadata['doc_files'] = get_all_files_with_relative_paths(doc_repo_path, type='doc')

        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        if file_type == 'code':
            _start_call_graph_rebuild(project_path)

        return {'status': 'success'}

    except Exception as e:
        print(f"Error during file upload: {e}")

        raise Exception(f"服务器处理文件上传时出错: {e}")


@bp.route('/project/upload-files', methods=['POST'])
def upload_files():
    try:
        project_path = request.form.get('path')
        file_type = request.form.get('fileType')  # 'doc' or 'code'
        # files = request.files.getlist('files')
        files = preprocess_files(request.files.getlist('files'))
        parseDocMethod = request.form.get('parseDocMethod')

        if not all([project_path, file_type, files]):
            return jsonify({"status": "error", "message": "请求参数不完整。"}), 400

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        if file_type == 'code':
            code_repo_path = metadata.get('code_repo')
            for file in files:
                # 保留包含中文的原始相对路径
                relative_path = file.filename.replace('\\', '/')

                # 安全检查：防止目录遍历攻击 (e.g., ../../secret.txt)
                # 1. 路径拼接后进行规范化
                dest_path = os.path.abspath(os.path.join(code_repo_path, relative_path))
                # 2. 确保目标路径仍然在 code_repo 目录内
                if not dest_path.startswith(os.path.abspath(code_repo_path)):
                    return jsonify({"status": "error", "message": f"检测到不安全的路径: {relative_path}"}), 400

                # 创建目标目录并保存文件
                dest_dir = os.path.dirname(dest_path)
                os.makedirs(dest_dir, exist_ok=True)
                file.save(dest_path)

            # 更新元数据
            metadata['code_files'] = get_all_files_with_relative_paths(code_repo_path, type='code')
            code_file_lines = {}
            total_loc = 0
            for f in metadata['code_files']:
                loc = count_lines_of_code(os.path.join(code_repo_path, f))
                code_file_lines[f] = loc
                total_loc += loc
            metadata['code_scale'] = total_loc
            metadata['code_file_lines'] = code_file_lines

        elif file_type == 'annotation':
            # 1. 既然 project_path 是根目录，我们在这里手动拼接 annotations
            target_dir = os.path.join(project_path, 'annotations')

            # 2. 如果文件夹不存在，自动创建
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)

            # 3. 保存文件
            for file in files:
                # 只取文件名，防止路径包含多余信息
                filename = os.path.basename(file.filename)
                save_dest = os.path.join(target_dir, filename)
                file.save(save_dest)
                print(f"标注文件已保存至: {save_dest}")

        elif file_type == 'doc':
            doc_repo_path = metadata.get('doc_repo')
            for file in files:
                # 直接使用原始文件名，仅取最后的文件名部分，天然防止了目录遍历
                filename = os.path.basename(file.filename)
                if filename.endswith(('.md', '.docx')):
                    doc_file_path = os.path.join(doc_repo_path, filename)
                    #print(doc_file_path)
                    # 去除文件名中的空格，重新保存
                    new_filename = process_word_file(doc_file_path)
                    if new_filename:
                        filename = new_filename
						#print("新保存的文件名！！！！！！！！！！！")
                        #print(doc_file_path)
                    doc_file_path = os.path.join(doc_repo_path, filename)
                    file.save(doc_file_path)
                    if filename.endswith('.docx'):
                        convert_docfile_to_markdown(doc_file_path, doc_repo_path, parseDocMethod)

            metadata['doc_files'] = get_all_files_with_relative_paths(doc_repo_path, type='doc')

        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)
        
        if file_type == 'code':
            _start_call_graph_rebuild(project_path)
        
        return jsonify({"status": "success"}), 200

    except Exception as e:
        print(f"Error during file upload: {e}")
        return jsonify({"status": "error", "message": f"服务器处理文件上传时出错: {e}"}), 500


def process_word_file(file_path):
    try:
        dir_name = os.path.dirname(file_path)
        original_name = os.path.basename(file_path)
        
        # 生成新文件名
        new_name = original_name.replace(" ", "_")  # 将空格替换为下划线
        new_path = os.path.join(dir_name, new_name)
        
        if new_path == file_path:
            print(f"文件名无需修改：{original_name}")
            return None
        
        #print(f"正在处理：{original_name} -> {new_name}")
    
    
        # 打开原文件
        #doc = Document(file_path)
        
        # 保存到新文件（这会创建一个新文件，而不是覆盖原文件）
        #doc.save(new_path)
        
        #print(f"成功保存为新文件：{new_name}")
        
        # 可选：如果确定新文件没问题，可以删除旧文件
        #os.remove(file_path) 
        #print("已删除原文件")
        return new_name
        
    except Exception as e:
        print(f"处理文件时出错：{e}")

        
        
@bp.route('/project/upload-folder', methods=['POST'])
def upload_folder():
    """处理文件夹上传功能"""
    try:
        now_str = project_now_str()
        # 获取上传的文件和文件夹名称
        # files = request.files.getlist('files')
        files = preprocess_files(request.files.getlist('files'))
        paths = request.form.getlist('paths')
        paths = [
            p[:-4] + '.docx' if p.lower().endswith('.doc') else p
            for p in paths
        ]
        folder_name = request.form.get('folderName')
        project_secret_level = request.form.get('projectSecretLevel')
        project_name = (request.form.get('projectName') or folder_name or '').strip()


        if not files or not folder_name:
            return jsonify({"status": "error", "message": "没有接收到文件或文件夹名称"}), 400

        # 确保testdata目录存在
        os.makedirs(TESTDATA_DIR, exist_ok=True)

        # 创建目标文件夹路径
        target_folder_path = os.path.join(TESTDATA_DIR, folder_name)
        # logger.info('target_folder_path',target_folder_path)
        timestamp = int(time.time())
        # 如果目标文件夹已存在，添加时间戳后缀
        if os.path.exists(target_folder_path):
            target_folder_path = os.path.join(TESTDATA_DIR, f"{folder_name}_{timestamp}")
        # 创建目标文件夹
        os.makedirs(target_folder_path, exist_ok=True)
        include_files = ['.py', '.c', '.cpp', '.h', '.hpp', '.java', '.html', '.vhd', '.v', '.sv', '.adb', '.ads']
        # 保存所有文件，保持目录结构
        for file, relative_path in zip(files, paths):
            # print(f'file:{file}, relative_path:{relative_path}')
            if not file.filename:
                continue
                
            filename = os.path.basename(file.filename)
            # 去除文件名中的空格，重新保存
            if filename.endswith(('.md', '.docx')):
 
                # 1. 分离目录路径和文件名
                dir_name, file_name = os.path.split(relative_path)
                # 2. 替换文件名中的空格
                new_file_name = file_name.replace(" ", "_")
                # 3. 拼接新路径
                relative_path = os.path.join(dir_name, new_file_name)
                #print("新保存的文件名！！！！！！！！！！！")
                #print(f"新路径：{relative_path}")
                
            # 移除文件夹名称前缀，获取相对路径
            if relative_path.startswith(folder_name + '/'):
                file_relative_path = relative_path[len(folder_name) + 1:]
            else:
                file_relative_path = relative_path

            # 构建完整的目标文件路径
            target_file_path = os.path.join(target_folder_path, file_relative_path)

            # 安全检查：确保目标路径在目标文件夹内
            target_file_path = os.path.abspath(target_file_path)
            if not target_file_path.startswith(os.path.abspath(target_folder_path)):
                return jsonify({"status": "error", "message": f"检测到不安全的路径: {relative_path}"}), 400

            # 创建目标目录
            target_dir = os.path.dirname(target_file_path)
            os.makedirs(target_dir, exist_ok=True)

            # 保存文件
            file.save(target_file_path)
            # ext = os.path.splitext(file.filename)[1].lower()
            # if ext in include_files:
            #     print(f'ext:{ext}, filename:{file.filename}')
            #     # ===== 文本文件：检测编码 -> 统一转 UTF-8 =====
            #     raw_bytes = file.read()
            #
            #     try:
            #         content = raw_bytes.decode('utf-8')
            #     except UnicodeDecodeError:
            #         content = raw_bytes.decode('gbk', errors='ignore')
            #
            #     content = content.replace('\r\n', '\n').replace('\r', '\n')
            #
            #     with open(target_file_path, 'w', encoding='utf-8') as f:
            #         f.write(content)
            #
            # else:
            #     # 其他文件
            #     # 保存文件
            #     file.save(target_file_path)

        sql = f"""
                insert into project(user_id,last_opened,name,path,create_time,update_time,project_secret_level) 
                values(%s, %s, %s, %s, %s, %s, %s);
                """
        # print('sql:', sql)
        params = (current_user.user_id, now_str, project_name, target_folder_path, now_str, now_str, project_secret_level)
        db = get_db()
        c = db.cursor()
        c.execute(sql, params)
        new_id = c.lastrowid

        return jsonify({
            "status": "success",
            "message": f"文件夹 '{folder_name}' 上传成功",
            "serverPath": target_folder_path,
            "folderName": os.path.basename(target_folder_path),
            "projectName": project_name,
            "new_id": new_id
        }), 200

    except Exception as e:
        print(f"Error during folder upload: {e}")
        traceback.print_exc()
        return jsonify({"status": "error", "message": f"文件夹上传失败: {str(e)}"}), 500


@bp.route('/project/file-remove', methods=['GET'])
def remove_file_content():
    """根据项目路径、相对路径和文件类型，删除文件或目录"""
    project_path = request.args.get('path')
    filename = request.args.get('filename')
    file_type = request.args.get('type') # 'doc' or 'code'
    node_type = request.args.get('node_type', 'file')

    if not project_path or not filename or file_type not in ('doc', 'code') or node_type not in ('file', 'directory'):
        return jsonify({"status": "error", "message": "参数错误"}), 400

    try:
        repo_name = 'doc_repo' if file_type == 'doc' else 'code_repo'
        repo_root = os.path.abspath(os.path.join(project_path, repo_name))
        target_rel_path = filename.replace('\\', '/').strip('/')
        target_abs_path = os.path.abspath(os.path.join(repo_root, target_rel_path))
        project_id = get_project_id_by_path(project_path)

        if os.path.commonpath([repo_root, target_abs_path]) != repo_root:
            return jsonify({"status": "error", "message": "非法路径"}), 400

        metadata_file = os.path.join(project_path, 'metadata.json')
        with open(metadata_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        metadata_key = 'doc_files' if file_type == 'doc' else 'code_files'
        matched_files = []
        if node_type == 'file':
            matched_files = [target_rel_path]
        else:
            prefix = f"{target_rel_path}/"
            matched_files = [item for item in metadata.get(metadata_key, []) if item == target_rel_path or item.startswith(prefix)]

        if node_type == 'directory' and not matched_files:
            return jsonify({"status": "error", "message": "目录下没有可删除文件"}), 404

        if not os.path.exists(target_abs_path):
            return jsonify({"status": "error", "message": "目标不存在"}), 404

        if node_type == 'directory':
            shutil.rmtree(target_abs_path)
        else:
            os.remove(target_abs_path)

        if file_type == 'doc':
            project_path_convert = os.path.join(project_path, 'doc_repo_converted')
            if os.path.exists(project_path_convert):
                for file_path in matched_files:
                    filefoldername = os.path.splitext(os.path.basename(file_path))[0]
                    for root, dirs, files in os.walk(project_path_convert, topdown=False):
                        if filefoldername in dirs:
                            folder_path = os.path.join(root, filefoldername)
                            try:
                                shutil.rmtree(folder_path)
                            except Exception as e:
                                print(f"删除失败: {folder_path}，错误: {e}")

            if project_id and matched_files:
                db = get_db()
                cur = db.cursor()
                placeholders = ','.join(['%s'] * len(matched_files))
                cur.execute(
                    f'DELETE FROM doc_blocks WHERE project_id=%s AND filename IN ({placeholders})',
                    tuple([project_id] + matched_files)
                )

            if "doc_files" in metadata:
                metadata["doc_files"] = [f for f in metadata["doc_files"] if f not in matched_files]
        else:
            if project_id and matched_files:
                db = get_db()
                cur = db.cursor()
                placeholders = ','.join(['%s'] * len(matched_files))
                cur.execute(
                    f'DELETE FROM code_blocks WHERE project_id=%s AND file IN ({placeholders})',
                    tuple([project_id] + matched_files)
                )

            if "code_files" in metadata:
                metadata["code_files"] = [f for f in metadata["code_files"] if f not in matched_files]
            if isinstance(metadata.get("code_file_lines"), dict):
                metadata["code_file_lines"] = {k: v for k, v in metadata["code_file_lines"].items() if k not in matched_files}

        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=4, ensure_ascii=False)

        message = '目录已递归删除' if node_type == 'directory' else '文件已删除'
        return jsonify({"status": "success", "message": message, "removed_files": matched_files}), 200

    except Exception as e:
        print(f"删除失败: {filename}，错误: {e}")
        return jsonify({"status": "error", "message": str(e)}), 404

        
DOC_PAGE_TARGET_CHARS = 12000
DOC_PAGE_MIN_CHARS = 8000
DOC_PAGE_MAX_CHARS = 18000
CODE_PAGE_LINES = 500


def _regularize_file_content(content, file_type):
    if content is None:
        return ''
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    content = content.replace('\u200b', '')
    if file_type == 'doc':
        content = re.sub(r'(?<=\S)\$\$(?=\S)', '$ $', content)
    return content


def _count_occurrences(text, token):
    if not text or not token:
        return 0
    count = 0
    start = 0
    while True:
        idx = text.find(token, start)
        if idx == -1:
            break
        count += 1
        start = idx + len(token)
    return count


def _build_doc_page_ranges(content):
    if not content:
        return []

    markers = ['\n# ', '\n## ', '\n### ', '\n#### ', '\n\n']
    pages = []
    start = 0

    def normalize_break_point(segment_start, candidate_end):
        safe_end = candidate_end
        segment = content[segment_start:safe_end]
        if _count_occurrences(segment, '```') % 2 != 0:
            closing_fence = content.find('```', safe_end)
            if closing_fence != -1:
                safe_end = closing_fence + 3

        if _count_occurrences(content[segment_start:safe_end], '$$') % 2 != 0:
            closing_math = content.find('$$', safe_end)
            if closing_math != -1:
                safe_end = closing_math + 2

        return min(safe_end, len(content))

    def pick_break_point(segment_start):
        remaining = len(content) - segment_start
        if remaining <= DOC_PAGE_MAX_CHARS:
            return len(content)

        min_end = min(len(content), segment_start + DOC_PAGE_MIN_CHARS)
        ideal_end = min(len(content), segment_start + DOC_PAGE_TARGET_CHARS)
        max_end = min(len(content), segment_start + DOC_PAGE_MAX_CHARS)

        def resolve_break(index, marker):
            if index < min_end:
                return None
            return index + (2 if marker == '\n\n' else 1)

        for marker in markers:
            backward_index = content.rfind(marker, segment_start, ideal_end + len(marker))
            break_point = resolve_break(backward_index, marker)
            if break_point is not None:
                return normalize_break_point(segment_start, break_point)

        for marker in markers:
            forward_index = content.find(marker, ideal_end)
            if forward_index != -1 and forward_index <= max_end:
                return normalize_break_point(segment_start, forward_index + (2 if marker == '\n\n' else 1))

        return normalize_break_point(segment_start, max_end)

    while start < len(content):
        end = pick_break_point(start)
        pages.append({'start': start, 'end': end})
        if end <= start:
            break
        start = end

    return pages


def _build_code_page_ranges(content):
    if not content:
        return []

    pages = []
    start = 0
    length = len(content)

    while start < length:
        pos = start
        for _ in range(CODE_PAGE_LINES):
            nl = content.find('\n', pos)
            if nl == -1:
                pos = length
                break
            pos = nl + 1

        end = pos if pos > start else length
        if end <= start:
            break
        pages.append({'start': start, 'end': end})
        start = end

    return pages


def _resolve_project_file_path(project_path, filename, file_type):
    metadata_file = os.path.join(project_path, 'metadata.json')
    with open(metadata_file, 'r', encoding='utf-8') as f:
        metadata = json.load(f)

    if file_type == 'code':
        repo_path = metadata.get('code_repo')
        return os.path.join(repo_path, filename)

    file_name_prefix = filename.split('.')[0]
    if filename.endswith('.md'):
        repo_path = metadata.get('doc_repo')
        file_path = os.path.join(repo_path, filename)
        if not os.path.exists(file_path):
            file_path = os.path.join(project_path, 'doc_repo_converted', file_name_prefix, file_name_prefix + '.md')
        return file_path

    return os.path.join(project_path, 'doc_repo_converted', file_name_prefix, file_name_prefix + '.md')


def _read_text_file_with_fallback(file_path):
    if not file_path or not os.path.exists(file_path):
        return None
    return read_source_file(file_path)


        
        
@bp.route('/project/file-content', methods=['GET'])
def get_file_content():
    """根据项目路径、文件名和文件类型获取文件内容"""
    project_path = request.args.get('path')
    page = request.args.get('page', type=int)
    filename = request.args.get('filename')
    file_type = request.args.get('type') # 'doc' or 'code'

    if not project_path or not filename or not file_type:
        return jsonify({"status": "error", "message": "缺少必要的参数"}), 400

    repo_map = {
        'doc': 'doc_repo',
        'code': 'code_repo'
    }

    if file_type not in repo_map:
        return jsonify({"status": "error", "message": "无效的文件类型"}), 400

    try:
        file_path = _resolve_project_file_path(project_path, filename, file_type)

        if not os.path.exists(file_path):
            print(file_path)
            return jsonify({"status": "error", "message": "文件未找到"}), 404

        content = _read_text_file_with_fallback(file_path)

        if page:
            normalized_content = _regularize_file_content(content, file_type)
            page_ranges = _build_doc_page_ranges(normalized_content) if file_type == 'doc' else _build_code_page_ranges(normalized_content)
            total_pages = max(len(page_ranges), 1)
            safe_page = min(max(page, 1), total_pages)
            if page_ranges:
                current_range = page_ranges[safe_page - 1]
                page_content = normalized_content[current_range['start']:current_range['end']]
                page_start = current_range['start']
                page_end = current_range['end']
            else:
                page_content = ''
                page_start = 0
                page_end = 0

            start_line = normalized_content[:page_start].count('\n') + 1 if page_start > 0 else 1

            return jsonify({
                "status": "success",
                "content": page_content,
                "pagination": {
                    "page": safe_page,
                    "pages": total_pages,
                    "total_pages": total_pages,
                    "page_ranges": page_ranges
                },
                "page_start": page_start,
                "page_end": page_end,
                "start_line": start_line
            }), 200

        return jsonify({"status": "success", "content": content}), 200

    except Exception as e:
        return jsonify({"status": "error", "message": f"读取文件内容时出错: {e}"}), 500



@bp.route('/api/requirement-decomposition', methods=['POST'])
def requirement_decomposition():
    """处理需求分解请求 - 保存到JSONL文件并同步数据库"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')

        if not project_path:
            return jsonify({'status':'error', 'message': '缺少项目路径'})

        annotations_dir = os.path.join(project_path, 'annotations')
        if not os.path.exists(annotations_dir):
            return jsonify({'status':'error', 'message': '标注目录不存在'})

        json_files = [f for f in os.listdir(annotations_dir) if f.endswith('.json')]
        if not json_files:
            return jsonify({'status':'error', 'message': '标注结果文件不存在'})

        annotation_file = os.path.join(annotations_dir, json_files[0])
        with open(annotation_file, 'r', encoding='utf-8') as f:
            annotation_data = json.load(f)

        # 解析docFiles建立映射
        doc_files = annotation_data.get('docFiles', [])
        doc_id_to_name = {doc['id']: doc['name'] for doc in doc_files}

        # 解析annotations构建需求点
        annotations = annotation_data.get('annotations', [])
        processed_count = 0

        req_blocks = []


        next_block_id = 1
        for annotation in annotations:
            doc_ranges = annotation.get('docRanges', [])
            category = annotation.get('category')


            # 获取文档ID
            document_id = doc_ranges[0].get('documentId')
            doc_name = doc_id_to_name.get(document_id)

            if not doc_name:
                continue

            # 构建需求块对象 (扁平化，每个docRange作为一个独立的块)
            for doc_range in doc_ranges:
                req_block = {
                    'id': next_block_id,
                    'name':category,
                    'type': category,
                    'filename': doc_name,
                    'documentId': doc_name,
                    'content': doc_range.get('content'),
                    'start': doc_range.get('start'),
                    'end': doc_range.get('end')
                }
                req_blocks.append(req_block)
                processed_count += 1
                next_block_id += 1

        replace_doc_blocks(project_path, req_blocks)

        return jsonify({
            'status': 'success',
            'message': '需求分解完成，结果已保存',
            'processedCount': processed_count
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/api/auto-markdown-split', methods=['POST'])
def auto_markdown_split():
    """处理自动Markdown分解请求 - 保存到JSONL文件并同步数据库"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')

        if not project_path:
            return jsonify({'status':'error', 'message': '缺少项目路径'})

        # 获取项目中的文档文件
        doc_repo_path = os.path.join(project_path, 'doc_repo')
        if not os.path.exists(doc_repo_path):
            return jsonify({'status':'error', 'message': '文档目录不存在'})
        
        annotations_dir = os.path.join(project_path, 'annotations')
        
        if os.path.exists(annotations_dir): 
            json_files = [f for f in os.listdir(annotations_dir) if f.endswith('.json')]
            if json_files:
            
                annotation_file = os.path.join(annotations_dir, json_files[0])
                with open(annotation_file, 'r', encoding='utf-8') as f:
                    annotation_data = json.load(f)

                # 解析docFiles建立映射
                doc_files = annotation_data.get('docFiles', [])
                doc_id_to_name = {doc['id']: doc['name'] for doc in doc_files}

                # 解析annotations构建需求点
                annotations = annotation_data.get('annotations', [])
                processed_count = 0

                req_blocks = []

                next_block_id = 1
                for annotation in annotations:
                    doc_ranges = annotation.get('docRanges', [])
                    category = annotation.get('category')


                    # 获取文档ID
                    document_id = doc_ranges[0].get('documentId')
                    doc_name = doc_id_to_name.get(document_id)

                    if not doc_name:
                        continue

                    # 构建需求块对象 (扁平化，每个docRange作为一个独立的块)
                    for doc_range in doc_ranges:
                        req_block = {
                            'id': next_block_id,
                            'name':category,
                            'type': category,
                            'filename': doc_name,
                            'documentId': doc_name,
                            'content': doc_range.get('content'),
                            'start': doc_range.get('start'),
                            'end': doc_range.get('end')
                        }
                        req_blocks.append(req_block)
                        processed_count += 1
                        next_block_id += 1

                replace_doc_blocks(project_path, req_blocks)
        
        else:
        
            # 查找所有markdown文件
            md_files = []
            for root, dirs, files in os.walk(doc_repo_path):
                for file in files:
                    if file.lower().endswith('.md'):
                        md_files.append(os.path.join(root, file))
            
            # 如果用户上传的不是markdown文件，就去格式转换后的文件夹中找对应的markdown文件
            if not md_files:
                # 获取项目中格式转换后的文档文件
                doc_repo_path = os.path.join(project_path, 'doc_repo_converted')
                if not os.path.exists(doc_repo_path):
                    return jsonify({'status':'error', 'message': '文档目录不存在'})

                # 查找所有markdown文件
                md_files = []
                for root, dirs, files in os.walk(doc_repo_path):
                    for file in files:
                        if file.lower().endswith('.md'):
                            md_files.append(os.path.join(root, file))
                if not md_files:
                    return jsonify({'status':'error', 'message': '未找到Markdown文档'})

            processed_count = 0
            req_blocks = []

            # 处理每个markdown文件
            next_block_id = 1
            for md_file in md_files:
                with open(md_file, 'r', encoding='utf-8') as f:
                    md_content = f.read()

                # 分解markdown内容
                doc_name = os.path.basename(md_file)
                blocks = chunk_markdown(doc_name, md_content)

                if not blocks:
                    continue


                # 构建需求块
                for block_info in blocks:
                    content = block_info['content']
                    if  '#' in content:
                        first_line = content.split('\n')[0]
                        chunk_name = first_line.lstrip('#')
                    else:
                        chunk_name = _compact_title_from_text(content, 24)

                    req_block = {
                        'id': next_block_id,
                        'name': chunk_name,
                        'type': block_info.get('type') or chunk_name,
                        'filename': doc_name,
                        'documentId': doc_name,
                        'content': block_info['content'],
                        'start': block_info['start'],
                        'end': block_info['end']
                    }
                    req_blocks.append(req_block)
                    processed_count += 1
                    next_block_id += 1

            replace_doc_blocks(project_path, req_blocks)

        return jsonify({
            'status': 'success',
            'message': '自动分解完成，结果已保存',
            'processedCount': processed_count
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/export-tasks', methods=['GET'])
def get_export_tasks():
    """
    获取导出记录列表(兼轮询查询)
    """
    project_id = request.args.get('project_id', type=int)

    db = get_db()
    cursor = db.cursor()

    try:
        cursor.execute("""
            SELECT id, task_id, status, filename, created_at, completed_at, error_msg
            FROM export_tasks
            WHERE project_id = %s AND export_type = 'reverse_requirement'
            ORDER BY created_at DESC
            LIMIT 20
        """, (project_id, ))

        tasks = []
        rows = cursor.fetchall()
        for row in rows:
            task = dict(row)
            for key in ['created_at', 'completed_at']:
                if isinstance(task[key], datetime):
                    task[key] = task[key].strftime('%Y-%m-%d %H:%M:%S')
            tasks.append(task)

        has_processing = any(r['status'] in ('pending', 'processing') for r in rows)

        return jsonify({'status': 'success', 'data': tasks, 'has_processing': has_processing})

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/export-issues-download', methods=['POST'])
def export_issues_download():
    """导出所有问题单到一个docx文件"""

    data = request.json
    issues = data.get('issues', [])
    form_data = data.get('formData', {})
    project_id = data.get('project_id')
    secret_level = data.get('secret_level')

    if not issues:
        return jsonify({'status': 'error', 'message': '没有问题单可导出'})

    temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
    # 生成docx文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"问题单导出_{form_data.get('issueId', 'BBB')}_{timestamp}" \
                    f"({SECRET_LEVEL_EXTENSION_MAP[secret_level]}).docx"
    docx_path = os.path.join(temp_dir, docx_filename)

    from tasks import export_issues_task
    task = export_issues_task.delay(data, docx_path)

    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("""
            INSERT INTO export_tasks (task_id, project_id, user_id, filename, url)
            VALUES (%s, %s, %s, %s, %s)
        """, (task.id, project_id, current_user.user_id, docx_filename, docx_path))

        return jsonify({
            'status': 'success',
            'task_id': task.id,
            'message': f'导出任务已创建',
        })

    except Exception as e:
        task.revoke(terminate=True, signal='SIGTERM')
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/download-file/<filename>', methods=['GET'])
def download_file(filename):
    """下载临时文件并在下载后删除，支持zip文件"""
    try:
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
        file_path = os.path.join(temp_dir, filename)

        if not os.path.exists(file_path):
            return jsonify({'status': 'error', 'message': '文件不存在'}), 404

        # 读取文件内容到内存
        with open(file_path, 'rb') as f:
            file_data = f.read()

        # 立即删除临时文件
        # try:
        #     os.remove(file_path)
        #     print(f"已删除临时文件: {file_path}")
        # except Exception as e:
        #     print(f"删除临时文件失败: {e}")

        # 根据文件类型设置MIME类型
        if filename.endswith('.zip'):
            mimetype = 'application/zip'
        elif filename.endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'text/plain'

        # 从内存返回文件
        return send_file(
            io.BytesIO(file_data),
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

#------------------------#
#     核心流程函数
#------------------------#
def get_abstracts_from_sqlite(project_id):
    """从project.db的abstracts表读取指定列数据"""
    try:
        # 连接SQLite数据库
        # conn = sqlite3.connect(db_path)
        conn = get_db_celery()
        #conn.execute('PRAGMA encoding = UTF-8')
        # 读取name、docRanges、codeRanges、GenReq、 GenMermaid 列所有数据
        query_sql = f"SELECT filename, abstract FROM abstracts where project_id={project_id}"
        # 直接用pandas读取SQL结果（简洁高效）
        # df = pd.read_sql(query_sql, conn)

        # 创建引擎
        db_url = f"mysql+pymysql://{DB_CONFIG['user']}:{DB_CONFIG['password']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"

        engine = create_engine(
            db_url,
            echo=False,
            pool_pre_ping=True
        )
        df = pd.read_sql(query_sql, engine)
        
        conn.close()
        print('读取摘要成功!!!!!!!!!!!')
        return df
    except sqlite3.Error as e:
        print(f"SQLite数据库读取失败：{e}")
        return pd.DataFrame()  # 返回空DataFrame
    except Exception as e:
        print(f"读取数据异常：{e}")
        return pd.DataFrame()


def generate_abstract(file_path):
    # 读取文件内容，保留原始行（包括空行）
    lines = []
    for line in read_source_file(file_path).splitlines():
        # 必须过滤被注释掉的代码，不然影响大模型理解和分析
        l_line = line.lstrip()
        is_code_comment = l_line.startswith('//') and not any(c in l_line[2:] for c in ('，', '。', '；'))
        if not is_code_comment:
            lines.append(l_line)
    # 调用 LLM 的代码文件摘要函数
    return query_codefile_abstract(lines)


def save_abstract_to_db(project_path, file, codefile_abstract, project_id, user_id):
    # 代码摘要写入数据库
    conn = get_db_celery()
    cur = conn.cursor()
    try:

        cur.execute(
            'INSERT INTO abstracts(filename,abstract,user_id,project_id,createdAt,updatedAt) '
            'VALUES (%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP) ',
            (
                file,
                codefile_abstract,
                user_id,
                project_id
            )
        )
        conn.commit()
        print(f'写入代码摘要数据，project_id={project_id}')
    except Exception as e:
        conn.rollback()
        logger.info("写入代码摘要数据失败")
        logger.info(e)
        return jsonify({"status": "error", "message": f"写入代码摘要数据失败: {e}"}), 500
    finally:
        conn.close()

    # conn.commit()
    # conn.close()

# 对齐 需求->代码 需要异步处理
# @bp.route('/api/get-code-abstract', methods=['GET'])
# def abstract_code_from_project():
#     """
#     代码摘要：摘要每个代码文件，存入数据库
#     """
#
#     project_path = request.args.get('projectPath')
#     project_id = request.args.get('project_id')
#     print('/api/get-code-abstract, project_id:', project_id)
#
#     if not project_path:
#         return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400
#
#     code_file_path = os.path.join(project_path, 'code_repo')
#     if not os.path.exists(code_file_path):
#         return jsonify({'status': 'success', 'data': []})
#
#     from tasks import abstract_code_from_project_task
#
#     task = abstract_code_from_project_task.delay(project_path, project_id, code_file_path, current_user.user_id)
#
#     return jsonify({'status': 'success', 'task_id': task.id})


@bp.route('/api/get-code-abstract', methods=['GET'])
def abstract_code_from_project():
    """
    代码摘要：摘要每个代码文件，存入数据库
    """

    try:
        project_path = request.args.get('projectPath')
        project_id = request.args.get('project_id')
        print('/api/get-code-abstract, project_id:', project_id)

        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400

        code_file_path = os.path.join(project_path, 'code_repo')
        if not os.path.exists(code_file_path):
            return jsonify({'status': 'success', 'data': []})

        # 读取数据库里的代码摘要
        # db_file = os.path.join(project_path, 'project.db')
        # if not os.path.exists(db_file):
        #     logger.info('未找到数据库文件')
        #     df = pd.DataFrame()
        # else:
            # 从SQLite读取数据
        df = get_abstracts_from_sqlite(project_id)
        
        # 排除无关文件夹/目录
        exclude_folders = ['.git', '.idea']
        # 基于文件名后缀，指定文件类型
        include_files = ['.py', '.c', '.cpp', '.h', '.hpp', '.java', '.html', '.vhd', '.v', '.sv', '.adb', '.ads']
        # 从FPGA代码中去除'.adb', '.ads'
        flag_FPGA = False
        adb_file_path_list = []
        adb_rel_path_list = []
        include_files_FPGA = ['.vhd', '.v', '.sv']
        include_files_Ada = ['.adb', '.ads']
        
        file_path_list = []
        rel_path_list = []
        # 遍历文件夹
        for root, dirs, files in os.walk(code_file_path):
            dirs[:] = [d for d in dirs if d not in exclude_folders]
            for file in files:
                if os.path.splitext(file)[1] in include_files:
                    if os.path.splitext(file)[1] in include_files_FPGA:
                        flag_FPGA = True
                    # 从FPGA代码中去除'.adb', '.ads'
                    if flag_FPGA and os.path.splitext(file)[1] in include_files_Ada:
                        continue
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, code_file_path)
                    if os.path.splitext(file)[1] in include_files_Ada:
                        adb_file_path_list.append(file_path)
                        adb_rel_path_list.append(rel_path)
                    else:
                        file_path_list.append(file_path)
                        rel_path_list.append(rel_path)
        if not flag_FPGA:
            file_path_list.extend(adb_file_path_list)      
            rel_path_list.extend(adb_rel_path_list)      
        
        # 遍历文件列表
        file_abstract = {}
        for i in range(0, len(file_path_list)):
            file_path = file_path_list[i]
            rel_path = rel_path_list[i]
            if not df.empty:
                # 先看数据库里有没有已经生成好的代码摘要
                row_data = df[df['filename'] == rel_path]

                # 数据库有该代码文件的摘要
                if not row_data.empty:
                    logger.info('数据库有该代码文件的摘要')
                    abstract_data = row_data['abstract'].values[0]
                    file_abstract[rel_path] = abstract_data
                # 数据库没有该代码文件的摘要
                else:
                    #file_path = os.path.join(root, file)
                    codefile_abstract = generate_abstract(file_path)
                    file_abstract[rel_path] = codefile_abstract

                    # save_abstract_to_db(project_path, file, codefile_abstract, project_id)
                    save_abstract_to_db(project_path, rel_path, codefile_abstract, project_id,
                                        current_user.user_id)

            # 数据库里代码摘要这张表是空的，需要新生成
            else:
                #file_path = os.path.join(root, file)
                codefile_abstract = generate_abstract(file_path)
                file_abstract[rel_path] = codefile_abstract
                # save_abstract_to_db(project_path, file, codefile_abstract, project_id)
                save_abstract_to_db(project_path, rel_path, codefile_abstract, project_id, current_user.user_id)

        # logger.info(file_abstract)
        # sys.exit()

        return jsonify({
            "status": "success",
            "data": file_abstract
        })
    except Exception as e:
        print(f"生成代码摘要失败: {str(e)}")
        traceback.print_exc()
        return jsonify({
            "status": "error",
            "message": f"生成代码摘要过程中出错: {str(e)}"
        }), 500


@bp.route('/api/get-code-abstract', methods=['POST'])
def abstract_code_from_project_post():
    data = request.json
    project_path = data.get('projectPath', '')
    project_id = data.get('project_id')
    abstract_length = data.get('abstract_length')
    all_align = data.get('all_align')
    user_id = current_user.user_id

    if not project_path:
        return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400

    code_file_path = os.path.join(project_path, 'code_repo')
    if not os.path.exists(code_file_path):
        return jsonify({'status': 'success', 'data': []})
    from tasks import abstract_code_from_project_task, align_requirement_to_project_task
    sig2 = align_requirement_to_project_task.s(data, user_id)
    sig2.freeze()
    task2_id = sig2.id
    sig1 = abstract_code_from_project_task.s(data, code_file_path, user_id)
    sig1.freeze()
    task1_id = sig1.id
    chain(sig1, sig2).apply_async()

    # 新增: 写入用户任务快照
    db = get_db()
    cursor = db.cursor()

    # 先清理该项目旧的任务快照(防止残留)
    cursor.execute("DELETE FROM user_task_snapshot WHERE user_id = %s AND project_id = %s",
                   (user_id, project_id))

    # 插入新的
    cursor.execute("""INSERT INTO user_task_snapshot 
                       (user_id, project_id, task_id, next_task_id, task_type, task1_total, task2_total, current_total,
                       state, title, is_running, task_category)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                   (
                       user_id, project_id, sig1.id, sig2.id, 'chain', abstract_length, all_align, abstract_length,
                       'PENDING', '代码摘要', 1, 'align'
                   ))

    return jsonify({'status': 'success', 'message': '链式任务已启动', 'task1_id': task1_id, 'task2_id': task2_id})


def get_project_abstract(project_id):
    """通过project_id获取所有摘要"""
    db = get_db_celery()
    cur = db.cursor()
    try:
        # logger.info(f'project_id==========={project_id}')
        cur.execute('select filename,abstract from abstracts where project_id=%s', (project_id,))
        rows = cur.fetchall()
        result = {row['filename']: row['abstract'] for row in rows}

    except Exception as e:
        logger.error(f'查询摘要失败:{str(e)}', exc_info=True)
        result = {}
    finally:
        db.close()
    # logger.info(f'获取摘要=============={result}')
    return result


# 对齐 需求->代码 需要异步处理
# @bp.route('/api/align-requirement-to-project', methods=['POST'])
# def align_requirement_to_project():
#     """
#     对齐功能：为单个doc需求点在项目中查找相关代码并返回codeRanges格式的结果
#     2026-01-21更新：为单个doc需求点在数据库中根据代码摘要检索相关代码文件，然后查找相关代码，返回codeRanges格式的结果
#     """
#     data = request.json
#     doc_ranges = data.get('docRanges', [])
#     # file_abstract = data.get('codeFileAbstract', {})
#     project_path = data.get('projectPath', '')
#     project_id = data.get('project_id')
#     alignment = data.get('alignment')
#     print('/api/align-requirement-to-project, project_id', project_id)
#     file_abstract = get_project_abstract(project_id)
#
#     from tasks import align_requirement_to_project_task
#     task = align_requirement_to_project_task.delay(project_path, doc_ranges, file_abstract, alignment, project_id,
#                                                    current_user.user_id)
#     return jsonify({'status': 'success', 'task_id': task.id})


# 需要异步处理
@bp.route('/api/align-requirement-to-project-addprompt', methods=['POST'])
def align_requirement_to_project_addprompt():
    """
    对齐功能：为单个doc需求点在项目中查找相关代码并返回codeRanges格式的结果
    2026-03-20更新：根据用户提示词优化输出结果
    """
    data = request.json
    doc_ranges = data.get('docRanges', [])
    abstract = data.get('codeFileAbstract', {})
    project_path = data.get('projectPath', '')
    codeRanges_aligned = data.get('codeRanges', [])
    userPrompt = data.get('userInputPrompt', [])
    project_id = data.get('project_id')
    user_id = getattr(current_user, 'user_id', None)
    print('/api/align-requirement-to-project-addprompt, project_id', project_id)
    # abstract = get_project_abstract(project_id)

    doc_name = doc_ranges[0]['filename']


    # 获取项目中所有代码文件
    code_repo_path = os.path.join(project_path, 'code_repo')
    all_files = get_all_files_with_relative_paths(code_repo_path, 'code')

    # 拼接所有docRanges的content作为requirement_text
    requirement_text = '\n\n'.join([doc_range.get('content', '') for doc_range in doc_ranges if doc_range.get('content')])
    if not requirement_text or not project_path:
        return jsonify({"status": "error", "message": "缺少需求内容或项目路径参数"}), 400

    # 如果有多个代码文件，执行检索代码摘要
    file_abstract = abstract.get('data', abstract) if isinstance(abstract, dict) else abstract
    if len(all_files) > 1:
        # 基于需求，利用大模型检索代码摘要，先定位代码文件
        # 调用llm
        file_name_list = query_codefile_from_abstract(requirement_text, file_abstract)
        
        # 解析异常，返回空列表时
        # 过滤掉含有乱码的代码摘要（作为被定位的代码文件防止遗漏），重新调用大模型定位代码文件
        FILE_MAX_LIMIT = ALIGN_FILE_ABSTRACT_BATCH_LIMIT
        if not file_name_list:
            filter_non_file, filter_file_abstract, file_cnt = filter_non_abstract_files(file_abstract)
            
            # 代码摘要数量小于阈值时，可以直接调用
            if file_cnt <=FILE_MAX_LIMIT:
                # 调用llm
                file_name_list = query_codefile_from_abstract(requirement_text, filter_file_abstract)
                #print(file_name_list)
            
            # 可能由于代码摘要过多，影响大模型分析理解而报错
            else:
                # 对代码摘要分批次处理
                file_name_list = []
                file_cnt = 0
                batch_file_abstract = {}
                for key, value in filter_file_abstract.items():
                    file_cnt += 1
                    batch_file_abstract[key] = value
                    if file_cnt >= FILE_MAX_LIMIT:
                        file_cnt = 0
                        batch_file_name_list = query_codefile_from_abstract(requirement_text, batch_file_abstract)
                        file_name_list += batch_file_name_list
            
            # 谨防遗漏，将有摘要是乱码的代码文件全部放入候选区
            if not file_name_list:
                file_name_list += filter_non_file 
            #print(file_name_list)
        
    # 如果只有一个代码文件，就不检索代码摘要
    else:
        file_name_list = all_files

    # 尝试初始化RAG引擎，以便在agent中使用
    # try:
        # rag_engine.initialize(project_path)
    # except Exception as e:
        # print(f"[Align] RAG initialize failed: {e}")

    try:
        code_ranges = []
        # 遍历经过代码摘要筛选的代码文件
        for file_name in file_name_list:
            # 假设 file_name 是一个字典，例如：{"filename": "main.py"}
            if isinstance(file_name, dict):
                file_name = file_name.get("filename", file_name.get("file", ""))

            # 确保 file_name 是字符串
            
            if not isinstance(file_name, (str, bytes, os.PathLike)):
                continue
            
            all_code_blocks = _get_or_build_code_blocks_for_file(project_path, file_name, project_id)
            if not all_code_blocks:
                continue

            # 调用对齐函数获取种子代码块
            seed_related_code = query_related_code_by_feedback(
                requirement_text,
                all_code_blocks,
                codeRanges_aligned,
                userPrompt,
                block_limit=50,
                user_id=user_id,
                project_path=project_path
            )
            
            try:
                seed_related_code = _trim_seed_related_code_items(seed_related_code)
                try:
                    seed_blocks = include_related_blocks(seed_related_code, all_code_blocks)
                except Exception as exc:
                    logger.info(f"related_id 扩展失败，使用基础匹配结果: {exc}")
                    seed_blocks = _match_related_items_to_code_blocks(seed_related_code, all_code_blocks)

                seed_blocks = _dedupe_code_blocks(seed_blocks)
                graph_blocks = _expand_code_blocks_with_call_graph_for_alignment(project_path, project_id, seed_blocks)
                candidate_blocks = _dedupe_code_blocks(seed_blocks + graph_blocks)

                final_blocks = []
                if graph_blocks:
                    reranked_code = query_related_code_graph_rerank(
                        requirement_text,
                        seed_blocks,
                        candidate_blocks
                    )
                    final_blocks = _match_related_items_to_code_blocks(reranked_code, candidate_blocks)

                if not final_blocks:
                    final_blocks = seed_blocks

                code_ranges.extend(_code_blocks_to_code_ranges(project_path, final_blocks))
                                
            except Exception as e:
                logger.error(f"add related_code failed: {e}", exc_info=True)
                
        #logger.info("对齐结果...................")
        #logger.info(code_ranges)
        #sys.exit()

        return jsonify({
            "status": "success",
            "codeRanges": code_ranges
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"对齐过程中出错: {str(e)}"
        }), 500



@bp.route('/api/generate-flowchart', methods=['POST'])
def generate_flowchart():
    try:
        data = request.get_json()
        code_content = data.get('codeContent')

        if not code_content:
            return jsonify({"status": "error", "message": "Missing code content"}), 400

        flowchart_code = query_flow_chart(code_content)

        return jsonify({
            "status": "success",
            "mermaidCode": flowchart_code,
            "flowchartCode": flowchart_code
        })

    except Exception as e:
        print(f"Error generating flowchart: {str(e)}")
        return jsonify({"status": "error", "message": f"Failed to generate flowchart: {str(e)}"}), 500

        
def fix_mermaid_syntax(mmd_code):
    """
    修复 LLM 生成的 Mermaid 代码中的语法错误：
    1. 给节点标签 [...] 中的内容加上双引号（如果包含特殊字符）。
    2. 确保边标签 |...| 格式正确。
    """
    if not mmd_code:
        return mmd_code

    # 1. 移除可能存在的多余空白或重复的 graph TD (如果它是独立的)
    lines = mmd_code.strip().split('\n')
    
    # 如果第一行是 graph TD，我们保留它，但确保后续内容正确
    # 这里我们主要处理节点标签 [...]
    
    fixed_lines = []
    for line in lines:
        # 正则匹配节点标签：[内容]
        # 匹配逻辑：找到 [ 和 ] 之间的内容
        def add_quotes(match):
            content = match.group(1)
            # 如果内容包含 =, -, /, (, ), 空格 等，必须加引号
            if re.search(r'[=\-/\(\)\s]', content):
                # 转义内容内部的双引号
                safe_content = content.replace('"', '\\"')
                return f'["{safe_content}"]'
            return match.group(0) # 不需要修改

        # 替换节点标签 [xxx] -> ["xxx"] (如果需要)
        # 注意：这可能会误伤边标签 |xxx|，所以要小心
        # 更安全的做法是只处理节点标签，不处理边标签
        
        # 使用正则替换节点标签
        # 匹配 [ 开头，中间内容，] 结尾，且不在 | ... | 内部
        # 简单的策略：匹配所有 [...] 并检查内容
        line = re.sub(r'\[([^\]]+)\]', add_quotes, line)
        
        fixed_lines.append(line)
    
    return '\n'.join(fixed_lines)      

    
# 改成从数据库查询
@bp.route('/api/generate-reverse-requirement', methods=['POST'])
def generate_reverse_requirement():
    try:
        data = request.get_json()
        requirement_content = data.get('requirementContent')
        code_content = data.get('codeContent')
        project_id = data.get('project_id')
        alignment_id = data.get('alignment_id')
        cache_only = bool(data.get('cacheOnly', False))
        force_regenerate = bool(data.get('forceRegenerate', False))

        if not project_id or not alignment_id:
            return jsonify({"status": "error", "message": "Missing project_id or alignment_id"}), 400
        if not cache_only and not code_content:
            return jsonify({"status": "error", "message": "Missing code content"}), 400

        db = get_db()
        cur = db.cursor()
        #cur.execute(f"select GenReq, GenMermaid from alignments where project_id={project_id} and id='{align_id}'")
        cur.execute(
            'SELECT GenReq, GenMermaid FROM alignments WHERE id=%s AND project_id=%s',
            (alignment_id, project_id)
        )
        row = cur.fetchone()
        has_cached_reverse = False
        if row is not None:
            has_cached_reverse = bool((row.get('GenReq') or '').strip()) or bool((row.get('GenMermaid') or '').strip())
        if has_cached_reverse and not force_regenerate:
            return jsonify({
                "status": "success",
                "cached": True,
                "generatedRequirement": row.get('GenReq'),
                "mermaidCode": row.get('GenMermaid')
            })
			
        if cache_only:
            return jsonify({
                "status": "success",
                "cached": False,
                "generatedRequirement": None,
                "mermaidCode": None
            })

        # 构建代码块列表，格式与现有函数兼容
        code_blocks = []
        if isinstance(code_content, list):
            for code_block in code_content:
                code_blocks.append({
                    'filename': code_block.get('filename', 'unknown'),
                    'content': code_block.get('content', '')
                })
        else:
            # 如果是字符串，创建单个代码块
            code_blocks.append({
                'filename': 'code',
                'content': code_content
            })

        # 调用LLM生成需求，传入参考需求内容
        generated_requirement = query_generated_requirement(code_blocks, requirement_content or "")

        # 调用LLM生成流程图
        mermaid_code = query_flow_chart(code_content if isinstance(code_content, str) else
                                       '\n\n'.join([block.get('content', '') for block in code_content]))

        # generated_requirement 是 LLM 返回的字符串
        # 关键步骤：处理流程图含有特殊字符无法渲染的问题
        #print(mermaid_code)
        mermaid_code = fix_mermaid_syntax(mermaid_code)
        #print(mermaid_code)
        
        
        cur.execute(
            'UPDATE alignments SET GenReq=%s, GenMermaid=%s, updatedAt=CURRENT_TIMESTAMP WHERE id=%s AND project_id=%s',
            (generated_requirement or '', mermaid_code or '', alignment_id, project_id)
        )
        
        
        
        return jsonify({
            "status": "success",
			"cached": False,
            "generatedRequirement": generated_requirement,
            "mermaidCode": mermaid_code
        })

    except Exception as e:
        print(f"Error generating reverse requirement: {str(e)}")
        return jsonify({"status": "error", "message": f"Failed to generate reverse requirement: {str(e)}"}), 500

 
    
# 将代码块数据写入对齐表，方便单个的纯代码审查
@bp.route('/api/add-code-data', methods=['POST'])
def add_code_data():
    """
    将代码块数据写入对齐表，用is_code_review=1来标识 纯代码审查的数据
    2026.6.9 更新：写入之前先查询alignments表有没有纯代码审查的数据，如果已存在就直接return
    """
    data = request.json
    project_id = data.get('project_id')
    code_blocks = data.get('code_blocks')

    db = get_db()
    c = db.cursor()
    c.execute(f'select count(*) count from alignments where project_id={project_id} and is_code_review=1')
    count = c.fetchone().get('count')
    if count > 0:
        return jsonify({"status": "success", "count": count, 'message': '已存在代码块数据'}), 200

    result = [{"doc_file": item['codeRanges'][0]['filename'], "alignment": item}
              for item in code_blocks]

    sql = "INSERT INTO alignments(id,user_id,project_id,name,reviewThoughts,docRanges,codeRanges," \
          "createdAt,updatedAt,is_code_review)" \
          "VALUES(%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP,1)"
    rows = [(item['alignment'].get('id'),
             current_user.user_id,
             project_id,
             item['alignment'].get('name'),
             item['alignment'].get('reviewThoughts'),
             json.dumps(item['alignment'].get('docRanges')),
             json.dumps(item['alignment'].get('codeRanges'))) for item in result]

    with db.cursor() as cursor:
        for i in range(0, len(rows), 500):
            batch = rows[i: i + 500]
            cursor.executemany(sql, batch)

    return jsonify({"status": "success", "message": "写入代码块成功"}), 200


@bp.route('/api/reverse_requirements', methods=['POST'])
def reverse_requirements():
    data = request.json
    project_id = data.get('project_id')
    selected_reverse = data.get('selected_reverse')
    user_id = current_user.user_id

    # sql = "SELECT * FROM alignments WHERE project_id=%s"
    params = [int(project_id)]

    align_types = [t for t in selected_reverse if t in ("req2code", "code2req")]
    has_code_review = "codeReview" in selected_reverse
    conditions = []

    if len(align_types) == 2:
        conditions.append("align_type IS NOT NULL")
    elif len(align_types) == 1:
        conditions.append("align_type = %s")
        params.append(align_types[0])

    if has_code_review:
        conditions.append("is_code_review = 1")

    where_extra = ""
    if conditions:
        if len(conditions) == 1:
            where_extra = " AND " + conditions[0]
        else:
            where_extra = " AND (" + " OR ".join(conditions) + ")"

    db = get_db()
    cursor = db.cursor()

    # 查询统计
    stats_sql = f"""
        SELECT
            COUNT(*) as total,
            SUM(CASE WHEN COALESCE(GenReq, '') = '' THEN 1 ELSE 0 END) as pending,
            SUM(CASE WHEN COALESCE(GenReq, '') <> '' THEN 1 ELSE 0 END) as done
        FROM alignments
        WHERE project_id = %s AND codeRanges != '[]' {where_extra}   
    """
    # print(stats_sql)
    cursor.execute(stats_sql, params)
    row = cursor.fetchone()

    if not row['total'] or not row['pending']:
        return jsonify({'status': 'error', 'message': '没有可以需求反生成的数据'})

    # 查询数据
    data_sql = f"SELECT * FROM alignments WHERE project_id = %s " \
               f"AND codeRanges != '[]' AND (GenReq is null OR GenReq = '') {where_extra}"
    cursor.execute(data_sql, params)
    alignments = cursor.fetchall()

    from tasks import gen_requirement_task
    task = gen_requirement_task.delay(alignments, int(row['total']), int(row['done']))

    # 先清理该项目旧的任务快照(防止残留)
    cursor.execute("DELETE FROM user_task_snapshot WHERE user_id = %s AND project_id = %s",
                   (user_id, project_id))

    # 插入新的
    cursor.execute("""INSERT INTO user_task_snapshot 
                               (user_id, project_id, task_id, task_type, task1_total, current_total,
                               state, title, is_running, task_category)
                               VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                   (
                       user_id, project_id, task.id, 'single', int(row['total']), int(row['total']),
                       'PENDING', '需求反生成', 1, 'reverse'
                   ))

    return jsonify({
        "status": "success",
        "task_id": task.id,
        "stats": {
            "total": row['total'],
            "pending": row['pending'],
            "generated": row['done'],
        }
    })


# 审查 需要异步处理
@bp.route('/api/review-alignment', methods=['POST'])
def review_alignment():
    data = request.json
    project_path = data.get('projectPath')
    # doc_file = data.get('docFile')
    # alignments = data.get('alignments')
    project_id = data.get('project_id')
    files = data.get('requirement_files')
    code_blocks = data.get('code_blocks')
    prompt_type = data.get('promptType', '')
    reviewed_count = data.get('reviewedCount', '')
    total_review_count = data.get('totalReviewCount', '')
    user_id = current_user.user_id

    has_requirement_files = isinstance(files, dict) and any(files.values())
    has_code_blocks = isinstance(code_blocks, list) and len(code_blocks) > 0
    if not all([project_path, project_id]) or not (has_requirement_files or has_code_blocks):
        return jsonify({"status": "error", "message": "Missing required parameters"}), 400

    from tasks import review_alignment_task
    task = review_alignment_task.delay(project_path, project_id, user_id, files, prompt_type, reviewed_count)

    # 新增: 写入用户任务快照
    db = get_db()
    cursor = db.cursor()

    # 先清理该项目旧的任务快照(防止残留)
    cursor.execute("DELETE FROM user_task_snapshot WHERE user_id = %s AND project_id = %s",
                   (user_id, project_id))

    # 插入新的
    cursor.execute("""INSERT INTO user_task_snapshot 
                           (user_id, project_id, task_id, task_type, task1_total, current_total,
                           state, title, is_running, task_category)
                           VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                   (
                       user_id, project_id, task.id, 'single', total_review_count, total_review_count,
                       'PENDING', '自动审查', 1, 'review'
                   ))

    return jsonify({"status": "success", "task_id": task.id})


# 需要异步处理
@bp.route('/api/review-alignment-addprompt', methods=['POST'])
def review_alignment_addprompt():
    from tasks import _normalize_kb_type_for_use, gen_requirement
    data = request.json
    project_path = data.get('projectPath')
    doc_file = data.get('docFile')
    alignment = data.get('alignment')
    userPrompt = data.get('userInputPrompt', [])
    project_id = data.get('project_id')
    prompt_type = data.get('promptType')

    if not all([project_path, project_id, alignment]) and not prompt_type:
        return jsonify({"status": "error", "message": "Missing required parameters"}), 400

    # 尝试初始化RAG引擎
    # try:
        # rag_engine.initialize()
    # except Exception as e:
        # print(f"[Review] RAG initialize failed: {e}")

    # 获取选定的 knowledge base
    selected_rule_kbs = []
    selected_issue_kbs = []
    try:
        metadata_file = os.path.join(project_path, 'metadata.json')
        if os.path.exists(metadata_file):
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
                selected_kbs = metadata.get('selected_kbs', [])
                selected_rule_kbs = [kb['name'] for kb in selected_kbs if
                                     _normalize_kb_type_for_use(kb.get('type')) == 'rule']
                selected_issue_kbs = [kb['name'] for kb in selected_kbs if
                                      _normalize_kb_type_for_use(kb.get('type')) == 'issue']
    except Exception as e:
        logger.error(str(e), exc_info=True)

    # 检索上下文
    retrieved_rules = []
    retrieved_issues = []

    doc_ranges = alignment.get('docRanges', [])
    code_ranges = alignment.get('codeRanges', [])
    reviewThoughts = alignment.get('reviewThoughts', [])

    # 构造查询文本
    query_text = ""
    if doc_ranges:
        query_text += doc_ranges[0].get('content', '') + "\n"
    if code_ranges:
        query_text += code_ranges[0].get('content', '')

    # 检索规则
    for kb_name in selected_rule_kbs:
        collection = rag_engine.get_collection('rule', kb_name)
        if collection:
            results = collection.query(query_texts=[query_text], n_results=3)
            if results and results['documents']:
                for doc in results['documents'][0]:
                    retrieved_rules.append(doc)

    # 检索问题单
    for kb_name in selected_issue_kbs:
        collection = rag_engine.get_collection('issue', kb_name)
        if collection:
            results = collection.query(query_texts=[query_text], n_results=3)
            if results and results['documents']:
                for doc in results['documents'][0]:
                    retrieved_issues.append(doc)

    # 1. 调用 agent 获取审查结果
    review_process, issue = query_review_result_by_feedback(
        doc_ranges,
        code_ranges,
        reviewThoughts,
        userPrompt,
        rules=retrieved_rules,
        issues=retrieved_issues,
        user_id=current_user.user_id,
        project_path=project_path,
        prompt_type=prompt_type)

    # 2. 更新对齐关系
    alignment['isReviewed'] = True
    alignment['reviewThoughts'] = review_process

    if isinstance(issue, list):
        issues_list = [x for x in issue if isinstance(x, dict)]
    elif isinstance(issue, dict):
        issues_list = [issue]
    else:
        issues_list = []

    # 需求反生成
    #generated_requirement, mermaid_code = gen_requirement(doc_ranges, code_ranges)
    generated_requirement = ""
    mermaid_code =""
    
    conn = get_db()
    cur = conn.cursor()

    try:
        cur.execute(
            'UPDATE alignments SET isReviewed=1, reviewThoughts=%s, GenReq=%s, GenMermaid=%s, updatedAt=CURRENT_TIMESTAMP '
            'WHERE id=%s and project_id=%s',
            (alignment.get('reviewThoughts') or '', generated_requirement or '', mermaid_code or '',
             alignment.get('id'), project_id)
        )

        if issues_list:
            cur.execute(f"SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%' and project_id={project_id}")
            used = set()
            for r in cur.fetchall():
                disp = r['displayId']
                if disp and disp.startswith('ISSUE-'):
                    try:
                        used.add(int(disp.split('-')[1]))
                    except Exception as e:
                        logger.error(str(e), exc_info=True)

            next_number = (max(used) + 1) if used else 1

            #brief_req = alignment.get('docRanges', [{}])[0].get('content', '') or ''
            #brief_code = alignment.get('codeRanges', [{}])[0].get('content', '') or ''

            doc_ranges_safe = alignment.get('docRanges', []) or []
            code_ranges_safe = alignment.get('codeRanges', []) or []
            brief_req = doc_ranges_safe[0].get('content', '') if doc_ranges_safe else ''
            brief_code = code_ranges_safe[0].get('content', '') if code_ranges_safe else ''
            related_doc_file = (
                doc_file
                or (doc_ranges_safe[0].get('filename', '') if doc_ranges_safe else '')
                or (code_ranges_safe[0].get('filename', '') if code_ranges_safe else '')
            )

            for one in issues_list:
                display_id = f"ISSUE-{next_number:03d}"
                next_number += 1

                severity = one.get('level') or one.get('severity')
                title = one.get('summary') or one.get('title') or ''
                content = one.get('description') or one.get('content') or ''
                status = one.get('status') or 'unconfirmed'

                cur.execute(
                    'INSERT INTO issues(user_id,project_id,displayId,alignmentId,severity,title,content,status,'
                    'relatedDocFile,relatedRequirementId,briefRequirement,briefCode,category,createdAt,updatedAt) '
                    'VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
                    (
                        current_user.user_id,
                        project_id,
                        display_id,
                        alignment.get('id'),
                        severity,
                        title,
                        content,
                        status,
                        related_doc_file if related_doc_file else alignment.get('codeRanges').get('filename'),
                        alignment.get('id'),
                        brief_req,
                        brief_code,
                        alignment.get('align_type') if alignment.get('align_type') else 'codeReview'
                    )
                )

    except Exception as e:
        conn.rollback()
        logger.error(f"审查失败 user_prompt2: {str(e)}", exc_info=True)
        return {"status": "error", "message": f"Failed to save review result: {str(e)}"}

    return {"status": "success", "createdIssues": len(issues_list)}


@bp.route('/api/clear-alignment-review', methods=['POST'])
def clear_alignment_review():
    """Clear review state for a single alignment and delete its related issues (DB)."""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        alignment_id = data.get('alignmentId')
        project_id = data.get('project_id')
        if not all([project_path, alignment_id]):
            return jsonify({"status": "error", "message": "Missing required parameters"}), 400
        # init_project_db(project_path)
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        cur.execute('UPDATE alignments SET isReviewed=0, reviewThoughts="" WHERE id=%s and project_id=%s',
                    (alignment_id, project_id))
        cur.execute('DELETE FROM issues WHERE alignmentId=%s and project_id=%s', (alignment_id, project_id))
        removed = cur.rowcount
        # conn.commit()
        # conn.close()
        return jsonify({"status": "success", "removedIssues": removed})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})




def get_filename_without_extension(filename):
    """去掉文件名的扩展名"""
    return os.path.splitext(filename)[0]


def _parse_csv_query_arg(value):
    if not value:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in str(value).split(',') if item.strip()]


def _safe_json_loads(value, default):
    try:
        return pyjson.loads(value or '[]')
    except Exception:
        return default


def _normalize_alignment_row(row):
    return {
        'id': row['id'],
        'name': row['name'],
        'isReviewed': bool(row['isReviewed']),
        'reviewThoughts': row['reviewThoughts'] or '',
        'docRanges': _safe_json_loads(row.get('docRanges'), []),
        'codeRanges': _safe_json_loads(row.get('codeRanges'), []),
        'isCodeReview': row.get('is_code_review', 0),
        'is_alignment': row.get('is_alignment', 0),
        'align_type': row['align_type']
    }


def _matches_range_file(ranges, target_file):
    if not target_file:
        return True
    if not isinstance(ranges, list):
        return False
    for one in ranges:
        if not isinstance(one, dict):
            continue
        if one.get('filename') == target_file or one.get('documentId') == target_file:
            return True
    return False


def _alignment_status(alignment):
    no_doc = not alignment.get('docRanges')
    no_code = not alignment.get('codeRanges')
    if no_doc or no_code:
        return 'unaligned'
    if alignment.get('isReviewed'):
        return 'reviewed'
    return 'unreviewed'


def _sort_alignments(items):
    def _sort_key(item):
        item_id = item.get('id') or ''
        is_auto = item_id.startswith('auto_')
        name = item.get('name') or ''
        if is_auto:
            start = 0
            if item.get('docRanges'):
                start = item['docRanges'][0].get('start', 0)
            return (1, 0, start, name)

        num_match = re.match(r'^(\d+)', name)
        if num_match:
            return (0, 0, int(num_match.group(1)), name)
        return (0, 1, name, item_id)

    return sorted(items, key=_sort_key)


def _paginate_items(items, page, page_size):
    total = len(items)
    safe_page_size = max(1, min(page_size, 500))
    total_pages = max(1, (total + safe_page_size - 1) // safe_page_size)
    safe_page = min(max(page, 1), total_pages)
    start = (safe_page - 1) * safe_page_size
    end = start + safe_page_size
    return {
        'items': items[start:end],
        'pagination': {
            'total': total,
            'page': safe_page,
            'page_size': safe_page_size,
            'total_pages': total_pages,
            'pages': total_pages
        }
    }


def _load_project_alignments(project_id):
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        "SELECT id,name,isReviewed,reviewThoughts,docRanges,codeRanges,is_code_review,is_alignment,align_type "
        "FROM alignments WHERE project_id=%s",
        (project_id,)
    )
    rows = cur.fetchall()
    return [_normalize_alignment_row(row) for row in rows]


def _filter_alignments(items, *, file_path=None, kind='doc', selected_doc_file=None,
                       selected_code_file=None, view_mode=None, status_filters=None,
                       include_code_review=None):
    result = items

    if include_code_review is not None:
        val = str(include_code_review).lower()
        if val in ('0', 'false', 'no'):
            # 排除所有 is_code_review=1 的对齐
            result = [item for item in result if not bool(item.get('isCodeReview'))]
        elif val == 'reviewed_only':
            # 排除 "is_code_review=1 且未审查" 的对齐（保留已审查的纯代码对齐）
            result = [item for item in result
                      if not item.get('isCodeReview') or item.get('isReviewed')]
        else:
            # 1/true/yes -> 包含所有
            pass

    if file_path:
        if kind == 'code':
            result = [item for item in result if _matches_range_file(item.get('codeRanges'), file_path)]
        else:
            result = [item for item in result if _matches_range_file(item.get('docRanges'), file_path)]

    if view_mode == 'current':
        result = [
            item for item in result
            if (selected_doc_file and _matches_range_file(item.get('docRanges'), selected_doc_file)) or
               (selected_code_file and _matches_range_file(item.get('codeRanges'), selected_code_file))
        ]
    elif view_mode in (None, '', 'all'):
        # “全部数据”模式下不要再叠加当前文件过滤，否则前端切换按钮看起来会失效。
        pass

    if status_filters:
        allowed = set(status_filters)
        result = [item for item in result if _alignment_status(item) in allowed]

    return _sort_alignments(result)


def _ranges_overlap(a_start, a_end, b_start, b_end):
    try:
        return max(float(a_start), float(b_start)) < min(float(a_end), float(b_end))
    except Exception:
        return False


def _line_ranges_overlap(a_start, a_end, b_start, b_end):
    try:
        return max(int(a_start), int(b_start)) <= min(int(a_end), int(b_end))
    except Exception:
        return False

@bp.route('/project/alignments', methods=['GET'])
def get_alignments():
    #print("request.args:", request.args)
    """获取项目对齐关系，支持分页和过滤"""
    project_path = request.args.get('path')
    file_path = request.args.get('file')
    kind = request.args.get('kind', 'doc')
    project_id = request.args.get('project_id')
    selected_doc_file = request.args.get('selected_doc_file')
    selected_code_file = request.args.get('selected_code_file')
    view_mode = request.args.get('view_mode')
    status_filters = _parse_csv_query_arg(request.args.get('status_filters'))
    include_code_review = request.args.get('include_code_review')
    page = request.args.get('page', type=int)
    page_size = request.args.get('page_size', type=int)
    align_type = request.args.get('align_type')
    print('get(`/project/alignments, project_id:', project_id)
    
    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径参数。"}), 400

    try:
        alignments = _load_project_alignments(project_id)
        filtered = _filter_alignments(
            alignments,
            file_path=file_path,
            kind=kind,
            selected_doc_file=selected_doc_file,
            selected_code_file=selected_code_file,
            view_mode=view_mode,
            status_filters=status_filters,
            include_code_review=include_code_review
        )

        if align_type:
            filtered = [item for item in filtered if item.get('align_type') == align_type]
        if page:
            paged = _paginate_items(filtered, page or 1, page_size)
            return jsonify({
                "status": "success",
                "data": paged['items'],
                "total": paged['pagination']['total'],
                "page": paged['pagination']['page'],
                "page_size": paged['pagination']['page_size']
            }), 200

        result = {alignment['id']: alignment for alignment in filtered}

        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:

        return jsonify({"status": "error", "message": f"读取对齐数据失败: {e}"}), 500

@bp.route('/project/alignments', methods=['POST'])
def add_alignment():
    """添加或更新对齐关系到数据库，并自动创建相关的块"""
    project_path = request.args.get('path')
    project_id = request.args.get('project_id')
    new_alignment = request.json
    if not project_path or not new_alignment or 'id' not in new_alignment:
        return jsonify({"status": "error", "message": "缺少项目路径或无效的对齐数据。"}), 400

    is_alignment = 0
    if new_alignment.get('is_alignment'):
        is_alignment = new_alignment.get('is_alignment')
    if new_alignment.get('docRanges') and new_alignment.get('codeRanges'):
        is_alignment = 1

    # --- 自动创建块的逻辑 ---
    try:
        # 1. 处理需求块
        doc_ranges = new_alignment.get('docRanges', [])
        if doc_ranges:
            append_missing_doc_blocks(project_path, doc_ranges)

        # 2. 处理代码块
        code_ranges = new_alignment.get('codeRanges', [])
        if code_ranges:
            append_missing_code_blocks(project_path, code_ranges)

    except Exception as e:
        print(f"Error auto-creating blocks: {e}")
        # 即使块创建失败，也不应该阻止对齐关系的保存，但最好记录日志
        pass

    generated_requirement = ''
    mermaid_code = ''

    # 反生成需求+流程图
    # try:
        # requirement_content = new_alignment.get('docRanges')
        # code_content = new_alignment.get('codeRanges')

        # # if not code_content:
        # #     return jsonify({"status": "error", "message": "Missing code content"}), 400

        # # 构建代码块列表，格式与现有函数兼容
        # code_blocks = []
        # if isinstance(code_content, list):
            # for code_block in code_content:
                # code_blocks.append({
                    # 'filename': code_block.get('filename', 'unknown'),
                    # 'content': code_block.get('content', '')
                # })
        # else:
            # # 如果是字符串，创建单个代码块
            # code_blocks.append({
                # 'filename': 'code',
                # 'content': code_content
            # })

        # # 调用LLM生成需求，传入参考需求内容
        # generated_requirement = query_generated_requirement(code_blocks, requirement_content or "")

        # # 调用LLM生成流程图
        # mermaid_code = query_flow_chart(code_content if isinstance(code_content, str) else
                                       # '\n\n'.join([block.get('content', '') for block in code_content]))


    # except Exception as e:
        # print(f"Error generating reverse requirement: {str(e)}")
        # logger.info(f"Error generating reverse requirement: {str(e)}")
        # #return jsonify({"status": "error", "message": f"Failed to generate reverse requirement: {str(e)}"}), 500

    try:
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        
        cur.execute(
            '''
            INSERT INTO alignments(id, user_id, project_id, name, isReviewed, reviewThoughts, docRanges, codeRanges, 
            GenReq, GenMermaid, createdAt, updatedAt, is_code_review, is_alignment, align_type) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, %s, %s, %s) 
            ON DUPLICATE KEY UPDATE 
                name = VALUES(name),
                isReviewed = VALUES(isReviewed),
                reviewThoughts = VALUES(reviewThoughts),
                docRanges = VALUES(docRanges),
                codeRanges = VALUES(codeRanges),
                GenReq = VALUES(GenReq),
                GenMermaid = VALUES(GenMermaid),
                updatedAt = CURRENT_TIMESTAMP,
                is_alignment = VALUES(is_alignment)
            ''',
            (
                new_alignment.get('id'),
                current_user.user_id,
                project_id,
                new_alignment.get('name'),
                1 if new_alignment.get('isReviewed') else 0,
                new_alignment.get('reviewThoughts') or '',
                json.dumps(new_alignment.get('docRanges') or []),
                json.dumps(code_ranges or []),
                generated_requirement or '',
                mermaid_code or '',
                0,
                is_alignment,
                new_alignment.get('align_type') or ('req2code' if new_alignment.get('docRanges') else 'code2req')
            )
        )

        
        # conn.commit()
        # conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        traceback.print_exc()
        logger.info("写入对齐数据失败")
        return jsonify({"status": "error", "message": f"写入对齐数据失败: {e}"}), 500



@bp.route('/project/alignment', methods=['DELETE'])
def delete_alignment():
    """删除一个对齐关系（数据库）"""
    project_path = request.args.get('path')
    alignment_id = request.args.get('id')
    project_id = request.args.get('project_id')
    if not all([project_path, alignment_id]):
        return jsonify({"status": "error", "message": "缺少项目路径或对齐ID参数。"}), 400
    try:
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        cur.execute('DELETE FROM alignments WHERE id=%s and project_id=%s', (alignment_id, project_id))
        # conn.commit()
        # conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"删除对齐项时出错: {e}"}), 500


@bp.route('/project/issues', methods=['GET'])
def get_issues():
    try:
        project_path = request.args.get('path')
        project_id = request.args.get('project_id')
        issue_page = request.args.get('issue_page', 1, type=int)
        issue_page_size = request.args.get('issue_page_size', 20, type=int)
        category = request.args.get('category', '')

        if issue_page < 1:
            issue_page = 1
        if issue_page_size < 1 or issue_page_size > 100:
            issue_page_size = 20

        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})

        offset = (issue_page - 1) * issue_page_size

        
        
        
        conn = get_db()
        cur = conn.cursor()

        where_clause = f"project_id={project_id} "
        if category:
            where_clause += f"AND category='{category}' "

        cur.execute(f"select count(*) as total from issues where {where_clause}")
        total = cur.fetchone()['total']

        cur.execute(
            """
            SELECT
                COUNT(*) AS total,
                SUM(CASE WHEN status = 'confirmed' THEN 1 ELSE 0 END) AS confirmed,
                SUM(CASE WHEN status = 'false_positive' THEN 1 ELSE 0 END) AS false_positive,
                SUM(CASE WHEN status IS NULL OR status = '' OR status = 'unconfirmed' THEN 1 ELSE 0 END) AS unchecked
            FROM issues
            WHERE project_id = %s
            """,
            (project_id,)
        )
        stats_row = cur.fetchone() or {}
        issue_status_stats = {
            'all': int(stats_row.get('total') or 0),
            'confirmed': int(stats_row.get('confirmed') or 0),
            'false_positive': int(stats_row.get('false_positive') or 0),
            'unchecked': int(stats_row.get('unchecked') or 0),
        }

        cur.execute(f'SELECT * FROM issues where {where_clause} ORDER BY id DESC LIMIT %s OFFSET %s',
                    (issue_page_size, offset))
        rows = cur.fetchall()
        issues = []

        
        for r in rows:
            issues.append({
                'id': str(r['id']),
                'displayId': r['displayId'],
                'alignmentId': r['alignmentId'],
                'level': r['severity'],
                'summary': r['title'] or '',
                'description': r['content'] or '',
                'status': r['status'] or 'unconfirmed',
                'relatedDocFile': r['relatedDocFile'] or '',
                'relatedRequirementId': r['relatedRequirementId'] or '',
                'briefRequirement': r['briefRequirement'] or '',
                'briefCode': r['briefCode'] or '',
                'createdDate': r['createdAt'],
                'updatedDate': r['updatedAt'],
                'category': r['category']
            })
        
        # conn.close()
        return jsonify({'status': 'success', 'data': issues,
                        'total': total, 'issue_page': issue_page, 'issue_page_size': issue_page_size,
                        'issue_status_stats': issue_status_stats})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


#没找到前端的调用地方，前端只有get put delete，先注释了
# @bp.route('/project/issues', methods=['POST'])
# def add_issue():
#     try:
#         project_path = request.args.get('path')
#         if not project_path:
#             return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
#         issue_data = request.json or {}
#         conn = get_db_conn(project_path)
#         cur = conn.cursor()
#         # 分配展示ID
#         if not issue_data.get('displayId'):
#             cur.execute("SELECT displayId FROM issues WHERE displayId LIKE 'ISSUE-%'")
#             used = set()
#             for r in cur.fetchall():
#                 disp = r['displayId']
#                 if disp and disp.startswith('ISSUE-'):
#                     try:
#                         used.add(int(disp.split('-')[1]))
#                     except Exception:
#                         pass
#             next_number = (max(used) + 1) if used else 1
#             issue_data['displayId'] = f"ISSUE-{next_number:03d}"
#         cur.execute(
#             'INSERT INTO issues(displayId,alignmentId,severity,title,content,status,createdAt,updatedAt) '
#             'VALUES (%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)',
#             (
#                 issue_data.get('displayId'),
#                 issue_data.get('alignmentId'),
#                 issue_data.get('level') or issue_data.get('severity'),
#                 issue_data.get('title'),
#                 issue_data.get('description') or issue_data.get('content'),
#                 issue_data.get('status')
#             )
#         )
#         conn.commit()
#         conn.close()
#         return jsonify({'status': 'success', 'message': '问题单添加成功'})
#     except Exception as e:
#         return jsonify({'status': 'error', 'message': str(e)})

@bp.route('/project/issues/<issue_id>', methods=['PUT'])
def update_issue(issue_id):
    try:
        project_path = request.args.get('path')
        project_id = request.args.get('project_id')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        if not project_id:
            return jsonify({'status': 'error', 'message': '缺少 project_id 参数'})
        issue_data = request.json or {}
        # init_project_db(project_path)
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        cur.execute('UPDATE issues SET severity=%s, title=%s, content=%s, status=%s, updatedAt=CURRENT_TIMESTAMP WHERE id=%s and project_id=%s', (
            issue_data.get('level') or issue_data.get('severity'),
            issue_data.get('summary'), # 前端传的是summary，不是title
            issue_data.get('description') or issue_data.get('content'),
            issue_data.get('status'),
            issue_id,
            project_id
        ))
        if cur.rowcount == 0:
            conn.rollback()
            return jsonify({'status': 'error', 'message': '未找到要更新的问题单'}), 404
        conn.commit()
        # conn.close()
        return jsonify({'status': 'success', 'message': '问题单更新成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/issues/batch-update', methods=['POST'])
def batch_update_issues():
    try:
        project_path = request.args.get('path')
        project_id = request.args.get('project_id')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        if not project_id:
            return jsonify({'status': 'error', 'message': '缺少 project_id 参数'})

        data = request.json or {}
        issue_ids = data.get('issue_ids', [])
        status = data.get('status')

        if not issue_ids:
            return jsonify({'status': 'error', 'message': '未提供要更新的问题单 ID'})
        if not status:
            return jsonify({'status': 'error', 'message': '未提供目标状态'})

        conn = get_db()
        cur = conn.cursor()

        format_strings = ','.join(['%s'] * len(issue_ids))
        sql = f'UPDATE issues SET status=%s, updatedAt=CURRENT_TIMESTAMP WHERE id IN ({format_strings}) AND project_id=%s'
        cur.execute(sql, (status, *issue_ids, project_id))
        updated_count = cur.rowcount
        conn.commit()

        return jsonify({
            'status': 'success',
            'message': f'已更新 {updated_count} 条问题单',
            'updated_count': updated_count
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/issues/<issue_id>', methods=['DELETE'])
def delete_issue(issue_id):
    try:
        project_path = request.args.get('path')
        project_id = request.args.get('project_id')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径参数'})
        # init_project_db(project_path)
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        cur.execute('DELETE FROM issues WHERE id=%s and project_id=%s', (issue_id, project_id))
        # conn.commit()
        # conn.close()
        return jsonify({'status': 'success', 'message': '问题单删除成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/api/clear-project-results', methods=['POST'])
def clear_project_results():
    """清空项目的所有结果：需求片段、对齐结果、审查结果、问题单结果"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        project_id = data.get('project_id')

        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        if not os.path.exists(project_path):
            return jsonify({'status': 'error', 'message': '项目路径不存在'})

        # init_project_db(project_path)
        # conn = get_db_conn(project_path)
        conn = get_db()
        conn = conn.cursor()
        conn.execute(f'DELETE FROM issues where project_id={project_id}')
        conn.execute(f'DELETE FROM alignments where project_id={project_id}')
        # conn.commit()
        # conn.close()

        return jsonify({
            'status': 'success',
            'message': '所有结果已清空'
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/api/clear-code-ranges', methods=['POST'])
def clear_code_ranges():
    """清空项目中的所有对齐关系(id_code_review = 0)"""
    data = request.json
    project_path = data.get('projectPath')
    project_id = data.get('project_id')

    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径参数"}), 400

    try:
        # init_project_db(project_path)
        # conn = get_db_conn(project_path)
        conn = get_db()

        cur = conn.cursor()
        cur.execute(f'DELETE FROM alignments where project_id={project_id} and is_code_review=0')
        # conn.commit()

        # finally:
            # conn.close()

        return jsonify({"status": "success", "message": "已清空项目的对齐关系"})

    except Exception as e:
        return jsonify({"status": "error", "message": f"清空对齐关系失败: {str(e)}"}), 500


@bp.route('/api/clear-review-results', methods=['POST'])
def clear_review_results():
    """清空项目的审查结果：问题单结果，并重置对齐关系的审查状态"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        project_id = data.get('project_id')
        review_type = data.get('reviewType')

        is_code_review = 1 if review_type else 0
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        if not os.path.exists(project_path):
            return jsonify({'status': 'error', 'message': '项目路径不存在'})

        #conn = get_db_conn(project_path)
        
        conn = get_db()
        cur = conn.cursor()
        cur.execute(f'DELETE FROM issues where project_id={project_id}')
        cur.execute(f'UPDATE alignments SET isReviewed=0, reviewThoughts="" where project_id={project_id} and is_code_review={is_code_review}')
        # try:
        #conn.execute(f'DELETE FROM issues where project_id={project_id}')
        #conn.execute(f'UPDATE alignments SET isReviewed=0, reviewThoughts="" where project_id={project_id}')
            # conn.commit()
        # finally:
        #     conn.close()

        return jsonify({
            'status': 'success',
            'message': '审查结果已清空'
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/project/issue/update', methods=['POST'])
def update_issue_content():
    data = request.json
    project_path = data.get('path')
    issue_id = data.get('issueId')
    new_description = data.get('description')
    new_status = data.get('status')
    new_level = data.get('level')  # 添加问题级别参数
    project_id = data.get('project_id')

    if not all([project_path, issue_id]):
        return jsonify({"status": "error", "message": "缺少项目路径或问题单ID"}), 400

    try:
        # conn = get_db_conn(project_path)
        conn = get_db()
        cur = conn.cursor()
        cur.execute('UPDATE issues SET content=%s, status=%s, severity=%s, updatedAt=CURRENT_TIMESTAMP WHERE id=%s and project_id=%s', (
            new_description,
            new_status,
            new_level,
            issue_id,
            project_id
        ))
        if cur.rowcount == 0:
            # conn.close()
            return jsonify({"status": "error", "message": "未找到指定ID的问题单"}), 404
        # conn.commit()
        # conn.close()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": f"更新问题单失败: {str(e)}"}), 500

@bp.route('/api/code-decomposition', methods=['POST'])
def code_decomposition():
    """使用调用图 tree-sitter 解析产物同步代码块和调用图。"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        code_repo_path = os.path.join(project_path, 'code_repo')
        if not os.path.exists(code_repo_path):
            return jsonify({'status': 'error', 'message': '代码目录不存在'})

        call_graph_result = build_project_call_graph(project_path)
        call_graph_metadata = call_graph_result.get('metadata') or get_call_graph_metadata(project_path)
        call_graph_status = call_graph_metadata.get('status')
        payload = call_graph_result.get('payload') or {}
        all_code_blocks = payload.get('code_blocks') or []

        if call_graph_status != 'ready':
            return jsonify({
                'status': 'error',
                'message': call_graph_result.get('message') or _call_graph_status_message(call_graph_metadata),
                'processedCount': 0,
                'call_graph_status': call_graph_status,
                'call_graph_message': call_graph_result.get('message') or _call_graph_status_message(call_graph_metadata),
            })

        replace_code_blocks(project_path, all_code_blocks)
        processed_count = len(all_code_blocks)

        return jsonify({
            'status': 'success',
            'message': '代码分解和调用图构建完成，结果已保存',
            'processedCount': processed_count,
            'call_graph_status': call_graph_status,
            'call_graph_message': call_graph_result.get('message') or _call_graph_status_message(call_graph_metadata),
        })
        
    except Exception as e:
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})

@bp.route('/api/call-graph/preview', methods=['POST'])
def call_graph_preview():
    try:
        data = request.get_json() or {}
        project_path = data.get('projectPath')
        project_id = data.get('project_id')
        file = data.get('file')
        start_line = _safe_int(data.get('startLine'))
        end_line = _safe_int(data.get('endLine'))
        max_depth = max(1, min(_safe_int(data.get('maxDepth'), 3), 8))
        start_mode = (data.get('startMode') or 'function').strip().lower()
        bidirectional = bool(data.get('bidirectional'))
        graph_direction = 'both' if bidirectional else 'forward'
        if start_mode not in {'function', 'line'}:
            start_mode = 'function'

        if not project_path or not file or start_line <= 0 or end_line <= 0:
            return jsonify({
                'status': 'error',
                'call_graph_status': 'unavailable',
                'message': '缺少调用图预览所需参数'
            }), 400
 
        ensure_result = ensure_project_call_graph(project_path)
        metadata = ensure_result.get('metadata') or get_call_graph_metadata(project_path)
        if not ensure_result.get('payload'):
            return jsonify({
                'status': 'error',
                'call_graph_status': metadata.get('status'),
                'message': ensure_result.get('message') or _call_graph_status_message(metadata),
            })
 
        resolved_project_id = resolve_project_id(project_path, project_id)
        if start_mode == 'line':
            called_functions = resolve_called_functions_in_line_range(
                project_path,
                file,
                start_line,
                end_line,
            )
            if not called_functions:
                return jsonify({
                    'status': 'error',
                    'call_graph_status': metadata.get('status'),
                    'query_mode': start_mode,
                    'message': '当前选中行内未命中可查询调用图的函数调用',
                })

            preview_bundle = query_multiple_function_graphs(
                project_path,
                [item.get('function_id') for item in called_functions],
                max_depth=max_depth,
                direction=graph_direction,
            )
            selection_code_range = build_selection_code_range(
                project_path,
                file,
                start_line,
                end_line,
            )
            selection_block = {
                'id': None,
                'name': selection_code_range.get('name'),
                'type': 'selection',
                'file': file,
                'filename': file,
                'range': [start_line, end_line],
                'codeRange': selection_code_range,
            }
            previews = []
            for index, item in enumerate(preview_bundle.get('previews') or []):
                preview_item = dict(item)
                preview_item['call_site_line'] = called_functions[index].get('call_site_line')
                preview_item['call_site_name'] = called_functions[index].get('raw_callee')
                previews.append(preview_item)

            return jsonify({
                'status': 'success',
                'call_graph_status': metadata.get('status'),
                'query_mode': start_mode,
                'message': (
                    f'调用图预览已生成，共 {len(previews)} 个起点'
                    if not ensure_result.get('built')
                    else f'调用图已构建并生成 {len(previews)} 个预览'
                ),
                'center_block': selection_block,
                'center_function': None,
                'selection_block': selection_block,
                'previews': previews,
                'mermaid_code': previews[0].get('mermaid_code') if previews else '',
                'code_ranges': [selection_code_range] + (preview_bundle.get('code_ranges') or []),
                'called_functions': [
                    {
                        'function_id': item.get('function_id'),
                        'call_site_line': item.get('call_site_line'),
                        'call_site_name': item.get('raw_callee'),
                        'function': {
                            'id': (item.get('function') or {}).get('id'),
                            'name': (item.get('function') or {}).get('name'),
                            'qualified_name': (item.get('function') or {}).get('qualified_name'),
                            'file': (item.get('function') or {}).get('file'),
                            'line': (item.get('function') or {}).get('line'),
                            'end_line': (item.get('function') or {}).get('end_line'),
                        }
                    }
                    for item in called_functions
                ],
            })

        function_id = resolve_code_block_to_function(
            project_path,
            resolved_project_id,
            file,
            start_line,
            end_line,
        )
        if not function_id:
            return jsonify({
                'status': 'error',
                'call_graph_status': metadata.get('status'),
                'message': '当前选区未命中可查询调用图的函数块',
            })
 
        preview = query_function_graph(
            project_path,
            function_id,
            max_depth=max_depth,
            direction=graph_direction,
        )
        center_block = find_first_intersecting_code_block(
            project_path,
            resolved_project_id,
            file,
            start_line,
            end_line,
        )
        if not center_block:
            center_function = preview['center_function']
            center_block = {
                'id': None,
                'name': center_function.get('qualified_name') or center_function.get('name'),
                'type': 'function',
                'file': center_function.get('file'),
                'filename': center_function.get('file'),
                'range': [
                    int(center_function.get('line') or 0),
                    int(center_function.get('end_line') or center_function.get('line') or 0),
                ],
            }
        return jsonify({
            'status': 'success',
            'call_graph_status': metadata.get('status'),
            'query_mode': start_mode,
            'message': '调用图预览已生成' if not ensure_result.get('built') else '调用图已构建并生成预览',
            'center_block': {
                'id': center_block.get('id'),
                'name': center_block.get('name'),
                'type': center_block.get('type'),
                'file': center_block.get('file') or center_block.get('filename'),
                'range': center_block.get('range'),
                'codeRange': code_block_to_code_range(project_path, center_block),
            },
            'center_function': {
                'id': preview['center_function'].get('id'),
                'name': preview['center_function'].get('name'),
                'qualified_name': preview['center_function'].get('qualified_name'),
                'file': preview['center_function'].get('file'),
                'line': preview['center_function'].get('line'),
                'end_line': preview['center_function'].get('end_line'),
                'signature': preview['center_function'].get('signature'),
            },
            'function_nodes': [
                {
                    'id': item.get('id'),
                    'name': item.get('name'),
                    'qualified_name': item.get('qualified_name'),
                    'file': item.get('file'),
                    'line': item.get('line'),
                    'end_line': item.get('end_line'),
                }
                for item in preview.get('nodes', [])
            ],
            'mermaid_code': preview.get('mermaid_code') or '',
            'code_ranges': preview.get('code_ranges', []),
            'previews': [
                {
                    'root_function': {
                        'id': preview['center_function'].get('id'),
                        'name': preview['center_function'].get('name'),
                        'qualified_name': preview['center_function'].get('qualified_name'),
                        'file': preview['center_function'].get('file'),
                        'line': preview['center_function'].get('line'),
                        'end_line': preview['center_function'].get('end_line'),
                        'signature': preview['center_function'].get('signature'),
                    },
                    'title': preview['center_function'].get('qualified_name') or preview['center_function'].get('name'),
                    'mermaid_code': preview.get('mermaid_code') or '',
                    'code_ranges': preview.get('code_ranges', []),
                    'reachable_function_ids': preview.get('reachable_function_ids', []),
                }
            ],
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'call_graph_status': 'failed',
            'message': str(e),
        }), 500        
        
# 没看到有调用的地方，先注释了
@bp.route('/project/alignment-by-id', methods=['GET'])
def get_alignment_by_id():
    project_path = request.args.get('path')
    alignment_id = request.args.get('id')
    if not project_path or not alignment_id:
        return jsonify({'status': 'error', 'message': '缺少项目路径或对齐ID'}), 400
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute('SELECT id,name,isReviewed,reviewThoughts,docRanges,codeRanges,is_code_review,is_alignment,align_type FROM alignments WHERE id=%s', (alignment_id,))
        r = cur.fetchone()
        if not r:
            return jsonify({'status': 'error', 'message': '未找到对齐关系'}), 404
        alignment = _normalize_alignment_row(r)
        return jsonify({'status': 'success', 'data': alignment})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})



@bp.route('/api/get-doc-blocks', methods=['GET'])
def get_doc_blocks():
    """获取需求分块列表"""
    try:
        project_path = request.args.get('projectPath')
        filename = request.args.get('filename')
        project_id = request.args.get('project_id')
        page = request.args.get('page', type=int)
        page_size = request.args.get('page_size', type=int)
        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        resolved_project_id = resolve_project_id(project_path, project_id)
        if page or page_size:
            doc_blocks, pagination = get_paginated_doc_blocks_by_project(
                project_path,
                resolved_project_id,
                filename=filename,
                page=page or 1,
                page_size=page_size
            )
        else:
            doc_blocks = get_doc_blocks_by_project(project_path, resolved_project_id, filename)
            pagination = None

        if resolved_project_id:
            alignments = _load_project_alignments(resolved_project_id)
            for block in doc_blocks:
                block_file = block.get('filename')
                block_start = block.get('start')
                block_end = block.get('end')
                matched_alignment = next((
                    alignment for alignment in alignments
                    if any(
                        (doc_range.get('documentId') == block_file or doc_range.get('filename') == block_file) and
                        _ranges_overlap(doc_range.get('start'), doc_range.get('end'), block_start, block_end)
                        for doc_range in alignment.get('docRanges', [])
                    )
                ), None)
                block['matchedAlignmentId'] = matched_alignment.get('id') if matched_alignment else None
                block['matchedAlignmentReviewed'] = bool(matched_alignment.get('isReviewed')) if matched_alignment else False

        if pagination is not None:
            return jsonify({'status': 'success', 'data': doc_blocks,
                            'total': pagination['total'], 'page': pagination['page'],
                            'page_size': pagination['page_size']})

        return jsonify({'status': 'success', 'data': doc_blocks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})

@bp.route('/api/get-code-blocks', methods=['GET'])
def get_code_blocks():
    """获取代码分块列表 (原始格式)"""
    try:
        project_path = request.args.get('projectPath')
        filename = request.args.get('filename') # 可选：按文件名筛选
        project_id = request.args.get('project_id')
        page = request.args.get('page', type=int)
        page_size = request.args.get('page_size', type=int)

        if not project_path:
            return jsonify({'status': 'error', 'message': '缺少项目路径'})

        resolved_project_id = resolve_project_id(project_path, project_id)
        if page or page_size:
            code_blocks, pagination = get_paginated_code_blocks_by_project(
                project_path,
                resolved_project_id,
                filename=filename,
                page=page or 1,
                page_size=page_size
            )
        else:
            code_blocks = get_code_blocks_by_project(project_path, resolved_project_id, filename)
            pagination = None
            
        if not code_blocks:
            code_block_base_path = os.path.join(project_path, 'code_block_repo')
            code_block_file_path = os.path.join(code_block_base_path, 'code_blocks.jsonl')

            if not os.path.exists(code_block_file_path):
                 return jsonify({'status': 'success', 'data': []})

            code_blocks = []
            with open(code_block_file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        try:
                            block = json.loads(line.strip())
                            # 如果提供了文件名，则进行筛选
                            if filename:
                                if block.get('file') == filename:
                                    code_blocks.append(block)
                            else:
                                code_blocks.append(block)
                        except json.JSONDecodeError:
                            continue

        if resolved_project_id:
            alignments = _load_project_alignments(resolved_project_id)
            for block in code_blocks:
                block_file = block.get('file')
                line_range = block.get('range') if isinstance(block.get('range'), list) else []
                start_line = line_range[0] if len(line_range) == 2 else block.get('startLine')
                end_line = line_range[1] if len(line_range) == 2 else block.get('endLine')

                matched_alignment = None
                matched_code_review_alignment = None
                for alignment in alignments:
                    has_overlap = any(
                        (code_range.get('documentId') == block_file or code_range.get('filename') == block_file) and
                        _line_ranges_overlap(code_range.get('startLine'), code_range.get('endLine'), start_line, end_line)
                        for code_range in alignment.get('codeRanges', [])
                    )
                    if not has_overlap:
                        continue
                    if not alignment.get('isCodeReview') and matched_alignment is None:
                        matched_alignment = alignment
                    if alignment.get('isCodeReview') and matched_code_review_alignment is None:
                        matched_code_review_alignment = alignment
                    if matched_alignment and matched_code_review_alignment:
                        break

                block['matchedAlignmentId'] = matched_alignment.get('id') if matched_alignment else None
                block['matchedAlignmentReviewed'] = bool(matched_alignment.get('isReviewed')) if matched_alignment else False
                block['matchedCodeReviewAlignmentId'] = matched_code_review_alignment.get('id') if matched_code_review_alignment else None
                block['matchedCodeReviewReviewed'] = bool(matched_code_review_alignment.get('isReviewed')) if matched_code_review_alignment else False

        if pagination is not None:
            return jsonify({'status': 'success', 'data': code_blocks, 'total': pagination['total'], 'page': pagination['page'],
                            'page_size': pagination['page_size']})

        return jsonify({'status': 'success', 'data': code_blocks})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


def _find_first_file_by_basename(base_dir: str, basename: str):
    if not base_dir or not basename or not os.path.isdir(base_dir):
        return None
    direct_path = os.path.join(base_dir, basename)
    if os.path.exists(direct_path):
        return direct_path
    for root, _, files in os.walk(base_dir):
        if basename in files:
            return os.path.join(root, basename)
    return None


def _read_text_file(file_path: str):
    if not file_path or not os.path.exists(file_path):
        return None
    return read_source_file(file_path)


def _read_doc_raw_content(project_path: str, filename: str):
    doc_repo_path = os.path.join(project_path, 'doc_repo')
    metadata_file = os.path.join(project_path, 'metadata.json')
    if os.path.exists(metadata_file):
        try:
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = pyjson.load(f)
            if metadata.get('doc_repo'):
                doc_repo_path = metadata.get('doc_repo')
        except Exception:
            pass

    if filename and filename.lower().endswith('.md'):
        file_path = _find_first_file_by_basename(doc_repo_path, os.path.basename(filename))
        return _read_text_file(file_path)

    if filename:
        prefix = filename.split('.')[0]
        converted_md = os.path.join(project_path, 'doc_repo_converted', prefix, prefix + '.md')
        converted_content = _read_text_file(converted_md)
        if converted_content is not None:
            return converted_content

        file_path = _find_first_file_by_basename(doc_repo_path, os.path.basename(filename))
        return _read_text_file(file_path)

    return None


def _offset_to_line_numbers(text: str, start: int, end: int):
    if text is None:
        return 1, 1
    n = len(text)
    try:
        s = max(0, min(int(start), n))
    except Exception:
        s = 0
    try:
        e = max(0, min(int(end), n))
    except Exception:
        e = s
    start_line = text.count('\n', 0, s) + 1
    end_line = text.count('\n', 0, e) + 1
    return start_line, end_line


def _line_range_to_char_offsets(text: str, start_line: int, end_line: int):
    if text is None:
        return 0, 0
    lines = text.splitlines(keepends=True)
    if not lines:
        return 0, 0
    s = max(1, int(start_line or 1))
    e = max(s, int(end_line or s))
    s = min(s, len(lines))
    e = min(e, len(lines))
    char_start = sum(len(line) for line in lines[:s - 1])
    char_end = sum(len(line) for line in lines[:e])
    return char_start, char_end


def _compact_title_from_text(text: str, max_len: int = 24):
    if not text:
        return '未命名'
    t = re.sub(r'\s+', ' ', text).strip()
    t = re.sub(r'[`*_>#-]+', '', t).strip()
    if len(t) <= max_len:
        return t or '未命名'
    return t[:max_len].rstrip() + '…'


def gen_req_chunks(project_path):
    """
    从数据库中的需求块生成需求对齐候选数据
    """
    chunks = []
    for block in get_doc_blocks_by_project(project_path):
        filename = block.get('filename') or block.get('documentId') or ''
        content = block.get('content') or ''
        start = block.get('start') if block.get('start') is not None else 0
        end = block.get('end') if block.get('end') is not None else 0

        raw_doc = _read_doc_raw_content(project_path, filename)
        start_line, end_line = _offset_to_line_numbers(raw_doc, start, end)

        doc_range = {
            'documentId': filename,
            'filename': filename,
            'start': start,
            'end': end,
            'content': content,
            'startLine': start_line,
            'endLine': end_line
        }

        chunk_id = f"auto_req_{uuid.uuid4().hex}"
        chunk_name = block.get('name')
        if not chunk_name:
            if '#' in content:
                first_line = content.split('\n')[0]
                chunk_name = first_line.lstrip('#')
            else:
                chunk_name = _compact_title_from_text(content, 24)

        chunks.append({
            'id': chunk_id,
            'name': chunk_name,
            'isReviewed': False,
            'reviewThoughts': '',
            'docRanges': [doc_range],
            'codeRanges': []
        })

    chunks.sort(key=lambda x: (
    (x.get('docRanges') or [{}])[0].get('filename') or '', (x.get('docRanges') or [{}])[0].get('start') or 0))

    return chunks


def query_by_project_id(project_id, align_type):
    db = get_db()
    c = db.cursor()

    c.execute(f"select codeRanges,docRanges,id,isReviewed,name,reviewThoughts,is_alignment from alignments "
              f"where project_id={project_id} and align_type='{align_type}'")
    rows = c.fetchall()
    return rows


def save_align_to_db(project_id, data, align_type):
    db = get_db()
    c = db.cursor()

    try:
        sql = """
                INSERT INTO alignments(id, user_id, project_id, name, isReviewed, reviewThoughts, docRanges, codeRanges, 
                                       createdAt, updatedAt, align_type) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, %s) 
              """
        values = []
        for item in data:
            values.append((
                item.get('id'),
                current_user.user_id,
                project_id,
                item.get('name'),
                0,
                item.get('reviewThoughts') or '',
                json.dumps(item.get('docRanges') or []),
                json.dumps(item.get('codeRanges') or []),
                align_type
            ))

        if values:
            c.executemany(sql, values)

    except Exception as e:
        db.rollback()
        raise e


@bp.route('/api/get-requirement-chunks', methods=['GET'])
def get_requirement_chunks():
    project_path = request.args.get('projectPath')
    project_id = request.args.get('projectId')
    need_save = False

    if not project_path:
        return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400

    if project_id:
        all_align = query_by_project_id(project_id, 'req2code')
        if len(all_align) != 0:
            no_align = []
            y_align_count = 0
            for item in all_align:
                item['docRanges'] = json.loads(item['docRanges'])
                item['codeRanges'] = json.loads(item['codeRanges'])
                if item['is_alignment'] == 0:
                    no_align.append(item)
                else:
                    y_align_count += 1

            return jsonify({'status': 'success', 'data': no_align,
                            'y_align_count': y_align_count,
                            'all_align_count': len(all_align),
                            'n_align_count': len(no_align)})

        need_save = True

    try:
        chunks = gen_req_chunks(project_path)
    except Exception as e:
        return jsonify({'status': 'error', 'message': f"文件处理失败:{str(e)}"})

    if need_save:
        try:
            save_align_to_db(project_id, chunks, 'req2code')
        except Exception as e:
            logger.info(f'对齐数据(req2code)入库失败:{e}')

    return jsonify({'status': 'success', 'data': chunks})



@bp.route('/api/add-block', methods=['POST'])
def add_block():
    """添加新的需求块或代码块"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        block_type = data.get('blockType')  # 'doc' or 'code'
        block_data = data.get('blockData')
        project_id = resolve_project_id(project_path, data.get('projectId') or data.get('project_id'))

        if not project_path or not block_type or not block_data:
            return jsonify({'status': 'error', 'message': '缺少必要参数'})
        if not project_id:
            return jsonify({'status': 'error', 'message': '未找到项目ID'}), 400

        if block_type == 'doc':
            filename = block_data.get('filename') or block_data.get('documentId') or ''
            start = _safe_int(block_data.get('start'))
            end = _safe_int(block_data.get('end'))
            content = block_data.get('content') or ''
            block_name = _resolve_doc_block_name(block_data)

            db = get_db()
            cur = db.cursor()
            cur.execute(
                'SELECT id FROM doc_blocks WHERE project_id=%s AND filename=%s AND start=%s AND end=%s LIMIT 1',
                (project_id, filename, start, end)
            )
            existing = cur.fetchone()
            if existing:
                return jsonify({'status': 'warning', 'message': '该需求块已存在'})

            block_id = block_data.get('id')
            if block_id in (None, ''):
                block_id = _get_next_block_id('doc_blocks', project_id)
                block_data['id'] = block_id

            block_data['filename'] = filename
            block_data['documentId'] = filename
            block_data['start'] = start
            block_data['end'] = end
            block_data['content'] = content
            block_data['name'] = block_name
            block_data['type'] = block_data.get('type') or block_data.get('name') or ''

            cur.execute(
                'INSERT INTO doc_blocks(project_id, id, name, filename, type, content, start, end, createdAt, updatedAt) '
                'VALUES (%s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)',
                (project_id, _safe_int(block_id), block_name, filename, block_data.get('type') or '', content, start, end)
            )
            return jsonify({'status': 'success', 'message': '需求块添加成功'})

        elif block_type == 'code':
            target_range = block_data.get('range', [])
            if not (isinstance(target_range, list) and len(target_range) == 2):
                return jsonify({'status': 'error', 'message': '代码块缺少有效范围'}), 400

            file_name = block_data.get('file') or block_data.get('filename') or block_data.get('documentId') or ''
            start_line = _safe_int(target_range[0])
            end_line = _safe_int(target_range[1])
            code_content = block_data.get('code') or block_data.get('content') or ''
            block_name = _resolve_code_block_name(block_data)

            db = get_db()
            cur = db.cursor()
            cur.execute(
                'SELECT id FROM code_blocks WHERE project_id=%s AND file=%s AND start_line=%s AND end_line=%s LIMIT 1',
                (project_id, file_name, start_line, end_line)
            )
            existing = cur.fetchone()
            if existing:
                return jsonify({'status': 'warning', 'message': '该代码块已存在'})

            block_id = block_data.get('id')
            if block_id in (None, ''):
                block_id = _get_next_block_id('code_blocks', project_id)
                block_data['id'] = block_id

            block_data['file'] = file_name
            block_data['filename'] = file_name
            block_data['documentId'] = file_name
            block_data['range'] = [start_line, end_line]
            block_data['startLine'] = start_line
            block_data['endLine'] = end_line
            block_data['code'] = code_content
            block_data['content'] = code_content
            block_data['name'] = block_name
            block_data['related_id'] = block_data.get('related_id') or []
            block_data['related_range'] = block_data.get('related_range') or {}

            cur.execute(
                'INSERT INTO code_blocks(project_id, id, name, file, start_line, end_line, type, code, related_id, related_range, createdAt, updatedAt) '
                'VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)',
                (
                    project_id,
                    _safe_int(block_id),
                    block_name,
                    file_name,
                    start_line,
                    end_line,
                    block_data.get('type') or '',
                    code_content,
                    pyjson.dumps(block_data.get('related_id') or [], ensure_ascii=False),
                    pyjson.dumps(block_data.get('related_range') or {}, ensure_ascii=False)
                )
            )
            return jsonify({'status': 'success', 'message': '代码块添加成功'})

        return jsonify({'status': 'error', 'message': '无效的块类型'})

    except Exception as e:
        print(f"Error adding block: {e}")
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/api/rename-block', methods=['POST'])
def rename_block():
    """重命名需求块或代码块"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        block_type = data.get('blockType')
        block_data = data.get('blockData') or {}
        new_name = (data.get('name') or '').strip()
        project_id = resolve_project_id(project_path, data.get('project_id') or data.get('projectId'))

        if not project_path or not block_type or not block_data:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400
        if not project_id:
            return jsonify({'status': 'error', 'message': '未找到项目ID'}), 400
        if not new_name:
            return jsonify({'status': 'error', 'message': '名称不能为空'}), 400

        db = get_db()
        cur = db.cursor()
        updated = 0

        if block_type == 'doc':
            filename = block_data.get('filename') or block_data.get('documentId') or ''
            start = _safe_int(block_data.get('start'))
            end = _safe_int(block_data.get('end'))
            if block_data.get('id') not in (None, ''):
                cur.execute(
                    'UPDATE doc_blocks SET name=%s, updatedAt=CURRENT_TIMESTAMP WHERE project_id=%s AND id=%s',
                    (new_name, project_id, _safe_int(block_data.get('id')))
                )
            else:
                cur.execute(
                    'UPDATE doc_blocks SET name=%s, updatedAt=CURRENT_TIMESTAMP WHERE project_id=%s AND filename=%s AND start=%s AND end=%s',
                    (new_name, project_id, filename, start, end)
                )
            updated = cur.rowcount
        elif block_type == 'code':
            line_range = block_data.get('range') if isinstance(block_data.get('range'), list) else []
            filename = block_data.get('file') or block_data.get('filename') or block_data.get('documentId') or ''
            start_line = _safe_int(line_range[0] if len(line_range) == 2 else block_data.get('startLine'))
            end_line = _safe_int(line_range[1] if len(line_range) == 2 else block_data.get('endLine'))
            if block_data.get('id') not in (None, ''):
                cur.execute(
                    'UPDATE code_blocks SET name=%s, updatedAt=CURRENT_TIMESTAMP WHERE project_id=%s AND id=%s',
                    (new_name, project_id, _safe_int(block_data.get('id')))
                )
            else:
                cur.execute(
                    'UPDATE code_blocks SET name=%s, updatedAt=CURRENT_TIMESTAMP WHERE project_id=%s AND file=%s AND start_line=%s AND end_line=%s',
                    (new_name, project_id, filename, start_line, end_line)
                )
            updated = cur.rowcount
        else:
            return jsonify({'status': 'error', 'message': '无效的块类型'}), 400

        if updated <= 0:
            return jsonify({'status': 'warning', 'message': '未找到要重命名的块'})
        return jsonify({'status': 'success', 'message': '重命名成功'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


# def save_code_block(project_path, project_id):
#     """
#     在add-block后面调用，取最后面新添加的block，存到alignments表，用来纯代码审查
#     """
#
#     try:
#         chunk = gen_code_chunks(project_path, True)
#     except Exception as e:
#         logger.info(f'文件处理失败:{e}')
#         return
#     db = get_db()
#     c = db.cursor()
#     try:
#         c.execute("INSERT INTO alignments(id,user_id,project_id,name,reviewThoughts,docRanges,codeRanges,createdAt,"
#                   "updatedAt,is_code_review) "
#                   "VALUES(%s,%s,%s,%s,%s,%s,%s,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP,1)",
#                   (chunk.get('id'),
#                    current_user.user_id,
#                    project_id,
#                    chunk.get('name'),
#                    chunk.get('reviewThoughts'),
#                    chunk.get('docRanges')),
#                   chunk.get('codeRanges'))
#     except Exception as e:
#         db.rollback()
#         logger.info(f'纯代码审查数据入库失败:{e}')


@bp.route('/api/delete-block', methods=['POST'])
def delete_block():
    """删除指定的需求块或代码块，并清理相关对齐关系"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        block_type = data.get('blockType')  # 'doc' or 'code'
        block_data = data.get('blockData')
        project_id = resolve_project_id(project_path, data.get('project_id') or data.get('projectId'))

        if not project_path or not block_type or not block_data:
            return jsonify({'status': 'error', 'message': '缺少必要参数'})

        # 1. 删除块数据
        deleted = False
        conn = get_db()
        cur = conn.cursor()
        if block_type == 'doc':
            target_filename = block_data.get('filename') or block_data.get('documentId')
            target_start = _safe_int(block_data.get('start'))
            target_end = _safe_int(block_data.get('end'))
            cur.execute(
                'DELETE FROM doc_blocks WHERE project_id=%s AND filename=%s AND start=%s AND end=%s',
                (project_id, target_filename, target_start, target_end)
            )
            deleted = cur.rowcount > 0

        elif block_type == 'code':
            target_range = block_data.get('range', [])
            target_file = block_data.get('file') or block_data.get('filename') or block_data.get('documentId')
            target_start = _safe_int(target_range[0] if isinstance(target_range, list) and len(target_range) == 2 else block_data.get('startLine'))
            target_end = _safe_int(target_range[1] if isinstance(target_range, list) and len(target_range) == 2 else block_data.get('endLine'))
            if block_data.get('id') not in (None, ''):
                cur.execute(
                    'DELETE FROM code_blocks WHERE project_id=%s AND id=%s',
                    (project_id, _safe_int(block_data.get('id')))
                )
            else:
                cur.execute(
                    'DELETE FROM code_blocks WHERE project_id=%s AND file=%s AND start_line=%s AND end_line=%s',
                    (project_id, target_file, target_start, target_end)
                )
            deleted = cur.rowcount > 0

        if not deleted:
            return jsonify({'status': 'warning', 'message': '未找到要删除的块'})

        # 2. 清理对齐关系
        cur.execute(f"SELECT * FROM alignments where project_id={project_id}")
        alignments = cur.fetchall()

        updates = []
        deletions = []

        for row in alignments:
            align_id = row['id']
            doc_ranges = json.loads(row['docRanges']) if row['docRanges'] else []
            code_ranges = json.loads(row['codeRanges']) if row['codeRanges'] else []

            modified = False

            if block_type == 'doc':
                # 过滤掉匹配的 docRange
                original_len = len(doc_ranges)
                doc_ranges = [
                    dr for dr in doc_ranges
                    if not (dr.get('filename') == block_data.get('filename') and
                            dr.get('start') == block_data.get('start') and
                            dr.get('end') == block_data.get('end'))
                ]
                if len(doc_ranges) < original_len:
                    modified = True

            elif block_type == 'code':
                original_len = len(code_ranges)
                target_range = block_data.get('range', [])
                target_file = block_data.get('file') or block_data.get('filename') or block_data.get('documentId')
                target_start = _safe_int(target_range[0] if isinstance(target_range, list) and len(target_range) == 2 else block_data.get('startLine'))
                target_end = _safe_int(target_range[1] if isinstance(target_range, list) and len(target_range) == 2 else block_data.get('endLine'))

                new_code_ranges = []
                for cr in code_ranges:
                    cr_file = cr.get('filename') or cr.get('documentId')
                    cr_start = _safe_int(cr.get('startLine'))
                    cr_end = _safe_int(cr.get('endLine'))

                    if (cr_file == target_file and
                        cr_start == target_start and
                        cr_end == target_end):
                        # Match found, remove it
                        pass
                    else:
                        new_code_ranges.append(cr)

                if len(new_code_ranges) < original_len:
                    code_ranges = new_code_ranges
                    modified = True

            if modified:
                # 检查是否为空
                if len(doc_ranges) == 0 or len(code_ranges) == 0:
                    deletions.append(align_id)
                else:
                    updates.append((json.dumps(doc_ranges, ensure_ascii=False),
                                    json.dumps(code_ranges, ensure_ascii=False),
                                    align_id))

        # 执行数据库更新
        for align_id in deletions:
            cur.execute("DELETE FROM alignments WHERE id = %s and project_id=%s", (align_id, project_id))

        for doc_r, code_r, align_id in updates:
            cur.execute("UPDATE alignments SET docRanges = %s, codeRanges = %s, updatedAt = CURRENT_TIMESTAMP "
                        "WHERE id = %s and project_id=%s",
                        (doc_r, code_r, align_id, project_id))

        # conn.commit()
        # conn.close()

        return jsonify({
            'status': 'success',
            'message': f'已删除块并更新了 {len(updates)} 个对齐关系，删除了 {len(deletions)} 个空对齐关系'
        })

    except Exception as e:
        print(f"Error deleting block: {e}")
        return jsonify({'status': 'error', 'message': str(e)})


@bp.route('/api/clear-alignment-target', methods=['POST'])
def clear_alignment_target():
    """清空单条对齐关系的目标侧块，用于删除对齐结果或重新对齐前准备"""
    try:
        data = request.get_json()
        project_path = data.get('projectPath')
        project_id = data.get('project_id')
        alignment_id = data.get('alignmentId')

        if not project_path or not alignment_id:
            return jsonify({'status': 'error', 'message': '缺少必要参数'}), 400

        resolved_project_id = resolve_project_id(project_path, project_id)
        if not resolved_project_id:
            return jsonify({'status': 'error', 'message': '未找到项目ID'}), 400

        conn = get_db()
        cur = conn.cursor()
        cur.execute(
            'SELECT id, align_type, docRanges, codeRanges FROM alignments WHERE id=%s AND project_id=%s',
            (alignment_id, resolved_project_id)
        )
        row = cur.fetchone()
        if not row:
            return jsonify({'status': 'error', 'message': '未找到对齐关系'}), 404

        align_type = row.get('align_type') or 'req2code'
        doc_ranges = pyjson.loads(row.get('docRanges') or '[]')
        code_ranges = pyjson.loads(row.get('codeRanges') or '[]')
        if align_type == 'code2req':
            doc_ranges = []
        else:
            code_ranges = []

        cur.execute(
            'UPDATE alignments SET docRanges=%s, codeRanges=%s, isReviewed=0, reviewThoughts="", '
            'GenReq="", GenMermaid="", is_alignment=0, updatedAt=CURRENT_TIMESTAMP '
            'WHERE id=%s AND project_id=%s',
            (
                pyjson.dumps(doc_ranges, ensure_ascii=False),
                pyjson.dumps(code_ranges, ensure_ascii=False),
                alignment_id,
                resolved_project_id
            )
        )
        cur.execute('DELETE FROM issues WHERE alignmentId=%s AND project_id=%s', (alignment_id, resolved_project_id))

        return jsonify({
            'status': 'success',
            'message': '已清空目标侧对齐结果',
            'data': {
                'id': alignment_id,
                'align_type': align_type,
                'docRanges': doc_ranges,
                'codeRanges': code_ranges,
                'isReviewed': False,
                'reviewThoughts': '',
                'is_alignment': 0
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


def gen_code_chunks(project_path, end=False):
    code_repo_path = os.path.join(project_path, 'code_repo')
    chunks = []
    for block in get_code_blocks_by_project(project_path):
        file_rel = block.get('file') or ''
        rng = block.get('range') or []
        if not (isinstance(rng, list) and len(rng) == 2):
            continue
        start_line = int(rng[0])
        end_line = int(rng[1])
        content = block.get('code') or block.get('content') or ''

        abs_code_path = os.path.join(code_repo_path, file_rel) if file_rel else None
        raw_code = _read_text_file(abs_code_path)
        char_start, char_end = _line_range_to_char_offsets(raw_code, start_line, end_line)

        chunk_name = f"{file_rel}:{start_line}-{end_line}" if file_rel else f"代码块:{start_line}-{end_line}"
        code_lines = [line.strip() for line in content.splitlines()]
        for code_line in code_lines:
            if not code_line:
                continue
            if code_line.startswith('//') or code_line.startswith('/*') or code_line.startswith('*'):
                continue
            chunk_name = code_line
            break

        code_range = {
            'name': chunk_name,
            'documentId': file_rel,
            'filename': file_rel,
            'start': char_start,
            'end': char_end,
            'content': content,
            'startLine': start_line,
            'endLine': end_line
        }

        chunk_id = f"auto_code_{uuid.uuid4().hex}"
        chunks.append({
            'id': chunk_id,
            'name': chunk_name,
            'isReviewed': False,
            'reviewThoughts': '',
            'docRanges': [],
            'codeRanges': [code_range]
        })

    if end:
        return chunks[-1] if chunks else None

    chunks.sort(key=lambda x: (
    (x.get('codeRanges') or [{}])[0].get('filename') or '', (x.get('codeRanges') or [{}])[0].get('startLine') or 0))

    return chunks


@bp.route('/api/get-code-chunks', methods=['GET'])
def get_code_chunks():
    project_path = request.args.get('projectPath')
    project_id = request.args.get('projectId')
    need_save = False

    if not project_path:
        return jsonify({'status': 'error', 'message': '缺少项目路径'}), 400
    if project_id:
        all_align = query_by_project_id(project_id, 'code2req')
        if len(all_align) != 0:
            no_align = []
            y_align_count = 0
            for item in all_align:
                item['docRanges'] = json.loads(item['docRanges'])
                item['codeRanges'] = json.loads(item['codeRanges'])
                if item['is_alignment'] == 0:
                    no_align.append(item)
                else:
                    y_align_count += 1

            return jsonify({'status': 'success', 'data': no_align,
                            'y_align_count': y_align_count,
                            'all_align_count': len(all_align)})

        need_save = True

    try:
        chunks = gen_code_chunks(project_path)
    except Exception as e:
        return jsonify({'status': 'error', 'message': f"文件处理失败:{str(e)}"})

    if need_save:
        try:
            save_align_to_db(project_id, chunks, 'code2req')
        except Exception as e:
            logger.info(f'对齐数据(code2req)入库失败:{e}')

    return jsonify({'status': 'success', 'data': chunks})


@bp.route('/api/align-code-to-requirement-for', methods=['POST'])
def align_code_to_requirements():
    """循环处理多个代码=>需求"""
    data = request.json
    code_blocks = data.get('codeBlocks', [])
    project_path = data.get('projectPath', '')
    project_id = data.get('project_id', '')
    y_align = data.get('y_align', '')
    all_align = data.get('all_align', '')
    user_id = current_user.user_id

    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径参数"}), 400

    from tasks import align_code_to_requirements_task
    task = align_code_to_requirements_task.delay(project_path, code_blocks, project_id, current_user.user_id, y_align)
    # 新增: 写入用户任务快照
    db = get_db()
    cursor = db.cursor()

    # 先清理该项目旧的任务快照(防止残留)
    cursor.execute("DELETE FROM user_task_snapshot WHERE user_id = %s AND project_id = %s",
                   (user_id, project_id))

    # 插入新的
    cursor.execute("""INSERT INTO user_task_snapshot 
                       (user_id, project_id, task_id, task_type, task1_total, current_total,
                       state, title, is_running, task_category)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                   (
                       user_id, project_id, task.id, 'single', all_align, all_align,
                       'PENDING', '自动对齐 (代码 → 需求)', 1, 'align'
                   ))


    return jsonify({"status": "success", "task_id": task.id})


# 需要异步处理
@bp.route('/api/align-code-to-requirement', methods=['POST'])
def align_code_to_requirement():
    """为单个代码块在项目中查找相关需求"""
    try:
        data = request.json
        code_ranges = data.get('codeRanges', [])
        project_path = data.get('projectPath', '')

        if not code_ranges or not project_path:
             return jsonify({"status": "error", "message": "缺少代码内容或项目路径参数"}), 400

        # 获取选定的 align 类型知识库
        selected_align_kbs = []
        try:
            metadata_file = os.path.join(project_path, 'metadata.json')
            if os.path.exists(metadata_file):
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    selected_kbs = metadata.get('selected_kbs', [])
                    selected_align_kbs = [kb['name'] for kb in selected_kbs if kb.get('type') in ('align', 'history_align')]
        except Exception:
            pass

        all_doc_blocks, all_original_doc_blocks, blocks_by_file = _get_doc_blocks_for_matching(project_path)
        if not all_doc_blocks:
            return jsonify({"status": "success", "docRanges": []})

        # 如果没有选择任何 align 知识库，使用原来的 LLM 逻辑
        if not selected_align_kbs:

            code_content = '\n\n'.join([code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

            # 调用LLM
            related_reqs = query_related_requirement(
                code_content,
                all_doc_blocks,
                block_limit=50,
                project_path=project_path
            )

            doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

            return jsonify({
                "status": "success",
                "docRanges": doc_ranges
            })

        # 如果选择了 align 知识库，使用 RAG 进行检索
        try:
            rag_engine.initialize() # 确保初始化
        except Exception as e:
            print(f"[Align] RAG initialize failed: {e}")

        code_content = '\n\n'.join([code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

        all_retrieved_items = []
        for kb_name in selected_align_kbs:
            # 检索 'align' 类型的知识库
            # 注意：align 知识库里存储的是历史对齐数据
            # 这里的检索策略可以是：用代码去搜相关的历史对齐，然后把历史对齐中的文档部分作为推荐

            # 暂时假设 rag_chroma 提供了 query 接口，如果没有需要添加
            # 这里先模拟直接调用 collection.query
            collection = rag_engine.get_collection('align', kb_name)
            if collection:
                results = collection.query(
                    query_texts=[code_content],
                    n_results=5 # Top 5 per KB
                )

                if results and results['documents']:
                    for i, doc in enumerate(results['documents'][0]):
                        meta = results['metadatas'][0][i]
                        # 历史对齐数据通常包含 code_text 和 doc_text (query_text)
                        # 我们需要提取其中的 doc 部分
                        # 在 build_from_json 中，document 是 doc_text, meta 中有 code_text
                        # 但我们现在是用 code 去搜 doc，所以 doc 正好是 document
                        all_retrieved_items.append({
                            'content': doc,
                            'score': results['distances'][0][i] if results['distances'] else 0,
                            'meta': meta
                        })

        # 对所有结果排序
        all_retrieved_items.sort(key=lambda x: x['score']) # distance 越小越好
        top_items = all_retrieved_items[:5]

        # 构造返回结果
        # 注意：这里返回的是参考的历史对齐文档内容，而不是当前项目中的具体需求块
        # 前端可能需要展示这些参考内容供用户选择，或者作为提示
        # 但目前的接口契约是返回 docRanges (当前项目的需求块)
        # 这是一个逻辑断层：历史对齐是"参考"，而不是"直接结果"
        # 如果要用历史对齐来辅助定位当前项目的需求，需要两步：
        # 1. 检索历史对齐 -> 得到相关的历史需求描述
        # 2. 用历史需求描述去匹配当前项目的需求块 (类似 query_related_requirement)

        history_doc_contents = [item['content'] for item in top_items]
        combined_history_content = "\n".join(history_doc_contents)

        # 使用历史需求内容 + 代码内容 共同作为 Query 去查询当前项目需求
        enhanced_query = f"Code:\n{code_content}\n\nRelated History Requirements:\n{combined_history_content}"

        # 调用LLM (使用增强后的 Query)
        related_reqs = query_related_requirement(
            enhanced_query,
            all_doc_blocks,
            block_limit=50,
            icl_examples=top_items,
            project_path=project_path
        )

        doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

        return jsonify({
            "status": "success",
            "docRanges": doc_ranges,
            "debug_info": "Used history alignment for enhancement"
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": f"对齐过程中出错: {str(e)}"}), 500
        

@bp.route('/api/align-code-to-requirement-addprompt', methods=['POST'])
def align_code_to_requirement_addprompt():
    """为单个代码块在项目中查找相关需求"""
    try:
        data = request.json
        code_ranges = data.get('codeRanges', [])
        project_path = data.get('projectPath', '')
        userPrompt = data.get('userInputPrompt', [])
        project_id = data.get('project_id')
        reqRanges_aligned = data.get('docRanges', [])

        if not code_ranges or not project_path:
             return jsonify({"status": "error", "message": "缺少代码内容或项目路径参数"}), 400

        # 获取选定的 align 类型知识库
        selected_align_kbs = []
        try:
            metadata_file = os.path.join(project_path, 'metadata.json')
            if os.path.exists(metadata_file):
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                    selected_kbs = metadata.get('selected_kbs', [])
                    selected_align_kbs = [kb['name'] for kb in selected_kbs if kb.get('type') in ('align', 'history_align')]
        except Exception:
            pass

        all_doc_blocks, all_original_doc_blocks, blocks_by_file = _get_doc_blocks_for_matching(project_path, project_id)
        if not all_doc_blocks:
            return jsonify({"status": "success", "docRanges": []})

        # 如果没有选择任何 align 知识库，使用原来的 LLM 逻辑
        if not selected_align_kbs:

            code_content = '\n\n'.join([code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

            # 调用LLM
            related_reqs = query_related_requirement_by_feedback(
                code_content,
                reqRanges_aligned,
                all_doc_blocks,
                userPrompt,
                block_limit=50,
                project_path=project_path
            )

            doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

            return jsonify({
                "status": "success",
                "docRanges": doc_ranges
            })

        # 如果选择了 align 知识库，使用 RAG 进行检索
        try:
            rag_engine.initialize() # 确保初始化
        except Exception as e:
            print(f"[Align] RAG initialize failed: {e}")

        code_content = '\n\n'.join([code_range.get('content', '') for code_range in code_ranges if code_range.get('content')])

        all_retrieved_items = []
        for kb_name in selected_align_kbs:
            # 检索 'align' 类型的知识库
            # 注意：align 知识库里存储的是历史对齐数据
            # 这里的检索策略可以是：用代码去搜相关的历史对齐，然后把历史对齐中的文档部分作为推荐

            # 暂时假设 rag_chroma 提供了 query 接口，如果没有需要添加
            # 这里先模拟直接调用 collection.query
            collection = rag_engine.get_collection('align', kb_name)
            if collection:
                results = collection.query(
                    query_texts=[code_content],
                    n_results=5 # Top 5 per KB
                )

                if results and results['documents']:
                    for i, doc in enumerate(results['documents'][0]):
                        meta = results['metadatas'][0][i]
                        # 历史对齐数据通常包含 code_text 和 doc_text (query_text)
                        # 我们需要提取其中的 doc 部分
                        # 在 build_from_json 中，document 是 doc_text, meta 中有 code_text
                        # 但我们现在是用 code 去搜 doc，所以 doc 正好是 document
                        all_retrieved_items.append({
                            'content': doc,
                            'score': results['distances'][0][i] if results['distances'] else 0,
                            'meta': meta
                        })

        # 对所有结果排序
        all_retrieved_items.sort(key=lambda x: x['score']) # distance 越小越好
        top_items = all_retrieved_items[:5]

        # 构造返回结果
        # 注意：这里返回的是参考的历史对齐文档内容，而不是当前项目中的具体需求块
        # 前端可能需要展示这些参考内容供用户选择，或者作为提示
        # 但目前的接口契约是返回 docRanges (当前项目的需求块)
        # 这是一个逻辑断层：历史对齐是"参考"，而不是"直接结果"
        # 如果要用历史对齐来辅助定位当前项目的需求，需要两步：
        # 1. 检索历史对齐 -> 得到相关的历史需求描述
        # 2. 用历史需求描述去匹配当前项目的需求块 (类似 query_related_requirement)

        history_doc_contents = [item['content'] for item in top_items]
        combined_history_content = "\n".join(history_doc_contents)

        # 使用历史需求内容 + 代码内容 共同作为 Query 去查询当前项目需求
        enhanced_query = f"Code:\n{code_content}\n\nRelated History Requirements:\n{combined_history_content}"

        # 调用LLM (使用增强后的 Query)
        related_reqs = query_related_requirement(
            enhanced_query,
            all_doc_blocks,
            block_limit=50,
            icl_examples=top_items,
            project_path=project_path
        )

        doc_ranges = _match_doc_ranges_from_related_reqs(related_reqs, blocks_by_file)

        return jsonify({
            "status": "success",
            "docRanges": doc_ranges,
            "debug_info": "Used history alignment for enhancement"
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": f"对齐过程中出错: {str(e)}"}), 500
        

# 1. 获取标注文件列表
@bp.route('/api/files/list-annotations', methods=['GET'])
def list_annotation_files():
    project_path = request.args.get('projectPath')

    if not project_path:
        return jsonify({"status": "error", "message": "缺少项目路径"})

    # 拼接 annotations 目录路径
    ann_dir = os.path.join(project_path, 'annotations')

    files = []
    # 检查目录是否存在
    if os.path.exists(ann_dir) and os.path.isdir(ann_dir):
        # 遍历目录，只获取 .json 文件
        files = [f for f in os.listdir(ann_dir) if f.endswith('.json')]
        # 可选：按修改时间排序，让最新的排在前面
        # files.sort(key=lambda x: os.path.getmtime(os.path.join(ann_dir, x)), reverse=True)

    return jsonify({"status": "success", "files": files})


# 2. 构建知识库 (调用 rag_chroma)
@bp.route('/api/rag/build', methods=['POST'])
@kb_perm_required(need="edit")
def build_rag_db():

    if request.content_type and 'multipart/form-data' in request.content_type:
        data = request.form
        uploaded_files = request.files.getlist('annotationFiles')
        uploaded_files = preprocess_files(request.files.getlist('annotationFiles'))
    else:
        data = request.json or {}
        uploaded_files = []

    project_path = data.get('projectPath') or PROJECT_ROOT # Default to system root if not provided
    kb_type = data.get('kbType', 'other')
    source_file_name = data.get('sourceFileName')

    # User provided KB Name, default to filename without extension
    kb_name = data.get('kbName')
    append_mode = data.get('append', False)
    if not kb_name or ' ' in kb_name:
        return jsonify({"status": "error", "message": "知识库名称不能为空，且不能包含空格"})


    try:
        rag_engine.initialize() # Ensure root exists
    except Exception as e:
        return jsonify({"status": "error", "message": f"初始化失败: {str(e)}"})

    # Resolve source file path robustly:
    # 1) try raw filename/path first (preserve spaces/Chinese),
    # 2) then try secure_filename fallback for local upload temp files.
    annotation_files = []
    temp_dir = os.path.join(PROJECT_ROOT, 'temp_uploads')

    if uploaded_files:
        for f in uploaded_files:
            if f and f.filename:
                safe_name = secure_filename(f.filename)
                target = os.path.join(temp_dir, safe_name)
                if os.path.exists(target):
                    name, ext = os.path.splitext(safe_name)
                    target = os.path.join(temp_dir, f"{name}_{int(time.time()*1000)}{ext}")
                f.save(target)
                annotation_files.append(target)
    else:
        # 服务器文件模式
        raw = data.get('annotationFiles')
        if isinstance(raw, str):
            annotation_files = [raw]
        elif isinstance(raw, list):
            annotation_files = raw

    if not annotation_files:
        return jsonify({'status': 'error', 'message': '参数缺失: annotationFile'})

    def normalize_kb_type(raw_type):
        mapping = {
            "rule": "coding_rule",
            "issue": "history_issue",
            "history_align": "align",
            "align": "align",
            "case": "typical_case",
            "coding_rule": "coding_rule",
            "history_issue": "history_issue",
            "typical_case": "typical_case",
            "checklist": "checklist",
            "other": "other"
        }
        return mapping.get((raw_type or "other").strip(), "other")

    def resolve_processing_type(raw_type):
        normalized = normalize_kb_type(raw_type)
        if normalized in ["coding_rule", "checklist"]:
            return "rule"
        if normalized == "history_issue":
            return "issue"
        if normalized == "align":
            return "align"
        return "other"

    def save_kb_metadata(k_type, k_name, total_count=0):
        try:
            target_type = normalize_kb_type(k_type)
            # 展平结构，直接存在根目录下
            kb_root = os.path.join(PROJECT_ROOT, "../rag_database", k_name)
            os.makedirs(kb_root, exist_ok=True)

            meta_file = os.path.join(kb_root, "metadata.json")
            if append_mode and os.path.exists(meta_file):
                # 追加模式：保留原有 create_time
                with open(meta_file, 'r', encoding='utf-8') as f:
                    meta = json.load(f)
                meta['doc_count'] = total_count
                meta['update_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            else:
                # 新建模式
                meta = {
                    "name": k_name,
                    "type": target_type,
                    "create_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "update_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "doc_count": total_count,
                    "source_file": normalized_source_name
                }
            with open(meta_file, 'w', encoding='utf-8') as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[KB] 保存元数据失败: {e}")

    # 循环处理每个文件 累加count
    total_count = 0
    results = []

    for idx, raw_annotation_file in enumerate(annotation_files):
        # === 找文件路径 (原逻辑) ===

        candidates = []
        if raw_annotation_file:
            if os.path.isabs(raw_annotation_file):
                candidates.append(raw_annotation_file)
            else:
                candidates.extend([
                    os.path.join(TESTDATA_DIR, raw_annotation_file),
                    os.path.join(temp_dir, raw_annotation_file),
                    os.path.join(PROJECT_ROOT, raw_annotation_file),
                    raw_annotation_file
                ])

        full_path = ""
        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                full_path = candidate
                break

        if not full_path:
            results.append({
                "file": raw_annotation_file,
                "status": "error",
                "message": "找不到文件",
            })

        try:
            json_data = None

            # Ensure kb_type is standardized for processing
            processing_type = resolve_processing_type(kb_type)

            # === Word 文档处理逻辑 ===
            if full_path.lower().endswith('.docx'):
                print(f"[RAG] 解析文档: {full_path}, 类型: {processing_type}")

                # A. 编程规则
                if processing_type == 'rule':
                    raw_rules = parse_programming_rules(full_path)
                    if not raw_rules:
                        doc_text = read_docx_text(full_path)
                        raw_rules = smart_parse_doc(doc_text, type='rule')
                    if raw_rules:
                        json_data = format_rules_for_rag(raw_rules)

                # B. 问题单
                elif processing_type == 'issue':
                    raw_issues = parse_issue_reports(full_path)
                    if not raw_issues:
                        doc_text = read_docx_text(full_path)
                        raw_issues = smart_parse_doc(doc_text, type='issue')
                    if raw_issues:
                        json_data = format_issues_for_rag(raw_issues)

                # C. 其他
                elif processing_type in ['case', 'other', 'align']:
                    doc_text = read_docx_text(full_path)
                    raw_data = smart_parse_doc(doc_text, type='rule')
                    if raw_data:
                        json_data = format_rules_for_rag(raw_data)

            # === JSON 逻辑 ===
            elif full_path.lower().endswith('.json'):
                with open(full_path, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)


            else:
                results.append({"file": raw_annotation_file, "status": "error", "message": "不支持的文件格式"})
                continue

            if not json_data:
                results.append({"file": raw_annotation_file, "status": "error", "message": "文档解析失败"})
                continue

            temp_json = full_path + '.parsed.json'
            with open(temp_json, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)

            is_append = append_mode or (idx > 0)
            normalized_source_name = re.sub(r'\s+', '_', source_file_name or os.path.basename(full_path))

            result = rag_engine.build_from_json(
                temp_json,
                kb_type=processing_type,
                kb_name=kb_name,
                append=is_append,
                source_file=normalized_source_name
            )
            if result.get("status") == "success":
                match = re.search(r'(\d+)', result.get("message", ""))
                count = int(match.group(1)) if match else 0
                total_count += count
                results.append({'file': raw_annotation_file, 'status': 'success', 'count': count})
                try:
                    os.remove(temp_json)
                except:
                    pass

            else:
                results.append({'file': raw_annotation_file, 'status': 'error', 'count': result.get('message')})

        except Exception as e:
            results.append({'file': raw_annotation_file, 'status': 'error', 'count': str(e)})

    if total_count > 0:
        save_kb_metadata(kb_type, kb_name, total_count)

    return jsonify({
        'status': 'success',
        'total_count': total_count,
        'results': results
    })


@bp.route('/api/rag/add_items', methods=['POST'])
@kb_perm_required(need="edit")
def add_items_to_kb():
    """专用于接收前端内存中的 items 数组，直接新建或追加到知识库"""
    data = request.json
    kb_name = data.get('kbName')
    kb_type = data.get('kbType', 'other')
    append_mode = data.get('append', False)
    items = data.get('items', [])

    if not kb_name or ' ' in kb_name:
        return jsonify({"status": "error", "message": "知识库名称不能为空，且不能包含空格"})
    if not items:
        return jsonify({"status": "warning", "message": "没有要入库的数据"})

    try:
        rag_engine.initialize()
        col_info = rag_engine._get_or_create_collection(kb_type, kb_name)
        if not col_info:
            return jsonify({"status": "error", "message": f"知识库 {kb_name} 初始化失败"})

        client = col_info['client']
        collection = col_info['collection']

        # 1. 新建模式下，清空原有集合
        if not append_mode:
            try:
                client.delete_collection(collection.name)
            except:
                pass
            try:
                col_info['collection'] = client.create_collection(
                    name=collection.name,
                    embedding_function=rag_engine.emb_fn,
                    metadata={"hnsw:space": "cosine"}
                )
            except Exception as e:
                print(f"[KB] 重置集合时发生警告(可忽略): {e}")

        collection = col_info['collection']

        # 2. 组装数据
        ids = []
        documents = []
        metadatas = []

        import uuid
        run_id = uuid.uuid4().hex[:6]

        for idx, item in enumerate(items):
            # 获取前端传来的ID，若无则自动生成
            doc_id = item.get('id') or f"direct_{run_id}_{idx}"
            content = item.get('content', '')
            if not content:
                continue

            # 元数据处理，确保全为基础类型以适应 ChromaDB 限制
            meta = {"source_type": "direct_import", 'source_file': '手动录入'}
            if 'full_data' in item and isinstance(item['full_data'], dict):
                for k, v in item['full_data'].items():
                    if isinstance(v, (str, int, float, bool)):
                        meta[k] = v
                    else:
                        # 复杂对象转为 JSON 字符串
                        meta[k] = pyjson.dumps(v, ensure_ascii=False)

            ids.append(doc_id)
            documents.append(content)
            metadatas.append(meta)

        # 3. 写入 ChromaDB
        if ids:
            collection.upsert(ids=ids, documents=documents, metadatas=metadatas)

        total_count = collection.count()
        
        # 4. 更新元数据 metadata.json (使用展平结构)
        target_type_map = {
            "rule": "coding_rule",
            "issue": "history_issue",
            "history_align": "align",
            "align": "align",
            "case": "typical_case",
            "coding_rule": "coding_rule",
            "history_issue": "history_issue",
            "typical_case": "typical_case",
            "checklist": "checklist",
            "other": "other"
        }
        target_type = target_type_map.get((kb_type or "other").strip(), "other")
        kb_root = os.path.join(PROJECT_ROOT, "../rag_database", kb_name)
        os.makedirs(kb_root, exist_ok=True)
        meta_file = os.path.join(kb_root, "metadata.json")

        if append_mode and os.path.exists(meta_file):
            with open(meta_file, 'r', encoding='utf-8') as f:
                meta = pyjson.load(f)
            meta['doc_count'] = total_count
            meta['update_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        else:
            meta = {
                "name": kb_name,
                "type": target_type,
                "create_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "update_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "doc_count": total_count,
                "source_file": "direct_import"
            }

        with open(meta_file, 'w', encoding='utf-8') as f:
            pyjson.dump(meta, f, ensure_ascii=False, indent=2)

        return jsonify({
            "status": "success",
            "message": f"入库完成！本次新增 {len(ids)} 条数据，当前库内总数: {total_count}",
            "total_count": total_count
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": f"入库失败: {str(e)}"}), 500

# ========================================================
# Knowledge Base API
# ========================================================

KB_TYPE_NORMALIZE_MAP = {
    "rule": "coding_rule",
    "issue": "history_issue",
    "history_align": "align",
    "align": "align",
    "case": "typical_case",
    "coding_rule": "coding_rule",
    "history_issue": "history_issue",
    "typical_case": "typical_case",
    "checklist": "checklist",
    "other": "other"
}

def normalize_kb_type(raw_type):
    return KB_TYPE_NORMALIZE_MAP.get((raw_type or "other").strip(), "other")

@bp.route('/api/list-testdata', methods=['GET'])
def list_testdata():
    """列出 testdata 目录下的所有文件"""
    if not os.path.exists(TESTDATA_DIR):
        return jsonify({"status": "success", "files": []})

    files = [f for f in os.listdir(TESTDATA_DIR) if os.path.isfile(os.path.join(TESTDATA_DIR, f)) and not f.startswith('.')]
    return jsonify({"status": "success", "files": files})


@bp.route('/api/kb/create', methods=['POST'])
def create_kb():
    """创建空知识库（先建库，再上传文件）"""
    data = request.json or {}
    kb_name = (data.get('name') or '').strip()
    kb_type = normalize_kb_type(data.get('type'))
    description = (data.get('description') or '').strip()
    security_level = (data.get('security_level') or '内部').strip()
    language = (data.get('language') or '中文').strip()
    parse_method = (data.get('parse_method') or '通用解析方法').strip()
    editors = data.get('editors') or []
    viewers = data.get('viewers') or []

    if not kb_name:
        return jsonify({"status": "error", "message": "知识库名称不能为空"})
    if ' ' in kb_name:
        return jsonify({"status": "error", "message": "知识库名称不能包含空格"})

    kb_root = os.path.join(PROJECT_ROOT, "../rag_database")
    os.makedirs(kb_root, exist_ok=True)
    kb_path = os.path.join(kb_root, kb_name)
    meta_file = os.path.join(kb_path, "metadata.json")
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    if os.path.exists(kb_path):
        return jsonify({"status": "error", "message": "知识库已存在，请更换名称"})

    try:
        os.makedirs(kb_path, exist_ok=True)
        metadata = {
            "name": kb_name,
            "type": kb_type,
            "description": description,
            "security_level": security_level,
            "language": language,
            "parse_method": parse_method,
            "owner": str(current_user.get_id()),
            "editors": [str(x) for x in editors],
            "viewers": [str(x) for x in viewers],
            "create_time": now,
            "update_time": now,
            "doc_count": 0,
            "source_file": ""
        }
        with open(meta_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "知识库创建成功", "kb": metadata})
    except Exception as e:
        return jsonify({"status": "error", "message": f"创建失败: {e}"})


@bp.route("/api/kb/update_permission", methods=["POST"])
@login_required
@kb_perm_required(need="owner")         # 只有创建者和管理员能改权限
def kb_update_permission(kb_meta=None):
    data = request.get_json()
    kb_name = data.get("name")
    editors = data.get("editors", [])   # 期望是数组
    viewers = data.get("viewers", [])

    kb_meta["editors"] = [str(x) for x in editors]
    kb_meta["viewers"] = [str(x) for x in viewers]
    kb_meta["update_time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    kb_root = current_app.config["KB_ROOT"]
    meta_file = os.path.join(kb_root, kb_name, "metadata.json")
    try:
        with open(meta_file, "w", encoding="utf-8") as f:
            json.dump(kb_meta, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "权限更新成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": f"权限更新失败：{e}"})


@bp.route('/api/list-kbs', methods=['GET'])
def list_kbs():
    kb_root = os.path.join(PROJECT_ROOT, "../rag_database")

    kbs = []
    if os.path.exists(kb_root):
        for kb_name in os.listdir(kb_root):
            kb_path = os.path.join(kb_root, kb_name)
            if not os.path.isdir(kb_path):
                continue

            kb_info = {
                "name": kb_name,
                "type": "other",
                "create_time": "",
                "doc_count": 0,
                "description": "",
                "security_level": "内部"
            }

            # 读取 metadata 决定类型
            meta_file = os.path.join(kb_path, "metadata.json")
            if os.path.exists(meta_file):
                try:
                    with open(meta_file, 'r', encoding='utf-8') as f:
                        meta = json.load(f)
                    kb_info.update(meta)
                except Exception as e:
                    print(f"[KB] 读取元数据失败 {kb_name}: {e}")

            if not kb_info["create_time"]:
                mtime = os.path.getmtime(kb_path)
                kb_info["create_time"] = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S')

            # ===== 新增：权限过滤开始 =====
            user_id = str(current_user.get_id())
            is_admin = current_user.role == 'admin'

            editors = [str(x) for x in kb_info.get("editors", [])]
            viewers = [str(x) for x in kb_info.get("viewers", [])]
            is_owner = is_admin or str(kb_info.get("owner")) == user_id
            # print(f'name:{user_id}, is_admin:{is_admin}, is_owner:{is_owner}, editors:{editors}, viewers:{viewers}')
            if not is_owner:
                # editors/viewers 都为空 = 公开库，所有人可见
                if editors or viewers:
                    if user_id not in editors and user_id not in viewers:
                        continue   # 无使用权限，跳过不返回

            # 顺带给前端返回操作权限，控制按钮显隐
            kb_info["is_owner"] = is_owner
            kb_info["can_edit"] = is_owner or (not editors and not viewers) or (user_id in editors)
            # ===== 新增：权限过滤结束 =====

            kbs.append(kb_info)

    kbs.sort(key=lambda x: x["create_time"], reverse=True)
    return jsonify({"status": "success", "kbs": kbs})


@bp.route('/api/kb/delete', methods=['POST'])
@kb_perm_required(need="owner")
def delete_kb():
    """删除知识库"""
    data = request.json
    kb_name = data.get('name')
    kb_type = data.get('type')

    if not kb_name or not kb_type:
        return jsonify({"status": "error", "message": "参数缺失"})

    cache_key = f"{kb_type}|{kb_name}"
    if hasattr(rag_engine, 'collections') and cache_key in rag_engine.collections:
        del rag_engine.collections[cache_key]
        print(f"[KB] 已清理 RAGEngine 内存缓存: {cache_key}")

    try:
        import chromadb.api.client
        import gc
        chromadb.api.client.SharedSystemClient.clear_system_cache()
        gc.collect()
        print("[KB] 已释放 ChromaDB 底层 SQLite 文件句柄")
    except Exception as e:
        print(f"[KB] 释放 ChromaDB 资源失败: {e}")

    kb_path = os.path.join(PROJECT_ROOT, "../rag_database", kb_name)

    if os.path.exists(kb_path):
        try:
            shutil.rmtree(kb_path)
            return jsonify({"status": "success", "message": "知识库已彻底删除"})
        except Exception as e:
            return jsonify({"status": "error", "message": f"删除文件夹失败: {e}"})
    else:
        return jsonify({"status": "error", "message": "知识库不存在"})


@bp.route('/api/kb/rename', methods=['POST'])
@kb_perm_required(need="owner")
def rename_kb():
    """重命名知识库"""
    data = request.json
    old_name = data.get('oldName')
    new_name = data.get('newName')

    if not old_name or not new_name:
        return jsonify({"status": "error", "message": "参数缺失"})

    base_path = os.path.join(PROJECT_ROOT, "../rag_database")
    old_path = os.path.join(base_path, old_name)
    new_path = os.path.join(base_path, new_name)

    if not os.path.exists(old_path):
        return jsonify({"status": "error", "message": "原知识库不存在"})

    if os.path.exists(new_path):
        return jsonify({"status": "error", "message": "新名称已存在"})

    try:
        os.rename(old_path, new_path)

        # Update metadata.json
        meta_file = os.path.join(new_path, "metadata.json")
        if os.path.exists(meta_file):
            with open(meta_file, 'r', encoding='utf-8') as f:
                meta = json.load(f)
            meta['name'] = new_name
            with open(meta_file, 'w', encoding='utf-8') as f:
                json.dump(meta, f, ensure_ascii=False, indent=2)

        return jsonify({"status": "success", "message": "重命名成功"})
    except Exception as e:
        return jsonify({"status": "error", "message": f"重命名失败: {e}"})


@bp.route('/api/kb/items', methods=['GET'])
def get_kb_items():
    """获取知识库条目"""
    kb_name = request.args.get('name')
    kb_type = request.args.get('type')
    limit = request.args.get('limit', 100, type=int)

    if not kb_name or not kb_type:
        return jsonify({"status": "error", "message": "参数缺失"})

    result = rag_engine.get_all_items(kb_type, kb_name, limit)
    return jsonify(result)


@bp.route('/api/kb/item/delete', methods=['POST'])
@kb_perm_required(need="edit")
def delete_kb_item():
    """删除知识库条目"""
    data = request.json
    kb_name = data.get('kbName')
    kb_type = data.get('kbType')
    item_id = data.get('itemId')

    if not kb_name or not kb_type or not item_id:
        return jsonify({"status": "error", "message": "参数缺失"})

    result = rag_engine.delete_item(kb_type, kb_name, item_id)

    if result.get('status') == 'success':
        # Update metadata count
        try:
            kb_path = os.path.join(PROJECT_ROOT, "../rag_database", kb_name)
            meta_file = os.path.join(kb_path, "metadata.json")

            if os.path.exists(meta_file):
                with open(meta_file, 'r', encoding='utf-8') as f:
                    meta = json.load(f)
                meta['doc_count'] = result.get('remaining', 0)
                with open(meta_file, 'w', encoding='utf-8') as f:
                    json.dump(meta, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[KB] 更新元数据失败: {e}")

    return jsonify(result)


@bp.route('/preview', methods=['POST'])
@kb_perm_required(need="edit")
def preview_file():
    doc_type = request.form.get('doc_type')
    use_server_file = request.form.get('use_server_file') == 'true'

    # 解析逻辑抽成函数
    def do_parse(target_path):
        # 调用解析逻辑
        preview_data = []

        if doc_type == 'rule':
            if target_path.endswith('.docx'):
                # 尝试规则解析
                rules = parse_programming_rules(target_path)
                if not rules:
                    # 兜底
                    text = read_docx_text(target_path)
                    rules = smart_parse_doc(text, type='rule')

                # 格式化为前端预览
                preview_data = rules

        elif doc_type == 'issue':
            if target_path.endswith('.docx'):
                issues = parse_issue_reports(target_path)
                if not issues:
                    text = read_docx_text(target_path)
                    issues = smart_parse_doc(text, type='issue')
                preview_data = issues

        elif doc_type == 'history_align':
            # 历史对齐通常是 JSON
            if target_path.endswith('.json'):
                with open(target_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # 简单提取一些预览信息
                    if "annotations" in data:
                        preview_data = [{"id": a.get("id"), "content": "历史对齐数据"} for a in data["annotations"][:10]]

        return preview_data

    try:
        if use_server_file:
            # 模式 A: 使用服务器上的 testdata 文件
            filename = request.form.get('filename')
            if not filename:
                return jsonify({"status": "error", "message": "未指定文件名"})

            target_path = os.path.join(TESTDATA_DIR, filename)
            if not os.path.exists(target_path):
                return jsonify({"status": "error", "message": "文件不存在"})

            # 调用解析，单文件直接返回
            preview_data = do_parse(target_path)
            return jsonify({"status": "success", "data": preview_data})

        else:
            # 模式 B: 上传文件
            # files = request.files.getlist('file')
            files = preprocess_files(request.files.getlist('file'))
            if not files or len(files) == 0:
                return jsonify({"status": "error", "message": "未上传文件"})

            # 保存到临时目录
            temp_dir = os.path.join(PROJECT_ROOT, 'temp_uploads')
            os.makedirs(temp_dir, exist_ok=True)

            all_results = []
            for idx, file in enumerate(files):
                if not file or file.filename == '':
                    continue

                filename = secure_filename(file.filename)
                target_path = os.path.join(temp_dir, filename)
                file.save(target_path)

                preview_data = do_parse(target_path)
                all_results.append({
                    "filename": file.filename,
                    "preview_data": preview_data
                })

            return jsonify({
                "status": "success",
                "data": all_results,
                "count": len(all_results)
            })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": str(e)})


def find_available_port(start_port):
    port = start_port
    while True:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            if s.connect_ex(('0.0.0.0', port)) != 0:
                return port
            port += 1


"""导出需求-代码匹配和审查结果"""

# ========== 新增：SQLite数据库操作函数 ==========
def get_alignments_from_sqlite(project_id):
    """从project.db的alignments表读取指定列数据"""
    try:
        # 连接SQLite数据库
        # conn = sqlite3.connect(db_path)
        
        conn = get_db()
        #conn = get_db_celery()
        
        # 读取name、docRanges、codeRanges三列所有数据
        query_sql = f"SELECT name, docRanges, codeRanges, GenReq, GenMermaid FROM alignments " \
                    f"where project_id={project_id} " \
                    f"AND (is_alignment = 1 or (is_code_review = 1 and isReviewed = 1))"
        # 直接用pandas读取SQL结果（简洁高效）
        # df = pd.read_sql(query_sql, conn)
        # 创建引擎
        db_url = f"mysql+pymysql://{DB_CONFIG['user']}:{DB_CONFIG['password']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"

        engine = create_engine(
            db_url,
            echo=False,
            pool_pre_ping=True
        )
        df = pd.read_sql(query_sql, engine)
        df.index = df.index + 1
        
        
        
        # conn.close()
        return df
    except sqlite3.Error as e:
        print(f"SQLite数据库读取失败：{e}")
        return pd.DataFrame()  # 返回空DataFrame
    except Exception as e:
        print(f"读取数据异常：{e}")
        return pd.DataFrame()


# ========== 导出结果：从SQLite读取数据，生成word文件 ==========
@bp.route('/project/export', methods=['GET'])
def export_project_results():
    """导出需求-代码匹配结果（从SQLite读取数据）"""
    # 1. 获取项目路径（用于定位project.db）
    project_path = request.args.get('path')
    project_id = request.args.get('project_id')
    secret_level = request.args.get('secret_level')
    if not project_path or not os.path.isdir(project_path):
        return jsonify({"status": "error", "message": "无效的项目路径。"}), 400

    # 2. 定位project.db文件（默认在项目路径根目录）
    # db_file = os.path.join(project_path, 'project.db')
    # if not os.path.exists(db_file):
        # return jsonify({"status": "error", "message": f"未找到数据库文件：{db_file}"}), 400
     
    # 3. 从SQLite读取数据
    df = get_alignments_from_sqlite(project_id)
    

    #logger.info(df.loc[2, "GenMermaid"])
    #sys.exit()

    if df.empty:
        return jsonify({"status": "warning", "message": "alignments表中暂无数据可导出"}), 200

    # 4、处理数据
    # 遍历行索引
    total_num = 0
    for idx in df.index:
        doc_data = df.loc[idx, "docRanges"]
        code_data = df.loc[idx, "codeRanges"]
        doc_data = json.loads(doc_data) #从string转成list
        code_data = json.loads(code_data)
        total_num += 1
        temp = []
        for doc in doc_data:
            temp.append(doc['content'])
        df.loc[idx, "docRanges"] = temp

        temp = []
        for code in code_data:
            temp.append(code['content'])
        df.loc[idx, "codeRanges"] = temp


    try:
        # 5. 生成并写入word文件
        template_path = os.path.join(os.path.dirname(__file__), '../templates', '需求表格.docx')
        # 创建临时目录存储文件
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir, exist_ok=True)

        # 生成docx文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"需求表格导出_{timestamp}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)

        # 检查是否提供了DOCX模板路径
        if template_path and os.path.exists(template_path):
            current_date = datetime.now().strftime("%Y%m%d")
            merged_doc = Document(template_path)

            # 处理第一个测试项表格作为基础文档
            merged_doc = Document(template_path)
            export_prompt = SECRET_LEVEL_MAP[secret_level]
            p = merged_doc.paragraphs[0].insert_paragraph_before(export_prompt) if merged_doc.paragraphs \
                else merged_doc.add_paragraph(export_prompt)
            p.alignment = 0
            run = p.runs[0]
            run.font.size = Pt(10)
            # run.font.color.rgb = RGBColor(128, 128, 128)
            number = 1 # 测试项编号
            category_id = 1 #章节号

            replacements = {}
            # 小标题名称
            replacements["6.6.1 XXXX功能"] = "6.6." + str(number) + " " + df.loc[number, "name"] + "功能"
            # 测试项名称
            replacements["AAAAA功能"] = df.loc[number, "name"]
            # 测试项标识
            replacements["T_FUNC"] = "T_FUNC" + str(number)
            # 追踪关系
            # replacements["BBBBB"] = "需求说明：" + str(category_id)
            pattern = r'^(\d+(?:\.\d+)*)\s*' # 匹配以数字开头，由数字和点组成，后接可选空格的字符串
            match = re.match(pattern, df.loc[number, "name"])
            if match:
                replacements["BBBBB"] = match.group(1)  # 返回编号，如 "3.2.1"
            else:
                replacements["BBBBB"] = "需求说明：" + str(category_id)
            
            # 需求描述
            replacements["CCCCC"] = "\n\n".join(["".join(sub_list) for sub_list in df.loc[number, "docRanges"]])
            # 生成需求
            replacements["DDDDD"] = df.loc[number, "GenReq"] if df.loc[number, "GenReq"] is not None else ""
            # 对齐代码
            replacements["EEEEE"] = "\n\n".join(["".join(sub_list) for sub_list in df.loc[number, "codeRanges"]])
            replacements["EEEEE"] = replacements["EEEEE"] if replacements["EEEEE"] is not None else ""
            # 流程图
            replacements["FFFFF"] = df.loc[number, "GenMermaid"] if df.loc[number, "GenMermaid"] is not None else ""

            # 替换第一个文档的占位符
            replace_text_in_docx(merged_doc, replacements, 'result')
            #logger.info(replacements)
            # 处理剩余的表单
            for i in range(1, total_num):
                # 添加分页符
                # merged_doc.add_page_break()
                if number == 1:
                    number += 1
                    category_id += 1

                # 为每个表单加载新的模板并填充
                temp_doc = Document(template_path)

                replacements = {}
                # 小标题名称
                replacements["6.6.1 XXXX功能"] = "6.6." + str(number) + " " + df.loc[number, "name"] + "功能"
                # 测试项名称
                replacements["AAAAA功能"] = df.loc[number, "name"]
                # 测试项标识
                replacements["T_FUNC"] = "T_FUNC" + str(number)
                # 追踪关系
                # replacements["BBBBB"] = "需求说明：" + str(category_id)
                pattern = r'^(\d+(?:\.\d+)*)\s*' # 匹配以数字开头，由数字和点组成，后接可选空格的字符串
                match = re.match(pattern, df.loc[number, "name"])
                if match:
                    replacements["BBBBB"] = match.group(1)  # 返回编号，如 "3.2.1"
                else:
                    replacements["BBBBB"] = "需求说明：" + str(category_id)
                # 需求描述
                replacements["CCCCC"] = "\n\n".join(["".join(sub_list) for sub_list in df.loc[number, "docRanges"]])
                # 生成需求
                replacements["DDDDD"] = df.loc[number, "GenReq"] if df.loc[number, "GenReq"] is not None else ""
                # 对齐代码
                replacements["EEEEE"] = "\n\n".join(["".join(sub_list) for sub_list in df.loc[number, "codeRanges"]])
                replacements["EEEEE"] = replacements["EEEEE"] if replacements["EEEEE"] is not None else ""
                # 流程图
                replacements["FFFFF"] = df.loc[number, "GenMermaid"] if df.loc[number, "GenMermaid"] is not None else ""

                number += 1
                category_id += 1

                # 替换模板中的占位符
                replace_text_in_docx(temp_doc, replacements, 'result')


                # 直接拼接填充好的页面内容到合并文档
                for element in temp_doc.element.body:
                    merged_doc.element.body.append(element)
                #logger.info(i)
                #logger.info(df.loc[number, "name"])

            # 输出文档内容，debug用
            #for paragraph in temp_doc.paragraphs:
            #    logger.info(paragraph.text)

            #for table in temp_doc.tables:
            #    for row in table.rows:
            #        for cell in row.cells:
            #            for paragraph in cell.paragraphs:
            #                logger.info(paragraph.text)

            # 保存合并后的文档
            #merged_doc.save(docx_path)
            try:
                merged_doc.save(docx_path)
                logger.info(f"文档可以被导出：{docx_path}")
                
            except Exception as e:
                # 打印详细错误（方便排查）
                logger.info(traceback.format_exc())
                logger.info(f"导出结果失败：{str(e)}")
                


            return send_file(
                docx_path,
                as_attachment=False,  # 配合前端自定义保存路径
                download_name=docx_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        else:
            # 使用excel格式导出（备用方案）
            # 重命名列（可选，让Excel列名更友好）
            df_renamed = df.rename(columns={
                'name': '需求块名称',
                'docRanges': '需求文档范围',
                'codeRanges': '代码范围'
            })
            # 生成并写入文件
            output = BytesIO()
            writer = pd.ExcelWriter(output, engine='openpyxl')
            df_renamed.to_excel(writer, sheet_name='对齐结果', index=False)
            writer.close()
            output.seek(0)  # 关键：重置文件指针

            # 返回Excel文件流
            return send_file(
                output,
                as_attachment=False,  # 配合前端自定义保存路径
                download_name=f'对齐结果_${pd.Timestamp.now().strftime("%Y%m%d%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )


    except Exception as e:
        logger.info(traceback.format_exc())
        logger.info(f"导出结果失败：{e}")
        return jsonify({"status": "error", "message": f"导出失败: {str(e)}"}), 500

# ========== 将生成的临时word文件删除 ==========
@bp.route('/project/delete-export-files', methods=['POST'])
def delete_export_project_results():
    try:        
        data = request.json
        filename = data.get('filename', '')
        # 删除文件（从文件系统中删除）
        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_exports')
        docx_path = os.path.join(temp_dir, filename)
        if os.path.exists(docx_path):
            os.remove(docx_path)
            #print(f'临时文件{filename}已删除')
            return jsonify({"status": "success", "message": f"已删除临时文件"}), 200    
            
        else:
            return jsonify({"status": "success", "message": f"临时文件不存在"}), 200  
        
    except Exception as e:
        logger.info(traceback.format_exc())
        logger.info(f"删除临时文件失败：{e}")
        return jsonify({"status": "error", "message": f"删除临时文件失败: {str(e)}"}), 500    
        
# 提示词设置对话框

# 保存用户自定义提示词的文件
from .prompt import ALIGN_PROMPT_TEMPLATE, ALIGN_REQ_PROMPT_TEMPLATE, REVIEW_PROMPT_TEMPLATE, \
    REVIEW_PROMPT_TEMPLATE_KBS, GENERATE_PROMPT_TEMPLATE, \
    ALIGN_PROMPT_TEMPLATE_ICL, RULE_EXTRACTION_PROMPT, ISSUE_EXTRACTION_PROMPT, \
    ABSTRACT_PROMPT_TEMPLATE, TOTAL_ABSTRACT_PROMPT_TEMPLATE, CODEFILE_PROMPT_TEMPLATE, TAB_MAP, DEFAULTS, \
    ALIGN_REQ_PROMPT_TEMPLATE_KBS, ALIGN_PROMPT_TEMPLATE_KBS, CODE_THINKING_PROMPT_TEMPLATE, \
    CODE_THINKING_PROMPT_TEMPLATE_KBS

# def load_user_prompt():
    # if os.path.exists(PROMPT_FILE):
        # with open(PROMPT_FILE, 'r') as f:
            # data = json.load(f)
            # return data.get('prompt', DEFAULT_PROMPT)
    # return DEFAULT_PROMPT

# @bp.route('/get_align_prompt')
# def get_align_prompt():
    # #prompt = load_user_prompt()
    # prompt = ALIGN_REQ_PROMPT_TEMPLATE
    # return jsonify({'prompt': prompt})

# 设置提示词相关

# 存储提示词（内存中，可替换为数据库或文件）
prompts = {
    'req-code-align': 'You are an expert in aligning AI responses to user intent...',
    'code-req-align': 'You are an expert in aligning AI responses to user intent...',
    'review': 'You are an expert in reviewing AI responses...'
}

# 默认提示词
default_prompts = {
    'req-code-align': ALIGN_PROMPT_TEMPLATE,
    'code-req-align': ALIGN_REQ_PROMPT_TEMPLATE,
    'review': REVIEW_PROMPT_TEMPLATE,
    'review-code': CODE_THINKING_PROMPT_TEMPLATE,
    'req-code-align-kbs': ALIGN_PROMPT_TEMPLATE_KBS,
    'code-req-align-kbs': ALIGN_REQ_PROMPT_TEMPLATE_KBS,
    'review-kbs': REVIEW_PROMPT_TEMPLATE_KBS,
    'review-code-kbs': CODE_THINKING_PROMPT_TEMPLATE_KBS
}

# 从文件加载
# def load_prompts():
#     global prompts
#     try:
#         with open('prompts.json', 'r') as f:
#             prompts = json.load(f)
#     except:
#         pass

# 保存提示词到文件
# def save_prompts():
#     with open('prompts.json', 'w') as f:
#         json.dump(prompts, f)


# 加载提示词（用于 /get_prompts）
def _get_prompt_table_columns(db):
    c = db.cursor()
    c.execute("SHOW COLUMNS FROM prompt")
    rows = c.fetchall() or []
    return {row.get('Field') for row in rows if row and row.get('Field')}


# 加载提示词（用于 /get_prompts）
@login_required
@bp.route('/get_prompts', methods=['GET'])
def get_prompts():
    db = get_db()
    c = db.cursor()
    existing_columns = _get_prompt_table_columns(db)

    requested_fields = ['Req2CodeAlign', 'Code2ReqAlign', 'review', 'reviewCode', 'Req2CodeAlignKbs', 'Code2ReqAlignKbs', 'reviewKbs', 'reviewCodeKbs']
    available_fields = [f for f in requested_fields if f in existing_columns]
    row = None
    if available_fields:
        sql = f"select {', '.join(available_fields)} from prompt where user_id=%s"
        c.execute(sql, (current_user.user_id,))
        row = c.fetchone()

    result = {
        'Req2CodeAlign': default_prompts['req-code-align'],
        'Code2ReqAlign': default_prompts['code-req-align'],
        'review': default_prompts['review'],
        'reviewCode': default_prompts['review-code'],
        'Req2CodeAlignKbs': default_prompts['req-code-align-kbs'],
        'Code2ReqAlignKbs': default_prompts['code-req-align-kbs'],
        'reviewKbs': default_prompts['review-kbs'],
        'reviewCodeKbs': default_prompts['review-code-kbs'],
    }
    if row:
        for k in requested_fields:
            if k in row and row.get(k):
                result[k] = row.get(k)

    # 兼容旧表结构（没有 KBS 字段时，前端也能拿到可展示值）
    if 'Req2CodeAlignKbs' not in existing_columns:
        result['Req2CodeAlignKbs'] = result['Req2CodeAlign']
    if 'Code2ReqAlignKbs' not in existing_columns:
        result['Code2ReqAlignKbs'] = result['Code2ReqAlign']
    if 'reviewKbs' not in existing_columns:
        result['reviewKbs'] = result['review']

    return jsonify(result)

# 保存提示词（根据 tab 和 content）
@login_required
@bp.route('/save_prompt', methods=['POST'])
def save_prompt():
    data = request.get_json()
    outer_tab = data.get('outerTab')
    tab = data.get('tab')
    content = data.get('content')

    field = TAB_MAP.get((outer_tab, tab))

    if not field:
        return jsonify({'success': False, 'message': 'Invalid tab'})

    db = get_db()
    c = db.cursor()
    existing_columns = _get_prompt_table_columns(db)

    if 'user_id' not in existing_columns:
        return jsonify({'success': False, 'message': 'prompt表缺少user_id字段'}), 500

    fallback_field_map = {
        'Req2CodeAlignKbs': 'Req2CodeAlign',
        'Code2ReqAlignKbs': 'Code2ReqAlign',
        'reviewKbs': 'review',
        'reviewCodeKbs': 'reviewCode'
    }
    effective_field = field
    if effective_field not in existing_columns:
        effective_field = fallback_field_map.get(field, field)

    if effective_field not in existing_columns:
        return jsonify({'success': False, 'message': f'prompt表缺少字段: {field}'}), 500

    supported_prompt_columns = [col for col in DEFAULTS.keys() if col in existing_columns]
    data = {'user_id': current_user.user_id}
    for col in supported_prompt_columns:
        data[col] = DEFAULTS[col]
    data[effective_field] = content

    columns = ['user_id'] + supported_prompt_columns
    placeholders = ['%s'] * len(columns)
    sql = f"""insert into prompt ({', '.join(columns)}) values ({', '.join(placeholders)})
            on duplicate key update {effective_field} = values({effective_field})"""
    params = [data[c] for c in columns]

    c.execute(sql, params)
    return jsonify({'success': True, 'message': 'prompt saved'})


# 恢复默认（根据 tab）
@bp.route('/restore_default', methods=['POST'])
def restore_default():
    data = request.get_json()
    tab = data.get('tab')

    if tab not in ['req-code-align', 'code-req-align', 'review', 'review-code',
                   'req-code-align-kbs', 'code-req-align-kbs', 'review-kbs', 'review-code-kbs']:
        return jsonify({'success': False, 'message': 'Invalid tab'})

    # 恢复默认
    prompts[tab] = default_prompts[tab]

    return jsonify({'default_prompt': prompts[tab]})

# 对齐提示词相关
@bp.route('/get_default_align_prompt')
def get_default_align_prompt():
    return jsonify({'default_prompt': ALIGN_REQ_PROMPT_TEMPLATE})

@bp.route('/save_align_prompt', methods=['POST'])
def save_align_prompt():
    data = request.get_json()
    prompt = data.get('prompt')
    if not prompt:
        return jsonify({'success': False, 'error': '提示词不能为空'})
    save_user_prompt(prompt)
    return jsonify({'success': True})

# 审查提示词相关
@bp.route('/get_default_review_prompt')
def get_default_review_prompt():
    return jsonify({'default_prompt': REVIEW_PROMPT_TEMPLATE})

@bp.route('/save_review_prompt', methods=['POST'])
def save_review_prompt():
    data = request.get_json()
    prompt = data.get('prompt')
    if not prompt:
        return jsonify({'success': False, 'error': '提示词不能为空'})
    save_user_prompt(prompt)
    return jsonify({'success': True})            

    
# 从数据库获取当前用户的所有项目
@bp.route('/welcome/get_user_projects', methods=['GET'])
def get_user_projects():
    user_id = current_user.user_id
    db = get_db_celery()
    c = db.cursor()
    c.execute(f'select name from project where user_id={user_id};')
    rows = c.fetchall()
    projects = [row['name'] for row in rows]
    db.close()
    paths = [os.path.join(TESTDATA_DIR, project) for project in projects]
    #print(projects)
    return jsonify({
            'status': 'success',
            'name': projects,
            'path': paths
        })
    
def contains_chinese(text):
    # 匹配中文字符的正则模式（包括常见汉字）
    # \u4e00-\u9fff 是基本汉字范围
    pattern = r'[\u4e00-\u9fff]'
    if text is None:
        return False
    if not isinstance(text, str):
        try:
            text = str(text)
        except:
            return False
    # 查找所有中文字符
    chinese_chars = re.findall(pattern, text)
    
    # 防止有乱码，如果中文字符数量 >= 10，则返回 True，否则 False
    return len(chinese_chars) >= 10

# 过滤代码摘要，输出代码摘要有乱码的文件名、无乱码的摘要字典、无乱码的摘要数量
def filter_non_abstract_files(file_abstract):
    filter_non_file = []
    filter_file_abstract = {}
    cnt = 0
    for key, value in file_abstract.items():
        if not isinstance(value, str):
            # print('************************************')
            # print(value)
            # print(type(value))
            # print('======================================')
            continue
        if not contains_chinese(value): 
            filter_non_file.append(key)
        else:
            filter_file_abstract[key] = value
            cnt += 1
            
    return filter_non_file, filter_file_abstract, cnt
    
# if __name__ == '__main__':
#     start_port = 5056
#     available_port = find_available_port(start_port)
#     app.run(host='0.0.0.0', port=available_port, debug=True)
