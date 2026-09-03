
"""
docx 案例集解析
依赖：pip install python-docx olefile
（olefile 用于把 Visio OLE 对象还原成 .vsd/.vsdx，不装也能跑，只是存成 .bin）
"""
import os
import re
import io
import traceback
import uuid
import olefile
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.document import Document as _Document

# ============ 可配置项 ============
CASE_HEADING_LEVEL = 3          # 案例标题级别（导航里的第3级）
SECTION_HEADING_LEVEL = 4       # 小节标题级别（1 概述 / 2 机理分析 ...）
DISCARD_SECTION_KEYWORD = '启示'  # 小节标题包含该词 → 整节丢弃
NS_OFFICE = 'urn:schemas-microsoft-com:office:office'


# ============ 基础工具 ============
def iter_block_items(parent):
    """按文档实际顺序遍历段落和表格"""
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def get_heading_level(para):
    """返回标题级别数字，非标题返回 None。兼容 'Heading 3' 和 '标题 3'"""
    name = (para.style.name or '').strip()
    m = re.match(r'^(?:heading|标题)\s*(\d+)$', name, re.IGNORECASE)
    return int(m.group(1)) if m else None


def extract_textbox_text(para):
    """提取段落中嵌入的文本框文字（python-docx 默认读不到文本框）"""
    lines = []
    for txbx in para._p.findall('.//' + qn('w:txbxContent')):
        for p in txbx.findall(qn('w:p')):
            line = ''.join(t.text or '' for t in p.findall('.//' + qn('w:t')))
            lines.append(line)
    return '\n'.join(lines).strip()


def table_to_text(table):
    """表格转纯文本（正文里的表格才会用到）"""
    return '\n'.join(
        ' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells)
        for row in table.rows
    )


def extract_from_ole10native(raw):
    """从 Ole10Native 流里抠原始文件： vsdx 是 zip(PK 头), 老 vsd 是 OLE 复合文件 (DOCF 头)"""
    idx = raw.find(b'PK\x03\x04')
    if idx != -1:
        return raw[idx:], '.vsdx'
    idx = raw.find(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')
    if idx != -1:
        return raw[idx:], '.vsd'
    return raw, '.bin'


def save_ole_blob(blob, save_dir, name_prefix):
    """保存 OLE 对象，Visio 对象尽量还原成 .vsd/.vsdx，否则存 .bin"""
    os.makedirs(save_dir, exist_ok=True)
    data, ext = blob, '.bin'
    try:

        bio = io.BytesIO(blob)
        if olefile.isOleFile(bio):
            ole = olefile.OleFileIO(bio)
            streams = ['/'.join(s) for s in ole.listdir()]
            # print('OLE streams:', streams)

            # 嵌入的 Visio 数据，就是完整 vsd
            if any('VisioDocument' in s or 'Visio Document' in s for s in streams):
                # 嵌入的 Visio 对象: 整个OLE 容器本身就是 vsd 格式, 直接存
                data, ext = blob, '.vsd'
            elif 'Package' in streams:           # 少数情况内嵌完整包
                pkg = ole.openstream('Package').read()
                data, ext = pkg, ('.vsdx' if pkg[:2] == b'PK' else '.bin')
            else:
                # Ole10Native 包装: 抠出里面的原始文件
                native = next((s for s in streams if s.endswith('Ole10Native')), None)
                if native:
                    data, ext = extract_from_ole10native(ole.openstream(native).read())
            ole.close()
    except ImportError:
        traceback.print_exc()
        pass  # 没装 olefile，按 .bin 原样保存
    except Exception as e:
        print(f'OLE 对象解析失败，按 .bin 保存: {e}')

    fname = f'{name_prefix}{ext}'
    with open(os.path.join(save_dir, fname), 'wb') as f:
        f.write(data)
    return fname


def extract_ole_objects(para, doc_part, save_dir, doc_tag, counter):
    """提取段落中的 OLE 嵌入对象（Visio），返回 (文件名列表, 新counter)"""
    saved = []
    for ole in para._p.findall('.//{%s}OLEObject' % NS_OFFICE):
        rid = ole.get(qn('r:id'))
        if not rid:
            continue
        try:
            blob = doc_part.rels[rid].target_part.blob
        except KeyError:
            continue
        progid = ole.get('ProgID', '')
        kind = 'visio' if 'visio' in progid.lower() else 'ole'
        fname = save_ole_blob(blob, save_dir, f'{doc_tag}_{kind}_{counter}')
        saved.append(fname)
        counter += 1
    return saved, counter


# ============ 主解析函数 ============
def parse_typical_cases(source, visio_dir='visio_files'):
    """
    解析案例集 docx
    source: 文件路径 或 文件流（Flask 上传直接传 file.stream）
    visio_dir: Visio 对象保存目录
    返回: [{'title': 案例标题, 'content': 内容, 'visio_files': [...]}, ...]
    """
    doc = Document(source)
    doc_tag = uuid.uuid4().hex[:6]   # 防止批量解析时 Visio 文件重名覆盖
    cases = []
    current = None              # 当前案例
    started = False             # 是否已进入第一个案例（之前的内容全部跳过：目录/封面/篇章标题）
    before_keywords = False     # 是否在 案例标题 ~ 关键词 之间
    skip_section = False        # 是否在"启示"小节里
    current_section = ''        # 当前小节标题文本
    ole_counter = 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block)

            # --- 案例标题：开启新案例 ---
            if level == CASE_HEADING_LEVEL:
                if current:
                    cases.append(current)
                current = {'title': block.text.strip(), 'parts': [], 'visio_files': []}
                started, before_keywords = True, True
                skip_section, current_section = False, ''
                continue

            if not started:
                continue  # 目录、封面、篇章大标题全部跳过

            # --- 小节标题 ---
            if level == SECTION_HEADING_LEVEL:
                sec = block.text.strip()
                before_keywords = False
                current_section = sec
                if DISCARD_SECTION_KEYWORD in sec:
                    skip_section = True        # 启示小节：整节丢弃
                    continue
                skip_section = False
                current['parts'].append(sec)   # 其他小节标题保留进 content
                continue

            # --- 其他级别标题（上篇/第一部分等）忽略 ---
            if level is not None:
                continue

            if skip_section:
                continue

            # --- 普通段落 ---
            text = block.text.strip()
            if text:
                if '【关键词】' in text:
                    before_keywords = False
                current['parts'].append(text)   # 简介、关键词、正文、图注都走这里

            # 文本框（代码段）：按所在小节命名为 违背示例 / 遵循示例
            tb_text = extract_textbox_text(block)
            if tb_text:
                if '机理分析' in current_section:
                    current['parts'].append('违背示例')
                elif '纠正措施' in current_section:
                    current['parts'].append('遵循示例')
                # 其他小节的文本框：不加标题，直接放内容
                current['parts'].append(tb_text)

            # Visio OLE 对象：保存文件 + content 留占位
            saved, ole_counter = extract_ole_objects(block, doc.part, visio_dir, doc_tag, ole_counter)
            for fname in saved:
                current['parts'].append(f'[Visio对象: {fname}]')
                current['visio_files'].append(fname)

        elif isinstance(block, Table):
            if not started or skip_section:
                continue
            if before_keywords:
                continue                        # 标题和关键词之间的表格不要
            current['parts'].append(table_to_text(block))

    if current:
        cases.append(current)

    for c in cases:
        c['content'] = '\n'.join(c.pop('parts'))
    return cases


def probe(file_path):
    doc = Document(file_path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            print(f'{p.style.name:<15} | {t[:40]}')


if __name__ == '__main__':

    # probe(r'C:\Users\Administrator\Desktop\知识库文档\案例\2024年示例(公开).docx')
    cases = parse_cases(r'C:\Users\Administrator\Desktop\知识库文档\案例\2024年示例(公开).docx', visio_dir='visio_files')
    # cases = parse_cases(r'C:\Users\Administrator\Desktop\知识库文档\案例\2025年示例(公开).docx', visio_dir='visio_files')
    print(type(cases), cases)
    print(f'共解析出 {len(cases)} 个案例')
    for i, c in enumerate(cases, 1):
        print(f'===== 案例{i}: {c["title"]} =====')
        print(c['content'])
        print('visio_files:', c['visio_files'])
        print()
