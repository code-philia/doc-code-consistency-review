
"""
docx 案例集解析
依赖：pip install python-docx olefile
（olefile 用于把 Visio OLE 对象还原成 .vsd/.vsdx，不装也能跑，只是存成 .bin）
"""
import os
import re
import io
import traceback
import uuid
import olefile
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.document import Document as _Document

# ============ 可配置项 ============
CASE_HEADING_LEVEL = 3          # 案例标题级别（导航里的第3级）
SECTION_HEADING_LEVEL = 4       # 小节标题级别（1 概述 / 2 机理分析 ...）
DISCARD_SECTION_KEYWORD = '启示'  # 小节标题包含该词 → 整节丢弃
NS_OFFICE = 'urn:schemas-microsoft-com:office:office'

# 备用逻辑关键词定义
CASE_KEYWORDS = ['案例', '问题', '事件', '事故', '典型', '实例']
SECTION_KEYWORDS = ['概述', '机理分析', '纠正措施', '启示和建议']

# ============ 基础工具 ============
def iter_block_items(parent):
    """按文档实际顺序遍历段落和表格"""
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)



def get_heading_level(para, doc=None):
    """
    返回标题级别数字，非标题返回 None。
    逻辑：
    1. 优先匹配样式名称 (Heading 3, 标题 3 等)。
    2. 如果样式名不匹配，检查段落的大纲级别 (Outline Level)。
       - 大纲级别 0 = 正文 (通常不视为标题，除非用户特意设为标题)
       - 大纲级别 1-9 = 对应 Level 1-9
       - 大纲级别 10+ = 正文 (None)
    """
    # 1. 优先检查样式名称
    name = (para.style.name or '').strip()
    m = re.match(r'^(?:heading|标题)\s*(\d+)$', name, re.IGNORECASE)
    if m and m is not None:     
        return int(m.group(1))
    
    # 2. 如果样式名不匹配，检查大纲级别
    if doc is None:
        return None

    # 获取样式对象
    style = para.style
    # 检查样式是否有定义 (有些内置样式可能没有独立的 style element)
    if style is None or style._element is None:
        return None

    # 解析样式定义的 XML
    style = para.style
    if not style or not style._element:
        return None

    style_elm = style._element
    style_name = (para.style.name or '').strip()
    # 打印调试信息：查看样式定义的 XML 结构
    #from lxml import etree
    #xml_str = etree.tostring(style_elm, encoding='utf-8').decode('utf-8')
    #print(f"样式 '{style_name}' 的 XML:\n{xml_str[:500]}...\n")

    # 在样式定义的 XML 中搜索 outlineLevel
    # 遍历所有子节点
    for elem in style_elm.iter():
        if elem.tag == qn('w:outlineLevel'):
            val_str = elem.get(qn('w:val'))
            if val_str:
                try:
                    level = int(val_str)
                    # 验证级别范围
                    if 1 <= level <= 9:
                        return level
                except ValueError:
                    pass

    # --- 策略 4: 兜底逻辑 ---
    if style_name == "RC标题":
        # 假设 RC标题 是一级标题
        return 1
    
    return None


def extract_textbox_text(para):
    """提取段落中嵌入的文本框文字（python-docx 默认读不到文本框）"""
    lines = []
    for txbx in para._p.findall('.//' + qn('w:txbxContent')):
        for p in txbx.findall(qn('w:p')):
            line = ''.join(t.text or '' for t in p.findall('.//' + qn('w:t')))
            lines.append(line)
    return '\n'.join(lines).strip()


def table_to_text(table):
    """表格转纯文本（正文里的表格才会用到）"""
    return '\n'.join(
        ' | '.join(cell.text.strip().replace('\n', ' ') for cell in row.cells)
        for row in table.rows
    )


def extract_from_ole10native(raw):
    """从 Ole10Native 流里抠原始文件： vsdx 是 zip(PK 头), 老 vsd 是 OLE 复合文件 (DOCF 头)"""
    idx = raw.find(b'PK\x03\x04')
    if idx != -1:
        return raw[idx:], '.vsdx'
    idx = raw.find(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1')
    if idx != -1:
        return raw[idx:], '.vsd'
    return raw, '.bin'


def save_ole_blob(blob, save_dir, name_prefix):
    """保存 OLE 对象，Visio 对象尽量还原成 .vsd/.vsdx，否则存 .bin"""
    os.makedirs(save_dir, exist_ok=True)
    data, ext = blob, '.bin'
    try:

        bio = io.BytesIO(blob)
        if olefile.isOleFile(bio):
            ole = olefile.OleFileIO(bio)
            streams = ['/'.join(s) for s in ole.listdir()]
            # print('OLE streams:', streams)

            # 嵌入的 Visio 数据，就是完整 vsd
            if any('VisioDocument' in s or 'Visio Document' in s for s in streams):
                # 嵌入的 Visio 对象: 整个OLE 容器本身就是 vsd 格式, 直接存
                data, ext = blob, '.vsd'
            elif 'Package' in streams:           # 少数情况内嵌完整包
                pkg = ole.openstream('Package').read()
                data, ext = pkg, ('.vsdx' if pkg[:2] == b'PK' else '.bin')
            else:
                # Ole10Native 包装: 抠出里面的原始文件
                native = next((s for s in streams if s.endswith('Ole10Native')), None)
                if native:
                    data, ext = extract_from_ole10native(ole.openstream(native).read())
            ole.close()
    except ImportError:
        traceback.print_exc()
        pass  # 没装 olefile，按 .bin 原样保存
    except Exception as e:
        print(f'OLE 对象解析失败，按 .bin 保存: {e}')

    fname = f'{name_prefix}{ext}'
    with open(os.path.join(save_dir, fname), 'wb') as f:
        f.write(data)
    return fname


def extract_ole_objects(para, doc_part, save_dir, doc_tag, counter):
    """提取段落中的 OLE 嵌入对象（Visio），返回 (文件名列表, 新counter)"""
    saved = []
    for ole in para._p.findall('.//{%s}OLEObject' % NS_OFFICE):
        rid = ole.get(qn('r:id'))
        if not rid:
            continue
        try:
            blob = doc_part.rels[rid].target_part.blob
        except KeyError:
            continue
        progid = ole.get('ProgID', '')
        kind = 'visio' if 'visio' in progid.lower() else 'ole'
        fname = save_ole_blob(blob, save_dir, f'{doc_tag}_{kind}_{counter}')
        saved.append(fname)
        counter += 1
    return saved, counter


# ============ 主解析函数 ============
def parse_typical_cases(source, visio_dir='visio_files'):
    """
    解析案例集 docx
    source: 文件路径 或 文件流（Flask 上传直接传 file.stream）
    visio_dir: Visio 对象保存目录
    返回: [{'title': 案例标题, 'content': 内容, 'visio_files': [...]}, ...]
    """
    doc = Document(source)
    doc_tag = uuid.uuid4().hex[:6]   # 防止批量解析时 Visio 文件重名覆盖
    cases = []
    current = None              # 当前案例
    started = False             # 是否已进入第一个案例（之前的内容全部跳过：目录/封面/篇章标题）
    before_keywords = False     # 是否在 案例标题 ~ 关键词 之间
    skip_section = False        # 是否在"启示"小节里
    current_section = ''        # 当前小节标题文本
    ole_counter = 1

    for block in iter_block_items(doc):

        if isinstance(block, Paragraph):
            level = get_heading_level(block)

            # --- 案例标题：开启新案例 ---
            if level == CASE_HEADING_LEVEL:
                if current:
                    cases.append(current)
                current = {'title': block.text.strip(), 'parts': [], 'visio_files': []}
                started, before_keywords = True, True
                skip_section, current_section = False, ''
                continue

            if not started:
                continue  # 目录、封面、篇章大标题全部跳过

            # --- 小节标题 ---
            if level == SECTION_HEADING_LEVEL:
                sec = block.text.strip()
                before_keywords = False
                current_section = sec
                if DISCARD_SECTION_KEYWORD in sec:
                    skip_section = True        # 启示小节：整节丢弃
                    continue
                skip_section = False
                current['parts'].append(sec)   # 其他小节标题保留进 content
                continue

            # --- 其他级别标题（上篇/第一部分等）忽略 ---
            if level is not None:
                continue

            if skip_section:
                continue

            # --- 普通段落 ---
            text = block.text.strip()
            if text:
                if '【关键词】' in text:
                    before_keywords = False
                current['parts'].append(text)   # 简介、关键词、正文、图注都走这里

            # 文本框（代码段）：按所在小节命名为 违背示例 / 遵循示例
            tb_text = extract_textbox_text(block)
            if tb_text:
                if '机理分析' in current_section:
                    current['parts'].append('违背示例')
                elif '纠正措施' in current_section:
                    current['parts'].append('遵循示例')
                # 其他小节的文本框：不加标题，直接放内容
                current['parts'].append(tb_text)

            # Visio OLE 对象：保存文件 + content 留占位
            saved, ole_counter = extract_ole_objects(block, doc.part, visio_dir, doc_tag, ole_counter)
            for fname in saved:
                current['parts'].append(f'[Visio对象: {fname}]')
                current['visio_files'].append(fname)

        elif isinstance(block, Table):
            if not started or skip_section:
                continue
            if before_keywords:
                continue                        # 标题和关键词之间的表格不要
            current['parts'].append(table_to_text(block))

    if current:
        cases.append(current)

    for c in cases:
        c['content'] = '\n'.join(c.pop('parts'))
    return cases


def probe(file_path):
    doc = Document(file_path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            print(f'{p.style.name:<15} | {t[:40]}')


# ============ 新解析函数：parse_typical_cases_new ============
def parse_typical_cases_new(source, visio_dir='visio_files_new'):
    """
    专门处理备用逻辑文档的解析函数
    逻辑：
    1. 先扫描文档，确认是否存在标准 Level 3 标题。
       - 如果存在：直接返回空列表（说明这是标准文档，应由 parse_typical_cases 处理）。
       - 如果不存在：继续判断是 B1 还是 B2 模式。
    2. 场景 B1 (全一级): Level 1 标题中，靠关键词区分案例和小节。
    3. 场景 B2 (一级案例+二级小节): Level 1 是案例，Level 2 是小节。
    """
    doc = Document(source)
    doc_tag = uuid.uuid4().hex[:6]
    
    # --- 第一步：探测文档结构 ---
    has_l3 = False
    has_l1 = False
    has_l2 = False

    # 快速扫描，只检查标题级别
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block, doc)
            if level == 3: has_l3 = True
            elif level == 1: has_l1 = True
            elif level == 2: has_l2 = True
    
    # 确定解析模式
    mode = None
    case_level = None
    section_level = None
    use_keywords = False

    if has_l1 and not has_l2:
        # 场景 B1: 只有 Level 1
        mode = 'B1_AllLevel1'
        case_level = 1
        section_level = 1
        use_keywords = True
    elif has_l1 and has_l2:
        # 场景 B2: Level 1 和 Level 2 共存
        mode = 'B2_L1Case_L2Section'
        case_level = 1
        section_level = 2
        use_keywords = False
    else:
        print(f"[{os.path.basename(source)}] 无法识别备用逻辑结构 (L1:{has_l1}, L2:{has_l2})")
        return []

    print(f"[{os.path.basename(source)}] 启用备用逻辑解析 -> 解析模式: {mode}")

    # --- 第二步：正式解析 ---
    cases = []
    current = None
    started = False
    before_keywords = False
    skip_section = False
    current_section = ''
    ole_counter = 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block, doc)
            text = block.text.strip()
            print(text,level)
            # --- 标题处理逻辑 ---
            if level is not None:
                is_case = False
                is_section = False
                
                if use_keywords and level == case_level:
                    # 场景 B1: 所有标题都是 L1，靠关键词区分
                    text_type = classify_text_type(text)
                    if text_type == 'case':
                        is_case = True
                    elif text_type == 'section':
                        is_section = True
                    
                else:
                    # 场景 B2: 直接按层级判断
                    if level == case_level:
                        is_case = True
                    elif level == section_level:
                        is_section = True
                
                print(text_type, is_case, is_section)
                # 处理案例标题
                if is_case:
                    print(current)
                    if current:
                        cases.append(current)
                    current = {'title': text, 'parts': [], 'visio_files': []}
                    started, before_keywords = True, True
                    skip_section, current_section = False, ''
                    continue
                
                # 处理小节标题
                if is_section:
                    if not started:
                        continue # 案例还没开始，忽略小节
                    before_keywords = False
                    current_section = text
                    if DISCARD_SECTION_KEYWORD in text:
                        skip_section = True
                        continue
                    skip_section = False
                    current['parts'].append(text)
                    continue
            
            # --- 非标题段落处理 ---
            if not started:
                continue
            
            if skip_section:
                continue

            # 普通段落
            if text:
                if '【关键词】' in text:
                    before_keywords = False
                if current:
                    current['parts'].append(text)

            # 文本框
            tb_text = extract_textbox_text(block)
            if tb_text and current:
                if '机理分析' in current_section:
                    current['parts'].append('违背示例')
                elif '纠正措施' in current_section:
                    current['parts'].append('遵循示例')
                current['parts'].append(tb_text)

            # Visio OLE
            saved, ole_counter = extract_ole_objects(block, doc.part, visio_dir, doc_tag, ole_counter)
            for fname in saved:
                if current:
                    current['parts'].append(f'[Visio对象: {fname}]')
                    current['visio_files'].append(fname)

        elif isinstance(block, Table):
            if not started or skip_section:
                continue
            if before_keywords:
                continue
            if current:
                current['parts'].append(table_to_text(block))

    if current:
        cases.append(current)

    for c in cases:
        c['content'] = '\n'.join(c.pop('parts'))
    
    return cases
            
 
# ============ 结构探测函数 (核心修改) ============
def detect_document_structure(doc):
    """
    扫描文档，判断标题层级结构，返回解析策略配置
    策略定义：
    1. 'standard': 存在 Level 3 (案例) 和 Level 4 (小节)
    2. 'level1_all': 最高级是 Level 1，且没有 Level 2/3。靠关键词区分案例和小节。
    3. 'level1_l2': 存在 Level 1 (案例) 和 Level 2 (小节)。
    """
    levels_found = set()
    has_l1 = False
    has_l2 = False
    has_l3 = False
    
    # 第一次扫描：统计存在的标准标题级别
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block)
            if level:
                levels_found.add(level)
                if level == 1: has_l1 = True
                if level == 2: has_l2 = True
                if level == 3: has_l3 = True
    
    # 策略判断
    if has_l3:
        # 标准模式：Level 3 = 案例，Level 4 = 小节
        return 'standard'
    
    if has_l1 and not has_l2 and not has_l3:
        # 场景 B1：所有标题都是 Level 1
        # 需要靠关键词区分：包含“案例/问题”的是案例，包含“概述/机理”的是小节
        return 'level1_all'
    
    if has_l1 and has_l2 and not has_l3:
        # 场景 B2：Level 1 = 案例，Level 2 = 小节
        return 'level1_l2'
    
    # 兜底：如果既没有 L3，也没有 L1，或者结构混乱
    return 'unknown'

# ============ 辅助函数：判断文本类型 (用于场景 B1) ============
def classify_text_type(text):
    """
    根据关键词判断文本是案例标题还是小节标题
    """
    text_lower = text.lower()
    
    # 优先判断小节 
    for kw in SECTION_KEYWORDS:
        if kw in text:
            return 'section'

    # 默认视为案例标题
    return 'case'
    
 
if __name__ == '__main__':
    # 测试
    file_path = r'命名方式战术组-问题类型-案例名称-姓名（公开）.docx'
    
    if os.path.exists(file_path):
        # 先探测结构
        doc = Document(file_path)
        mode = detect_document_structure(doc)
        print(f"=== 检测到文档结构模式: {mode} ===")
        
        if  mode== 'standard':
            cases = parse_typical_cases(file_path, visio_dir='visio_files')
        else:
            cases = parse_typical_cases_new(file_path, visio_dir='visio_files')
            
        print(f'共解析出 {len(cases)} 个案例')
        for i, c in enumerate(cases, 1):
            print(f'===== 案例{i}: {c["title"]} =====')
            print(c['content'][:300] + '...')
            print('visio_files:', c['visio_files'])
            print("-" * 30)
    else:
        print(f"文件未找到：{file_path}")